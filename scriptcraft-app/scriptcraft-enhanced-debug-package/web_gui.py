#!/usr/bin/env python3
"""
ScriptCraft Web GUI - Clean Console Capture Version
Fixed implementation that displays real-time EnhancedAutoGenSystem messages in browser
"""

# CRITICAL: Add repo root to path FIRST (before any other imports)
import re
import os
import json
import time
import queue
import uuid
import asyncio
import threading
import logging
from datetime import datetime
from io import BytesIO
from flask import Flask, render_template, request, jsonify, Response, make_response, send_file
import sys
from pathlib import Path

# Add paths for imports - prioritize local scriptcraft-app directory
current_dir = Path(__file__).parent
sys.path.insert(0, str(current_dir))  # Prioritize local directory FIRST
print(f"✅ Prioritized local dir in sys.path: {current_dir}")

REPO_ROOT = Path(__file__).parent.parent
if str(REPO_ROOT) not in sys.path:
    # Add parent as backup (position 1, not 0)
    sys.path.insert(1, str(REPO_ROOT))
    print(f"✅ Added to sys.path: {REPO_ROOT}")


VERSION = "15.34-quotes-progress-mapping"

# Set Google API key at module level to ensure it's available
# This ensures the API key is set before any thumbnail generation
if "GOOGLE_API_KEY" not in os.environ:
    os.environ["GOOGLE_API_KEY"] = "AIzaSyAiFFlgDokz-s4U8UrV73Fhdnl8Ukx2jCM"
    print("🔑 Set GOOGLE_API_KEY environment variable from default")
else:
    print(f"🔑 Using GOOGLE_API_KEY from environment")


# Set up logging
logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Paths already added above - no need to duplicate

# Import text processing functions

app = Flask(__name__)
progress_streams = {}
results = {}
running_tasks = {}


class ConsoleCapture:
    """Captures console output and redirects to progress streamer"""

    def __init__(self, streamer):
        self.streamer = streamer
        self.original_stdout = sys.stdout
        self.buffer = []
        self.long_operation_start = None
        self.last_chapter_revision = None

    def write(self, message):
        """Write to both original stdout and the progress streamer"""
        if not message.strip():
            return

        # Debug: Add marker to identify when write method is called
        if "TEST 5" in message or "Minimal rapid" in message:
            self.original_stdout.write(f"🔍 WRITE DEBUG: {message.strip()}\n")
            self.original_stdout.flush()

        # Write to original stdout first
        self.original_stdout.write(message)
        self.original_stdout.flush()

        # Filter Azure SDK noise that can cause blocking
        if self._should_filter_azure_sdk_logging(message):
            return

        # Extract progress and send to web interface
        progress = self._extract_progress(message)
        if progress is not None and self.streamer:
            # Only send updates when we have explicit progress values
            try:
                self.streamer.send_update(message.strip(), progress)
            except Exception as e:
                self.original_stdout.write(
                    f"⚠️ Stream update failed: {e}\n")
                self.original_stdout.flush()
        # Don't send messages without progress - they would default to 50%

    def _track_long_operations(self, message):
        """Track when long operations start for timing purposes"""
        import time

        # Track when chapter revision starts
        if "Reviewing Chapter" in message and "revision failed" not in message:
            chapter_match = re.search(r'Chapter (\d+)', message)
            if chapter_match:
                chapter_num = int(chapter_match.group(1))
                self.last_chapter_revision = {
                    'chapter': chapter_num,
                    'start_time': time.time()
                }

        # Track other long operations
        elif any(phrase in message for phrase in [
            "Script Reviewer", "Script review", "Reviewing", "Revising"
        ]) and "failed" not in message and "completed" not in message:
            if self.long_operation_start is None:
                self.long_operation_start = time.time()

    def _enhance_message_with_timing(self, primary_text, fallback_text):
        """Add timing information to progress messages"""
        import time

        msg = primary_text if len(primary_text) > 10 else fallback_text

        # Add elapsed time for chapter revisions
        if self.last_chapter_revision and "revised" in msg:
            elapsed = time.time() - self.last_chapter_revision['start_time']
            if elapsed > 30:  # Only show timing after 30 seconds
                chapter_num = self.last_chapter_revision['chapter']
                mins, secs = divmod(int(elapsed), 60)
                if mins > 0:
                    msg = f"{msg} (Chapter {chapter_num} took {mins}m {secs}s)"
                else:
                    msg = f"{msg} (Chapter {chapter_num} took {secs}s)"

        # Add elapsed time for other long operations
        elif self.long_operation_start and any(phrase in msg for phrase in [
            "reviewing", "revising", "processing"
        ]):
            elapsed = time.time() - self.long_operation_start
            if elapsed > 60:  # Show timing after 1 minute
                mins, secs = divmod(int(elapsed), 60)
                msg = f"{msg} (running {mins}m {secs}s)"

        return msg

    def _should_filter_initialization_message(self, message):
        """Filter out initialization messages not relevant to script work"""
        initialization_patterns = [
            "TournamentAgentClient initialized",
            "TournamentAgent model client created",
            "AITipsAgentClient initialized",
            "AITipsAgent model client created",
            "📡 AUTOGEN SYSTEM TRACE: Initializing agents",
            "🤖 AUTOGEN SYSTEM TRACE: Creating model clients",
            "🎯 AUTOGEN SYSTEM TRACE: System fully initialized"
        ]

        return any(pattern in message for pattern in initialization_patterns)

    def _should_filter_azure_sdk_logging(self, message):
        """Filter out Azure AI SDK internal logging that can block operations"""
        # Only filter very specific Azure SDK internal messages that cause blocking
        blocking_patterns = [
            "INFO - ManagedIdentityCredential will use IMDS",
            "INFO - Request URL:",
            "Request method:",
            "Request headers:",
            "No body was attached to the request",
            "A body is sent with the request",
            "INFO - Response status:",
            "Response headers:",
            "INFO - DefaultAzureCredential acquired a token",
            "User-Agent': 'azsdk-python-identity",
            "User-Agent': 'AIProjectClient azsdk-python-ai-agents",
            "Authorization': 'REDACTED'",
            "x-ms-client-request-id':",
            "Content-Type': 'application/json'",
            "openai-processing-ms':",
            "Date': 'Fri, 26 Sep 2025",
            "No environment configuration found."
        ]

        # Keep important workflow messages even if they contain Azure terms
        important_workflow_patterns = [
            "STEP 1:", "STEP 2:", "STEP 3:", "STEP 4:",
            "Topic enhanced", "Chapter", "Script", "Writing", "Reviewing", "Revising",
            "completed", "revised", "Assembling", "WORKFLOW", "DEBUG:"
        ]

        # Don't filter if it's an important workflow message
        if any(pattern in message for pattern in important_workflow_patterns):
            return False

        return any(pattern in message for pattern in blocking_patterns)

    def _extract_progress(self, message):
        """
        Extract progress based on actual workflow messages.

        Progress Distribution (matches actual workflow timeline):
        - 0-10%: Initialization (system startup)
        - 10-20%: Topic Enhancement (STEP 1)
        - 20-30%: Start Chapter Writing (STEP 2 begins)
        - 30-55%: Chapter Writing Progress (based on completed chapters)
        - 55-65%: Script Review (STEP 3)
        - 65-80%: Chapter-by-Chapter Review Progress
        - 80-83%: YouTube Metadata
        - 83-86%: B-roll Search Terms
        - 86-89%: Demo Packages
        - 89-90%: B-roll Terms (second generation)
        - 90-100%: Thumbnail Generation & Completion
        """
        import re

        # ========== INITIALIZATION (0-10%) ==========
        if "Starting script creation" in message:
            return 1
        elif "Initializing script creation system" in message:
            return 3
        elif "Console capture test" in message:
            return 5
        elif "Creating EnhancedAutoGenSystem" in message:
            return 6
        elif "Initializing Enhanced AutoGen" in message:
            return 7
        elif "All 6 agents initialized successfully" in message:
            return 10

        # ========== STEP 1: TOPIC ENHANCEMENT (10-20%) ==========
        elif "STEP 1: Topic Enhancement" in message or "Topic Enhancement & Chapter Planning" in message:
            return 12
        elif "Running Script-Topic-Assistant-Agent" in message:
            return 14
        elif "Topic enhanced" in message and "chars" in message:
            return 20

        # Chapter extraction debug and breakdown display (20-21%)
        elif "DEBUG: Found" in message and "chapter matches" in message:
            return 20  # Show debug info right after topic enhanced
        elif "Raw Match" in message or "Skipped:" in message or "Trimming from" in message:
            return 20  # Show all debug messages
        elif "Chapter Planning Complete [21%]" in message:
            return 21
        elif "📋 CHAPTER BREAKDOWN:" in message:
            return 21
        elif "Chapter" in message and ":" in message and "/" not in message and "completed" not in message:
            # Matches "Chapter 1: Title" but not "Chapter 1/5" or "Chapter 1 completed"
            return 21

        # ========== STEP 2: CHAPTER WRITING (20-55%) ==========
        elif "STEP 2: Chapter-by-Chapter Script Writing" in message:
            return 22
        elif "Fetching style references" in message:
            return 24
        elif "Retrieved transcript style references" in message:
            return 26

        # Filter out "Writing Chapter X/Y" - don't update progress when starting
        elif "Writing Chapter" in message and "/" in message:
            return None  # Don't change progress while starting to write

        # Track individual chapter completion ONLY when done (26% to 55%)
        elif "Chapter" in message and "completed" in message and "chars" in message:
            match = re.search(r'Chapter (\d+)/(\d+)', message)
            if match:
                completed_num, total = int(match.group(1)), int(match.group(2))
                # Progress from 26% to 55%
                chapter_progress = int(26 + (completed_num * 29 / total))
                return min(chapter_progress, 55)
            return 40

        elif "Combined all" in message and "chapters into full script" in message:
            return 55

        # ========== STEP 3: SCRIPT REVIEW (55-80%) ==========
        elif "STEP 3: Script Review" in message:
            return 57
        elif "Running Script-Review-Agent" in message and "Script-Review-Agent" in message:
            return 59
        elif "Script reviewed" in message and "chars" in message:
            return 62
        elif "Retrying with chapter-by-chapter revision" in message:
            return 64

        # Filter out "Reviewing Chapter X/Y" - don't update progress when starting
        elif "Reviewing Chapter" in message and "/" in message:
            return None  # Don't change progress while starting review

        # Reviewer debug output (show what reviewer returned)
        elif "🔍 DEBUG - Reviewer Response Preview" in message:
            return 65  # Show during review phase
        elif "Contains '" in message and "': " in message:
            return 65  # Show debug checks

        # Track chapter-by-chapter review progress ONLY when complete (65% to 78%)
        elif "Chapter" in message and "revised" in message and "chars" in message:
            match = re.search(r'Chapter (\d+)/(\d+)', message)
            if match:
                completed_num, total = int(match.group(1)), int(match.group(2))
                # Progress from 65% to 78%
                review_progress = int(65 + (completed_num * 13 / total))
                return min(review_progress, 78)
            return 70

        elif "Assembling" in message and "revised chapters" in message:
            return 79
        elif "All" in message and "chapters revised individually" in message:
            return 80

        # Review feedback summary
        elif "📋 REVIEW FEEDBACK SUMMARY" in message:
            return 80
        elif "Total chapters reviewed:" in message:
            return 80
        elif "Full review feedback saved to:" in message:
            return 80

        # ========== STEP 3.5: QUOTES & STATISTICS (80-81%) ==========
        elif "STEP 3.5: Quotes & Statistics Generation" in message:
            return 80
        elif "Quotes & Statistics generated" in message and "chars" in message:
            return 81
        elif "Quotes/Stats generation failed" in message:
            return 81
        elif "Exception in Quotes/Stats generation" in message:
            return 81
        elif "Continuing workflow without quotes/stats" in message:
            return 81
        elif "Quotes & Statistics section inserted" in message:
            return 81

        # ========== STEP 4: HOOK & SUMMARY (81-83%) ==========
        elif "STEP 4: Hook & Summary Generation" in message:
            return 81
        elif "Hook generated" in message and "chars" in message:
            return 82
        elif "Summary generated" in message and "chars" in message:
            return 82
        elif "Hook prepended to script" in message:
            return 82
        elif "Summary appended to script" in message:
            return 83

        # ========== POST-PROCESSING (83-90%) ==========
        elif "SEQUENTIAL WORKFLOW COMPLETED" in message:
            return 83
        elif "Script enhanced with formatting" in message:
            return 83

        # YouTube Metadata (83-85%)
        elif "Generating YouTube upload metadata" in message:
            return 83
        elif "YouTube metadata generated" in message:
            return 85

        # B-roll Search Terms (85-87%)
        elif "Generating B-roll search terms" in message:
            return 85
        elif "B-roll table generated" in message:
            return 87

        # Demo Packages (87-90%)
        elif "Creating demo packages" in message:
            return 88
        elif "Demo packages generated" in message:
            return 90

        # ========== THUMBNAIL GENERATION (90-100%) ==========
        elif "Generating" in message and "thumbnail" in message:
            return 92
        elif "Generated" in message and "thumbnails" in message:
            return 98
        elif "Script creation completed" in message:
            return 100

        return None  # Don't change progress for unrecognized messages

    def flush(self):
        self.original_stdout.flush()

    def fileno(self):
        """Return file descriptor for compatibility"""
        return self.original_stdout.fileno()

    def isatty(self):
        """Return whether this is a terminal"""
        return self.original_stdout.isatty()

    def readable(self):
        """Return whether file is readable"""
        return False

    def writable(self):
        """Return whether file is writable"""
        return True

    def seekable(self):
        """Return whether file is seekable"""
        return False

    def __getattr__(self, name):
        """Delegate unknown attributes to original stdout"""
        return getattr(self.original_stdout, name)


class ProgressStreamer:
    """Handles Server-Sent Events progress streaming"""

    def __init__(self, session_id: str):
        self.session_id = session_id
        self.queue = queue.Queue()
        self.result = None
        self.done = False

    def send_update(self, message: str, progress: int = None):
        """Send a progress update"""
        # Don't send updates with no progress - they would reset to default
        if progress is None:
            return

        update = {
            "message": message,
            "progress": progress,
            "timestamp": time.time(),
            "done": progress == 100
        }

        try:
            logger.info(
                f"🔍 send_update: putting message in queue for session {self.session_id}")
            self.queue.put(json.dumps(update))
            logger.info(f"🔍 send_update: message queued successfully")
            if progress == 100:
                logger.info(
                    f"🔍 send_update: marking done=True for session {self.session_id}")
                self.done = True
            logger.info(f"🔍 send_update: method completed successfully")
        except Exception as e:
            logger.error(f"Failed to send update: {e}")


async def process_script_creation(session_id, topic, audience, tone,
                                  video_length, production_type, goals,
                                  quick_test=False, checkboxes=None):
    """Clean script creation with only console capture"""
    logger.info(f"🎬 SCRIPT CREATION STARTED: session={session_id}")
    if quick_test:
        logger.info("⚡ QUICK TEST: 1-chapter mode enabled")

    # Default to script only if no checkboxes provided
    if checkboxes is None:
        checkboxes = {"script": True}

    try:
        if session_id not in progress_streams:
            logger.error(f"❌ Session {session_id} not found!")
            return

        streamer = progress_streams[session_id]

        # Initial progress
        streamer.send_update("🚀 Initializing script creation system...", 5)

        # Import the system
        from linedrive_azure.agents.enhanced_autogen_system import (
            EnhancedAutoGenSystem
        )

        # Set up console capture
        console_capture = ConsoleCapture(streamer)
        original_stdout = sys.stdout

        # Note: Removed background progress updater thread as it was
        # interfering with message-based progress tracking

        try:
            # Redirect stdout to capture console output
            sys.stdout = console_capture

            # Send a test message to verify streaming works
            streamer.send_update(
                "🧪 Console capture test - starting script creation", 5)

            # Initialize system (this will be captured)
            print("🔧 Creating EnhancedAutoGenSystem...")
            system = EnhancedAutoGenSystem(verbose=True)
            print("🔧 All 6 agents initialized successfully")

            # Run the workflow with timeout to prevent hanging
            print("🚀 Starting script workflow with 20-minute timeout...")

            # DEBUG: Add detailed logging for hang investigation
            print("🔍 DEBUG: About to call async workflow method")
            print(f"📋 Checkboxes received: {checkboxes}")
            print(f"🎯 quick_test: {quick_test}, topic: {topic[:50]}...")
            print(
                f"🔍 DEBUG: Current thread: {threading.current_thread().name}")
            print(
                f"🔍 DEBUG: ConsoleCapture active: {isinstance(sys.stdout, ConsoleCapture)}")

            try:
                print("🔍 DEBUG: Calling asyncio.wait_for with async method")
                result = await asyncio.wait_for(
                    system.run_complete_script_workflow_sequential(
                        script_topic=topic,
                        topic_description="",
                        audience=audience,
                        tone=tone,
                        script_length=video_length,
                        max_chapters=1 if quick_test else 8,
                    ),
                    # 20 minutes - accounts for Script Writer (~5min) + Script Review (~12min)
                    timeout=1200
                )
                print("🔍 DEBUG: Async workflow completed successfully")

            except asyncio.TimeoutError:
                print("🔍 DEBUG: AsyncIO timeout occurred")
                print("⚠️ Workflow timeout - using fallback approach")
                streamer.send_update(
                    "⚠️ Script creation taking longer than expected, "
                    "trying alternative approach...", 85
                )

                # Fallback: Create a simpler script without revision
                result = {
                    "success": False,
                    "error": "Workflow timeout - chapter revision taking too long",
                    "timeout": True
                }
            except asyncio.TimeoutError:
                print("⚠️ Workflow timeout - using fallback approach")
                streamer.send_update(
                    "⚠️ Script creation taking longer than expected, "
                    "trying alternative approach...", 85
                )

                # Fallback: Create a simpler script without revision
                result = {
                    "success": False,
                    "error": "Workflow timeout - chapter revision taking too long",
                    "timeout": True
                }

        finally:
            # Always restore stdout
            sys.stdout = original_stdout

        # Debug: Check what we actually got
        print(f"🔍 DEBUG: Result type: {type(result)}")
        if isinstance(result, dict):
            print(f"🔍 DEBUG: Result keys: {list(result.keys())}")
        else:
            print(f"🔍 DEBUG: Result value: {result}")
            # Convert string result to error format
            result = {"success": False, "error": str(result)}

        if result.get("success"):
            # EXACT COPY OF WORKING CONSOLE UI LOGIC
            print(f"\n🎉 COMPLETE 4-AGENT WORKFLOW FINISHED!")
            print("=" * 60)
            print(f"📋 Topic: {topic}")
            print(f"👥 Audience: {audience}")
            print(f"💬 Tone: conversational and educational")
            print(f"⏱️ Length: {video_length}")
            print("-" * 60)
            print(f"✅ Topic Enhanced by Topic Assistant")
            print(f"✅ Script Created by Script Writer")
            print(f"✅ Script Reviewed by Script Reviewer")
            print(f"✅ Sequential Workflow Completed Successfully")

            # Add introduction to script content (EXACT COPY)
            introduction = """Hi, I'm Roz's AI Digital Twin. She is a high-tech sales leader, bonafide tech nerd and busy Mom of two really incredible kids. She has spent her career in high-tech working to this moment and beyond and we are here to guide you through it. This is, AI with Roz.

---

"""

            # Combine introduction with script content (EXACT COPY)
            raw_script_content = introduction + result["script_content"]

            # Enhanced script with bold tool formatting (EXACT COPY - NO TIMEOUTS!)
            from console_ui.text_processing import (
                enhance_script_with_bold_tools,
                extract_tool_links_and_info,
            )

            enhanced_script_content = enhance_script_with_bold_tools(
                raw_script_content)
            print("✅ Script enhanced with bold tool formatting")
            # Force progress update for enhancement step
            streamer.send_update("✅ Script enhanced with formatting", 98)

            # Use enhanced script as final content (EXACT COPY)
            final_script_content = enhanced_script_content

            # Extract tool links for YouTube description (EXACT COPY - NO TIMEOUTS!)
            tool_links = extract_tool_links_and_info(final_script_content)
            tool_count = len(tool_links.splitlines())
            print(f"📊 Extracted {tool_count} tool references")
            # Force progress update for tool extraction
            streamer.send_update(
                f"📊 Extracted {tool_count} tool references", 99)

            # Add tool links to the script content for YouTube description (EXACT COPY)
            final_script_with_tools = final_script_content + "\n\n" + "=" * 60 + "\n"
            final_script_with_tools += "� YOUTUBE VIDEO DESCRIPTION\n"
            final_script_with_tools += "=" * 60 + "\n"
            final_script_with_tools += (
                "Copy the section below for your YouTube video description:\n\n"
            )
            final_script_with_tools += f"🎥 {topic}\n\n"
            final_script_with_tools += tool_links
            final_script_with_tools += "\n\n🔔 Don't forget to SUBSCRIBE for more AI tools and productivity tips!"
            final_script_with_tools += "\n� What tools would you like to see featured next? Drop a comment below!"

            # Store result (SIMPLIFIED)
            # Initialize variables that will be used later
            youtube_upload_details = None
            curl_commands = None  # Initialize curl commands variable
            edl_content = None  # Initialize EDL content variable
            edl_filename = None  # Initialize EDL filename variable

            # NEW: Generate B-roll Table using AI Agent (skip unless checkbox checked)
            print(
                f"🔍 DEBUG: Checking broll - checkbox: {checkboxes.get('broll', False)}, quick_test: {quick_test}")
            if not quick_test and checkboxes.get("broll", False):
                print("✅ DEBUG: B-roll condition TRUE - calling agent")
                try:
                    print("\n📊 Generating B-roll Search Terms Table...")
                    streamer.send_update(
                        "📊 Generating B-roll search terms...", 97.5)

                    try:
                        from linedrive_azure.agents import ScriptBRollAgentClient
                    except ImportError as import_err:
                        print(f"⚠️ B-roll agent not available: {import_err}")
                        raise

                    broll_agent = ScriptBRollAgentClient()
                    print("✅ B-roll agent initialized")

                    # Generate B-roll table WITH timecodes for EDL export
                    broll_result = broll_agent.generate_broll_table_with_timecodes(
                        script_content=final_script_content,
                        script_title=topic,
                        words_per_minute=150,  # Adjust based on speaking pace
                        timeout=180
                    )

                    if broll_result.get("success", False):
                        broll_table = broll_result.get("table", "")
                        parsed_data = broll_result.get("parsed_data", [])

                        print(
                            f"✅ B-roll table generated ({len(broll_table)} characters, {len(parsed_data)} entries)")
                        streamer.send_update(
                            "✅ B-roll table generated", 98.5)

                        # Append B-roll table to script
                        broll_section = f"\n\n{'=' * 80}\n"
                        broll_section += "# 🎬 B-ROLL SEARCH TERMS TABLE WITH TIMECODES\n"
                        broll_section += f"{'=' * 80}\n\n"
                        broll_section += broll_table

                        final_script_with_tools += broll_section
                        print("✅ B-roll table appended to script")

                        # Generate EDL file for DaVinci Resolve
                        edl_content = None
                        edl_filename = None
                        if parsed_data:
                            try:
                                import os
                                import time

                                # Create filename with timestamp
                                timestamp = time.strftime("%Y%m%d_%H%M%S")
                                edl_filename = f"broll_markers_{timestamp}.edl"

                                edl_result = broll_agent.create_edl_markers(
                                    broll_data=parsed_data,
                                    output_file=edl_filename,
                                    frame_rate='24'
                                )

                                if edl_result.get("success"):
                                    print(
                                        f"✅ EDL file created: {edl_filename} ({edl_result.get('marker_count')} markers)")
                                    streamer.send_update(
                                        f"✅ EDL file created: {edl_filename}", 98.7)

                                    # Read EDL content for frontend display
                                    try:
                                        with open(edl_filename, 'r', encoding='utf-8') as f:
                                            edl_content = f.read()
                                        print(
                                            f"✅ EDL content read ({len(edl_content)} bytes)")
                                    except Exception as read_error:
                                        print(
                                            f"⚠️ Could not read EDL file: {read_error}")

                                    # Add EDL info to script
                                    edl_section = f"\n\n{'=' * 80}\n"
                                    edl_section += "# 📋 DAVINCI RESOLVE EDL MARKERS\n"
                                    edl_section += f"{'=' * 80}\n\n"
                                    edl_section += f"**EDL File Generated:** `{edl_filename}`\n\n"
                                    edl_section += f"**Total Markers:** {edl_result.get('marker_count')}\n\n"
                                    edl_section += "**Import Instructions:**\n"
                                    edl_section += "1. Open DaVinci Resolve\n"
                                    edl_section += "2. Go to File > Import > Timeline Markers from EDL\n"
                                    edl_section += f"3. Select the file: `{edl_filename}`\n"
                                    edl_section += "4. Markers will be placed at the specified timecodes\n"
                                    edl_section += "5. Each marker shows the B-roll search term and description\n\n"

                                    final_script_with_tools += edl_section
                                    print("✅ EDL instructions appended to script")
                                else:
                                    print(
                                        f"⚠️ EDL creation failed: {edl_result.get('error')}")

                            except Exception as edl_error:
                                print(f"⚠️ EDL generation error: {edl_error}")
                        else:
                            print("⚠️ No parsed data available for EDL generation")
                    else:
                        print(
                            f"⚠️ B-roll table generation failed: {broll_result.get('error')}")

                except Exception as broll_error:
                    print(f"⚠️ B-roll table generation error: {broll_error}")
            else:
                print("⚡ Quick test: Skipping B-roll table generation")

            # NEW: Generate B-roll Images using Gemini (only if broll_images checkbox checked)
            broll_images = None
            if not quick_test and checkboxes.get("broll_images", False) and broll_table:
                try:
                    print("\n🎨 Generating B-roll Images...")
                    streamer.send_update("🎨 Generating B-roll images...", 98.8)

                    from tools.media.broll_image_generator import BRollImageGenerator

                    broll_gen = BRollImageGenerator()
                    print("✅ B-roll image generator initialized")

                    # Generate images from ALL entries in the B-roll table
                    # max_images now refers to max ENTRIES (each gets 3 variations)
                    # Set to None to generate ALL entries with 3 variations each
                    image_results = broll_gen.generate_all_broll_images(
                        broll_table=broll_table,
                        script_title=topic,
                        # Generate all entries (matching script processing workflow)
                        max_images=None
                    )

                    if image_results.get("success"):
                        broll_images = image_results.get("images", [])
                        entries_count = image_results.get("total_entries", 0)
                        variations = image_results.get(
                            "variations_per_entry", 3)
                        print(
                            f"✅ Generated {len(broll_images)} B-roll images ({entries_count} entries × {variations} variations)")
                        streamer.send_update(
                            f"✅ Generated {len(broll_images)} B-roll images ({entries_count} entries × {variations} variations)", 99)
                    else:
                        print(
                            f"⚠️ B-roll image generation failed: {image_results.get('error')}")

                except Exception as broll_img_error:
                    print(
                        f"⚠️ B-roll image generation error: {broll_img_error}")

            # NEW: Generate HeyGen Ready Section (only if checkbox checked)
            if checkboxes.get("heygen", False):
                try:
                    print("\n🎬 Generating HeyGen Ready section...")
                    from console_ui.text_processing import extract_heygen_host_script

                    heygen_script = extract_heygen_host_script(
                        final_script_content)

                    if heygen_script:
                        heygen_section = f"\n\n{'=' * 80}\n"
                        heygen_section += "# 🎬 HEYGEN READY SCRIPT\n"
                        heygen_section += f"{'=' * 80}\n\n"
                        heygen_section += heygen_script

                        final_script_with_tools += heygen_section
                        print(
                            f"✅ HeyGen section added ({len(heygen_script)} characters, {len(heygen_script.split())} words)")

                        # Generate curl commands only if curl checkbox is checked
                        print(
                            f"🔍 DEBUG: curl checkbox = {checkboxes.get('curl', False)}")
                        if checkboxes.get("curl", False):
                            try:
                                from console_ui.text_processing import generate_heygen_curl_commands

                                print("\n🔨 Generating HeyGen curl commands...")
                                curl_commands = generate_heygen_curl_commands(
                                    final_script_with_tools,
                                    topic  # script_title
                                )

                                if curl_commands:
                                    # Store curl commands separately (don't append to script)
                                    print(
                                        f"✅ Generated {curl_commands.count('curl --location')} curl commands")
                                else:
                                    print("⚠️ No curl commands generated")
                                    curl_commands = None

                            except Exception as curl_error:
                                print(
                                    f"⚠️ Curl generation error: {curl_error}")
                                curl_commands = None
                    else:
                        print("⚠️ No host dialogue found for HeyGen section")

                except Exception as heygen_error:
                    print(
                        f"⚠️ HeyGen section generation error: {heygen_error}")

            # NEW: Generate Demo Packages (skip unless checkbox checked)
            demo_packages = None
            if not quick_test and checkboxes.get("demo", False):
                try:
                    print("\n🎥 Generating Demo Packages...")
                    streamer.send_update("🎥 Creating demo packages...", 96)

                    from linedrive_azure.agents.openai_demo_agent_client import OpenAIDemoAgentClient

                    demo_client = OpenAIDemoAgentClient()
                    demo_result = demo_client.generate_demo_packages(
                        final_script_content,
                        max_tokens=12000,
                        audience=audience
                    )

                    if demo_result.get("success"):
                        demo_packages = demo_result["response"]
                        print(
                            f"✅ Demo packages created ({len(demo_packages)} characters)")
                        streamer.send_update("✅ Demo packages generated", 97)

                        # Extract HeyGen content from demos
                        from console_ui.text_processing import (
                            extract_demo_heygen_content,
                            format_demo_steps_plain
                        )
                        demo_heygen_content = extract_demo_heygen_content(
                            demo_packages)

                        if demo_heygen_content:
                            print(
                                f"✅ Demo HeyGen section prepared ({len(demo_heygen_content)} characters)")

                        # Format demo packages - remove numbers/bullets from steps
                        formatted_demo_packages = format_demo_steps_plain(
                            demo_packages)

                        # Append demo packages to final script
                        demo_section = f"\n\n{'=' * 80}\n"
                        demo_section += "# 🎥 DEMO PACKAGES\n"
                        demo_section += f"{'=' * 80}\n\n"
                        demo_section += formatted_demo_packages

                        final_script_with_tools += demo_section
                        print("✅ Demo packages appended to script")
                    else:
                        print(
                            f"⚠️ Demo generation failed: {demo_result.get('error', 'Unknown error')}")

                except Exception as demo_error:
                    print(f"⚠️ Demo generation error: {demo_error}")
            else:
                print("⚡ Quick test: Skipping demo package generation")

            # Store the original script before flow analysis
            original_script_before_flow = final_script_with_tools
            flow_analysis_report = None

            # NEW: Flow & Repetition Analysis (only if checkbox checked)
            # IMPORTANT: This runs AFTER all content is assembled
            if checkboxes.get("flow_analysis", False):
                try:
                    print("\n🔄 Analyzing script for repetition and flow...")
                    streamer.send_update(
                        "🔄 Analyzing script for repetition and flow...", 96)

                    from linedrive_azure.agents import ScriptRepeatAndFlowAgentClient

                    flow_agent = ScriptRepeatAndFlowAgentClient()
                    flow_result = flow_agent.analyze_and_improve_flow(
                        script_content=final_script_with_tools,
                        script_title=topic,
                        target_audience=audience,
                        timeout=300
                    )

                    logger.info(
                        f"📊 Flow agent result: success={flow_result.get('success')}, keys={list(flow_result.keys())}")

                    if flow_result.get("success", False):
                        improved_script = flow_result.get(
                            "improved_script", "")
                        repetition_analysis = flow_result.get(
                            "repetition_analysis", "")
                        flow_improvements = flow_result.get(
                            "flow_analysis", "")

                        logger.info(
                            f"📊 Improved script length: {len(improved_script)}, rep analysis: {len(repetition_analysis)}, flow: {len(flow_improvements)}")

                        if improved_script:
                            # Build analysis report
                            analysis_section = f"\n\n{'=' * 80}\n"
                            analysis_section += "# 🔄 FLOW & REPETITION ANALYSIS\n"
                            analysis_section += f"{'=' * 80}\n\n"

                            if repetition_analysis:
                                analysis_section += "## Repetition Analysis\n\n"
                                analysis_section += f"{repetition_analysis}\n\n"

                            if flow_improvements:
                                analysis_section += "## Flow Improvements\n\n"
                                analysis_section += f"{flow_improvements}\n\n"

                            # Store analysis for separate display
                            flow_analysis_report = analysis_section

                            # Update final_script_with_tools to improved version (WITHOUT analysis appended)
                            final_script_with_tools = improved_script

                            streamer.send_update(
                                f"✅ Flow analysis complete - script improved ({len(improved_script)} chars)",
                                97
                            )
                            print(
                                f"✅ Flow analysis complete - script improved ({len(improved_script)} chars)")
                        else:
                            logger.warning(
                                "⚠️ Flow agent returned empty improved script")
                            print("⚠️ Flow agent returned empty improved script")
                    else:
                        logger.warning(
                            f"⚠️ Flow analysis failed - result: {flow_result}")
                        print(f"⚠️ Flow analysis failed")

                except Exception as flow_error:
                    logger.error(f"❌ Flow analysis error: {flow_error}")
                    print(f"⚠️ Flow analysis error: {flow_error}")
            else:
                print("⚡ Skipping flow & repetition analysis (not selected)")

            # NEW: Generate YouTube Upload Details (AFTER flow analysis)
            print(
                f"🔍 DEBUG: Checking youtube_details - checkbox: {checkboxes.get('youtube_details', False)}, quick_test: {quick_test}")
            if not quick_test and checkboxes.get("youtube_details", False):
                print("✅ DEBUG: YouTube Details condition TRUE - calling agent")
                try:
                    print("\n📺 Generating YouTube Upload Details...")
                    streamer.send_update(
                        "📺 Generating YouTube upload metadata...", 97)

                    from linedrive_azure.agents import YouTubeUploadDetailsAgentClient

                    youtube_agent = YouTubeUploadDetailsAgentClient()
                    youtube_result = youtube_agent.generate_upload_details(
                        script_content=final_script_content,
                        script_title=topic,
                        target_audience=audience,
                        video_length=video_length,
                        primary_keywords=None,
                        channel_focus=None,
                        timeout=180
                    )

                    if youtube_result.get("success", False):
                        youtube_upload_details = youtube_result.get(
                            "upload_details", "")
                        print(
                            f"✅ YouTube upload details generated ({len(youtube_upload_details)} characters)")
                        streamer.send_update(
                            "✅ YouTube metadata generated", 98)

                        # Append YouTube details to script
                        youtube_section = f"\n\n{'=' * 80}\n"
                        youtube_section += "# 📺 YOUTUBE UPLOAD DETAILS\n"
                        youtube_section += f"{'=' * 80}\n\n"
                        youtube_section += youtube_upload_details

                        final_script_with_tools += youtube_section
                        print("✅ YouTube details appended to script")
                    else:
                        print(
                            f"⚠️ YouTube generation failed: {youtube_result.get('error')}")

                except Exception as youtube_error:
                    print(f"⚠️ YouTube upload details error: {youtube_error}")
            else:
                print("⚡ Quick test: Skipping YouTube details generation")

            # NEW: Generate Emotional Thumbnails (only if checkbox checked)
            thumbnail_results = None
            if checkboxes.get("thumbnails", False):
                try:
                    print("\n" + "="*70)
                    print("🖼️  THUMBNAIL GENERATION STARTING")
                    print("="*70)
                    print(f"📝 Topic: {topic}")
                    print(
                        f"📄 Script length: {len(final_script_content)} chars")
                    print(
                        f"🎥 YouTube details: {'Available' if youtube_upload_details else 'None'}")

                    streamer.send_update(
                        "🖼️ Generating 6 thumbnail variations...", 98)

                    from tools.media.emotional_thumbnail_generator import (
                        EmotionalThumbnailGenerator,
                    )

                    print(f"\n🔧 Initializing EmotionalThumbnailGenerator...")
                    # Let generator get API key from environment (same as test page)
                    api_key = os.getenv(
                        "GOOGLE_API_KEY", "AIzaSyAiFFlgDokz-s4U8UrV73Fhdnl8Ukx2jCM")
                    print(
                        f"   Environment API key: {api_key[:20]}... (length: {len(api_key)})")

                    # No api_key param - use environment
                    thumbnail_gen = EmotionalThumbnailGenerator()
                    print(f"✅ Generator initialized successfully")
                    print(f"   Template: {thumbnail_gen.template_path}")
                    print(f"   Output: {thumbnail_gen.output_dir}")
                    print(
                        f"   API Key set: {'Set' if thumbnail_gen.api_key else 'MISSING'}")
                    print(
                        f"   API Key matches: {thumbnail_gen.api_key == api_key}")

                    print(f"\n🎬 Calling generate_all_thumbnails()...")
                    thumbnail_results = thumbnail_gen.generate_all_thumbnails(
                        script_title=topic,
                        script_content=final_script_content,
                        youtube_upload_details=youtube_upload_details
                    )

                    print(f"\n🔍 THUMBNAIL GENERATION RESULTS:")
                    print(f"   Type: {type(thumbnail_results)}")
                    print(f"   Is None: {thumbnail_results is None}")

                    if thumbnail_results:
                        print(f"   Keys: {list(thumbnail_results.keys())}")
                        if thumbnail_results.get("variations"):
                            variations = thumbnail_results["variations"]
                            print(f"   Variations count: {len(variations)}")
                            print(
                                f"\n✅ Generated {len(variations)} thumbnail variations")

                            # Show details of each variation
                            for i, var in enumerate(variations, 1):
                                print(
                                    f"      #{i}: {var.get('emotion')} - {var.get('filename')}")

                            streamer.send_update(
                                f"✅ Generated {len(variations)} thumbnails", 99
                            )
                        else:
                            print("   ⚠️ No 'variations' key in results")
                            print(f"   Full results: {thumbnail_results}")
                    else:
                        print("   ⚠️ thumbnail_results is None")

                    print("="*70 + "\n")

                except Exception as thumb_error:
                    print("\n" + "="*70)
                    print("❌ THUMBNAIL GENERATION ERROR")
                    print("="*70)
                    print(f"Error type: {type(thumb_error).__name__}")
                    print(f"Error message: {thumb_error}")
                    import traceback
                    print("\n📋 Full traceback:")
                    traceback.print_exc()
                    print("="*70 + "\n")

            streamer.result = {
                "success": True,
                "enhanced_script": final_script_with_tools,
                "demo_packages": demo_packages,
                "youtube_details": youtube_upload_details,
                "thumbnail_results": thumbnail_results,
                "broll_images": broll_images,
                "word_count": len(final_script_with_tools.split()),
                "reading_time": f"{len(final_script_with_tools.split()) // 150} min",
                "audience": audience,
                "production_type": production_type,
                "comparison_file": result.get("comparison_file") if isinstance(result, dict) else None,
                "chapter_comparisons": result.get("chapter_comparisons") if isinstance(result, dict) else None,
                "flow_original_script": original_script_before_flow if checkboxes.get("flow_analysis") else None,
                "flow_improved_script": final_script_with_tools if checkboxes.get("flow_analysis") and flow_analysis_report else None,
                "flow_analysis_report": flow_analysis_report,
                "edl_content": edl_content,
                "edl_filename": edl_filename,
                "curl_commands": curl_commands,
            }

            # Debug: Log what's in the result
            print(
                f"🔍 DEBUG: Result object curl_commands = {curl_commands is not None} ({len(curl_commands) if curl_commands else 0} chars)")
            print(f"🔍 DEBUG: Result keys: {list(streamer.result.keys())}")

            # Send final completion message (SIMPLE)
            print("🎯 Sending completion message...")
            print(f"DEBUG: About to send completion for session {session_id}")
            print(
                f"DEBUG: Streamer done status before completion: {streamer.done}")

            # Send the completion message
            streamer.send_update("✅ Script creation complete!", 100)

            print("✅ Completion message sent successfully")
            print(
                f"DEBUG: Streamer done status after completion: {streamer.done}")
            print("🎉 Script creation completed successfully!")

        else:
            error_msg = result.get("error", "Unknown workflow error")
            streamer.send_update(f"❌ Script creation failed: {error_msg}", -1)
            streamer.result = {"success": False, "error": error_msg}
            print(f"❌ Script creation failed: {error_msg}")
            logger.error(
                f"❌ Script creation failed for session {session_id}: {error_msg}")

    except Exception as e:
        error_msg = f"Error: {str(e)}"
        logger.error(f"💥 Exception during script creation: {error_msg}")
        if session_id in progress_streams:
            streamer = progress_streams[session_id]
            streamer.send_update(f"❌ {error_msg}", -1)
            streamer.result = {"success": False, "error": error_msg}


async def process_existing_script(
    session_id, script_content, audience, tone,
    video_length, checkboxes, heygen_template_id="",
    heygen_api_key=""
):
    """Process existing script with progress updates"""
    import re

    logger.info(f"📝 SCRIPT PROCESSING STARTED: session={session_id}")

    if session_id not in progress_streams:
        logger.error(f"❌ Session {session_id} not found!")
        return

    streamer = progress_streams[session_id]

    try:
        # Initialize EDL variables
        edl_content = None
        edl_filename = None

        # Calculate total steps based on checkboxes
        total_steps = sum([
            checkboxes.get("hook_summary", False),
            checkboxes.get("youtube_details", False),
            checkboxes.get("broll", False),
            checkboxes.get("broll_images", False),
            checkboxes.get("heygen", False),
            checkboxes.get("curl", False),
            checkboxes.get("demo", False),
            checkboxes.get("thumbnails", False),
            checkboxes.get("flow_analysis", False)
        ])

        if total_steps == 0:
            streamer.send_update("❌ No processing options selected", -1)
            streamer.result = {"success": False,
                               "error": "No options selected"}
            return

        completed_steps = 0
        progress_per_step = 90 / total_steps  # Reserve 10% for completion

        streamer.send_update("🚀 Initializing script processing...", 5)

        # Extract title from script
        script_title = "Untitled Script"
        logger.info("📰 Extracting title from script content...")

        title_match = re.search(r'^#\s+(.+)$', script_content, re.MULTILINE)
        if title_match:
            script_title = title_match.group(1).strip()
            logger.info(f"📰 ✅ Found title: '{script_title}'")
        else:
            lines = script_content.split('\n')
            for line in lines[:10]:
                line = line.strip()
                if not line or line.startswith(('#', '*', '-', '[')):
                    continue
                if all(c in '_=-~' for c in line):
                    continue
                if len(line) > 150:
                    continue
                script_title = line
                logger.info(f"📰 ✅ Using line as title: '{script_title}'")
                break

        # Clean the script
        streamer.send_update("🧹 Cleaning script formatting...", 10)
        cleaned_script = script_content

        # Remove separator lines
        cleaned_script = re.sub(r'^[_\-=~]+\s*$', '', cleaned_script,
                                flags=re.MULTILINE)
        cleaned_script = re.sub(r'\n{3,}', '\n\n', cleaned_script)

        # Remove emojis
        cleaned_script = re.sub(r'🎬|📺|📊|🎥|🎯|💡|✨|🚀', '', cleaned_script)

        # Remove markdown formatting
        cleaned_script = re.sub(r'\*\*([^*]+)\*\*', r'\1', cleaned_script)
        cleaned_script = re.sub(r'\*([^*]+)\*', r'\1', cleaned_script)
        cleaned_script = re.sub(r'__([^_]+)__', r'\1', cleaned_script)
        cleaned_script = re.sub(r'_([^_]+)_', r'\1', cleaned_script)

        # Apply clean formatting
        cleaned_script = re.sub(r'^Host:\s*$', r'**Host:**\n\n',
                                cleaned_script, flags=re.MULTILINE)
        cleaned_script = re.sub(r'^VISUAL CUE:\s*(.+)$', r'\n**VISUAL CUE:** \1',
                                cleaned_script, flags=re.MULTILINE)

        # Remove the title from the beginning of cleaned_script to avoid duplication
        # since we'll prepend it separately with proper formatting
        if script_title and script_title in cleaned_script:
            # Remove markdown title format (# Title)
            cleaned_script = re.sub(
                r'^#\s+' + re.escape(script_title) + r'\s*\n*',
                '',
                cleaned_script,
                count=1,
                flags=re.MULTILINE
            )
            # Also remove plain title at the start
            if cleaned_script.strip().startswith(script_title):
                cleaned_script = cleaned_script.strip()[
                    len(script_title):].lstrip()

        # Initialize output variables (will build final_output after all sections are generated)
        final_output = ""
        hook_text = ""
        summary_text = ""
        flow_analysis = ""
        heygen_script = None
        curl_commands = None
        hook_result = {}  # Initialize to empty dict to prevent undefined variable errors

        # Generate Hook & Summary if requested
        if checkboxes.get("hook_summary", False):
            streamer.send_update("🎯 Generating Hook & Summary...", 15)
            try:
                from linedrive_azure.agents import HookAndSummaryAgentClient

                hook_agent = HookAndSummaryAgentClient()
                hook_result = hook_agent.generate_hook_and_summary(
                    script_content=cleaned_script,
                    script_title=script_title,
                    target_audience=audience,
                    tone=tone,
                    video_length=video_length,
                    timeout=180
                )

                if hook_result.get("success", False):
                    hook_text = hook_result.get("hook", "")
                    summary_text = hook_result.get("summary", "")
                    flow_analysis = hook_result.get("flow_analysis", "")
                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        f"✅ Hook & Summary generated "
                        f"({len(hook_text)} + {len(summary_text)} chars)",
                        int(progress)
                    )
                else:
                    logger.warning(f"⚠️ Hook & summary failed")
            except Exception as e:
                logger.error(f"❌ Hook & summary error: {e}")

        # Build script with hooks and summary
        hook1_text = ""
        hook2_text = ""
        hook3_text = ""
        opening_statement = ""

        if hook_result.get("success", False):
            # Extract all three hooks from the result
            # Use hook1 or fallback to hook (for backward compatibility)
            hook1_text = hook_result.get("hook1", hook_result.get("hook", ""))
            hook2_text = hook_result.get("hook2", "")
            hook3_text = hook_result.get("hook3", "")
            opening_statement = hook_result.get("opening_statement", "")

            # DEBUG: Log what we actually got
            logger.info(
                f"🔍 DEBUG: hook_result keys: {list(hook_result.keys())}")
            logger.info(
                f"🔍 DEBUG: hook1_text length: {len(hook1_text) if hook1_text else 0}")
            logger.info(
                f"🔍 DEBUG: hook2_text length: {len(hook2_text) if hook2_text else 0}")
            logger.info(
                f"🔍 DEBUG: hook3_text length: {len(hook3_text) if hook3_text else 0}")

            logger.info(
                f"📌 Adding {3 if hook3_text else (2 if hook2_text else 1)} hook option(s) to output")

            final_output += "**🎬 OPENING HOOK OPTIONS**\n\n"
            final_output += "*Choose one of these three hooks for your video:*\n\n"

            # Add Hook Option 1
            final_output += "**OPTION 1:**\n\n"
            final_output += f"**Host:**\n\n{hook1_text}\n\n"

            # Add Hook Option 2 if available
            if hook2_text:
                final_output += "---\n\n**OPTION 2:**\n\n"
                final_output += f"**Host:**\n\n{hook2_text}\n\n"

            # Add Hook Option 3 if available
            if hook3_text:
                final_output += "---\n\n**OPTION 3:**\n\n"
                final_output += f"**Host:**\n\n{hook3_text}\n\n"

            final_output += "---\n\n"

            # Add opening statement after hooks
            if opening_statement:
                final_output += "**📺 OPENING STATEMENT**\n\n"
                final_output += f"**Host:**\n\n{opening_statement}\n\n"
                final_output += "---\n\n"

        # Add script title with line break and bold formatting
        final_output += f"\n**{script_title}**\n\n"

        final_output += cleaned_script
        logger.info(
            f"📌 After adding cleaned_script, final_output length: {len(final_output)} chars")

        if summary_text:
            logger.info(
                f"📌 Adding summary to output ({len(summary_text)} chars)")
            final_output += "\n\n---\n\n**📝 CONCLUSION/SUMMARY**\n\n"
            final_output += f"**Host:**\n\n{summary_text}\n\n"

        if flow_analysis:
            logger.info(
                f"📌 Adding flow analysis to output ({len(flow_analysis)} chars)")
            final_output += f"\n\n{'=' * 80}\n# 📊 FLOW ANALYSIS\n"
            final_output += f"{'=' * 80}\n\n{flow_analysis}"

        # Generate YouTube Upload Details if requested
        if checkboxes.get("youtube_details", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("📺 Generating YouTube Upload Details...",
                                 int(current_progress))
            try:
                from linedrive_azure.agents import YouTubeUploadDetailsAgentClient

                youtube_agent = YouTubeUploadDetailsAgentClient()
                youtube_result = youtube_agent.generate_upload_details(
                    script_content=cleaned_script,
                    script_title=script_title,
                    target_audience=audience,
                    video_length=video_length,
                    timeout=180
                )

                if youtube_result.get("success", False):
                    youtube_details = youtube_result.get("upload_details", "")
                    youtube_section = f"\n\n{'=' * 80}\n"
                    youtube_section += "# 📺 YOUTUBE UPLOAD DETAILS\n"
                    youtube_section += f"{'=' * 80}\n\n{youtube_details}"
                    final_output += youtube_section
                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        f"✅ YouTube details generated ({len(youtube_details)} chars)",
                        int(progress)
                    )
            except Exception as e:
                logger.error(f"❌ YouTube details error: {e}")

        # Generate B-roll Table if requested
        broll_table = None
        if checkboxes.get("broll", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("📊 Generating B-roll Search Terms Table...",
                                 int(current_progress))
            try:
                from linedrive_azure.agents import ScriptBRollAgentClient

                broll_agent = ScriptBRollAgentClient()
                broll_result = broll_agent.generate_broll_table_with_timecodes(
                    script_content=cleaned_script,
                    script_title=script_title,
                    words_per_minute=150,
                    # Increased to 5 minutes for larger tables (40-60+ rows)
                    timeout=300
                )

                if broll_result.get("success", False):
                    broll_table = broll_result.get("table", "")
                    parsed_data = broll_result.get("parsed_data", [])

                    broll_section = f"\n\n{'=' * 80}\n"
                    broll_section += "# 📊 B-ROLL SEARCH TERMS TABLE\n"
                    broll_section += f"{'=' * 80}\n\n{broll_table}"
                    final_output += broll_section

                    # Generate EDL file
                    if parsed_data:
                        try:
                            import time
                            timestamp = time.strftime("%Y%m%d_%H%M%S")
                            edl_filename = f"broll_markers_{timestamp}.edl"

                            edl_result = broll_agent.create_edl_markers(
                                broll_data=parsed_data,
                                output_file=edl_filename,
                                frame_rate='24'
                            )

                            if edl_result.get("success"):
                                # Read EDL content for frontend display
                                try:
                                    with open(edl_filename, 'r', encoding='utf-8') as f:
                                        edl_content = f.read()
                                    logger.info(
                                        f"✅ EDL content read ({len(edl_content)} bytes)")
                                except Exception as read_error:
                                    logger.error(
                                        f"⚠️ Could not read EDL file: {read_error}")

                                edl_info = f"\n\n**EDL File:** `{edl_filename}`"
                                edl_info += f" ({edl_result.get('marker_count')}"
                                edl_info += " markers)\n"
                                final_output += edl_info
                        except Exception as edl_error:
                            logger.error(f"❌ EDL error: {edl_error}")

                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        f"✅ B-roll table generated ({len(broll_table)} chars)",
                        int(progress)
                    )
                else:
                    logger.warning(f"⚠️ B-roll agent returned success=False")
                    streamer.send_update(
                        "⚠️ B-roll table generation failed", int(current_progress))
            except Exception as e:
                logger.error(f"❌ B-roll table error: {e}")
                import traceback
                traceback.print_exc()
                streamer.send_update(
                    f"❌ B-roll table error: {str(e)}", int(current_progress))

        # Generate B-roll Images using Gemini (ONLY if broll_images checkbox checked AND broll_table was successfully generated)
        broll_images = None
        if checkboxes.get("broll_images", False) and broll_table:
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("🎨 Generating B-roll images with AI...",
                                 int(current_progress))
            try:
                from tools.media.broll_image_generator import BRollImageGenerator

                broll_gen = BRollImageGenerator()

                # Generate images from the B-roll table
                # max_images now refers to max ENTRIES (each gets 3 variations)
                # Set to None to generate ALL entries, or set a limit if needed
                image_results = broll_gen.generate_all_broll_images(
                    broll_table=broll_table,
                    script_title=script_title,
                    max_images=None  # Generate all entries with 3 variations each
                )

                if image_results.get("success", False):
                    broll_images = image_results.get("images", [])
                    entries_count = image_results.get("total_entries", 0)
                    variations = image_results.get("variations_per_entry", 3)
                    logger.info(
                        f"✅ Generated {len(broll_images)} B-roll images ({entries_count} entries × {variations} variations)")
                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        f"✅ Generated {len(broll_images)} B-roll images ({entries_count} entries × {variations} variations)",
                        int(progress)
                    )
                else:
                    logger.warning(
                        f"⚠️ B-roll images generation failed: {image_results.get('error')}")
            except Exception as e:
                logger.error(f"❌ B-roll images error: {e}")
                import traceback
                traceback.print_exc()

        # Generate HeyGen Ready Script if requested
        if checkboxes.get("heygen", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("🎬 Generating HeyGen Ready section...",
                                 int(current_progress))
            try:
                from console_ui.text_processing import extract_heygen_host_script

                heygen_script = extract_heygen_host_script(final_output)
                if heygen_script:
                    heygen_section = f"\n\n{'=' * 80}\n"
                    heygen_section += "# 🎬 HEYGEN READY SCRIPT\n"
                    heygen_section += f"{'=' * 80}\n\n{heygen_script}"
                    final_output += heygen_section
                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        f"✅ HeyGen script extracted ({len(heygen_script)} chars)",
                        int(progress)
                    )
            except Exception as e:
                logger.error(f"❌ HeyGen section error: {e}")
                heygen_script = None

        # Generate curl commands if requested
        if checkboxes.get("curl", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("🚀 Generating HeyGen curl commands...",
                                 int(current_progress))
            try:
                from console_ui.text_processing import (
                    extract_heygen_host_script,
                    generate_heygen_curl_commands
                )

                if not heygen_script:
                    heygen_script = extract_heygen_host_script(final_output)

                if heygen_script:
                    heygen_with_header = (
                        f"# 🎬 HEYGEN READY SCRIPT\n{'=' * 80}\n\n{heygen_script}"
                    )
                    curl_commands = generate_heygen_curl_commands(
                        heygen_with_header, script_title
                    )
                    if curl_commands:
                        # Store curl commands separately (don't append to script)
                        num_commands = curl_commands.count('curl --location')
                        completed_steps += 1
                        progress = 15 + (completed_steps * progress_per_step)
                        streamer.send_update(
                            f"✅ Generated {num_commands} curl commands",
                            int(progress)
                        )
                    else:
                        curl_commands = None
            except Exception as e:
                logger.error(f"❌ Curl generation error: {e}")
                curl_commands = None

        # Generate Demo Package if requested
        if checkboxes.get("demo", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("🔧 Generating demo package...",
                                 int(current_progress))
            try:
                from linedrive_azure.agents.openai_demo_agent_client import (
                    OpenAIDemoAgentClient
                )

                demo_agent = OpenAIDemoAgentClient()
                demo_result = demo_agent.generate_demo_package(
                    script_content=script_content,
                    script_title=script_title,
                    timeout=180
                )

                if demo_result.get("success", False):
                    demo_packages = demo_result.get("demo_packages", "")
                    demo_section = f"\n\n{'=' * 80}\n# 🔧 DEMO PACKAGE\n"
                    demo_section += f"{'=' * 80}\n\n{demo_packages}"
                    final_output += demo_section
                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        f"✅ Demo package generated ({len(demo_packages)} chars)",
                        int(progress)
                    )
            except Exception as e:
                logger.error(f"❌ Demo package error: {e}")

        # Generate Thumbnails if requested
        if checkboxes.get("thumbnails", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("🖼️ Generating emotional thumbnails...",
                                 int(current_progress))
            try:
                from linedrive_azure.agents import EmotionalThumbnailAgentClient

                thumb_agent = EmotionalThumbnailAgentClient()
                thumb_result = thumb_agent.generate_thumbnail_variations(
                    script_title=script_title,
                    script_content=cleaned_script,
                    target_audience=audience,
                    timeout=120
                )

                if thumb_result.get("success", False):
                    thumbnails = thumb_result.get("thumbnails", "")
                    thumb_section = f"\n\n{'=' * 80}\n"
                    thumb_section += "# 🖼️ EMOTIONAL THUMBNAIL VARIATIONS\n"
                    thumb_section += f"{'=' * 80}\n\n{thumbnails}"
                    final_output += thumb_section
                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        "✅ Emotional thumbnails generated",
                        int(progress)
                    )
            except Exception as e:
                logger.error(f"❌ Thumbnail generation error: {e}")

        # Generate Flow & Repetition Analysis if requested
        # IMPORTANT: This should run AFTER all other content is assembled
        if checkboxes.get("flow_analysis", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("🔄 Analyzing script for repetition and flow...",
                                 int(current_progress))
            try:
                from linedrive_azure.agents import ScriptRepeatAndFlowAgentClient

                flow_agent = ScriptRepeatAndFlowAgentClient()
                flow_result = flow_agent.analyze_and_improve_flow(
                    script_content=final_output,
                    script_title=script_title,
                    target_audience=audience,
                    timeout=300
                )

                logger.info(
                    f"📊 Flow agent result: success={flow_result.get('success')}, keys={list(flow_result.keys())}")

                if flow_result.get("success", False):
                    # Replace final_output with the improved version
                    improved_script = flow_result.get("improved_script", "")
                    repetition_analysis = flow_result.get(
                        "repetition_analysis", "")
                    flow_improvements = flow_result.get("flow_analysis", "")

                    logger.info(
                        f"📊 Improved script length: {len(improved_script)}, rep analysis: {len(repetition_analysis)}, flow: {len(flow_improvements)}")

                    if improved_script:
                        # Add analysis section first
                        analysis_section = f"\n\n{'=' * 80}\n"
                        analysis_section += "# 🔄 FLOW & REPETITION ANALYSIS\n"
                        analysis_section += f"{'=' * 80}\n\n"

                        if repetition_analysis:
                            analysis_section += "## Repetition Analysis\n\n"
                            analysis_section += f"{repetition_analysis}\n\n"

                        if flow_improvements:
                            analysis_section += "## Flow Improvements\n\n"
                            analysis_section += f"{flow_improvements}\n\n"

                        # Replace the script with improved version and add analysis at end
                        final_output = improved_script + analysis_section

                        completed_steps += 1
                        progress = 15 + (completed_steps * progress_per_step)
                        streamer.send_update(
                            f"✅ Flow analysis complete - script improved ({len(improved_script)} chars)",
                            int(progress)
                        )
                    else:
                        logger.warning(
                            "⚠️ Flow agent returned empty improved script")
                else:
                    logger.warning(
                        f"⚠️ Flow analysis failed - result: {flow_result}")
            except Exception as e:
                logger.error(f"❌ Flow analysis error: {e}")

        # Complete
        streamer.send_update("✅ Script processing complete!", 100)
        streamer.result = {
            "success": True,
            "enhanced_script": final_output,  # SSE handler expects "enhanced_script" key
            "script": final_output,  # Also include "script" for compatibility
            "script_title": script_title,
            "edl_content": edl_content,
            "edl_filename": edl_filename,
            "curl_commands": curl_commands,
            "broll_images": broll_images,
        }
        logger.info("✅ Script processing completed successfully")

    except Exception as e:
        error_msg = f"Error: {str(e)}"
        logger.error(f"💥 Exception during script processing: {error_msg}")
        streamer.send_update(f"❌ {error_msg}", -1)
        streamer.result = {"success": False, "error": error_msg}


async def test_mode_simulation(session_id):
    """Quick test mode - simulates script creation in 10 seconds"""
    import time
    import asyncio

    logger.info(f"⚡ Starting test mode simulation for {session_id}")

    if session_id not in progress_streams:
        logger.error(f"❌ No progress stream for session {session_id}")
        return

    streamer = progress_streams[session_id]

    try:
        # Simulate the 7-chapter workflow with fast updates
        test_chapters = [
            "Introduction: Setting the Stage",
            "Chapter 1: The Problem",
            "Chapter 2: Understanding the Context",
            "Chapter 3: Exploring Solutions",
            "Chapter 4: Implementation Strategy",
            "Chapter 5: Real-World Applications",
            "Conclusion: Looking Forward"
        ]

        # Send initial message
        streamer.send_update("🚀 Starting test script creation...", 5)
        await asyncio.sleep(1)

        # Simulate writing each chapter (1 second per chapter)
        for i, chapter in enumerate(test_chapters):
            progress = 10 + (i * 12)  # Distribute progress from 10-94%
            streamer.send_update(f"📝 Writing {chapter}...", progress)
            await asyncio.sleep(1)

            # Show completion of chapter
            char_count = 1000 + (i * 200)  # Fake character counts
            streamer.send_update(f"✅ {chapter} completed ({char_count} chars)",
                                 progress + 2)
            await asyncio.sleep(0.5)

        # Final processing steps
        streamer.send_update("✨ Enhancing script formatting...", 96)
        await asyncio.sleep(1)

        streamer.send_update("🔍 Adding tool recommendations...", 98)
        await asyncio.sleep(1)

        # Create a simple test result
        test_script = """Welcome to AI with Roz, I'm Roz's AI Clone. 

This is a test script created in fast mode for debugging purposes.

Chapter 1: The Problem
This chapter discusses the main challenges...

Chapter 2: Understanding the Context  
Here we explore the broader implications...

[Continue with remaining chapters...]

===============================
TOOLS
===============================

• OpenAI GPT-4 - AI writing assistant
• Canva - Design and presentation tools
• YouTube Studio - Video optimization"""

        # Store test result
        streamer.result = {
            "success": True,
            "enhanced_script": test_script,
            "word_count": len(test_script.split()),
            "reading_time": f"{len(test_script.split()) // 150} min",
            "audience": "test audience",
            "production_type": "test",
        }

        # Send completion message
        streamer.send_update("✅ Test script creation complete!", 100)
        logger.info(f"✅ Test mode completed for session {session_id}")

    except Exception as e:
        error_msg = f"Test mode error: {str(e)}"
        logger.error(f"❌ Test mode error for {session_id}: {error_msg}")
        streamer.send_update(f"❌ {error_msg}", -1)
        streamer.result = {"success": False, "error": error_msg}


@app.route("/test-sse")
def test_sse():
    """Test SSE completion behavior without running full script creation"""
    return render_template("test_sse.html")


@app.route("/test-progress/<session_id>")
def test_progress_stream(session_id):
    """Test SSE endpoint that simulates script creation completion"""
    import time

    def generate():
        try:
            # Simulate progress messages
            messages = [
                ("🚀 Starting test...", 10),
                ("📝 Processing step 1...", 30),
                ("📝 Processing step 2...", 60),
                ("✅ Almost done...", 90),
                ("✅ Test complete!", 100)
            ]

            for message, progress in messages:
                update = {
                    "message": message,
                    "progress": progress,
                    "timestamp": time.time(),
                    "done": progress == 100
                }

                logger.info(f"🔄 TEST SSE sending: {json.dumps(update)}")
                yield f"data: {json.dumps(update)}\n\n"

                if progress == 100:
                    logger.info(f"🎯 TEST SSE completion message sent")
                    # Don't send closing message - client will close connection
                    break

                time.sleep(1)  # 1 second between messages

        except Exception as e:
            logger.error(f"TEST SSE error: {e}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
        finally:
            logger.info(f"🔚 TEST SSE stream ended")

    response = Response(generate(), mimetype="text/event-stream")
    response.headers['Cache-Control'] = 'no-cache'
    response.headers['Connection'] = 'keep-alive'
    response.headers['X-Accel-Buffering'] = 'no'
    return response


@app.route("/")
def index():
    """Serve the main ScriptCraft web interface"""
    return render_template("index.html", version=VERSION)


@app.route("/test")
def test_page():
    """Serve the test page for debugging SSE streaming"""
    return render_template("test.html")


@app.route("/test-capture")
def test_capture_page():
    """Test page to debug console capture and threading issues"""
    return """
    <!DOCTYPE html>
    <html>
    <head><title>Console Capture Test</title></head>
    <body>
        <h1>Console Capture Test Harness</h1>
        <button onclick="runTest()">Run Console Capture Test</button>
        <div id="progress"></div>
        <div id="result"></div>
        <script>
        function runTest() {
            fetch('/test-capture-process', {method: 'POST'})
            .then(r => r.json())
            .then(data => {
                document.getElementById('progress').innerHTML = 'Test started...';
                watchProgress(data.session_id);
            });
        }
        
        function watchProgress(sessionId) {
            const eventSource = new EventSource('/progress/' + sessionId);
            eventSource.onmessage = function(event) {
                const data = JSON.parse(event.data);
                document.getElementById('progress').innerHTML = 
                    `Progress: ${data.progress}% - ${data.message}`;
                if (data.done) {
                    eventSource.close();
                    document.getElementById('result').innerHTML = 'Test completed!';
                }
            };
        }
        </script>
    </body>
    </html>
    """


@app.route('/test-capture-process', methods=['POST'])
def test_capture_process():
    """Test the console capture mechanism without full workflow"""
    session_id = str(uuid.uuid4())
    progress_streams[session_id] = ProgressStreamer(session_id)

    # Start the test in background
    thread = threading.Thread(
        target=run_test_capture,
        args=(session_id,),
        daemon=True
    )
    thread.start()
    running_tasks[session_id] = thread

    return jsonify({'session_id': session_id, 'status': 'started'})


def run_test_capture(session_id):
    """Test console capture and the exact point where it hangs"""
    import time
    import sys

    logger.info(f"🧪 TEST CAPTURE STARTED: session={session_id}")

    try:
        if session_id not in progress_streams:
            logger.error(f"❌ Session {session_id} not found!")
            return

        streamer = progress_streams[session_id]

        # Set up console capture exactly like the real workflow
        console_capture = ConsoleCapture(streamer)
        original_stdout = sys.stdout

        try:
            # Redirect stdout to capture console output
            sys.stdout = console_capture

            # Test 1: Basic console output
            print("🧪 TEST 1: Basic console output")
            streamer.send_update("🧪 Basic console output test", 10)
            time.sleep(1)

            # Test 2: Import text processing (where it hangs)
            print("🧪 TEST 2: Importing text processing modules")
            streamer.send_update("🧪 Testing text processing import", 20)

            try:
                from console_ui.text_processing import (
                    enhance_script_with_bold_tools,
                    extract_tool_links_and_info,
                )
                print("✅ Text processing modules imported successfully")
                streamer.send_update("✅ Text processing import success", 30)
            except Exception as e:
                print(f"❌ Text processing import failed: {e}")
                streamer.send_update(
                    f"❌ Text processing import failed: {e}", 30)
                raise

            # Test 3: Call enhance_script_with_bold_tools (where it might hang)
            print("🧪 TEST 3: Testing enhance_script_with_bold_tools")
            streamer.send_update("🧪 Testing script enhancement", 40)

            test_script = "This is a test script with **bold** and [ChatGPT](https://chatgpt.com)."

            try:
                enhanced_script = enhance_script_with_bold_tools(test_script)
                print(f"✅ Script enhanced: {len(enhanced_script)} chars")
                streamer.send_update("✅ Script enhancement success", 50)
            except Exception as e:
                print(f"❌ Script enhancement failed: {e}")
                streamer.send_update(f"❌ Script enhancement failed: {e}", 50)
                raise

            # Test 4: Call extract_tool_links_and_info
            print("🧪 TEST 4: Testing extract_tool_links_and_info")
            streamer.send_update("🧪 Testing tool link extraction", 60)

            try:
                tool_links = extract_tool_links_and_info(enhanced_script)
                print(f"✅ Tool links extracted: {len(tool_links)} chars")
                streamer.send_update("✅ Tool link extraction success", 70)
            except Exception as e:
                print(f"❌ Tool link extraction failed: {e}")
                streamer.send_update(f"❌ Tool link extraction failed: {e}", 70)
                raise

            # Test 5: Minimal rapid output test to isolate deadlock
            print("🧪 TEST 5: Minimal rapid output test")
            streamer.send_update("🧪 Starting minimal rapid test", 80)

            # Test just 3 rapid outputs with debugging
            for i in range(3):
                print(f"   📝 Before output {i+1}")
                streamer.send_update(f"🧪 Processing output {i+1}/3", 85 + i*2)
                print(f"   📝 After output {i+1}")
                time.sleep(0.5)  # Longer sleep to prevent issues

            print("✅ Minimal rapid test completed")
            logger.info(
                "🔍 DEBUG: About to send minimal rapid test success update")
            streamer.send_update("✅ Minimal rapid test success", 90)
            logger.info("🔍 DEBUG: Minimal rapid test success update sent")

            # Final test
            print("🧪 ALL TESTS COMPLETED SUCCESSFULLY!")
            logger.info("🔍 DEBUG: About to send all tests completed update")
            streamer.send_update("🧪 All tests completed!", 100)
            logger.info("🔍 DEBUG: All tests completed update sent")

            # Store success result
            logger.info("🔍 DEBUG: About to set streamer.result")
            streamer.result = {
                "success": True,
                "message": "All console capture tests passed",
                "enhanced_script": enhanced_script,
                "tool_links": tool_links
            }
            logger.info("🔍 DEBUG: streamer.result set successfully")

        finally:
            # Always restore stdout
            logger.info("🔍 DEBUG: Entering finally block")
            sys.stdout = original_stdout
            logger.info("🔍 DEBUG: stdout restored, function should exit now")

    except Exception as e:
        error_msg = f"Test failed: {str(e)}"
        logger.error(f"💥 Test exception: {error_msg}")
        if session_id in progress_streams:
            streamer = progress_streams[session_id]
            streamer.send_update(f"❌ {error_msg}", -1)
            streamer.result = {"success": False, "error": error_msg}


@app.route("/test_progress/<session_id>/<int:progress>/<message>")
def test_progress(session_id, progress, message):
    """Test endpoint to manually send progress updates"""
    if session_id in progress_streams:
        streamer = progress_streams[session_id]
        streamer.send_update(message, progress)
        logger.info(f"🧪 Test progress sent: {message} -> {progress}%")
        return jsonify({
            "success": True,
            "message": f"Sent: {message} -> {progress}%"
        })
    return jsonify({
        "success": False,
        "error": "Session not found"
    }), 404


@app.route("/test_create")
def test_create():
    """Create a test session with hardcoded progress messages"""
    import uuid
    import time
    import threading

    # Generate session ID
    session_id = str(uuid.uuid4())

    # Create progress streamer
    progress_streams[session_id] = ProgressStreamer(session_id)

    logger.info(f"🧪 Starting test session: {session_id}")

    def send_test_messages():
        streamer = progress_streams[session_id]
        time.sleep(1)

        # Send a series of test messages
        test_messages = [
            (10, "🚀 Starting test..."),
            (25, "📝 Writing Chapter 1/7: Test Chapter"),
            (40, "✅ Chapter 1 completed (1234 chars)"),
            (55, "📝 Writing Chapter 2/7: Another Test Chapter"),
            (70, "✅ Chapter 2 completed (2345 chars)"),
            (85, "🔍 Reviewing content..."),
            (100, "✅ Test completed successfully!")
        ]

        for progress, message in test_messages:
            streamer.send_update(message, progress)
            logger.info(f"🧪 Sent: {message} -> {progress}%")
            time.sleep(2)  # 2 second delay between messages

    # Start test in background
    thread = threading.Thread(target=send_test_messages, daemon=True)
    running_tasks[session_id] = thread
    thread.start()

    return jsonify({
        "success": True,
        "session_id": session_id,
        "message": "Test session started",
        "test_url": f"/progress/{session_id}"
    })


@app.route("/debug/stdout")
def debug_stdout():
    """Debug endpoint to check stdout status"""
    return jsonify({
        "stdout_type": str(type(sys.stdout)),
        "is_console_capture": isinstance(sys.stdout, ConsoleCapture),
        "active_sessions": len(progress_streams),
        "running_tasks": len(running_tasks)
    })


@app.route("/create", methods=["POST"])
def create():
    """Create a new script with progress streaming"""
    # Get form data
    data = request.get_json() or {}
    topic = data.get("topic", "AI in daily life")
    audience = data.get("audience", "general")
    tone = data.get("tone", "professional")
    video_length = data.get("video_length", "medium")
    production_type = data.get("production_type", "standard")
    goals = data.get("goals", "educational")
    test_mode = data.get("test_mode", False)  # Quick test mode flag
    quick_test = data.get("quick_test", False)  # 1-chapter quick test

    # NEW: Get checkbox selections
    checkboxes = data.get("checkboxes", {})
    # If "All" is checked, enable everything
    if checkboxes.get("all", False):
        checkboxes = {
            "script": True,
            "hook_summary": True,
            "youtube_details": True,
            "heygen": True,
            "curl": True,
            "broll": True,
            "broll_images": True,
            "thumbnails": True,
            "demo": True,
            "flow_analysis": True,
            "all": True
        }
    # If no checkboxes provided, default to script only
    elif not checkboxes:
        checkboxes = {"script": True}

    # Generate session ID
    session_id = str(uuid.uuid4())

    # Create progress streamer
    progress_streams[session_id] = ProgressStreamer(session_id)

    logger.info(f"🎬 Starting script creation: {session_id}")
    logger.info(f"📝 Topic: {topic}, Audience: {audience}, Tone: {tone}")
    logger.info(f"📋 Checkboxes: {checkboxes}")
    if quick_test:
        logger.info(f"⚡ QUICK TEST MODE: 1 chapter only")

    # Start script creation in background thread
    def run_script_creation():
        try:
            if test_mode:
                # Quick test mode - simulate progress and complete in 10 seconds
                logger.info(f"⚡ Running in test mode for session {session_id}")
                asyncio.run(test_mode_simulation(session_id))
            else:
                # Override video_length for quick test
                actual_length = "1 minute (150-200 words)" if quick_test else video_length
                asyncio.run(process_script_creation(
                    session_id, topic, audience, tone, actual_length,
                    production_type, goals, quick_test, checkboxes
                ))
        except Exception as e:
            logger.error(f"❌ Script creation error: {e}")
            if session_id in progress_streams:
                progress_streams[session_id].send_update(
                    f"❌ Error: {str(e)}", 100
                )

    thread = threading.Thread(target=run_script_creation, daemon=True)
    running_tasks[session_id] = thread
    thread.start()

    return jsonify({
        "success": True,
        "session_id": session_id,
        "message": "Script creation started"
    })


@app.route("/progress/<session_id>")
def progress_stream(session_id):
    """Server-Sent Events endpoint for progress updates"""
    logger.info(f"🌊 SSE connection established for session {session_id}")

    if session_id not in progress_streams:
        return "Session not found", 404

    streamer = progress_streams[session_id]

    def generate():
        try:
            while True:  # Changed from while not streamer.done
                try:
                    # Get update with timeout
                    update = streamer.queue.get(timeout=30)
                    logger.info(f"🔄 SSE sending: {update}")

                    # Parse the update to check if it's the completion message
                    update_data = json.loads(update)

                    # If this is the completion message (progress=100), add the script content
                    if update_data.get("done", False) or update_data.get("progress") == 100:
                        logger.info(
                            f"🎯 Processing completion message for session {session_id}")

                        # Create a completion message with status and script
                        if streamer.result and streamer.result.get("success"):
                            # Prepare thumbnail data for frontend
                            thumbnail_results = streamer.result.get(
                                "thumbnail_results", {})

                            print("\n" + "="*70)
                            print("📦 PREPARING COMPLETION MESSAGE")
                            print("="*70)
                            print(f"Session: {session_id}")
                            print(
                                f"Thumbnail results type: {type(thumbnail_results)}")
                            print(
                                f"Thumbnail results is None: {thumbnail_results is None}")

                            thumbnails = []
                            if thumbnail_results and thumbnail_results.get("variations"):
                                print(f"✅ Found variations in thumbnail_results")
                                variations = thumbnail_results["variations"]
                                print(f"   Count: {len(variations)}")

                                for idx, variation in enumerate(variations, 1):
                                    print(f"\n   Processing variation #{idx}:")
                                    print(
                                        f"      Keys: {list(variation.keys())}")
                                    print(
                                        f"      Emotion: {variation.get('emotion')}")
                                    print(
                                        f"      Text: {variation.get('text')}")
                                    print(
                                        f"      Filename: {variation.get('filename')}")

                                    thumbnails.append({
                                        "emotion": variation["emotion"],
                                        "text": variation["text"],
                                        # Generator returns filename directly
                                        "filename": variation["filename"]
                                    })

                                print(
                                    f"\n✅ Prepared {len(thumbnails)} thumbnail objects for frontend")
                            else:
                                print(
                                    f"⚠️ No variations found in thumbnail_results")
                                if thumbnail_results:
                                    print(
                                        f"   Available keys: {list(thumbnail_results.keys())}")

                            print("="*70 + "\n")

                            # DEBUG: Check chapter_comparisons data
                            chapter_comps = streamer.result.get(
                                "chapter_comparisons")
                            print(
                                f"🔍 DEBUG: chapter_comparisons type: {type(chapter_comps)}")
                            if chapter_comps:
                                print(
                                    f"🔍 DEBUG: chapter_comparisons length: {len(chapter_comps)}")
                                print(
                                    f"🔍 DEBUG: First comparison keys: {list(chapter_comps[0].keys()) if len(chapter_comps) > 0 else 'N/A'}")
                            else:
                                print(
                                    f"🔍 DEBUG: chapter_comparisons is None or empty")

                            completion_msg = {
                                "status": "complete",
                                "message": update_data.get("message", "✅ Script creation complete!"),
                                "script": streamer.result.get("enhanced_script", ""),
                                "script_title": streamer.result.get("script_title", "Untitled Script"),
                                "demo_packages": streamer.result.get("demo_packages"),
                                "word_count": streamer.result.get("word_count", 0),
                                "reading_time": streamer.result.get("reading_time", "N/A"),
                                "thumbnails": thumbnails,
                                "comparison_file": streamer.result.get("comparison_file"),
                                "chapter_comparisons": streamer.result.get("chapter_comparisons"),
                                "flow_original_script": streamer.result.get("flow_original_script"),
                                "flow_improved_script": streamer.result.get("flow_improved_script"),
                                "flow_analysis_report": streamer.result.get("flow_analysis_report"),
                                "edl_content": streamer.result.get("edl_content"),
                                "edl_filename": streamer.result.get("edl_filename"),
                                "curl_commands": streamer.result.get("curl_commands"),
                                "broll_images": streamer.result.get("broll_images")
                            }

                            print("\n" + "="*70)
                            print("📤 SENDING COMPLETION MESSAGE TO FRONTEND")
                            print("="*70)
                            print(f"Status: {completion_msg['status']}")
                            print(
                                f"Script length: {len(completion_msg['script'])} chars")
                            print(
                                f"Thumbnails array length: {len(completion_msg['thumbnails'])}")
                            if completion_msg['thumbnails']:
                                print("\nThumbnails being sent:")
                                for i, t in enumerate(completion_msg['thumbnails'], 1):
                                    print(
                                        f"   #{i}: {t['emotion']} - {t['filename']}")
                            else:
                                print("⚠️ No thumbnails in completion message!")

                            # Debug B-roll images
                            broll_imgs = completion_msg.get('broll_images')
                            if broll_imgs:
                                print(
                                    f"\n🎨 B-roll images being sent: {len(broll_imgs)} images")
                                for i, img in enumerate(broll_imgs, 1):
                                    print(
                                        f"   #{i}: {img.get('search_term')} - {img.get('filename')}")
                            else:
                                print(
                                    "\n⚠️ No B-roll images in completion message!")

                            print("="*70 + "\n")

                            logger.info(
                                f"✅ Sending completion with script ({len(completion_msg.get('script', ''))} chars) and {len(thumbnails)} thumbnails")
                        else:
                            # Error case
                            error_msg = streamer.result.get(
                                "error", "Unknown error") if streamer.result else "No result available"
                            completion_msg = {
                                "status": "error",
                                "error": error_msg
                            }
                            logger.error(
                                f"❌ Sending error completion: {error_msg}")

                        yield f"data: {json.dumps(completion_msg)}\n\n"
                        logger.info(
                            f"🎯 SSE completion message sent for session {session_id}")
                        break
                    else:
                        # Regular progress update
                        progress_msg = {
                            "status": "progress",
                            "message": update_data.get("message", "Processing..."),
                            "progress": update_data.get("progress", 0)
                        }
                        yield f"data: {json.dumps(progress_msg)}\n\n"

                except queue.Empty:
                    # Only check done status if queue is empty
                    if streamer.done:
                        logger.info(
                            f"🏁 SSE ending - streamer marked done for session {session_id}")
                        break
                    # Send keepalive
                    logger.debug(f"💓 SSE keepalive for session {session_id}")
                    yield f"data: {json.dumps({'status': 'keepalive'})}\n\n"
                    continue

        except Exception as e:
            # Suppress "Broken pipe" errors - these happen when client disconnects
            import errno
            if isinstance(e, (BrokenPipeError, IOError)) and getattr(e, 'errno', None) == errno.EPIPE:
                logger.debug(
                    f"🔌 Client disconnected (broken pipe) for session {session_id}")
            else:
                logger.error(f"SSE error for session {session_id}: {e}")
                yield f"data: {json.dumps({'status': 'error', 'error': str(e)})}\n\n"
        finally:
            logger.info(f"🔚 SSE stream ended for session {session_id}")

    response = Response(generate(), mimetype="text/event-stream")
    response.headers['Cache-Control'] = 'no-cache'
    response.headers['Connection'] = 'keep-alive'
    response.headers['X-Accel-Buffering'] = 'no'  # Disable nginx buffering
    return response


@app.route("/result/<session_id>")
def get_result(session_id):
    """Get final result"""
    logger.info(f"🔍 Result requested for session {session_id}")

    if session_id in progress_streams:
        streamer = progress_streams[session_id]
        logger.info(
            f"📊 Session {session_id} found. Result available: {bool(streamer.result)}")
        logger.info(f"📊 Session {session_id} done status: {streamer.done}")

        if streamer.result:
            logger.info(f"✅ Result found for session {session_id}")
            result_size = len(str(streamer.result)) if streamer.result else 0
            logger.info(f"📏 Result size: {result_size} chars")
            return jsonify(streamer.result)
        else:
            logger.warning(
                f"⚠️ No result available yet for session {session_id}")
            return jsonify({"error": "Result not ready yet, please try again"}), 202
    else:
        logger.error(f"❌ Session {session_id} not found in progress_streams")
        logger.info(f"📋 Available sessions: {list(progress_streams.keys())}")

    return jsonify({"error": "Result not available"}), 404


@app.route("/process-script", methods=["POST"])
def process_script():
    """Process existing script to generate YouTube details, B-roll, HeyGen, etc."""
    logger.info("=== PROCESS SCRIPT REQUEST RECEIVED ===")

    try:
        data = request.json
        script_content = data.get("script", "")
        # Note: topic field is only used for script CREATION, not processing
        audience = data.get("audience", "general audience")
        video_length = data.get("videoLength", "5-8 minutes")
        tone = data.get("tone", "informative")
        checkboxes = data.get("checkboxes", {})
        heygen_template_id = data.get("heygen_template_id", "")
        heygen_api_key = data.get("heygen_api_key", "")

        if not script_content.strip():
            logger.warning("❌ Empty script received")
            return jsonify({"success": False, "error": "Please provide a script to process"})

        logger.info(f"📋 Checkboxes: {checkboxes}")
        if heygen_template_id:
            logger.info(f"🎬 HeyGen Template: {heygen_template_id}")

        # Generate session ID
        session_id = str(uuid.uuid4())

        # Create progress streamer
        progress_streams[session_id] = ProgressStreamer(session_id)

        logger.info(f"🎬 Starting script processing: {session_id}")

        # Start script processing in background thread
        def run_script_processing():
            try:
                asyncio.run(process_existing_script(
                    session_id, script_content, audience, tone,
                    video_length, checkboxes, heygen_template_id,
                    heygen_api_key
                ))
            except Exception as e:
                logger.error(f"❌ Script processing error: {e}")
                if session_id in progress_streams:
                    progress_streams[session_id].send_update(
                        f"❌ Error: {str(e)}", -1
                    )

        thread = threading.Thread(target=run_script_processing, daemon=True)
        running_tasks[session_id] = thread
        thread.start()

        return jsonify({
            "success": True,
            "session_id": session_id,
            "message": "Script processing started"
        })

    except Exception as e:
        error_msg = f"Script processing request failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({"success": False, "error": error_msg})


@app.route("/export_markdown", methods=["POST"])
def export_markdown():
    """Export script content as Markdown file"""
    logger.info("=== MARKDOWN EXPORT REQUEST RECEIVED ===")

    try:
        data = request.json
        script_content = data.get("script_content", "")
        title = data.get("title", "AI_Script")

        if not script_content:
            logger.warning("❌ No script content provided for Markdown export")
            return jsonify(
                {"success": False, "error": "No script content to export"}
            )

        logger.info(f"📝 Exporting script to Markdown: {title}")

        # Clean up title for filename
        safe_title = "".join(
            c for c in title if c.isalnum() or c in (' ', '-', '_')
        ).strip()
        safe_title = safe_title.replace(' ', '_')

        # Create Markdown content with proper formatting
        markdown_content = f"# {title}\n\n"
        markdown_content += f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n"
        markdown_content += "---\n\n"
        markdown_content += script_content

        # Create in-memory file
        markdown_bytes = markdown_content.encode('utf-8')
        markdown_io = BytesIO(markdown_bytes)

        logger.info(
            f"✅ Markdown file created successfully ({len(markdown_bytes)} bytes)"
        )

        # Send file
        response = send_file(
            markdown_io,
            mimetype='text/markdown',
            as_attachment=True,
            download_name=f'{safe_title}.md'
        )

        return response

    except Exception as e:
        error_msg = f"Markdown export failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({"success": False, "error": error_msg})


@app.route("/download_comparison", methods=["POST"])
def download_comparison():
    """Download the chapter comparison report file"""
    logger.info("=== COMPARISON FILE DOWNLOAD REQUEST RECEIVED ===")

    try:
        data = request.json
        file_path = data.get("file_path", "")

        if not file_path:
            logger.warning("❌ No file path provided for comparison download")
            return jsonify(
                {"success": False, "error": "No file path provided"}
            ), 400

        # Security: Ensure the file is in the current directory and is a comparison file
        import os
        file_name = os.path.basename(file_path)
        if not file_name.startswith("chapter_comparison_") or not file_name.endswith(".md"):
            logger.warning(f"❌ Invalid file name: {file_name}")
            return jsonify(
                {"success": False, "error": "Invalid comparison file"}
            ), 400

        # Check if file exists
        if not os.path.exists(file_path):
            logger.warning(f"❌ File not found: {file_path}")
            return jsonify(
                {"success": False, "error": "Comparison file not found"}
            ), 404

        logger.info(f"📊 Sending comparison file: {file_name}")

        # Send the file
        return send_file(
            file_path,
            mimetype='text/markdown',
            as_attachment=True,
            download_name=file_name
        )

    except Exception as e:
        error_msg = f"Comparison file download failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({"success": False, "error": error_msg}), 500


@app.route("/export_heygen_curl", methods=["POST"])
def export_heygen_curl():
    """Extract and export HeyGen curl commands from script"""
    logger.info("=== HEYGEN CURL EXPORT REQUEST RECEIVED ===")

    try:
        data = request.json
        script_content = data.get("script_content", "")

        if not script_content:
            logger.warning(
                "❌ No script content provided for HeyGen curl export")
            return jsonify({"error": "No script content to export"}), 400

        # Extract HeyGen section
        heygen_section_start = script_content.find("# 🎬 HEYGEN READY SCRIPT")

        if heygen_section_start == -1:
            logger.warning("❌ No HeyGen section found in script")
            return jsonify({"error": "No HeyGen section found in this script. Make sure the script has been generated with HeyGen content."}), 404

        # Find curl commands section (try multiple markers)
        curl_markers = [
            "# 🚀 HEYGEN API CURL COMMANDS",
            "CURL COMMANDS:",
            "# HEYGEN API CURL COMMANDS"
        ]

        curl_start = -1
        marker_found = None
        for marker in curl_markers:
            pos = script_content.find(marker, heygen_section_start)
            if pos != -1:
                curl_start = pos
                marker_found = marker
                break

        if curl_start == -1:
            logger.warning("❌ No curl commands found in HeyGen section")
            msg = "No curl commands found in the HeyGen section."
            return jsonify({"error": msg}), 404

        # Extract curl commands (skip the marker line, get content after)
        curl_commands = script_content[curl_start:]

        # Curl commands now run to the end of the script (after section order fix)
        # No need to look for Demo Packages section as it comes BEFORE HeyGen now
        curl_commands = curl_commands.strip()

        if not curl_commands:
            logger.warning("❌ Curl commands section is empty")
            return jsonify({"error": "HeyGen curl commands section is empty."}), 404

        logger.info(
            f"✅ Extracted {len(curl_commands)} characters of curl commands")

        # Create in-memory text file
        curl_bytes = curl_commands.encode('utf-8')
        curl_io = BytesIO(curl_bytes)

        # Send file
        response = send_file(
            curl_io,
            mimetype='text/plain',
            as_attachment=True,
            download_name='heygen_curl_commands.txt'
        )

        logger.info("✅ HeyGen curl commands exported successfully")
        return response

    except Exception as e:
        error_msg = f"HeyGen curl export failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({"error": error_msg}), 500


@app.route("/test_word")
def test_word_page():
    """Test page for Word export functionality"""
    return render_template("test_word_export.html")


@app.route("/export_word", methods=["POST"])
def export_word():
    """Export script content to Word document"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No data provided"})

        script_content = data.get("script_content", "")
        title = data.get("title", "AI Script")

        if not script_content:
            return jsonify({"success": False, "error": "No script content to export"})

        logger.info(
            f"📄 Exporting script to Word: {len(script_content)} characters")

        # Import Word processing libraries directly
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            logger.info("✅ Word processing libraries imported successfully")
        except ImportError as e:
            logger.error(f"❌ Failed to import docx libraries: {e}")
            return jsonify({"success": False, "error": "python-docx not available"})

        # Create temporary file path
        import tempfile
        temp_dir = Path(tempfile.gettempdir())
        temp_file = temp_dir / f"script_{uuid.uuid4().hex[:8]}.docx"

        # Convert script to Word document directly
        try:
            logger.info("📝 Creating Word document...")

            # Create a new Word document
            doc = Document()

            # Process the script content line by line
            lines = script_content.split('\n')
            logger.info(f"🔍 Processing {len(lines)} lines")

            i = 0
            while i < len(lines):
                line = lines[i].rstrip()

                # Skip empty lines
                if not line:
                    i += 1
                    continue

                # Handle markdown tables - detect header row with pipes
                if '|' in line:
                    # Look ahead to check if next line is a separator
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].rstrip()
                        # Check if next line is a table separator (pipes, dashes, colons, spaces)
                        if '|' in next_line and all(c in '|-: \t' for c in next_line.replace('|', '')):
                            # This is a markdown table!
                            logger.info(f"📊 Detected table at line {i}")

                            # Collect header row
                            header_line = line
                            i += 1  # Skip to separator line
                            i += 1  # Skip separator, move to first data row

                            # Collect all data rows (lines with pipes)
                            data_lines = []
                            while i < len(lines):
                                data_line = lines[i].rstrip()
                                if '|' in data_line:
                                    data_lines.append(data_line)
                                    i += 1
                                else:
                                    break

                            # Parse the table
                            # Split on pipes and clean up
                            header_cells = [cell.strip()
                                            for cell in header_line.split('|')]
                            # Remove empty cells from start/end if table has leading/trailing pipes
                            if header_cells and not header_cells[0]:
                                header_cells = header_cells[1:]
                            if header_cells and not header_cells[-1]:
                                header_cells = header_cells[:-1]

                            # Parse data rows
                            data_rows = []
                            for data_line in data_lines:
                                cells = [cell.strip()
                                         for cell in data_line.split('|')]
                                # Remove empty cells from start/end
                                if cells and not cells[0]:
                                    cells = cells[1:]
                                if cells and not cells[-1]:
                                    cells = cells[:-1]
                                if cells:  # Only add non-empty rows
                                    data_rows.append(cells)

                            # Create Word table
                            if header_cells and data_rows:
                                logger.info(
                                    f"📊 Creating Word table: {len(header_cells)} columns, {len(data_rows)} rows")

                                table = doc.add_table(
                                    rows=1 + len(data_rows), cols=len(header_cells))
                                table.style = 'Light Grid Accent 1'

                                # Add headers
                                for col_idx, header in enumerate(header_cells):
                                    if col_idx < len(table.rows[0].cells):
                                        cell = table.rows[0].cells[col_idx]
                                        cell.text = header
                                        # Make header bold
                                        for paragraph in cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.bold = True

                                # Add data rows
                                for row_idx, row_data in enumerate(data_rows, start=1):
                                    for col_idx, cell_text in enumerate(row_data):
                                        if col_idx < len(header_cells) and row_idx < len(table.rows):
                                            table.rows[row_idx].cells[col_idx].text = cell_text

                                # Add spacing after table
                                doc.add_paragraph()
                                logger.info(f"✅ Table created successfully")
                            else:
                                logger.warning(
                                    f"⚠️ Table parsing failed: headers={len(header_cells)}, rows={len(data_rows)}")

                            continue

                # Handle horizontal rules (---)
                elif line.strip() in ['---', '___', '***']:
                    doc.add_paragraph('_' * 50)
                    i += 1
                    continue

                # Handle headers (lines starting with #)
                elif line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('#').strip()

                    if level == 1:  # Main title
                        paragraph = doc.add_heading(text, level=1)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif level == 2:  # Chapter/Section
                        paragraph = doc.add_heading(text, level=2)
                    else:  # Sub-sections
                        paragraph = doc.add_heading(text, level=3)
                    i += 1
                    continue

                # Handle bold text (**text**)
                elif '**' in line:
                    paragraph = doc.add_paragraph()
                    parts = line.split('**')
                    for idx, part in enumerate(parts):
                        if idx % 2 == 0:  # Normal text
                            if part:
                                paragraph.add_run(part)
                        else:  # Bold text
                            if part:
                                run = paragraph.add_run(part)
                                run.bold = True
                    i += 1
                    continue

                # Handle visual cues (lines with 🎬 emoji or [Visual:])
                elif ('🎬' in line or '[Visual:' in line or
                      (line.startswith('[') and ']' in line)):
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.italic = True
                    paragraph.space_before = Pt(6)
                    paragraph.space_after = Pt(6)
                    i += 1
                    continue

                # Regular paragraph
                else:
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(line)
                    i += 1

            # Save document to temporary file
            doc.save(str(temp_file))
            logger.info(f"💾 Word document saved to: {temp_file}")

            if not temp_file.exists():
                return jsonify({"success": False, "error": "Failed to create Word document"})

            # Read the file and return as binary response
            with open(temp_file, 'rb') as f:
                file_content = f.read()

            # Clean up temp file
            temp_file.unlink()

            logger.info(f"✅ Word export successful: {len(file_content)} bytes")

            # Return file as binary response
            response = make_response(file_content)
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename="{title}.docx"'
            return response

        except Exception as e:
            logger.error(f"❌ Word conversion failed: {e}")
            return jsonify({"success": False, "error": f"Word conversion failed: {str(e)}"})

    except Exception as e:
        error_msg = f"Export to Word failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({"success": False, "error": error_msg})


@app.route("/test-thumbnails")
def test_thumbnails():
    """Test page to view thumbnail gallery with existing thumbnails"""
    try:
        thumbnail_dir = Path.home() / "Dev" / "Thumbnails"

        # Get all PNG files
        if thumbnail_dir.exists():
            png_files = list(thumbnail_dir.glob("*.png"))
            # Take the most recent 6 files
            png_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
            png_files = png_files[:6]

            # Create thumbnail data
            thumbnails = []
            emotions = ["ANGRY", "SHOCKED", "SCARED",
                        "EXCITED", "SKEPTICAL", "DETERMINED"]
            for i, file_path in enumerate(png_files):
                thumbnails.append({
                    "emotion": emotions[i] if i < len(emotions) else f"VARIANT_{i+1}",
                    "text": f"Test thumbnail {i+1}",
                    "filename": file_path.name
                })

            return render_template("test_thumbnails.html", thumbnails=thumbnails)
        else:
            return "Thumbnail directory not found", 404
    except Exception as e:
        return f"Error: {str(e)}", 500


@app.route("/thumbnail-test")
def thumbnail_test():
    """Interactive thumbnail generation test page"""
    return render_template("thumbnail_test.html")


@app.route("/generate-test-thumbnails", methods=["POST"])
def generate_test_thumbnails():
    """Generate thumbnails for testing the gallery"""
    try:
        data = request.json or {}
        topic = data.get("topic", "AI Testing")

        print(f"\n{'='*70}")
        print(f"🎬 GENERATING TEST THUMBNAILS")
        print(f"📝 Topic: {topic}")
        print(f"{'='*70}\n")

        from tools.media.emotional_thumbnail_generator import (
            EmotionalThumbnailGenerator,
        )

        thumbnail_gen = EmotionalThumbnailGenerator()
        print(f"✅ Generator initialized")
        print(f"   Template: {thumbnail_gen.template_path}")
        print(f"   Output: {thumbnail_gen.output_dir}")

        thumbnail_results = thumbnail_gen.generate_all_thumbnails(
            script_title=topic,
            script_content="Test script content",
            youtube_upload_details=None
        )

        if thumbnail_results and thumbnail_results.get("variations"):
            variations = thumbnail_results["variations"]
            print(f"\n✅ Generated {len(variations)} thumbnails successfully!")

            return jsonify({
                "success": True,
                "thumbnails": variations,
                "count": len(variations)
            })
        else:
            print(f"\n⚠️ No thumbnails generated")
            return jsonify({
                "success": False,
                "error": "No thumbnails generated",
                "results": thumbnail_results
            })

    except Exception as e:
        print(f"\n❌ Error generating test thumbnails: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route("/thumbnails/<filename>")
def serve_thumbnail(filename):
    """Serve generated thumbnail images"""
    try:
        thumbnail_dir = Path.home() / "Dev" / "Thumbnails"
        file_path = thumbnail_dir / filename

        if file_path.exists() and file_path.suffix.lower() in ['.png', '.jpg', '.jpeg']:
            return send_file(file_path, mimetype='image/png')
        else:
            return jsonify({"error": "Thumbnail not found"}), 404
    except Exception as e:
        logger.error(f"Error serving thumbnail: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/broll-images/<filename>")
def serve_broll_image(filename):
    """Serve generated B-roll images"""
    try:
        broll_dir = Path.home() / "Dev" / "brollimages"
        file_path = broll_dir / filename

        if file_path.exists() and file_path.suffix.lower() in ['.png', '.jpg', '.jpeg', '.webp']:
            # Determine mimetype based on extension
            mimetype = 'image/png'
            if file_path.suffix.lower() in ['.jpg', '.jpeg']:
                mimetype = 'image/jpeg'
            elif file_path.suffix.lower() == '.webp':
                mimetype = 'image/webp'

            return send_file(file_path, mimetype=mimetype)
        else:
            return jsonify({"error": "B-roll image not found"}), 404
    except Exception as e:
        logger.error(f"Error serving B-roll image: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/create_resolve_project", methods=["POST"])
def create_resolve_project():
    """Create a new DaVinci Resolve project with automated setup"""
    logger.info("=== DAVINCI RESOLVE PROJECT CREATION REQUEST ===")

    try:
        data = request.json
        script_title = data.get("script_title", "AI_Video_Project")
        edl_content = data.get("edl_content")
        edl_filename = data.get("edl_filename")

        logger.info(f"📊 Script Title: {script_title}")
        logger.info(f"📄 EDL Filename: {edl_filename}")
        logger.info(
            f"📄 EDL Content Length: {len(edl_content) if edl_content else 0}")

        # Save EDL content to temp file if provided
        edl_file_path = None
        if edl_content:
            import tempfile
            # Create temp file with .edl extension
            temp_fd, edl_file_path = tempfile.mkstemp(suffix='.edl', text=True)
            try:
                with os.fdopen(temp_fd, 'w') as f:
                    f.write(edl_content)
                logger.info(f"💾 Saved EDL to temp file: {edl_file_path}")
            except Exception as write_error:
                logger.error(f"❌ Failed to write EDL temp file: {write_error}")
                edl_file_path = None

        # Import the DaVinci Resolve API module
        from davinci_resolve_api import create_resolve_project as create_project

        # Create the project with EDL file path
        result = create_project(script_title, edl_file_path)

        # Clean up temp file
        if edl_file_path and os.path.exists(edl_file_path):
            try:
                os.unlink(edl_file_path)
                logger.info(f"🗑️ Deleted temp EDL file: {edl_file_path}")
            except Exception as cleanup_error:
                logger.warning(
                    f"⚠️ Failed to delete temp EDL file: {cleanup_error}")

        if result.get("success"):
            logger.info(f"✅ Project created: {result.get('project_name')}")
            return jsonify(result)
        else:
            logger.error(f"❌ Project creation failed: {result.get('error')}")
            return jsonify(result), 500

    except ImportError as import_error:
        error_msg = f"Failed to import DaVinci Resolve API: {str(import_error)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500
    except Exception as e:
        error_msg = f"Unexpected error: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/check_aroll_videos", methods=["POST"])
def check_aroll_videos():
    """Check if aRoll folder has any videos"""
    logger.info("=== CHECKING AROLL VIDEOS ===")

    try:
        data = request.json
        script_title = data.get("script_title", "")

        if not script_title:
            return jsonify({
                "success": False,
                "error": "No script title provided"
            }), 400

        # Sanitize title for folder name
        import re
        safe_title = re.sub(r'[<>:"/\\|?*]', '_', script_title)
        safe_title = safe_title.replace(' ', '_')

        # Build path to aroll folder
        base_path = os.path.expanduser("~/Dev/Videos/Edited/Final")
        aroll_path = os.path.join(base_path, safe_title, "aroll")

        logger.info(f"📁 Script title: '{script_title}'")
        logger.info(f"📁 Safe title: '{safe_title}'")
        logger.info(f"📁 Checking aroll path: {aroll_path}")
        logger.info(f"📁 Path exists: {os.path.exists(aroll_path)}")

        if not os.path.exists(aroll_path):
            # List what's actually in the base path to help debug
            if os.path.exists(base_path):
                logger.info(f"📁 Base path exists, folders found:")
                for folder in os.listdir(base_path):
                    logger.info(f"   - {folder}")
            else:
                logger.info(f"📁 Base path does not exist: {base_path}")

            logger.info(f"📁 aRoll folder does not exist yet")
            return jsonify({
                "success": True,
                "video_count": 0,
                "videos": []
            })

        # Count .mp4 files in aRoll folder
        video_files = [f for f in os.listdir(aroll_path) if f.endswith('.mp4')]

        logger.info(f"✅ Found {len(video_files)} video(s) in aRoll folder:")
        for vf in video_files:
            logger.info(f"   - {vf}")

        return jsonify({
            "success": True,
            "video_count": len(video_files),
            "videos": video_files
        })

    except Exception as e:
        error_msg = f"Error checking aRoll videos: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/create_resolve_with_videos", methods=["POST"])
def create_resolve_with_videos():
    """Create a DaVinci Resolve project with aRoll videos added to timeline"""
    logger.info("=== DAVINCI RESOLVE PROJECT CREATION WITH VIDEOS ===")

    try:
        data = request.json
        script_title = data.get("script_title", "AI_Video_Project")
        edl_content = data.get("edl_content")
        edl_filename = data.get("edl_filename")

        logger.info(f"📊 Script Title: {script_title}")

        # Sanitize title for folder name
        import re
        safe_title = re.sub(r'[<>:"/\\|?*]', '_', script_title)
        safe_title = safe_title.replace(' ', '_')

        # Build path to aroll folder
        base_path = os.path.expanduser("~/Dev/Videos/Edited/Final")
        aroll_path = os.path.join(base_path, safe_title, "aroll")

        # Get list of video files
        video_files = []
        if os.path.exists(aroll_path):
            video_files = [f for f in os.listdir(
                aroll_path) if f.endswith('.mp4')]
            logger.info(f"📹 Found {len(video_files)} video(s) in aroll folder")

        # Sort videos by chapter/part order
        def parse_chapter_info(filename):
            """Extract chapter and part numbers from filename like heygen_Ch1p1_abc123.mp4"""
            import re
            # Try to match pattern: Ch{X}p{Y} where X is chapter, Y is part
            match = re.search(r'Ch(\d+)p(\d+)', filename, re.IGNORECASE)
            if match:
                chapter_num = int(match.group(1))
                part_num = int(match.group(2))
                return (chapter_num, part_num)

            # Try to match just Chapter_{X}
            match = re.search(r'Chapter[_\s]+(\d+)', filename, re.IGNORECASE)
            if match:
                chapter_num = int(match.group(1))
                return (chapter_num, 0)

            # Default to end of list
            return (999, 999)

        # Sort videos by chapter order
        sorted_videos = sorted(video_files, key=parse_chapter_info)
        logger.info(f"📋 Sorted video order: {sorted_videos}")

        # Save EDL content to temp file if provided
        edl_file_path = None
        if edl_content:
            import tempfile
            temp_fd, edl_file_path = tempfile.mkstemp(suffix='.edl', text=True)
            try:
                with os.fdopen(temp_fd, 'w') as f:
                    f.write(edl_content)
                logger.info(f"💾 Saved EDL to temp file: {edl_file_path}")
            except Exception as write_error:
                logger.error(f"❌ Failed to write EDL temp file: {write_error}")
                edl_file_path = None

        # Import the DaVinci Resolve API module (force reload to get latest changes)
        import importlib
        import sys
        if 'davinci_resolve_api' in sys.modules:
            logger.info("🔄 Reloading davinci_resolve_api module...")
            import davinci_resolve_api
            importlib.reload(davinci_resolve_api)
            from davinci_resolve_api import create_resolve_project_with_videos
            logger.info("✅ Module reloaded successfully")
        else:
            logger.info(
                "📦 Loading davinci_resolve_api module for first time...")
            from davinci_resolve_api import create_resolve_project_with_videos

        logger.info(f"🎬 Calling create_resolve_project_with_videos...")
        logger.info(f"   Script title: {script_title}")
        logger.info(f"   EDL file: {edl_file_path}")
        logger.info(f"   aRoll path: {aroll_path}")
        logger.info(f"   Videos to add: {len(sorted_videos)}")

        # Create the project with videos
        result = create_resolve_project_with_videos(
            script_title,
            edl_file_path,
            aroll_path,
            sorted_videos
        )

        # Clean up temp file
        if edl_file_path and os.path.exists(edl_file_path):
            try:
                os.unlink(edl_file_path)
                logger.info(f"🗑️ Deleted temp EDL file: {edl_file_path}")
            except Exception as cleanup_error:
                logger.warning(
                    f"⚠️ Failed to delete temp EDL file: {cleanup_error}")

        if result.get("success"):
            logger.info(
                f"✅ Project created with {len(sorted_videos)} videos: {result.get('project_name')}")
            return jsonify(result)
        else:
            logger.error(f"❌ Project creation failed: {result.get('error')}")
            return jsonify(result), 500

    except ImportError as import_error:
        error_msg = f"Failed to import DaVinci Resolve API: {str(import_error)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500
    except Exception as e:
        error_msg = f"Unexpected error: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/execute_curl", methods=["POST"])
def execute_curl():
    """Execute a curl command and return the result"""
    logger.info("=== CURL COMMAND EXECUTION REQUEST ===")

    try:
        data = request.json
        curl_command = data.get("curl_command", "")

        if not curl_command:
            return jsonify({
                "success": False,
                "error": "No curl command provided"
            }), 400

        logger.info(
            f"📤 Executing curl command (first 100 chars): {curl_command[:100]}...")

        # Execute the curl command using subprocess
        import subprocess
        import shlex

        # Parse the curl command
        # Remove leading/trailing whitespace and newlines
        curl_command = curl_command.strip()

        # Execute the command
        try:
            result = subprocess.run(
                curl_command,
                shell=True,
                capture_output=True,
                text=True,
                timeout=30
            )

            # Parse the response
            if result.returncode == 0:
                # Try to parse JSON response
                try:
                    import json
                    response_data = json.loads(result.stdout)

                    # Extract job ID if available (HeyGen specific)
                    job_id = None
                    if isinstance(response_data, dict):
                        job_id = response_data.get('data', {}).get('video_id') or \
                            response_data.get('video_id') or \
                            response_data.get('job_id') or \
                            response_data.get('id')

                    logger.info(f"✅ Curl command executed successfully")
                    if job_id:
                        logger.info(f"📋 Job ID: {job_id}")

                    return jsonify({
                        "success": True,
                        "response": response_data,
                        "job_id": job_id,
                        "raw_output": result.stdout
                    })
                except json.JSONDecodeError:
                    # Not JSON, return raw output
                    logger.info(f"✅ Curl command executed (non-JSON response)")
                    return jsonify({
                        "success": True,
                        "raw_output": result.stdout
                    })
            else:
                error_msg = result.stderr or "Command failed with no error message"
                logger.error(f"❌ Curl command failed: {error_msg}")
                return jsonify({
                    "success": False,
                    "error": error_msg,
                    "return_code": result.returncode
                }), 500

        except subprocess.TimeoutExpired:
            logger.error("❌ Curl command timed out after 30 seconds")
            return jsonify({
                "success": False,
                "error": "Command timed out after 30 seconds"
            }), 500

    except Exception as e:
        error_msg = f"Error executing curl command: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/setup_project", methods=["POST"])
def setup_project():
    """Set up project folder structure and copy template files"""
    logger.info("=== PROJECT SETUP REQUEST ===")

    try:
        data = request.json
        script_title = data.get("script_title", "")
        script_content = data.get("script_content", "")

        if not script_title:
            return jsonify({
                "success": False,
                "error": "No script title provided"
            }), 400

        import os
        import shutil
        from pathlib import Path

        # Create safe directory name
        safe_title = "".join(
            c for c in script_title
            if c.isalnum() or c in (' ', '-', '_')
        ).rstrip()
        safe_title = safe_title.replace(' ', '_')

        # Define paths
        base_path = Path.home() / "Dev" / "Videos" / "Edited" / "Final"
        project_path = base_path / safe_title
        template_path = Path.home() / "Dev" / "Davinci" / "Template"

        logger.info(f"📁 Creating project folder: {project_path}")

        # Create project directory (remove if exists)
        if project_path.exists():
            logger.info(f"⚠️ Removing existing project: {project_path}")
            shutil.rmtree(project_path)

        project_path.mkdir(parents=True, exist_ok=True)

        # Copy template files
        if template_path.exists():
            logger.info(f"📋 Copying template from: {template_path}")
            file_count = sum(
                1 for _ in template_path.rglob('*') if _.is_file())
            logger.info(f"📦 Found {file_count} files to copy...")

            copied_count = 0
            for item in template_path.iterdir():
                dest = project_path / item.name
                if item.is_dir():
                    # Count files in directory
                    dir_files = sum(1 for _ in item.rglob('*') if _.is_file())
                    logger.info(
                        f"📁 Copying directory: {item.name} ({dir_files} files)..."
                    )
                    shutil.copytree(item, dest, dirs_exist_ok=True)
                    copied_count += dir_files
                else:
                    shutil.copy2(item, dest)
                    copied_count += 1

                # Log progress every 100 files
                if copied_count % 100 == 0:
                    logger.info(
                        f"📊 Progress: {copied_count}/{file_count} files copied")

            logger.info(f"✅ Template files copied ({copied_count} total)")
        else:
            logger.warning(
                f"⚠️ Template path not found: {template_path}"
            )

        # Create script subdirectory
        script_dir = project_path / "script"
        script_dir.mkdir(exist_ok=True)

        # Copy B-roll images to images folder if they exist
        broll_images_dir = Path.home() / "Dev" / "brollimages"
        images_dir = project_path / "images"

        if broll_images_dir.exists():
            images_dir.mkdir(exist_ok=True)

            # Get all image files from broll images directory
            image_files = list(broll_images_dir.glob("*.png")) + \
                list(broll_images_dir.glob("*.jpg")) + \
                list(broll_images_dir.glob("*.jpeg")) + \
                list(broll_images_dir.glob("*.webp"))

            if image_files:
                logger.info(
                    f"📸 Found {len(image_files)} B-roll images to copy")
                copied_count = 0

                for image_file in image_files:
                    try:
                        dest_file = images_dir / image_file.name
                        shutil.copy2(image_file, dest_file)
                        copied_count += 1
                    except Exception as copy_error:
                        logger.warning(
                            f"⚠️ Could not copy {image_file.name}: {copy_error}")

                logger.info(
                    f"✅ Copied {copied_count} B-roll images to project")
            else:
                logger.info("ℹ️ No B-roll images found to copy")
        else:
            logger.info(
                f"ℹ️ B-roll images directory not found: {broll_images_dir}")

        # Save script as Word document if content provided
        if script_content:
            import sys
            import asyncio

            # Import word processing from console_ui
            sys.path.insert(
                0, str(Path(__file__).parent.parent / "console_ui")
            )
            from word_processing import convert_markdown_to_word

            script_file = script_dir / f"{safe_title}.docx"
            logger.info(f"📝 Creating script document: {script_file}")

            # Convert and save (handle async function)
            try:
                asyncio.run(convert_markdown_to_word(
                    markdown_content=script_content,
                    output_file_path=str(script_file),
                    template_path=None,
                    title=script_title
                ))
                logger.info(f"✅ Script saved: {script_file}")
            except Exception as word_error:
                logger.error(f"❌ Word document error: {word_error}")
                # Continue anyway - project still set up

        return jsonify({
            "success": True,
            "project_path": str(project_path),
            "script_dir": str(script_dir)
        })

    except Exception as e:
        error_msg = f"Error setting up project: {str(e)}"
        logger.error(f"❌ {error_msg}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/check_video_status", methods=["POST"])
def check_video_status():
    """Check HeyGen video generation status and download when complete"""
    logger.info("=== VIDEO STATUS CHECK REQUEST ===")

    try:
        data = request.json
        video_id = data.get("video_id", "")
        api_key = data.get("api_key", "")

        if not video_id:
            return jsonify({
                "success": False,
                "error": "No video ID provided"
            }), 400

        if not api_key:
            return jsonify({
                "success": False,
                "error": "No API key provided"
            }), 400

        logger.info(f"📹 Checking status for video ID: {video_id}")

        import requests

        # Check video status
        base_url = "https://api.heygen.com/v1/video_status.get"
        video_status_url = f"{base_url}?video_id={video_id}"
        headers = {
            "X-Api-Key": api_key
        }

        try:
            response = requests.get(
                video_status_url, headers=headers, timeout=10
            )
            response_data = response.json()

            if response.status_code != 200:
                error_msg = response_data.get("message", "Unknown error")
                logger.error(f"❌ HeyGen API error: {error_msg}")
                return jsonify({
                    "success": False,
                    "error": f"HeyGen API error: {error_msg}"
                }), 500

            status = response_data.get("data", {}).get("status", "unknown")
            logger.info(f"📊 Video status: {status}")

            result = {
                "success": True,
                "status": status,
                "video_id": video_id
            }

            if status == "completed":
                video_url = response_data.get("data", {}).get("video_url")
                thumb_url = response_data.get("data", {}).get(
                    "thumbnail_url"
                )

                if video_url:
                    result["video_url"] = video_url
                    result["thumbnail_url"] = thumb_url
                    logger.info(f"✅ Video completed! URL: {video_url}")
                else:
                    msg = "⚠️ Video marked as completed but no URL"
                    logger.warning(msg)

            elif status == "failed":
                error = response_data.get("data", {}).get(
                    "error", "Unknown error"
                )
                result["error"] = error
                logger.error(f"❌ Video generation failed: {error}")

            return jsonify(result)

        except requests.RequestException as req_error:
            error_msg = f"Network error: {str(req_error)}"
            logger.error(f"❌ {error_msg}")
            return jsonify({
                "success": False,
                "error": error_msg
            }), 500

    except Exception as e:
        error_msg = f"Error checking video status: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/download_heygen_video", methods=["POST"])
def download_heygen_video():
    """Download completed HeyGen video to project folder"""
    logger.info("=== VIDEO DOWNLOAD REQUEST ===")

    try:
        data = request.json
        video_url = data.get("video_url", "")
        video_id = data.get("video_id", "")
        chapter_name = data.get("chapter_name", "video")
        script_title = data.get("script_title", "")

        if not video_url:
            return jsonify({
                "success": False,
                "error": "No video URL provided"
            }), 400

        logger.info(f"⬇️ Downloading video: {video_id}")

        import requests
        from pathlib import Path

        try:
            # Download video content
            video_response = requests.get(video_url, timeout=60)

            if video_response.status_code != 200:
                status_code = video_response.status_code
                logger.error(f"❌ Failed to download video: {status_code}")
                return jsonify({
                    "success": False,
                    "error": f"Failed to download video: {status_code}"
                }), 500

            # Create safe names
            safe_chars = (' ', '-', '_')
            safe_chapter_name = "".join(
                c for c in chapter_name
                if c.isalnum() or c in safe_chars
            ).rstrip()

            safe_title = "".join(
                c for c in script_title
                if c.isalnum() or c in safe_chars
            ).rstrip().replace(' ', '_') if script_title else "default"

            filename = f"heygen_{safe_chapter_name}_{video_id[:8]}.mp4"

            # Save to project aroll folder if script title provided
            if script_title:
                base_path = Path.home() / "Dev" / "Videos" / "Edited"
                base_path = base_path / "Final"
                project_path = base_path / safe_title
                aroll_path = project_path / "aroll"

                # Create project folders if they don't exist
                aroll_path.mkdir(parents=True, exist_ok=True)
                broll_path = project_path / "broll"
                broll_path.mkdir(parents=True, exist_ok=True)
                images_path = project_path / "images"
                images_path.mkdir(parents=True, exist_ok=True)

                # Save video file
                video_file = aroll_path / filename
                with open(video_file, 'wb') as f:
                    f.write(video_response.content)

                logger.info(
                    f"✅ Video saved to project: {video_file}"
                )

                return jsonify({
                    "success": True,
                    "filename": filename,
                    "path": str(video_file)
                })
            else:
                # No script title - return as download
                from io import BytesIO
                video_buffer = BytesIO(video_response.content)
                video_buffer.seek(0)

                logger.info(
                    f"✅ Video downloaded successfully: {filename}"
                )

                return send_file(
                    video_buffer,
                    mimetype='video/mp4',
                    as_attachment=True,
                    download_name=filename
                )

        except requests.RequestException as req_error:
            error_msg = f"Network error: {str(req_error)}"
            logger.error(f"❌ {error_msg}")
            return jsonify({
                "success": False,
                "error": error_msg
            }), 500

    except Exception as e:
        error_msg = f"Error downloading video: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/test-comparisons")
def test_comparisons():
    """Test endpoint to preview the chapter comparisons tab without running full workflow"""

    # Mock chapter comparison data
    mock_comparisons = [
        {
            'chapter_num': 1,
            'original': '''Welcome to our channel! Today we're diving into an amazing topic that will change how you think about productivity.

Let's start with the basics and understand why this matters.''',
            'revised': '''Welcome to our channel! Today we're exploring a game-changing approach to productivity that actually works.

Let's dive into the fundamentals and discover why this is so important for your success.''',
            'feedback': '''Reviewer Feedback:
- Opening hook needs more impact
- Changed "amazing topic" to "game-changing approach" for stronger language
- Adjusted second paragraph for better flow'''
        },
        {
            'chapter_num': 2,
            'original': '''The first key principle is understanding your goals. You need to know what you want to achieve.

Without clear goals, you'll struggle to make progress.''',
            'revised': '''The first crucial principle is crystal-clear goal setting. You must define exactly what success looks like.

Without specific, measurable goals, you're just wandering without direction.''',
            'feedback': '''Reviewer Feedback:
- Made language more specific and actionable
- Changed "understanding" to "crystal-clear goal setting"
- Strengthened the consequence statement'''
        },
        {
            'chapter_num': 3,
            'original': '''Now let's talk about time management. It's important to schedule your day properly.

Make sure you allocate time for important tasks.''',
            'revised': '''Here's where time management becomes your superpower. Block out your day strategically, not randomly.

Protect prime hours for your most critical, high-impact work.''',
            'feedback': '''Reviewer Feedback:
- Added energy to the opening with "superpower"
- Changed passive "schedule" to active "block out"
- Made the advice more specific and actionable'''
        }
    ]

    # Mock thumbnails for testing
    mock_thumbnails = [
        {
            'id': 'test-thumb-1',
            'url': 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMTI4MCIgaGVpZ2h0PSI3MjAiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+PHJlY3Qgd2lkdGg9IjEyODAiIGhlaWdodD0iNzIwIiBmaWxsPSIjMWU0MGFmIi8+PHRleHQgeD0iNTAlIiB5PSI1MCUiIGZvbnQtZmFtaWx5PSJBcmlhbCIgZm9udC1zaXplPSI4MCIgZmlsbD0id2hpdGUiIHRleHQtYW5jaG9yPSJtaWRkbGUiIGR5PSIuM2VtIj5UZXN0IFRodW1ibmFpbCAxPC90ZXh0Pjwvc3ZnPg==',
            'style': 'bold'
        }
    ]

    # Render the main template
    return render_template('index.html',
                           test_mode=True,
                           test_comparisons=mock_comparisons,
                           test_thumbnails=mock_thumbnails)


if __name__ == "__main__":
    # Use environment variable for port, default to 8080 for local, 5007 for container
    port = int(os.environ.get("PORT", "8080"))

    logger.info("🎬 Starting ScriptCraft Web GUI (Universal Version)...")
    logger.info("📦 Version: %s", VERSION)
    logger.info("🚀 Server starting on http://0.0.0.0:%d", port)

    if port == 5007:
        logger.info("🐳 Running in Azure Container mode")
    else:
        logger.info("💻 Running in Local mode")

    app.run(host="0.0.0.0", port=port, debug=False)
