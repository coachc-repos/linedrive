#!/usr/bin/env python3
"""
ScriptCraft Web GUI - Enhanced Debug Version
This version includes comprehensive debugging for container deployment
"""

from flask import Flask, render_template, request, jsonify, send_file, make_response
from flask_cors import CORS
import threading
import webbrowser
import time
import sys
import os
import logging
import datetime
import tempfile
import asyncio
import traceback
from pathlib import Path

# Set up logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# Add the parent directory to Python path to find linedrive_azure
sys.path.insert(0, str(Path(__file__).parent))

# Import the agent client
try:
    from linedrive_azure.agents.script_polisher_agent_client import (
        ScriptPolisherAgentClient,
    )

    logger.info("✅ ScriptPolisherAgentClient imported successfully")
except ImportError as e:
    logger.error(f"❌ Failed to import ScriptPolisherAgentClient: {e}")
    ScriptPolisherAgentClient = None

# Create Flask app
app = Flask(__name__)
CORS(app)


def convert_markdown_to_word(markdown_content, title="Script"):
    """Convert markdown to Word document with formatting"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        logger.info(f"🔄 Converting {len(markdown_content)} chars to Word doc")

        # Create document
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(16)
        title_run.bold = True
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add blank line
        doc.add_paragraph()

        # Process markdown content line by line
        lines = markdown_content.split("\n")
        current_list_level = 0

        for line in lines:
            if not line.strip():
                doc.add_paragraph()
                current_list_level = 0
                continue

            # Headers
            if line.startswith("# "):
                p = doc.add_paragraph()
                run = p.add_run(line[2:])
                run.font.size = Pt(14)
                run.bold = True
                current_list_level = 0
            elif line.startswith("## "):
                p = doc.add_paragraph()
                run = p.add_run(line[3:])
                run.font.size = Pt(12)
                run.bold = True
                current_list_level = 0
            elif line.startswith("### "):
                p = doc.add_paragraph()
                run = p.add_run(line[4:])
                run.font.size = Pt(11)
                run.bold = True
                current_list_level = 0
            # Bullet points
            elif line.strip().startswith("- "):
                p = doc.add_paragraph(line.strip()[2:])
                p.style = "List Bullet"
                current_list_level = 1
            elif line.strip().startswith("* "):
                p = doc.add_paragraph(line.strip()[2:])
                p.style = "List Bullet"
                current_list_level = 1
            # Numbered lists
            elif any(line.strip().startswith(f"{i}. ") for i in range(1, 21)):
                content = line.strip()
                for i in range(1, 21):
                    if content.startswith(f"{i}. "):
                        content = content[len(f"{i}. ") :]
                        break
                p = doc.add_paragraph(content)
                p.style = "List Number"
                current_list_level = 1
            # Regular text
            else:
                # Handle bold and italic
                if "**" in line or "*" in line:
                    p = doc.add_paragraph()
                    parts = line.split("**")
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            if "*" in part:
                                italic_parts = part.split("*")
                                for j, italic_part in enumerate(italic_parts):
                                    run = p.add_run(italic_part)
                                    if j % 2 == 1:
                                        run.italic = True
                            else:
                                p.add_run(part)
                        else:
                            run = p.add_run(part)
                            run.bold = True
                else:
                    doc.add_paragraph(line)
                current_list_level = 0

        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)

        logger.info("✅ Word document created successfully")
        return temp_file.name

    except ImportError:
        logger.error("❌ python-docx not installed")
        raise Exception(
            "python-docx not installed. Install with: pip install python-docx"
        )
    except Exception as e:
        logger.error(f"❌ Word conversion failed: {e}")
        raise


class ScriptCraftEngine:
    """Main engine for script polishing with comprehensive debugging"""

    def __init__(self):
        self.agent_client = None
        self._init_agent()

    def _init_agent(self):
        """Initialize the agent client with debugging"""
        try:
            logger.info("🚀 Starting script polish - initializing agent")

            if not ScriptPolisherAgentClient:
                logger.warning("⚠️  No agent client available - using stub")
                return

            self.agent_client = ScriptPolisherAgentClient()
            logger.info("✅ ScriptPolisherAgentClient initialized successfully")

        except Exception as e:
            logger.error(f"❌ Agent initialization failed: {e}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            self.agent_client = None

    def polish_script(
        self, script_content, audience="General", production_type="General"
    ):
        """Polish script with enhanced debugging"""
        try:
            logger.info(f"🚀 Polish request - Length: {len(script_content)}")

            if self.agent_client:
                return self._real_polish_script(
                    script_content, audience, production_type
                )
            else:
                logger.warning("⚠️  Using stub polish (no agent available)")
                return self._stub_polish_script(
                    script_content, audience, production_type
                )

        except Exception as e:
            logger.error(f"❌ Script polishing failed: {e}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                "success": False,
                "error": str(e),
                "message": "Script polishing failed",
                "debug_info": {
                    "error_type": type(e).__name__,
                    "traceback": traceback.format_exc(),
                },
            }

    def _real_polish_script(self, script_content, audience, production_type):
        """Use real agent for polishing"""
        try:
            result = self.agent_client.polish_script(
                script_content=script_content,
                audience=audience,
                production_type=production_type,
            )

            logger.info("📋 Agent response received:")
            logger.info(f"   Type: {type(result)}")
            logger.info(f"   Success: {result.get('success', 'Not specified')}")

            if result.get("success"):
                # The real agent returns response in "response" field
                enhanced_script = result.get("response", "")
                logger.info(f"   Response field content: {repr(enhanced_script[:200])}")
                logger.info(
                    f"✅ AI enhancement complete - Length: {len(enhanced_script)}"
                )

                # Debug: Log the first 500 characters
                logger.info(f"   Original (first 500): {script_content[:500]}...")
                logger.info(f"   Enhanced (first 500): {enhanced_script[:500]}...")

                # Quality check
                if enhanced_script == script_content:
                    logger.warning("⚠️  WARNING: Enhanced script identical to original!")
                elif len(enhanced_script) < len(script_content) * 1.1:
                    logger.warning("⚠️  WARNING: Enhanced script minimally changed!")
                else:
                    logger.info("✅ Script appears significantly enhanced")

                return {
                    "success": True,
                    "message": enhanced_script,
                    "original_length": len(script_content),
                    "enhanced_length": len(enhanced_script),
                    "timestamp": datetime.datetime.now().isoformat(),
                }
            else:
                error_msg = result.get("error", "Unknown error")
                logger.error(f"❌ Agent returned error: {error_msg}")
                return {
                    "success": False,
                    "error": error_msg,
                    "message": "Agent processing failed",
                }

        except Exception as e:
            logger.error(f"❌ Real polish failed: {e}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return self._stub_polish_script(script_content, audience, production_type)

    def _stub_polish_script(self, script_content, audience, production_type):
        """Fallback stub polishing when agent is not available"""
        try:
            logger.info("🔄 Using stub polish (offline mode)")

            enhanced_script = f"""# Enhanced Script for {audience}
Production Type: {production_type}
Timestamp: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

---

{script_content}

---

**Note: This script was processed in offline mode. For full AI enhancement, ensure Azure AI Projects connection is configured.**
"""

            logger.info(f"✅ Stub polish complete - Length: {len(enhanced_script)}")

            return {
                "success": True,
                "message": enhanced_script,
                "original_length": len(script_content),
                "enhanced_length": len(enhanced_script),
                "timestamp": datetime.datetime.now().isoformat(),
                "mode": "stub",
            }

        except Exception as e:
            logger.error(f"❌ Even stub polish failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "message": "Complete processing failure",
            }


# Create the engine instance
scriptcraft = ScriptCraftEngine()


@app.route("/")
def index():
    """Main page"""
    return render_template("index.html")


@app.route("/create", methods=["POST"])
def create():
    """Create/polish a script with comprehensive debugging"""
    debug_info = {
        "request_received": True,
        "timestamp": datetime.datetime.now().isoformat(),
    }

    try:
        # Parse JSON request with validation
        try:
            data = request.get_json()
            debug_info["json_parsing"] = "SUCCESS"
            debug_info["data_keys"] = list(data.keys()) if data else None
        except Exception as json_error:
            debug_info["json_parsing"] = f"FAILED: {str(json_error)}"
            return (
                jsonify(
                    {
                        "success": False,
                        "error": "Invalid JSON request",
                        "debug_info": debug_info,
                    }
                ),
                400,
            )

        if not data:
            debug_info["data_validation"] = "FAILED: No data"
            return (
                jsonify(
                    {
                        "success": False,
                        "error": "No data received",
                        "debug_info": debug_info,
                    }
                ),
                400,
            )

        # Extract parameters with validation - handle both 'script' and 'topic' parameters
        script_content = data.get("script", "").strip()
        topic = data.get("topic", "").strip()

        # Use topic if script is empty (for script creation)
        if not script_content and topic:
            script_content = topic

        audience = data.get("audience", "General")
        production_type = data.get(
            "productionType", data.get("production_type", "General")
        )
        duration = data.get("duration", "")
        description = data.get("description", "")

        debug_info["parameters"] = {
            "script_length": len(script_content),
            "topic_length": len(topic),
            "audience": audience,
            "production_type": production_type,
            "duration": duration,
            "description": description,
            "script_preview": script_content[:100] if script_content else None,
            "using_topic": bool(topic and not data.get("script")),
        }

        if not script_content:
            debug_info["validation"] = "FAILED: Empty script/topic"
            return jsonify(
                {
                    "success": False,
                    "error": "Please enter a topic or script content",
                    "debug_info": debug_info,
                }
            )

        debug_info["validation"] = "SUCCESS"

        # Polish the script with comprehensive error handling
        try:
            logger.info(f"🎯 Processing request for {audience}/{production_type}")
            result = scriptcraft.polish_script(
                script_content, audience, production_type
            )
            debug_info["polish_attempt"] = "COMPLETED"
            debug_info["polish_result_type"] = type(result).__name__
            debug_info["polish_success"] = result.get("success", False)

        except Exception as polish_error:
            debug_info["polish_attempt"] = f"FAILED: {str(polish_error)}"
            debug_info["polish_traceback"] = traceback.format_exc()
            return (
                jsonify(
                    {
                        "success": False,
                        "error": f"Script processing failed: {str(polish_error)}",
                        "debug_info": debug_info,
                    }
                ),
                500,
            )

        # Process result with debugging
        if result.get("success"):
            debug_info["response_processing"] = "SUCCESS"
            return jsonify(
                {
                    "success": True,
                    "message": result.get("message", ""),
                    "originalLength": result.get("original_length", 0),
                    "enhancedLength": result.get("enhanced_length", 0),
                    "mode": result.get("mode", "agent"),
                    "debug_info": debug_info,
                }
            )
        else:
            debug_info["response_processing"] = "FAILED"
            debug_info["error_from_result"] = result.get("error", "Unknown")
            return (
                jsonify(
                    {
                        "success": False,
                        "error": result.get("error", "Processing failed"),
                        "debug_info": debug_info,
                    }
                ),
                500,
            )

    except Exception as e:
        debug_info["unexpected_error"] = str(e)
        debug_info["unexpected_traceback"] = traceback.format_exc()
        logger.error(f"❌ Unexpected error in create endpoint: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return (
            jsonify(
                {
                    "success": False,
                    "error": f"Unexpected error: {str(e)}",
                    "debug_info": debug_info,
                }
            ),
            500,
        )


@app.route("/polish", methods=["POST"])
def polish():
    """Polish an existing script with comprehensive debugging"""
    debug_info = {
        "request_received": True,
        "timestamp": datetime.datetime.now().isoformat(),
    }

    try:
        # Parse JSON request with validation
        try:
            data = request.get_json()
            debug_info["json_parsing"] = "SUCCESS"
            debug_info["data_keys"] = list(data.keys()) if data else None
        except Exception as json_error:
            debug_info["json_parsing"] = f"FAILED: {str(json_error)}"
            return (
                jsonify(
                    {
                        "success": False,
                        "error": "Invalid JSON request",
                        "debug_info": debug_info,
                    }
                ),
                400,
            )

        if not data:
            debug_info["data_validation"] = "FAILED: No data"
            return (
                jsonify(
                    {
                        "success": False,
                        "error": "No data received",
                        "debug_info": debug_info,
                    }
                ),
                400,
            )

        # Extract parameters for polish endpoint
        script_content = data.get("script", "").strip()
        audience = data.get("audience", "General")
        production_type = data.get(
            "production_type", data.get("productionType", "General")
        )

        debug_info["parameters"] = {
            "script_length": len(script_content),
            "audience": audience,
            "production_type": production_type,
            "script_preview": script_content[:100] if script_content else None,
        }

        if not script_content:
            debug_info["validation"] = "FAILED: Empty script"
            return jsonify(
                {
                    "success": False,
                    "error": "Please enter a script to polish",
                    "debug_info": debug_info,
                }
            )

        debug_info["validation"] = "SUCCESS"

        # Polish the script with comprehensive error handling
        try:
            logger.info(f"🔄 Polishing script for {audience}/{production_type}")
            result = scriptcraft.polish_script(
                script_content, audience, production_type
            )
            debug_info["polish_attempt"] = "COMPLETED"
            debug_info["polish_result_type"] = type(result).__name__
            debug_info["polish_success"] = result.get("success", False)

        except Exception as polish_error:
            debug_info["polish_attempt"] = f"FAILED: {str(polish_error)}"
            debug_info["polish_traceback"] = traceback.format_exc()
            return (
                jsonify(
                    {
                        "success": False,
                        "error": f"Script polishing failed: {str(polish_error)}",
                        "debug_info": debug_info,
                    }
                ),
                500,
            )

        # Process result with debugging
        if result.get("success"):
            debug_info["response_processing"] = "SUCCESS"

            # Calculate additional metrics for polish endpoint
            enhanced_script = result.get("message", "")
            word_count = len(enhanced_script.split()) if enhanced_script else 0
            reading_time = (
                f"{max(1, word_count // 200)} min" if word_count > 0 else "< 1 min"
            )

            return jsonify(
                {
                    "success": True,
                    "enhanced_script": enhanced_script,
                    "word_count": word_count,
                    "reading_time": reading_time,
                    "originalLength": result.get(
                        "original_length", len(script_content)
                    ),
                    "enhancedLength": result.get(
                        "enhanced_length", len(enhanced_script)
                    ),
                    "mode": result.get("mode", "agent"),
                    "debug_info": debug_info,
                }
            )
        else:
            debug_info["response_processing"] = "FAILED"
            debug_info["error_from_result"] = result.get("error", "Unknown")
            return (
                jsonify(
                    {
                        "success": False,
                        "error": result.get("error", "Processing failed"),
                        "debug_info": debug_info,
                    }
                ),
                500,
            )

    except Exception as e:
        debug_info["unexpected_error"] = str(e)
        debug_info["unexpected_traceback"] = traceback.format_exc()
        logger.error(f"❌ Unexpected error in polish endpoint: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return (
            jsonify(
                {
                    "success": False,
                    "error": f"Unexpected error: {str(e)}",
                    "debug_info": debug_info,
                }
            ),
            500,
        )


@app.route("/export_word", methods=["POST"])
def export_word():
    """Export polished script to Word document"""
    try:
        data = request.get_json()
        if not data or "content" not in data:
            return jsonify({"success": False, "error": "No content provided"})

        content = data["content"]
        title = data.get("title", "Polished Script")

        # Convert to Word
        word_file_path = convert_markdown_to_word(content, title)

        # Send file
        response = send_file(
            word_file_path,
            as_attachment=True,
            download_name=f"{title.replace(' ', '_')}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # Clean up temp file after sending
        @response.call_on_close
        def cleanup():
            try:
                os.unlink(word_file_path)
            except:
                pass

        return response

    except Exception as e:
        logger.error(f"Word export failed: {e}")
        return jsonify({"success": False, "error": f"Word export failed: {str(e)}"})


@app.route("/load_file", methods=["POST"])
def load_file():
    """Load a file from the file system"""
    try:
        from flask import request

        if "file" not in request.files:
            return jsonify({"success": False, "error": "No file selected"})

        file = request.files["file"]
        if file.filename == "":
            return jsonify({"success": False, "error": "No file selected"})

        # Read file content
        content = file.read().decode("utf-8")
        filename = file.filename

        # Get file stats
        file_size = len(content)
        line_count = content.count("\n") + 1

        return jsonify(
            {
                "success": True,
                "content": content,
                "filename": filename,
                "file_size": file_size,
                "line_count": line_count,
            }
        )

    except Exception as e:
        return jsonify({"success": False, "error": f"Failed to load file: {str(e)}"})


@app.route("/debug")
def debug_system():
    """Debug endpoint to check system status"""
    debug_info = {}

    try:
        # Check ScriptPolisherAgentClient import and initialization
        debug_info["imports"] = {
            "ScriptPolisherAgentClient": (
                "OK" if ScriptPolisherAgentClient else "FAILED"
            ),
            "current_directory": os.getcwd(),
            "python_path": sys.path[:3],  # First 3 entries
        }

        # Check environment variables
        env_conn_str = os.getenv("AZURE_AI_PROJECTS_CONNECTION_STRING")
        debug_info["environment"] = {
            "AZURE_AI_PROJECTS_CONNECTION_STRING": (
                "SET" if env_conn_str else "MISSING"
            ),
            "PORT": os.getenv("PORT", "Not set"),
            "Environment Variables Count": len(os.environ),
        }

        # Test agent initialization
        try:
            test_agent = ScriptPolisherAgentClient()
            debug_info["agent_init"] = "SUCCESS"

            # Test agent availability
            try:
                # This should be a minimal test
                debug_info["agent_test"] = "AGENT_CREATED"
            except Exception as agent_test_error:
                debug_info["agent_test"] = f"AGENT_TEST_ERROR: {str(agent_test_error)}"

        except Exception as init_error:
            debug_info["agent_init"] = f"INIT_ERROR: {str(init_error)}"
            debug_info["agent_init_traceback"] = traceback.format_exc()

        # Check module locations
        try:
            import scriptpolisher_agent_client

            debug_info["module_location"] = {
                "scriptpolisher_module": scriptpolisher_agent_client.__file__,
                "working_directory": os.getcwd(),
            }
        except ImportError as import_error:
            debug_info["module_location"] = f"IMPORT_ERROR: {str(import_error)}"

        # System info
        debug_info["system"] = {
            "python_version": sys.version,
            "platform": sys.platform,
            "executable": sys.executable,
        }

        return jsonify(
            {
                "success": True,
                "debug_info": debug_info,
                "timestamp": datetime.datetime.now().isoformat(),
            }
        )

    except Exception as e:
        return (
            jsonify(
                {
                    "success": False,
                    "error": str(e),
                    "traceback": traceback.format_exc(),
                    "partial_debug_info": debug_info,
                }
            ),
            500,
        )


# Template creation function
def create_template():
    """Create the HTML template"""
    template_dir = Path(__file__).parent / "templates"
    template_dir.mkdir(exist_ok=True)

    html_content = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>ScriptCraft - AI Script Polisher</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }
        
        .container {
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            border-radius: 15px;
            box-shadow: 0 20px 40px rgba(0, 0, 0, 0.1);
            overflow: hidden;
        }
        
        .header {
            background: linear-gradient(135deg, #4CAF50 0%, #45a049 100%);
            color: white;
            padding: 30px;
            text-align: center;
        }
        
        .header h1 {
            font-size: 2.5em;
            margin-bottom: 10px;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.3);
        }
        
        .header p {
            font-size: 1.2em;
            opacity: 0.9;
        }
        
        .main-content {
            padding: 40px;
        }
        
        .form-group {
            margin-bottom: 20px;
        }
        
        .form-group label {
            display: block;
            font-weight: bold;
            margin-bottom: 8px;
            color: #333;
            font-size: 1.1em;
        }
        
        .form-row {
            display: flex;
            gap: 20px;
            margin-bottom: 20px;
        }
        
        .form-row .form-group {
            flex: 1;
        }
        
        select {
            width: 100%;
            padding: 12px;
            border: 2px solid #ddd;
            border-radius: 8px;
            font-size: 16px;
            background: white;
        }
        
        select:focus {
            border-color: #4CAF50;
            outline: none;
        }
        
        textarea {
            width: 100%;
            padding: 15px;
            border: 2px solid #ddd;
            border-radius: 8px;
            font-family: 'Monaco', 'Courier New', monospace;
            font-size: 14px;
            line-height: 1.5;
            resize: vertical;
            min-height: 200px;
        }
        
        textarea:focus {
            border-color: #4CAF50;
            outline: none;
        }
        
        .buttons {
            display: flex;
            gap: 15px;
            margin: 20px 0;
        }
        
        .btn {
            padding: 15px 30px;
            border: none;
            border-radius: 8px;
            font-size: 16px;
            font-weight: bold;
            cursor: pointer;
            transition: all 0.3s ease;
        }
        
        .btn-primary {
            background: linear-gradient(135deg, #4CAF50 0%, #45a049 100%);
            color: white;
        }
        
        .btn-primary:hover {
            transform: translateY(-2px);
            box-shadow: 0 5px 15px rgba(76, 175, 80, 0.3);
        }
        
        .btn-secondary {
            background: #6c757d;
            color: white;
        }
        
        .btn-secondary:hover {
            background: #5a6268;
        }
        
        .btn:disabled {
            background: #ccc;
            cursor: not-allowed;
            transform: none;
        }
        
        .progress {
            display: none;
            margin: 20px 0;
            padding: 20px;
            background: #f8f9fa;
            border-radius: 8px;
            text-align: center;
        }
        
        .progress-bar {
            width: 100%;
            height: 4px;
            background: #e9ecef;
            border-radius: 2px;
            overflow: hidden;
            margin-top: 10px;
        }
        
        .progress-bar-fill {
            height: 100%;
            background: linear-gradient(90deg, #4CAF50, #45a049);
            animation: progress 2s ease-in-out infinite;
        }
        
        @keyframes progress {
            0% { transform: translateX(-100%); }
            100% { transform: translateX(100%); }
        }
        
        .result {
            margin-top: 30px;
            padding: 25px;
            background: #f8f9fa;
            border-radius: 10px;
            border-left: 4px solid #4CAF50;
        }
        
        .result h3 {
            color: #4CAF50;
            margin-bottom: 15px;
            font-size: 1.3em;
        }
        
        .result textarea {
            background: white;
            min-height: 300px;
        }
        
        .stats {
            display: flex;
            gap: 20px;
            margin: 15px 0;
            padding: 10px;
            background: white;
            border-radius: 5px;
        }
        
        .stats span {
            font-size: 0.9em;
            color: #666;
        }
        
        .file-upload {
            margin-bottom: 20px;
            padding: 20px;
            border: 2px dashed #ddd;
            border-radius: 8px;
            text-align: center;
            background: #fafafa;
        }
        
        .file-upload input[type="file"] {
            margin: 10px 0;
        }
        
        .file-upload p {
            color: #666;
            margin: 5px 0;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>✨ ScriptCraft</h1>
            <p>AI-Powered Script Enhancement & Polishing</p>
        </div>
        
        <div class="main-content">
            <form id="scriptForm">
                <div class="file-upload">
                    <h3>📁 Load Script from File</h3>
                    <input type="file" id="fileInput" accept=".txt,.md,.docx" onchange="loadFile(event)">
                    <p>Supported formats: .txt, .md, .docx</p>
                    <p>Or paste/type your script below</p>
                </div>
                
                <div class="form-row">
                    <div class="form-group">
                        <label for="audience">🎯 Target Audience:</label>
                        <select id="audience" name="audience">
                            <option value="General">General Audience</option>
                            <option value="Corporate">Corporate/Business</option>
                            <option value="Academic">Academic</option>
                            <option value="Creative">Creative/Entertainment</option>
                            <option value="Technical">Technical</option>
                            <option value="Youth">Youth/Students</option>
                        </select>
                    </div>
                    
                    <div class="form-group">
                        <label for="productionType">🎬 Production Type:</label>
                        <select id="productionType" name="productionType">
                            <option value="General">General Script</option>
                            <option value="Video">Video Production</option>
                            <option value="Podcast">Podcast</option>
                            <option value="Presentation">Presentation</option>
                            <option value="Article">Article/Blog</option>
                            <option value="Speech">Speech/Talk</option>
                        </select>
                    </div>
                </div>
                
                <div class="form-group">
                    <label for="scriptInput">📝 Your Script:</label>
                    <textarea id="scriptInput" name="script" 
                              placeholder="Enter your script here... or use the sample button below to see how it works!"></textarea>
                </div>
                
                <div class="buttons">
                    <button type="button" class="btn btn-primary" onclick="polishScript()">
                        ✨ Polish My Script
                    </button>
                    <button type="button" class="btn btn-secondary" onclick="loadSample()">
                        📄 Load Sample
                    </button>
                    <button type="button" class="btn btn-secondary" onclick="clearForm()">
                        🗑️ Clear
                    </button>
                </div>
            </form>
            
            <div class="progress" id="progress">
                <h4>🤖 AI is enhancing your script...</h4>
                <p>This may take 30-60 seconds for best results</p>
                <div class="progress-bar">
                    <div class="progress-bar-fill"></div>
                </div>
            </div>
            
            <div id="result" style="display: none;"></div>
        </div>
    </div>
    
    <script>
        async function polishScript() {
            const scriptContent = document.getElementById('scriptInput').value.trim();
            const audience = document.getElementById('audience').value;
            const productionType = document.getElementById('productionType').value;
            
            if (!scriptContent) {
                alert('Please enter a script to polish!');
                return;
            }
            
            // Show progress
            document.getElementById('progress').style.display = 'block';
            document.getElementById('result').style.display = 'none';
            
            // Disable button
            const button = document.querySelector('.btn-primary');
            button.disabled = true;
            button.textContent = '🤖 Processing...';
            
            try {
                const response = await fetch('/create', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json',
                    },
                    body: JSON.stringify({
                        script: scriptContent,
                        audience: audience,
                        productionType: productionType
                    })
                });
                
                const data = await response.json();
                
                // Hide progress
                document.getElementById('progress').style.display = 'none';
                
                if (data.success) {
                    showResult(data.message, {
                        originalLength: data.originalLength || 0,
                        enhancedLength: data.enhancedLength || 0,
                        mode: data.mode || 'agent'
                    });
                } else {
                    showError(data.error || 'Unknown error occurred', data.debug_info);
                }
            } catch (error) {
                document.getElementById('progress').style.display = 'none';
                showError('Network error: ' + error.message);
            } finally {
                // Re-enable button
                button.disabled = false;
                button.textContent = '✨ Polish My Script';
            }
        }
        
        function showResult(polishedScript, stats) {
            const resultDiv = document.getElementById('result');
            const mode = stats.mode === 'stub' ? ' (Offline Mode)' : '';
            
            resultDiv.innerHTML = `
                <h3>✨ Script Enhanced Successfully${mode}!</h3>
                <div class="stats">
                    <span>📏 Original: ${stats.originalLength} chars</span>
                    <span>📈 Enhanced: ${stats.enhancedLength} chars</span>
                    <span>📊 Growth: ${(((stats.enhancedLength - stats.originalLength) / stats.originalLength) * 100).toFixed(1)}%</span>
                </div>
                <textarea readonly>${polishedScript}</textarea>
                <div class="buttons">
                    <button type="button" class="btn btn-primary" onclick="exportToWord()">
                        📄 Export to Word
                    </button>
                    <button type="button" class="btn btn-secondary" onclick="saveResult()">
                        💾 Save as File
                    </button>
                </div>
            `;
            resultDiv.className = 'result';
            resultDiv.style.display = 'block';
        }
        
        function showError(error, debugInfo) {
            const resultDiv = document.getElementById('result');
            let debugSection = '';
            
            if (debugInfo) {
                debugSection = `
                    <details style="margin-top: 15px;">
                        <summary style="cursor: pointer; color: #666;">🔍 Debug Information</summary>
                        <pre style="background: #f1f1f1; padding: 10px; margin-top: 10px; border-radius: 5px; font-size: 12px; overflow-x: auto;">${JSON.stringify(debugInfo, null, 2)}</pre>
                    </details>
                `;
            }
            
            resultDiv.innerHTML = `
                <h3 style="color: #d32f2f;">❌ Error</h3>
                <p style="color: #666; margin: 10px 0;">${error}</p>
                ${debugSection}
            `;
            resultDiv.className = 'result';
            resultDiv.style.display = 'block';
        }
        
        async function exportToWord() {
            const textarea = document.querySelector('#result textarea');
            if (!textarea) {
                alert('No enhanced script to export.');
                return;
            }
            
            try {
                const response = await fetch('/export_word', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json',
                    },
                    body: JSON.stringify({
                        content: textarea.value,
                        title: 'Enhanced Script'
                    })
                });
                
                if (response.ok) {
                    const blob = await response.blob();
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = 'enhanced_script.docx';
                    document.body.appendChild(a);
                    a.click();
                    document.body.removeChild(a);
                    window.URL.revokeObjectURL(url);
                } else {
                    const error = await response.json();
                    alert('Export failed: ' + error.error);
                }
            } catch (error) {
                alert('Export failed: ' + error.message);
            }
        }
        
        function clearForm() {
            document.getElementById('scriptInput').value = '';
            document.getElementById('audience').value = 'General';
            document.getElementById('productionType').value = 'General';
            document.getElementById('result').style.display = 'none';
        }
        
        function loadSample() {
            const sample = `# The Future of Work: Embracing AI as Your Professional Partner

## Introduction

Hello, and welcome to today's discussion about one of the most important topics of our time: the future of work in the age of artificial intelligence.

## The Current Landscape

Many workers fear that AI will eliminate their roles entirely. However, research shows a more nuanced reality. AI excels at automating routine tasks, freeing humans to focus on creative problem-solving, emotional intelligence, and strategic thinking.

## Building Future-Ready Skills

Focus on developing skills that complement AI rather than compete with it:

1. **Complex Problem Solving**: AI provides data, humans provide context
2. **Emotional Intelligence**: Relationships remain purely human domains  
3. **Creative Thinking**: Innovation requires human intuition
4. **Adaptability**: Learning and pivoting quickly becomes crucial

## Conclusion

The future belongs to professionals who view AI as a powerful assistant rather than a threat.`;
            
            document.getElementById('scriptInput').value = sample;
        }
        
        async function loadFile(event) {
            const file = event.target.files[0];
            if (!file) return;
            
            const formData = new FormData();
            formData.append('file', file);
            
            try {
                const response = await fetch('/load_file', {
                    method: 'POST',
                    body: formData
                });
                
                const data = await response.json();
                
                if (data.success) {
                    document.getElementById('scriptInput').value = data.content;
                    
                    // Show file info
                    const resultDiv = document.getElementById('result');
                    resultDiv.innerHTML = `
                        <h3>📁 File Loaded Successfully!</h3>
                        <div class="stats">
                            <span>📄 File: ${data.filename}</span>
                            <span>📏 Size: ${data.file_size} characters</span>
                            <span>📋 Lines: ${data.line_count}</span>
                        </div>
                        <p>✅ Content loaded into the script input area above.</p>
                    `;
                    resultDiv.className = 'result';
                    resultDiv.style.display = 'block';
                } else {
                    alert('Error loading file: ' + data.error);
                }
            } catch (error) {
                alert('Failed to load file: ' + error.message);
            }
            
            // Reset file input
            event.target.value = '';
        }
        
        function saveResult() {
            const resultDiv = document.getElementById('result');
            const textarea = resultDiv.querySelector('textarea');
            
            if (!textarea) {
                alert('No polished script to save. Please polish a script first.');
                return;
            }
            
            const content = textarea.value;
            const filename = `polished_script_${new Date().toISOString().slice(0,19).replace(/:/g, '-')}.md`;
            
            // Create download link
            const blob = new Blob([content], { type: 'text/markdown' });
            const url = window.URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url;
            a.download = filename;
            document.body.appendChild(a);
            a.click();
            document.body.removeChild(a);
            window.URL.revokeObjectURL(url);
            
            alert(`✅ Script saved as: ${filename}`);
        }
    </script>
</body>
</html>"""

    with open(template_dir / "index.html", "w") as f:
        f.write(html_content)


def start_web_gui():
    """Start the web-based GUI"""
    print("🌐 Starting ScriptCraft Web GUI...")

    # Create template
    # create_template()  # COMMENTED OUT - stops overwriting manual edits

    # Start Flask in a separate thread
    def run_flask():
        app.run(debug=False, host="0.0.0.0", port=5007, use_reloader=False)

    flask_thread = threading.Thread(target=run_flask, daemon=True)
    flask_thread.start()

    # Wait a moment for Flask to start
    time.sleep(2)

    # Open browser
    url = "http://127.0.0.1:5007"
    print(f"🚀 Opening ScriptCraft in your browser: {url}")
    webbrowser.open(url)

    print("✅ ScriptCraft Web GUI is running!")
    print("🌐 Access it at: http://127.0.0.1:5007")
    print("🛑 Press Ctrl+C to stop the server")

    try:
        # Keep the main thread alive
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        print("\n👋 ScriptCraft Web GUI stopped")


if __name__ == "__main__":
    start_web_gui()
