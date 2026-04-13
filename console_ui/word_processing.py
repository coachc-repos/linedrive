#!/usr/bin/env python3
"""
Word document processing utilities for LineDrive Console
Handles conversion from markdown to Word documents with templates
"""

import asyncio
import re
from pathlib import Path
from typing import Optional
from utils import get_word_template_path


async def convert_markdown_to_word(
    markdown_content: str,
    output_file_path: str,
    template_path: Optional[Path] = None,
    title: Optional[str] = None,
) -> bool:
    """Convert markdown content to Word document with enhanced formatting or template"""

    try:
        # Dynamic imports to avoid issues if packages aren't installed
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.shared import OxmlElement, qn

        print(f"🔍 Word conversion: Starting with {len(markdown_content)} chars")

        # Get template path if not provided
        if not template_path:
            template_path = get_word_template_path()

        if template_path:
            print(f"📄 Using Word template: {template_path}")

        print("🔍 Word conversion: Imports successful")
        print("🔍 Word conversion: Creating document...")

        # Create document from template or blank
        try:
            if template_path and template_path.exists():
                # Try to load template directly first
                print(f"✅ Loading template from: {template_path}")
                try:
                    doc = Document(str(template_path))
                    print(
                        "🔄 Template format detected, creating document with template styles..."
                    )
                    print(
                        "✅ Document created - template styles will be applied during content processing"
                    )
                except ValueError as ve:
                    # Handle DOTX template content type issue
                    if "is not a Word file" in str(ve) and "template.main+xml" in str(
                        ve
                    ):
                        print(
                            f"⚠️ DOTX template format detected, using workaround for python-docx compatibility..."
                        )
                        # Create a blank document and manually load styles later
                        doc = Document()
                        print(
                            "✅ Blank document created (DOTX template styles not fully supported)"
                        )
                    else:
                        raise ve
                except Exception as template_error:
                    print(
                        f"⚠️ Template loading failed ({template_error}), creating blank document"
                    )
                    doc = Document()
            else:
                print("📄 No template found, creating blank document")
                doc = Document()

            print("🔍 Word conversion: Enhanced document styles created")

            # Clear existing content if loading from template
            if template_path and template_path.exists():
                # Clear the document content but keep the styles
                for paragraph in doc.paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)
                    p._p = p._element = None

        except Exception as doc_creation_error:
            print(f"❌ Document creation failed: {doc_creation_error}")
            return False

        # Process markdown content
        lines = markdown_content.split("\n")
        print(f"🔍 Word conversion: Processing {len(lines)} lines")

        for line in lines:
            line = line.rstrip()

            # Skip empty lines
            if not line.strip():
                doc.add_paragraph()
                continue

            # Handle headers - enhanced detection
            if line.startswith("# "):
                # H1 - Main title
                title_text = line[2:].strip()
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(title_text)
                run.font.size = Pt(18)
                run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            elif line.startswith("## "):
                # H2 - Chapter/Section headers
                header_text = line[3:].strip()
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(header_text)
                run.font.size = Pt(14)
                run.font.bold = True
                paragraph.space_before = Pt(12)
                paragraph.space_after = Pt(6)
                continue

            elif line.startswith("### "):
                # H3 - Subsection headers
                header_text = line[4:].strip()
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(header_text)
                run.font.size = Pt(12)
                run.font.bold = True
                paragraph.space_before = Pt(8)
                paragraph.space_after = Pt(4)
                continue

            # Handle "Heading: Title" format headers
            elif line.startswith("Heading:"):
                # Extract title after "Heading:"
                header_text = line[8:].strip()  # Remove "Heading:" prefix
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(header_text)
                run.font.size = Pt(16)
                run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.space_before = Pt(12)
                paragraph.space_after = Pt(8)
                print(f"✅ Formatted heading: {header_text}")
                continue

            # Handle other common header patterns
            elif re.match(r"^[A-Z][^:]*:", line) and len(line.split()) <= 8:
                # Detect "Title:" pattern headers (but not long sentences with colons)
                header_text = line.rstrip(":").strip()
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(header_text)
                run.font.size = Pt(14)
                run.font.bold = True
                paragraph.space_before = Pt(10)
                paragraph.space_after = Pt(6)
                print(f"✅ Formatted section header: {header_text}")
                continue

            # Handle bold text (**text**)
            elif "**" in line:
                paragraph = doc.add_paragraph()
                parts = re.split(r"\*\*(.*?)\*\*", line)

                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        # Regular text
                        if part.strip():
                            paragraph.add_run(part)
                    else:
                        # Bold text
                        run = paragraph.add_run(part)
                        run.font.bold = True
                continue

            # Handle special formatting markers
            elif line.startswith("**") and line.endswith("**"):
                # Entire line is bold
                text = line[2:-2].strip()
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(text)
                run.font.bold = True
                continue

            elif (
                line.startswith("*")
                and line.endswith("*")
                and not line.startswith("**")
            ):
                # Italic text
                text = line[1:-1].strip()
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(text)
                run.font.italic = True
                continue

            # Handle visual cues (lines with 🎬 emoji or [Visual:])
            elif (
                "🎬" in line
                or "[Visual:" in line
                or line.startswith("[")
                and "]" in line
            ):
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line)
                run.font.italic = True
                run.font.color.rgb = None  # Use theme color
                paragraph.space_before = Pt(6)
                paragraph.space_after = Pt(6)
                continue

            # Regular paragraph
            else:
                paragraph = doc.add_paragraph()
                paragraph.add_run(line)

        print(f"🔄 Saving Word document to: {output_file_path}")
        doc.save(output_file_path)

        # Check file size
        output_path = Path(output_file_path)
        if output_path.exists():
            file_size = output_path.stat().st_size
            print(
                f"✅ Word document created successfully: {output_file_path} ({file_size} bytes)"
            )
            return True
        else:
            print(f"❌ Word document was not created: {output_file_path}")
            return False

    except ImportError as e:
        print(f"❌ Word conversion failed: Missing required package ({e})")
        print("💡 Install with: pip install python-docx")
        return False

    except Exception as e:
        print(f"❌ Word conversion failed: {e}")
        return False
