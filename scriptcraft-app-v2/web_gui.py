#!/usr/bin/env python3
"""
ScriptCraft Web GUI - Clean Console Capture Version
Fixed implementation that displays real-time EnhancedAutoGenSystem messages in browser
"""

# CRITICAL: Add repo root to path FIRST (before any other imports)
import queue as _queue
import threading as _threading
import re
import os
import json
import time
import queue
import uuid
import asyncio
import threading
import logging
from datetime import datetime
from io import BytesIO
from flask import Flask, render_template, request, jsonify, Response, make_response, send_file, stream_with_context
import requests
import sys
from pathlib import Path
from typing import Optional

# Add paths for imports - prioritize local scriptcraft-app directory
current_dir = Path(__file__).parent
sys.path.insert(0, str(current_dir))  # Prioritize local directory FIRST
print(f"✅ Prioritized local dir in sys.path: {current_dir}")

REPO_ROOT = Path(__file__).parent.parent
if str(REPO_ROOT) not in sys.path:
    # Add parent as backup (position 1, not 0)
    sys.path.insert(1, str(REPO_ROOT))
    print(f"✅ Added to sys.path: {REPO_ROOT}")


VERSION = "15.34-quotes-progress-mapping"

# Verify Google API key availability for thumbnail generation.
if "GOOGLE_API_KEY" in os.environ and os.environ.get("GOOGLE_API_KEY"):
    print("🔑 Using GOOGLE_API_KEY from environment")
else:
    print("⚠️ GOOGLE_API_KEY is not set; thumbnail generation will fail until provided")


# Set up logging
logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

SCRIPTCRAFT_SETTINGS_PATH = Path.home() / ".scriptcraft" / \
    "web_gui_settings.json"
DEFAULT_OUTPUT_PARENT = Path.home() / "Dev" / "Videos" / "Edited" / "Final"


def _load_scriptcraft_settings() -> dict:
    try:
        if SCRIPTCRAFT_SETTINGS_PATH.exists():
            with open(SCRIPTCRAFT_SETTINGS_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception as e:
        logger.warning(f"⚠️ Could not load ScriptCraft settings: {e}")
    return {}


def _save_scriptcraft_settings(settings: dict) -> bool:
    try:
        SCRIPTCRAFT_SETTINGS_PATH.parent.mkdir(parents=True, exist_ok=True)
        with open(SCRIPTCRAFT_SETTINGS_PATH, "w", encoding="utf-8") as f:
            json.dump(settings, f, indent=2)
        return True
    except Exception as e:
        logger.error(f"❌ Could not save ScriptCraft settings: {e}")
        return False


# On startup, hydrate environment variables from saved settings so any module
# that reads os.getenv(...) (e.g. GOOGLE_API_KEY for Gemini) picks them up
# even when the user only set them through the API Keys modal.
def _hydrate_env_from_settings() -> None:
    try:
        s = _load_scriptcraft_settings()
        mapping = {
            "google_api_key": "GOOGLE_API_KEY",
            "heygen_api_key": "HEYGEN_API_KEY",
            "heygen_voice_id": "HEYGEN_VOICE_ID",
            "grok_api_key": "XAI_API_KEY",
        }
        for setting_key, env_var in mapping.items():
            val = (s.get(setting_key) or "").strip()
            if val and not os.environ.get(env_var):
                os.environ[env_var] = val
    except Exception as e:
        logger.warning(f"⚠️ Could not hydrate env from saved settings: {e}")


_hydrate_env_from_settings()


def _get_output_parent_dir() -> Path:
    """Return configured output directory as-is, or Final root as fallback."""
    try:
        s = _load_scriptcraft_settings()
        raw = (s.get("output_dir") or "").strip()
        base = Path(raw).expanduser().resolve(
        ) if raw else DEFAULT_OUTPUT_PARENT
    except Exception:
        base = DEFAULT_OUTPUT_PARENT
    if base.exists() and not base.is_dir():
        logger.warning(
            f"⚠️ Configured output path is not a directory: {base}. Using default.")
        base = DEFAULT_OUTPUT_PARENT
    base.mkdir(parents=True, exist_ok=True)
    return base


def _get_output_base_dir() -> Optional[Path]:
    """Back-compat helper: returns the output parent directory."""
    try:
        return _get_output_parent_dir()
    except Exception:
        return None


def _get_run_output_dir(script_title: str, create: bool = True) -> Path:
    """Resolve run directory for this execution.

    Rules:
    - If configured output_dir is the bare Final root, derive Final/<safe_title>.
    - If configured output_dir is any subdirectory/custom path, use it exactly.
    """
    configured_dir = _get_output_parent_dir()
    default_root = DEFAULT_OUTPUT_PARENT.resolve()

    try:
        configured_is_default_root = configured_dir.resolve() == default_root
    except Exception:
        configured_is_default_root = False

    if configured_is_default_root:
        run_dir = configured_dir / _safe_title_for_paths(script_title)
    else:
        run_dir = configured_dir

    if create:
        run_dir.mkdir(parents=True, exist_ok=True)
    return run_dir


def _copy_to_output_subfolder(src, subfolder: str, script_title: Optional[str] = None) -> Optional[Path]:
    """Copy a source file into {output_base}/{subfolder}/ if a base dir is configured.

    Returns the destination path on success; None if no base dir, src missing, or copy fails.
    Best-effort: errors are logged but never raised so generation flows aren't disrupted.
    """
    try:
        base = _get_run_output_dir(
            script_title, create=True) if script_title else _get_output_parent_dir()
        if not base:
            return None
        if not src:
            return None
        src_path = Path(src).expanduser()
        if not src_path.exists() or not src_path.is_file():
            return None
        dst_dir = base / subfolder
        # If the file already lives inside the destination folder, nothing to do.
        try:
            if src_path.resolve().parent == dst_dir.resolve():
                return src_path
        except Exception:
            pass
        dst_dir.mkdir(parents=True, exist_ok=True)
        dst = dst_dir / src_path.name
        # Skip identical re-copies
        if dst.exists() and dst.stat().st_size == src_path.stat().st_size:
            return dst
        import shutil as _shutil
        _shutil.copy2(src_path, dst)
        return dst
    except Exception as e:
        logger.warning(f"⚠️ Could not auto-copy '{src}' → {subfolder}/: {e}")
        return None


def _save_curl_commands_live(curl_commands, script_title: str) -> Optional[Path]:
    """Write curl commands to {output_base}/curls/<safe_title>_curls.sh as soon as they are generated.

    Returns the destination path on success, or None if no base dir / nothing to write / failure.
    """
    try:
        if not curl_commands:
            return None
        base = _get_run_output_dir(script_title, create=True)
        safe_title = re.sub(r'[^A-Za-z0-9._-]+', '_',
                            (script_title or 'script').strip()).strip('._') or 'script'
        curl_dir = base / "curls"
        curl_dir.mkdir(parents=True, exist_ok=True)
        if isinstance(curl_commands, str):
            dst = curl_dir / f"{safe_title}_curls.sh"
            dst.write_text(curl_commands, encoding="utf-8")
        else:
            import json as _json
            dst = curl_dir / f"{safe_title}_curls.json"
            dst.write_text(_json.dumps(
                curl_commands, indent=2), encoding="utf-8")
        return dst
    except Exception as e:
        logger.warning(f"⚠️ Live curl save failed: {e}")
        return None


def _sync_generated_media_to_project(project_path: Path) -> dict:
    """Copy generated media into DaVinci project folders.

    - Grok videos: {project}/broll -> {project}/bRoll
    - Generated images: fallback copy from ~/Dev/brollimages -> {project}/images
      (only used when project/images does not already contain generated files)
    """
    import shutil as _shutil

    project_path = Path(project_path)
    broll_src = project_path / "broll"
    broll_dst = project_path / "bRoll"
    images_dst = project_path / "images"

    broll_dst.mkdir(parents=True, exist_ok=True)
    images_dst.mkdir(parents=True, exist_ok=True)

    videos_copied = 0
    images_copied = 0

    # Sync Grok videos into DaVinci bRoll folder.
    if broll_src.exists() and broll_src.is_dir():
        for src in broll_src.rglob('*'):
            if not src.is_file():
                continue
            if src.suffix.lower() not in {'.mp4', '.mov', '.m4v'}:
                continue
            dst = broll_dst / src.name
            try:
                if dst.exists() and dst.stat().st_size == src.stat().st_size:
                    continue
                _shutil.copy2(src, dst)
                videos_copied += 1
            except Exception as copy_err:
                logger.warning(
                    f"⚠️ Could not copy Grok video {src.name}: {copy_err}")

    # If images are already generated into project/images, leave them as-is.
    local_generated_images = []
    for pattern in ('*.png', '*.jpg', '*.jpeg', '*.webp'):
        local_generated_images.extend(images_dst.glob(pattern))

    # Back-compat fallback: copy from legacy global brollimages folder.
    if len(local_generated_images) == 0:
        legacy_images_src = Path.home() / "Dev" / "brollimages"
        if legacy_images_src.exists() and legacy_images_src.is_dir():
            for src in legacy_images_src.iterdir():
                if not src.is_file() or src.suffix.lower() not in {'.png', '.jpg', '.jpeg', '.webp'}:
                    continue
                dst = images_dst / src.name
                try:
                    if dst.exists() and dst.stat().st_size == src.stat().st_size:
                        continue
                    _shutil.copy2(src, dst)
                    images_copied += 1
                except Exception as copy_err:
                    logger.warning(
                        f"⚠️ Could not copy legacy image {src.name}: {copy_err}")

    return {
        "videos_copied": videos_copied,
        "images_copied": images_copied,
        "broll_source": str(broll_src),
        "broll_target": str(broll_dst),
        "images_target": str(images_dst),
    }


# Paths already added above - no need to duplicate

# Import text processing functions

app = Flask(__name__)
# Allow large audio/video uploads (Whisper transcription, etc.) — 2 GB cap.
app.config["MAX_CONTENT_LENGTH"] = 2 * 1024 * 1024 * 1024

# ---------------------------------------------------------------------------
# Optional token-based auth gate.
#
# When env APP_AUTH_TOKEN is set, every request OUTSIDE the public allowlist
# must supply the token via one of:
#   - X-App-Token: <token>      header
#   - ?app_token=<token>        query string
#   - app_token=<token>         cookie  (set by GET /login?token=...)
#
# When unset, the app is open (current behavior — used for local dev).
#
# Public paths (always allowed) live in PUBLIC_PREFIXES below. The video
# share page /videos and its read-only API stay public; the upload endpoint
# is explicitly protected even though it shares the same prefix.
# ---------------------------------------------------------------------------
APP_AUTH_TOKEN = os.environ.get("APP_AUTH_TOKEN", "").strip()
PUBLIC_PREFIXES = (
    "/videos",
    "/api/finished-videos",   # GET list (and local-mode file stream)
    "/static/",
    "/favicon.ico",
    "/login",
    "/healthz",
)
PROTECTED_OVERRIDES = ("/api/finished-videos/upload",)


def _request_app_token() -> str:
    return (
        request.headers.get("X-App-Token", "").strip()
        or request.args.get("app_token", "").strip()
        or request.cookies.get("app_token", "").strip()
    )


@app.before_request
def _enforce_app_auth_token():
    if not APP_AUTH_TOKEN:
        return None
    p = request.path or ""
    # Protect upload endpoint first (it lives under a public prefix).
    if any(p.startswith(x) for x in PROTECTED_OVERRIDES):
        if _request_app_token() != APP_AUTH_TOKEN:
            return jsonify({"success": False, "error": "Unauthorized"}), 401
        return None
    # Public allowlist.
    if any(p == x or p.startswith(x) for x in PUBLIC_PREFIXES):
        return None
    # Everything else: require the token.
    if _request_app_token() != APP_AUTH_TOKEN:
        if p.startswith("/api/"):
            return jsonify({"success": False, "error": "Unauthorized"}), 401
        from flask import redirect as _redirect
        return _redirect("/login")
    return None


@app.route("/login", methods=["GET"])
def app_login_page():
    """Set the app_token cookie via ?token=... and redirect to /."""
    from flask import redirect as _redirect
    token = (request.args.get("token") or "").strip()
    if APP_AUTH_TOKEN and token == APP_AUTH_TOKEN:
        resp = make_response(_redirect("/"))
        resp.set_cookie(
            "app_token", token,
            max_age=30 * 24 * 3600,
            httponly=True,
            samesite="Lax",
            secure=request.is_secure,
        )
        return resp
    html = """<!doctype html><html><head><meta charset="utf-8"><title>Sign in</title>
<style>
body{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;background:#0b1220;color:#e5e7eb;display:flex;align-items:center;justify-content:center;height:100vh;margin:0}
form{background:#111827;padding:24px;border-radius:10px;border:1px solid #1f2937;display:flex;gap:8px;flex-direction:column;min-width:300px}
input,button{padding:8px 10px;border-radius:6px;border:1px solid #334155;background:#0f172a;color:#e5e7eb;font-size:14px}
button{cursor:pointer;background:#1d4ed8;border-color:#1d4ed8}
h1{margin:0 0 6px;font-size:16px}
.hint{color:#94a3b8;font-size:12px;margin-top:4px}
</style></head><body>
<form method="get" action="/login">
<h1>🔒 Sign in</h1>
<input type="password" name="token" placeholder="Access token" autofocus />
<button type="submit">Continue</button>
<div class="hint">The /videos share page is always public.</div>
</form></body></html>"""
    return Response(html, mimetype="text/html")


progress_streams = {}
results = {}
running_tasks = {}
# Per-session cancel events for long-running, cancellable steps.
broll_image_cancel_events: "dict[str, _threading.Event]" = {}
thumbnail_cancel_events: "dict[str, _threading.Event]" = {}
grok_video_cancel_events: "dict[str, _threading.Event]" = {}
grok_video_streams: "dict[str, _queue.Queue]" = {}
grok_video_results: "dict[str, dict]" = {}
audio_render_streams: "dict[str, _queue.Queue]" = {}
audio_render_results: "dict[str, dict]" = {}
transcribe_streams: "dict[str, _queue.Queue]" = {}
transcribe_results: "dict[str, dict]" = {}
thumbnail_test_streams: "dict[str, _queue.Queue]" = {}
thumbnail_test_results: "dict[str, dict]" = {}
# Background video-upload/transcode jobs. Keyed by job_id (uuid4 hex).
# Shape: {state, percent, duration, out_time, error, video, filename, blob_name, started, finished}
video_upload_jobs: "dict[str, dict]" = {}
video_upload_jobs_lock = _threading.Lock()


@app.route('/api/proxy')
def proxy_external_page():
    """
    Lightweight HTML proxy used by the Shutterstock split-pane preview.

    Some sites (e.g. shutterstock.com) refuse to render in iframes via X-Frame-Options
    or Content-Security-Policy. We fetch the page server-side and strip those headers
    so the result can be embedded in our preview pane. We also inject a <base> tag so
    relative URLs (CSS/JS/images) continue to resolve against the original origin.

    Only http/https URLs from a small allow-list are proxied to avoid being abused as
    an open redirector or generic SSRF surface.
    """
    from urllib.parse import urlparse, urljoin

    target = (request.args.get('url') or '').strip()
    if not target:
        return Response("Missing url parameter", status=400)

    parsed = urlparse(target)
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        return Response("Invalid url", status=400)

    allowed_hosts = {
        "www.shutterstock.com",
        "shutterstock.com",
    }
    if parsed.netloc not in allowed_hosts:
        return Response("Host not allowed for proxy", status=403)

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
    }

    try:
        upstream = requests.get(target, headers=headers,
                                timeout=15, allow_redirects=True)
    except Exception as fetch_err:
        return Response(f"Proxy fetch failed: {fetch_err}", status=502)

    content_type = upstream.headers.get(
        "Content-Type", "text/html; charset=utf-8")
    body = upstream.content

    # Only rewrite HTML responses; pass through other content types untouched.
    if "text/html" in content_type.lower():
        try:
            text = body.decode(upstream.encoding or "utf-8", errors="replace")
            base_href = f"{parsed.scheme}://{parsed.netloc}/"
            base_tag = f'<base href="{base_href}">'
            # Insert <base> right after <head> if present, otherwise prepend.
            if re.search(r"<head[^>]*>", text, re.IGNORECASE):
                text = re.sub(r"(<head[^>]*>)", r"\1" + base_tag,
                              text, count=1, flags=re.IGNORECASE)
            else:
                text = base_tag + text
            # Strip any inline CSP meta tags that would block embedded scripts/styles.
            text = re.sub(
                r"<meta[^>]+http-equiv=['\"]Content-Security-Policy['\"][^>]*>",
                "",
                text,
                flags=re.IGNORECASE,
            )
            body = text.encode("utf-8")
            content_type = "text/html; charset=utf-8"
        except Exception:
            pass

    resp = Response(body, status=upstream.status_code,
                    content_type=content_type)
    # Explicitly drop framing-related headers so the iframe can render the response.
    for hdr in ("X-Frame-Options", "Content-Security-Policy", "Content-Security-Policy-Report-Only"):
        if hdr in resp.headers:
            del resp.headers[hdr]
    return resp


def _extract_broll_table_from_script(script_text: str) -> str:
    """
    Find a B-roll markdown table already embedded in a script.

    The B-roll agent emits tables shaped like:
        | Timecode | Search Term | Description | Scene Context |
        |----------|-------------|-------------|---------------|
        | 00:00:05 | ...         | ...         | ...           |

    Some scripts only have the 3-column variant (Search Term | Description | Scene Context).
    Returns the contiguous block of pipe-prefixed lines that contains the header, or "" if
    no plausible B-roll table is present.
    """
    if not script_text:
        return ""

    lines = script_text.splitlines()
    n = len(lines)
    header_idx = -1
    for i, ln in enumerate(lines):
        s = ln.strip()
        if not s.startswith("|"):
            continue
        low = s.lower()
        # Header heuristic: pipe row that names recognizable B-roll columns.
        if ("search term" in low) or ("timecode" in low and "description" in low):
            header_idx = i
            break

    if header_idx < 0:
        return ""

    # Walk forward collecting contiguous table rows (allow blank lines between rows? no — agents emit them solid).
    end_idx = header_idx
    for j in range(header_idx, n):
        if lines[j].strip().startswith("|"):
            end_idx = j
        else:
            break

    block = "\n".join(lines[header_idx:end_idx + 1]).strip()
    # Require at least one data row beyond header + separator.
    data_rows = [
        ln for ln in block.splitlines()
        if ln.strip().startswith("|") and not ln.strip().startswith("|---")
        and "search term" not in ln.lower() and "timecode" not in ln.lower()
    ]
    if not data_rows:
        return ""
    return block


@app.after_request
def add_no_cache_headers(response):
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response


EMOTION_BY_VARIATION = {
    1: "ANGRY/FRUSTRATED",
    2: "SHOCKED/SURPRISED",
    3: "SCARED/WORRIED",
    4: "EXCITED/ENERGETIC",
    5: "SKEPTICAL/DOUBTFUL",
    6: "DETERMINED/INTENSE",
}


def _safe_title_for_paths(script_title: str) -> str:
    """Normalize script title to filesystem-safe folder name used by generators."""
    normalized = (script_title or '').strip()
    normalized = re.sub(r'^\s*#\s*', '', normalized)
    normalized = re.sub(r'^\s*Direct\s+Video\s*-\s*', '',
                        normalized, flags=re.IGNORECASE)
    safe_title = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '', normalized)
    safe_title = re.sub(r'\s+', '_', safe_title).strip('_.')
    return safe_title or "untitled_script"


def _extract_script_title_for_output(script_text: str, fallback: str = "Untitled Script") -> str:
    """Extract a stable script title from generated markdown content."""
    text = script_text or ""

    # HIGHEST priority: explicit "Title: ..." line (plain or **Title:** Foo).
    # Wins over per-chapter Heading: lines and any H1/Chapter-1 fallback.
    title_match = re.search(
        r'^[ \t]*\**[ \t]*Title[ \t]*\**[ \t]*:[ \t]*\**[ \t]*(.+?)[ \t]*\**[ \t]*$',
        text,
        re.MULTILINE | re.IGNORECASE,
    )
    if title_match:
        title = title_match.group(1).strip().strip('*').strip()
        if title:
            return title

    direct_video_match = re.search(
        r'^\s*#\s*Direct\s+Video\s*-\s*(.+)$',
        text,
        re.MULTILINE | re.IGNORECASE,
    )
    if direct_video_match:
        return direct_video_match.group(1).strip() or fallback

    heading_match = re.search(r'^\s*Heading:\s*(.+)$',
                              text, re.MULTILINE | re.IGNORECASE)
    if heading_match:
        return heading_match.group(1).strip() or fallback

    h1_match = re.search(r'^#\s+(.+)$', text, re.MULTILINE)
    if h1_match:
        title = h1_match.group(1).strip()
        title = re.sub(r'^\s*Direct\s+Video\s*-\s*',
                       '', title, flags=re.IGNORECASE)
        return title or fallback

    for raw_line in text.splitlines()[:20]:
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith(("#", "*", "-", "[")):
            continue
        if all(c in "_=-~" for c in line):
            continue
        if len(line) > 150:
            continue
        return line

    return fallback


async def _save_script_md_and_docx(script_title: str, script_content: str) -> tuple[Optional[str], Optional[str]]:
    """Save script outputs as .md and .docx in {run_folder}/Script/."""
    run_dir = _get_run_output_dir(script_title, create=True)
    script_dir = run_dir / "Script"
    script_dir.mkdir(parents=True, exist_ok=True)
    safe_title = _safe_title_for_paths(script_title)

    md_path = script_dir / f"{safe_title}.md"
    md_path.write_text(script_content or "", encoding="utf-8")

    docx_path = script_dir / f"{safe_title}.docx"
    try:
        sys.path.insert(0, str(Path(__file__).parent.parent / "console_ui"))
        from word_processing import convert_markdown_to_word

        await convert_markdown_to_word(
            markdown_content=script_content or "",
            output_file_path=str(docx_path),
            template_path=None,
            title=script_title,
        )
        return str(md_path), str(docx_path)
    except Exception as e:
        logger.warning(f"⚠️ Could not save DOCX output: {e}")
        return str(md_path), None


def _guess_emotion_from_filename(filename: str) -> str:
    """Infer variation emotion label from filename pattern like *_v2_*.png."""
    match = re.search(r'_v(\d+)_', filename)
    if not match:
        return "EXISTING THUMBNAIL"
    return EMOTION_BY_VARIATION.get(int(match.group(1)), "EXISTING THUMBNAIL")


def _collect_all_thumbnail_entries(script_title: str, thumbnail_results: dict | None = None) -> list[dict]:
    """Return all thumbnails present in the script's thumbnails directory."""
    thumbnail_results = thumbnail_results or {}
    generated_variations = thumbnail_results.get("variations") or []
    output_dir_hint = thumbnail_results.get("output_dir")

    thumbnail_dir = None
    if output_dir_hint:
        thumbnail_dir = Path(output_dir_hint)
    else:
        thumbnail_dir = _get_run_output_dir(
            script_title, create=False) / "thumbnails"

    generated_by_filename = {}
    for variation in generated_variations:
        filename = variation.get("filename")
        if not filename and variation.get("filepath"):
            filename = Path(variation.get("filepath")).name
        if filename:
            generated_by_filename[filename] = variation

    if not thumbnail_dir.exists() or not thumbnail_dir.is_dir():
        return []

    image_paths = sorted(
        [
            p for p in thumbnail_dir.iterdir()
            if p.is_file() and p.suffix.lower() in {".png", ".jpg", ".jpeg"}
        ],
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )

    thumbnails = []
    for image_path in image_paths:
        variation = generated_by_filename.get(image_path.name, {})
        thumbnails.append({
            "emotion": variation.get("emotion") or variation.get("mood") or _guess_emotion_from_filename(image_path.name),
            "text": variation.get("text") or variation.get("thumbnail_text") or "Existing thumbnail",
            "filename": image_path.name,
        })

    return thumbnails


def _extract_thumbnail_hook_text_options(hook_result: dict | None = None, script_text: str = "") -> list[str]:
    """Extract up to three thumbnail hook text options from hook agent result or script text blocks."""
    hook_result = hook_result or {}

    options: list[str] = []

    direct_options = hook_result.get("thumbnail_hook_text_options") or []
    if isinstance(direct_options, list):
        for option in direct_options:
            text = (option or "").strip().strip('"').strip()
            if text and text not in options:
                options.append(text)

    direct_single = (hook_result.get("thumbnail_hook_text")
                     or "").strip().strip('"').strip()
    if direct_single and direct_single not in options:
        options.append(direct_single)

    if options:
        return options[:3]

    def _add_option(text: str):
        cleaned = (text or "").strip().strip('"').strip()
        if cleaned and cleaned not in options:
            options.append(cleaned)

    raw_candidates = [
        hook_result.get("full_response", ""),
        hook_result.get("raw_response", ""),
        script_text,
    ]

    for raw in raw_candidates:
        if not raw:
            continue

        numbered = re.findall(
            r'THUMBNAIL_HOOK_TEXT_(\d+)\s*:\s*"?([^"\n]+)"?',
            raw,
            flags=re.IGNORECASE,
        )
        if numbered:
            for _, text in sorted(numbered, key=lambda item: int(item[0])):
                _add_option(text)
            if options:
                return options[:3]

        match = re.search(
            r'THUMBNAIL_HOOK_TEXT\s*:\s*"?([^"\n]+)"?',
            raw,
            flags=re.IGNORECASE,
        )
        if match:
            _add_option(match.group(1))
            if options:
                return options[:3]

    return options[:3]


def _extract_thumbnail_hook_text(hook_result: dict | None = None, script_text: str = "") -> str:
    """Extract the first available thumbnail hook text option."""
    options = _extract_thumbnail_hook_text_options(
        hook_result=hook_result,
        script_text=script_text,
    )
    return options[0] if options else ""


class ConsoleCapture:
    """Captures console output and redirects to progress streamer"""

    def __init__(self, streamer):
        self.streamer = streamer
        self.original_stdout = sys.stdout
        self.buffer = []
        self.long_operation_start = None
        self.last_chapter_revision = None

    def write(self, message):
        """Write to both original stdout and the progress streamer"""
        if not message.strip():
            return

        # Debug: Add marker to identify when write method is called
        if "TEST 5" in message or "Minimal rapid" in message:
            self.original_stdout.write(f"🔍 WRITE DEBUG: {message.strip()}\n")
            self.original_stdout.flush()

        # Write to original stdout first
        self.original_stdout.write(message)
        self.original_stdout.flush()

        # Filter Azure SDK noise that can cause blocking
        if self._should_filter_azure_sdk_logging(message):
            return

        # Extract progress and send to web interface
        progress = self._extract_progress(message)
        if progress is not None and self.streamer:
            # Only send updates when we have explicit progress values
            try:
                self.streamer.send_update(message.strip(), progress)
            except Exception as e:
                self.original_stdout.write(
                    f"⚠️ Stream update failed: {e}\n")
                self.original_stdout.flush()
        # Don't send messages without progress - they would default to 50%

    def _track_long_operations(self, message):
        """Track when long operations start for timing purposes"""
        import time

        # Track when chapter revision starts
        if "Reviewing Chapter" in message and "revision failed" not in message:
            chapter_match = re.search(r'Chapter (\d+)', message)
            if chapter_match:
                chapter_num = int(chapter_match.group(1))
                self.last_chapter_revision = {
                    'chapter': chapter_num,
                    'start_time': time.time()
                }

        # Track other long operations
        elif any(phrase in message for phrase in [
            "Script Reviewer", "Script review", "Reviewing", "Revising"
        ]) and "failed" not in message and "completed" not in message:
            if self.long_operation_start is None:
                self.long_operation_start = time.time()

    def _enhance_message_with_timing(self, primary_text, fallback_text):
        """Add timing information to progress messages"""
        import time

        msg = primary_text if len(primary_text) > 10 else fallback_text

        # Add elapsed time for chapter revisions
        if self.last_chapter_revision and "revised" in msg:
            elapsed = time.time() - self.last_chapter_revision['start_time']
            if elapsed > 30:  # Only show timing after 30 seconds
                chapter_num = self.last_chapter_revision['chapter']
                mins, secs = divmod(int(elapsed), 60)
                if mins > 0:
                    msg = f"{msg} (Chapter {chapter_num} took {mins}m {secs}s)"
                else:
                    msg = f"{msg} (Chapter {chapter_num} took {secs}s)"

        # Add elapsed time for other long operations
        elif self.long_operation_start and any(phrase in msg for phrase in [
            "reviewing", "revising", "processing"
        ]):
            elapsed = time.time() - self.long_operation_start
            if elapsed > 60:  # Show timing after 1 minute
                mins, secs = divmod(int(elapsed), 60)
                msg = f"{msg} (running {mins}m {secs}s)"

        return msg

    def _should_filter_initialization_message(self, message):
        """Filter out initialization messages not relevant to script work"""
        initialization_patterns = [
            "TournamentAgentClient initialized",
            "TournamentAgent model client created",
            "AITipsAgentClient initialized",
            "AITipsAgent model client created",
            "📡 AUTOGEN SYSTEM TRACE: Initializing agents",
            "🤖 AUTOGEN SYSTEM TRACE: Creating model clients",
            "🎯 AUTOGEN SYSTEM TRACE: System fully initialized"
        ]

        return any(pattern in message for pattern in initialization_patterns)

    def _should_filter_azure_sdk_logging(self, message):
        """Filter out Azure AI SDK internal logging that can block operations"""
        # Only filter very specific Azure SDK internal messages that cause blocking
        blocking_patterns = [
            "INFO - ManagedIdentityCredential will use IMDS",
            "INFO - Request URL:",
            "Request method:",
            "Request headers:",
            "No body was attached to the request",
            "A body is sent with the request",
            "INFO - Response status:",
            "Response headers:",
            "INFO - DefaultAzureCredential acquired a token",
            "User-Agent': 'azsdk-python-identity",
            "User-Agent': 'AIProjectClient azsdk-python-ai-agents",
            "Authorization': 'REDACTED'",
            "x-ms-client-request-id':",
            "Content-Type': 'application/json'",
            "openai-processing-ms':",
            "Date': 'Fri, 26 Sep 2025",
            "No environment configuration found."
        ]

        # Keep important workflow messages even if they contain Azure terms
        important_workflow_patterns = [
            "STEP 1:", "STEP 2:", "STEP 3:", "STEP 4:",
            "Topic enhanced", "Chapter", "Script", "Writing", "Reviewing", "Revising",
            "completed", "revised", "Assembling", "WORKFLOW", "DEBUG:"
        ]

        # Don't filter if it's an important workflow message
        if any(pattern in message for pattern in important_workflow_patterns):
            return False

        return any(pattern in message for pattern in blocking_patterns)

    def _extract_progress(self, message):
        """
        Extract progress based on actual workflow messages.

        Progress Distribution (matches actual workflow timeline):
        - 0-10%: Initialization (system startup)
        - 10-20%: Topic Enhancement (STEP 1)
        - 20-30%: Start Chapter Writing (STEP 2 begins)
        - 30-55%: Chapter Writing Progress (based on completed chapters)
        - 55-65%: Script Review (STEP 3)
        - 65-80%: Chapter-by-Chapter Review Progress
        - 80-83%: YouTube Metadata
        - 83-86%: B-roll Search Terms
        - 86-89%: Demo Packages
        - 89-90%: B-roll Terms (second generation)
        - 90-100%: Thumbnail Generation & Completion
        """
        import re

        # ========== INITIALIZATION (0-10%) ==========
        if "Starting script creation" in message:
            return 1
        elif "Initializing script creation system" in message:
            return 3
        elif "Console capture test" in message:
            return 5
        elif "Creating EnhancedAutoGenSystem" in message:
            return 6
        elif "Initializing Enhanced AutoGen" in message:
            return 7
        elif "All 6 agents initialized successfully" in message:
            return 10

        # ========== STEP 1: TOPIC ENHANCEMENT (10-20%) ==========
        elif "STEP 1: Topic Enhancement" in message or "Topic Enhancement & Chapter Planning" in message:
            return 12
        elif "Running Script-Topic-Assistant-Agent" in message:
            return 14
        elif "Topic enhanced" in message and "chars" in message:
            return 20

        # Chapter extraction debug and breakdown display (20-21%)
        elif "DEBUG: Found" in message and "chapter matches" in message:
            return 20  # Show debug info right after topic enhanced
        elif "Raw Match" in message or "Skipped:" in message or "Trimming from" in message:
            return 20  # Show all debug messages
        elif "Chapter Planning Complete [21%]" in message:
            return 21
        elif "📋 CHAPTER BREAKDOWN:" in message:
            return 21
        elif "Chapter" in message and ":" in message and "/" not in message and "completed" not in message:
            # Matches "Chapter 1: Title" but not "Chapter 1/5" or "Chapter 1 completed"
            return 21

        # ========== STEP 2: CHAPTER WRITING (20-55%) ==========
        elif "STEP 2: Chapter-by-Chapter Script Writing" in message:
            return 22
        elif "Fetching style references" in message:
            return 24
        elif "Retrieved transcript style references" in message:
            return 26

        # Filter out "Writing Chapter X/Y" - don't update progress when starting
        elif "Writing Chapter" in message and "/" in message:
            return None  # Don't change progress while starting to write

        # Track individual chapter completion ONLY when done (26% to 55%)
        elif "Chapter" in message and "completed" in message and "chars" in message:
            match = re.search(r'Chapter (\d+)/(\d+)', message)
            if match:
                completed_num, total = int(match.group(1)), int(match.group(2))
                # Progress from 26% to 55%
                chapter_progress = int(26 + (completed_num * 29 / total))
                return min(chapter_progress, 55)
            return 40

        elif "Combined all" in message and "chapters into full script" in message:
            return 55

        # ========== STEP 3: SCRIPT REVIEW (55-80%) ==========
        elif "STEP 3: Script Review" in message:
            return 57
        elif "Running Script-Review-Agent" in message and "Script-Review-Agent" in message:
            return 59
        elif "Script reviewed" in message and "chars" in message:
            return 62
        elif "Retrying with chapter-by-chapter revision" in message:
            return 64

        # Filter out "Reviewing Chapter X/Y" - don't update progress when starting
        elif "Reviewing Chapter" in message and "/" in message:
            return None  # Don't change progress while starting review

        # Reviewer debug output (show what reviewer returned)
        elif "🔍 DEBUG - Reviewer Response Preview" in message:
            return 65  # Show during review phase
        elif "Contains '" in message and "': " in message:
            return 65  # Show debug checks

        # Track chapter-by-chapter review progress ONLY when complete (65% to 78%)
        elif "Chapter" in message and "revised" in message and "chars" in message:
            match = re.search(r'Chapter (\d+)/(\d+)', message)
            if match:
                completed_num, total = int(match.group(1)), int(match.group(2))
                # Progress from 65% to 78%
                review_progress = int(65 + (completed_num * 13 / total))
                return min(review_progress, 78)
            return 70

        elif "Assembling" in message and "revised chapters" in message:
            return 79
        elif "All" in message and "chapters revised individually" in message:
            return 80

        # Step 3 failure/skip paths (previously invisible to the UI)
        elif "Script Review failed or timed out" in message:
            return 58
        elif "FALLBACK: Skipping review" in message:
            return 58
        elif "revision failed, using original" in message:
            return 70
        elif "Script Review skipped" in message:
            return 80

        # Review feedback summary
        elif "📋 REVIEW FEEDBACK SUMMARY" in message:
            return 80
        elif "Total chapters reviewed:" in message:
            return 80
        elif "Full review feedback saved to:" in message:
            return 80

        # ========== STEP 3.5: QUOTES & STATISTICS (80-81%) ==========
        elif "STEP 3.5: Quotes & Statistics Generation" in message:
            return 80
        elif "Quotes & Statistics generated" in message and "chars" in message:
            return 81
        elif "Quotes/Stats generation failed" in message:
            return 81
        elif "Exception in Quotes/Stats generation" in message:
            return 81
        elif "Continuing workflow without quotes/stats" in message:
            return 81
        elif "Quotes & Statistics section inserted" in message:
            return 81

        # ========== STEP 4: HOOK & SUMMARY (81-83%) ==========
        elif "STEP 4: Hook & Summary Generation" in message:
            return 81
        elif "Hook generated" in message and "chars" in message:
            return 82
        elif "Summary generated" in message and "chars" in message:
            return 82
        elif "Hook prepended to script" in message:
            return 82
        elif "Summary appended to script" in message:
            return 83

        # ========== POST-PROCESSING (83-90%) ==========
        elif "SEQUENTIAL WORKFLOW COMPLETED" in message:
            return 83
        elif "Script enhanced with formatting" in message:
            return 83

        # YouTube Metadata (83-85%)
        elif "Generating YouTube upload metadata" in message:
            return 83
        elif "YouTube metadata generated" in message:
            return 85

        # B-roll Search Terms (85-87%)
        elif "Generating B-roll search terms" in message:
            return 85
        elif "B-roll table generated" in message:
            return 87

        # Demo Packages (87-90%)
        elif "Creating demo packages" in message:
            return 88
        elif "Demo packages generated" in message:
            return 90

        # ========== THUMBNAIL GENERATION (90-100%) ==========
        elif "Generating" in message and "thumbnail" in message:
            return 92
        elif "Generated" in message and "thumbnails" in message:
            return 98
        elif "Script creation completed" in message:
            return 100

        return None  # Don't change progress for unrecognized messages

    def flush(self):
        self.original_stdout.flush()

    def fileno(self):
        """Return file descriptor for compatibility"""
        return self.original_stdout.fileno()

    def isatty(self):
        """Return whether this is a terminal"""
        return self.original_stdout.isatty()

    def readable(self):
        """Return whether file is readable"""
        return False

    def writable(self):
        """Return whether file is writable"""
        return True

    def seekable(self):
        """Return whether file is seekable"""
        return False

    def __getattr__(self, name):
        """Delegate unknown attributes to original stdout"""
        return getattr(self.original_stdout, name)


class ProgressStreamer:
    """Handles Server-Sent Events progress streaming"""

    def __init__(self, session_id: str):
        self.session_id = session_id
        self.queue = queue.Queue()
        self.result = None
        self.done = False

    def send_update(self, message: str, progress: int = None):
        """Send a progress update"""
        # Don't send updates with no progress - they would reset to default
        if progress is None:
            return

        update = {
            "message": message,
            "progress": progress,
            "timestamp": time.time(),
            "done": progress == 100
        }

        try:
            logger.info(
                f"🔍 send_update: putting message in queue for session {self.session_id}")
            self.queue.put(json.dumps(update))
            logger.info(f"🔍 send_update: message queued successfully")
            if progress == 100:
                logger.info(
                    f"🔍 send_update: marking done=True for session {self.session_id}")
                self.done = True
            logger.info(f"🔍 send_update: method completed successfully")
        except Exception as e:
            logger.error(f"Failed to send update: {e}")


async def process_script_creation(session_id, topic, audience, tone,
                                  video_length, production_type, goals,
                                  quick_test=False, checkboxes=None,
                                  heygen_template_id="", heygen_api_key="",
                                  heygen_voice_id="", grok_api_key=""):
    """Clean script creation with only console capture"""
    logger.info(f"🎬 SCRIPT CREATION STARTED: session={session_id}")
    if quick_test:
        logger.info("⚡ QUICK TEST: 1-chapter mode enabled")

    # Default to script only if no checkboxes provided
    if checkboxes is None:
        checkboxes = {"script": True}

    try:
        if session_id not in progress_streams:
            logger.error(f"❌ Session {session_id} not found!")
            return

        streamer = progress_streams[session_id]
        resolved_script_title = topic
        run_output_dir = _get_run_output_dir(
            resolved_script_title, create=True)

        # Initial progress
        streamer.send_update("🚀 Initializing script creation system...", 5)

        # Import the system
        from linedrive_azure.agents.enhanced_autogen_system import (
            EnhancedAutoGenSystem
        )

        # Set up console capture
        console_capture = ConsoleCapture(streamer)
        original_stdout = sys.stdout

        # Note: Removed background progress updater thread as it was
        # interfering with message-based progress tracking

        try:
            # Redirect stdout to capture console output
            sys.stdout = console_capture

            # Send a test message to verify streaming works
            streamer.send_update(
                "🧪 Console capture test - starting script creation", 5)

            # Initialize system (this will be captured)
            print("🔧 Creating EnhancedAutoGenSystem...")
            system = EnhancedAutoGenSystem(verbose=True)
            print("🔧 All 6 agents initialized successfully")

            # Run the workflow with timeout to prevent hanging
            print("🚀 Starting script workflow with 20-minute timeout...")

            # DEBUG: Add detailed logging for hang investigation
            print("🔍 DEBUG: About to call async workflow method")
            print(f"📋 Checkboxes received: {checkboxes}")
            print(f"🎯 quick_test: {quick_test}, topic: {topic[:50]}...")
            print(
                f"🔍 DEBUG: Current thread: {threading.current_thread().name}")
            print(
                f"🔍 DEBUG: ConsoleCapture active: {isinstance(sys.stdout, ConsoleCapture)}")

            try:
                print("🔍 DEBUG: Calling asyncio.wait_for with async method")
                result = await asyncio.wait_for(
                    system.run_complete_script_workflow_sequential(
                        script_topic=topic,
                        topic_description="",
                        audience=audience,
                        tone=tone,
                        script_length=video_length,
                        max_chapters=1 if quick_test else 8,
                    ),
                    # 20 minutes - accounts for Script Writer (~5min) + Script Review (~12min)
                    timeout=1200
                )
                print("🔍 DEBUG: Async workflow completed successfully")

            except asyncio.TimeoutError:
                print("🔍 DEBUG: AsyncIO timeout occurred")
                print("⚠️ Workflow timeout - using fallback approach")
                streamer.send_update(
                    "⚠️ Script creation taking longer than expected, "
                    "trying alternative approach...", 85
                )

                # Fallback: Create a simpler script without revision
                result = {
                    "success": False,
                    "error": "Workflow timeout - chapter revision taking too long",
                    "timeout": True
                }
            except asyncio.TimeoutError:
                print("⚠️ Workflow timeout - using fallback approach")
                streamer.send_update(
                    "⚠️ Script creation taking longer than expected, "
                    "trying alternative approach...", 85
                )

                # Fallback: Create a simpler script without revision
                result = {
                    "success": False,
                    "error": "Workflow timeout - chapter revision taking too long",
                    "timeout": True
                }

        finally:
            # Always restore stdout
            sys.stdout = original_stdout

        # Debug: Check what we actually got
        print(f"🔍 DEBUG: Result type: {type(result)}")
        if isinstance(result, dict):
            print(f"🔍 DEBUG: Result keys: {list(result.keys())}")
        else:
            print(f"🔍 DEBUG: Result value: {result}")
            # Convert string result to error format
            result = {"success": False, "error": str(result)}

        if result.get("success"):
            # EXACT COPY OF WORKING CONSOLE UI LOGIC
            print(f"\n🎉 COMPLETE 4-AGENT WORKFLOW FINISHED!")
            print("=" * 60)
            print(f"📋 Topic: {topic}")
            print(f"👥 Audience: {audience}")
            print(f"💬 Tone: conversational and educational")
            print(f"⏱️ Length: {video_length}")
            print("-" * 60)
            print(f"✅ Topic Enhanced by Topic Assistant")
            print(f"✅ Script Created by Script Writer")
            print(f"✅ Script Reviewed by Script Reviewer")
            print(f"✅ Sequential Workflow Completed Successfully")

            # Add introduction to script content (EXACT COPY)
            introduction = """Hi, I'm Roz's AI Digital Twin. She is a high-tech sales leader, bonafide tech nerd and busy Mom of two really incredible kids. She has spent her career in high-tech working to this moment and beyond and we are here to guide you through it. This is, AI with Roz.

---

"""

            # Combine introduction with script content (EXACT COPY)
            raw_script_content = introduction + result["script_content"]

            # Enhanced script with bold tool formatting (EXACT COPY - NO TIMEOUTS!)
            from console_ui.text_processing import (
                enhance_script_with_bold_tools,
                extract_tool_links_and_info,
            )

            enhanced_script_content = enhance_script_with_bold_tools(
                raw_script_content)
            print("✅ Script enhanced with bold tool formatting")
            # Force progress update for enhancement step
            streamer.send_update("✅ Script enhanced with formatting", 98)

            # Use enhanced script as final content (EXACT COPY)
            final_script_content = enhanced_script_content

            # Use the user's requested topic for folder/file naming.
            # If topic is empty, fall back to a title extracted from the generated content.
            extracted_script_title = _extract_script_title_for_output(
                final_script_content, fallback=topic)
            resolved_script_title = (
                topic or extracted_script_title or "Untitled Script").strip()
            run_output_dir = _get_run_output_dir(
                resolved_script_title, create=True)
            print(f"📁 Using run folder: {run_output_dir}")

            # Extract tool links for YouTube description (EXACT COPY - NO TIMEOUTS!)
            tool_links = extract_tool_links_and_info(final_script_content)
            tool_count = len(tool_links.splitlines())
            print(f"📊 Extracted {tool_count} tool references")
            # Force progress update for tool extraction
            streamer.send_update(
                f"📊 Extracted {tool_count} tool references", 99)

            # Add tool links to the script content for YouTube description (EXACT COPY)
            final_script_with_tools = final_script_content + "\n\n" + "=" * 60 + "\n"
            final_script_with_tools += "� YOUTUBE VIDEO DESCRIPTION\n"
            final_script_with_tools += "=" * 60 + "\n"
            final_script_with_tools += (
                "Copy the section below for your YouTube video description:\n\n"
            )
            final_script_with_tools += f"🎥 {topic}\n\n"
            final_script_with_tools += tool_links
            final_script_with_tools += "\n\n🔔 Don't forget to SUBSCRIBE for more AI tools and productivity tips!"
            final_script_with_tools += "\n� What tools would you like to see featured next? Drop a comment below!"

            # Store result (SIMPLIFIED)
            # Initialize variables that will be used later
            youtube_upload_details = None
            curl_commands = None  # Initialize curl commands variable
            edl_content = None  # Initialize EDL content variable
            edl_filename = None  # Initialize EDL filename variable
            broll_table = None
            broll_rows = []

            # NEW: Generate B-roll Table using AI Agent (skip unless checkbox checked)
            print(
                f"🔍 DEBUG: Checking broll - checkbox: {checkboxes.get('broll', False)}, quick_test: {quick_test}")
            if not quick_test and checkboxes.get("broll", False):
                print("✅ DEBUG: B-roll condition TRUE - calling agent")
                try:
                    print("\n📊 Generating B-roll Search Terms Table...")
                    streamer.send_update(
                        "📊 Generating B-roll search terms...", 97.5)

                    try:
                        from linedrive_azure.agents import ScriptBRollAgentClient
                    except ImportError as import_err:
                        print(f"⚠️ B-roll agent not available: {import_err}")
                        raise

                    broll_agent = ScriptBRollAgentClient()
                    print("✅ B-roll agent initialized")

                    # Generate B-roll table WITH timecodes for EDL export
                    broll_result = broll_agent.generate_broll_table_with_timecodes(
                        script_content=final_script_content,
                        script_title=topic,
                        words_per_minute=150,  # Adjust based on speaking pace
                        timeout=180
                    )

                    if broll_result.get("success", False):
                        broll_table = broll_result.get("table", "")
                        parsed_data = broll_result.get("parsed_data", [])
                        broll_rows = parsed_data

                        print(
                            f"✅ B-roll table generated ({len(broll_table)} characters, {len(parsed_data)} entries)")
                        streamer.send_update(
                            "✅ B-roll table generated", 98.5)

                        # Append B-roll table to script
                        broll_section = f"\n\n{'=' * 80}\n"
                        broll_section += "# 🎬 B-ROLL SEARCH TERMS TABLE WITH TIMECODES\n"
                        broll_section += f"{'=' * 80}\n\n"
                        broll_section += broll_table

                        final_script_with_tools += broll_section
                        print("✅ B-roll table appended to script")

                        # Generate EDL file for DaVinci Resolve
                        edl_content = None
                        edl_filename = None
                        if parsed_data:
                            try:
                                import time

                                # Create filename with timestamp
                                timestamp = time.strftime("%Y%m%d_%H%M%S")
                                edl_filename = f"broll_markers_{timestamp}.edl"
                                edl_dir = run_output_dir / "MDL"
                                edl_dir.mkdir(parents=True, exist_ok=True)
                                edl_path = edl_dir / edl_filename

                                edl_result = broll_agent.create_edl_markers(
                                    broll_data=parsed_data,
                                    output_file=str(edl_path),
                                    frame_rate='24'
                                )

                                if edl_result.get("success"):
                                    print(
                                        f"✅ EDL file created: {edl_filename} ({edl_result.get('marker_count')} markers)")
                                    streamer.send_update(
                                        f"✅ EDL file created: {edl_filename}", 98.7)

                                    # Read EDL content for frontend display
                                    try:
                                        with open(edl_path, 'r', encoding='utf-8') as f:
                                            edl_content = f.read()
                                        print(
                                            f"✅ EDL content read ({len(edl_content)} bytes)")
                                    except Exception as read_error:
                                        print(
                                            f"⚠️ Could not read EDL file: {read_error}")

                                    # Add EDL info to script
                                    edl_section = f"\n\n{'=' * 80}\n"
                                    edl_section += "# 📋 DAVINCI RESOLVE EDL MARKERS\n"
                                    edl_section += f"{'=' * 80}\n\n"
                                    edl_section += f"**EDL File Generated:** `{edl_filename}`\n\n"
                                    edl_section += f"**Total Markers:** {edl_result.get('marker_count')}\n\n"
                                    edl_section += "**Import Instructions:**\n"
                                    edl_section += "1. Open DaVinci Resolve\n"
                                    edl_section += "2. Go to File > Import > Timeline Markers from EDL\n"
                                    edl_section += f"3. Select the file: `{edl_filename}`\n"
                                    edl_section += "4. Markers will be placed at the specified timecodes\n"
                                    edl_section += "5. Each marker shows the B-roll search term and description\n\n"

                                    final_script_with_tools += edl_section
                                    print("✅ EDL instructions appended to script")
                                else:
                                    print(
                                        f"⚠️ EDL creation failed: {edl_result.get('error')}")

                            except Exception as edl_error:
                                print(f"⚠️ EDL generation error: {edl_error}")
                        else:
                            print("⚠️ No parsed data available for EDL generation")
                    else:
                        print(
                            f"⚠️ B-roll table generation failed: {broll_result.get('error')}")

                except Exception as broll_error:
                    print(f"⚠️ B-roll table generation error: {broll_error}")
            else:
                print("⚡ Quick test: Skipping B-roll table generation")

            # NEW: Generate B-roll Images using Gemini (only if broll_images checkbox checked)
            broll_images = None
            if not quick_test and checkboxes.get("broll_images", False) and broll_table:
                try:
                    print("\n🎨 Generating B-roll Images...")
                    streamer.send_update("🎨 Generating B-roll images...", 98.8)

                    from tools.media.broll_image_generator import BRollImageGenerator

                    broll_gen = BRollImageGenerator()
                    print("✅ B-roll image generator initialized")

                    def _create_broll_progress(message, current, total, image_info=None):
                        try:
                            streamer.send_update(message, 99)
                        except Exception:
                            pass
                        try:
                            if image_info and image_info.get("success"):
                                src = image_info.get(
                                    "filename") or image_info.get("filepath")
                                dst = _copy_to_output_subfolder(
                                    src, "images", script_title=resolved_script_title)
                                if dst:
                                    streamer.send_update(
                                        f"💾 Saved → {dst}", 99)
                        except Exception as copy_err:
                            logger.warning(
                                f"⚠️ live copy of B-roll image failed: {copy_err}")

                    # Generate images from ALL entries in the B-roll table
                    # max_images now refers to max ENTRIES (each gets 3 variations)
                    # Set to None to generate ALL entries with 3 variations each
                    _broll_out_dir = run_output_dir / "images"
                    image_results = broll_gen.generate_all_broll_images(
                        broll_table=broll_table,
                        script_title=topic,
                        # Generate all entries (matching script processing workflow)
                        max_images=None,
                        progress_callback=_create_broll_progress,
                        output_dir=_broll_out_dir,
                    )

                    if image_results.get("success"):
                        broll_images = image_results.get("images", [])
                        entries_count = image_results.get("total_entries", 0)
                        variations = image_results.get(
                            "variations_per_entry", 3)
                        print(
                            f"✅ Generated {len(broll_images)} B-roll images ({entries_count} entries × {variations} variations)")
                        streamer.send_update(
                            f"✅ Generated {len(broll_images)} B-roll images ({entries_count} entries × {variations} variations)", 99)
                    else:
                        print(
                            f"⚠️ B-roll image generation failed: {image_results.get('error')}")

                except Exception as broll_img_error:
                    print(
                        f"⚠️ B-roll image generation error: {broll_img_error}")

            # NEW: Defer Grok AI generation until user selects rows from B-roll table
            grok_videos = []
            if not quick_test and checkboxes.get("grok_videos", False):
                if not broll_table:
                    # Prefer reusing a B-roll table already embedded in the script before regenerating one.
                    try:
                        extracted_table = _extract_broll_table_from_script(
                            final_script_content)
                    except Exception:
                        extracted_table = ""
                    if extracted_table:
                        broll_table = extracted_table
                        print(
                            f"📋 Grok Videos: reusing existing B-roll table from script ({len(broll_table)} chars)")
                        streamer.send_update(
                            "📋 Grok Videos: using existing B-roll table from script", 98
                        )
                if not broll_table:
                    print(
                        "🤖 Grok Videos: broll_table not available, auto-generating...")
                    streamer.send_update(
                        "🤖 Grok Videos: generating B-roll table first...", 98)
                    try:
                        from linedrive_azure.agents import ScriptBRollAgentClient
                        broll_agent = ScriptBRollAgentClient()
                        broll_result = broll_agent.generate_broll_table_with_timecodes(
                            script_content=final_script_content,
                            script_title=topic,
                            words_per_minute=150,
                            timeout=300
                        )
                        if broll_result.get("success", False):
                            broll_table = broll_result.get("table", "")
                            broll_rows = broll_result.get("parsed_data", [])
                            print(
                                f"✅ Auto-generated broll_table ({len(broll_table)} chars) for Grok Videos")
                        else:
                            print(
                                "⚠️ Auto broll_table generation failed for Grok Videos")
                    except Exception as auto_broll_err:
                        print(f"❌ Auto broll_table error: {auto_broll_err}")

                if broll_table:
                    total_rows = len(broll_rows) if broll_rows else 0
                    streamer.send_update(
                        f"⏸️ Select B-roll rows for Grok generation ({total_rows} available)",
                        99
                    )

            # NEW: Generate HeyGen Ready Section (only if checkbox checked)
            if checkboxes.get("heygen", False):
                try:
                    print("\n🎬 Generating HeyGen Ready section...")
                    from console_ui.text_processing import extract_heygen_host_script

                    heygen_script = extract_heygen_host_script(
                        final_script_content)

                    # Fallback: if no Host: markers found, use the script content directly
                    if not heygen_script:
                        print(
                            "⚠️ No Host: markers found - using full script for HeyGen section")
                        heygen_script = final_script_content

                    if heygen_script:
                        heygen_section = f"\n\n{'=' * 80}\n"
                        heygen_section += "# 🎬 HEYGEN READY SCRIPT\n"
                        heygen_section += f"{'=' * 80}\n\n"
                        heygen_section += heygen_script

                        final_script_with_tools += heygen_section
                        print(
                            f"✅ HeyGen section added ({len(heygen_script)} characters, {len(heygen_script.split())} words)")

                        # Generate curl commands only if curl checkbox is checked
                        print(
                            f"🔍 DEBUG: curl checkbox = {checkboxes.get('curl', False)}")
                        if checkboxes.get("curl", False):
                            try:
                                from console_ui.text_processing import generate_heygen_curl_commands

                                print("\n🔨 Generating HeyGen curl commands...")
                                curl_commands = generate_heygen_curl_commands(
                                    final_script_with_tools,
                                    topic,  # script_title
                                    heygen_api_key,
                                    heygen_template_id,
                                    heygen_voice_id
                                )

                                if curl_commands:
                                    # Store curl commands separately (don't append to script)
                                    print(
                                        f"✅ Generated {curl_commands.count('curl --request POST')} curl commands")
                                    _live_curl_dst = _save_curl_commands_live(
                                        curl_commands, resolved_script_title)
                                    if _live_curl_dst:
                                        print(
                                            f"💾 Saved curl commands → {_live_curl_dst}")
                                else:
                                    print("⚠️ No curl commands generated")
                                    curl_commands = None

                            except Exception as curl_error:
                                print(
                                    f"⚠️ Curl generation error: {curl_error}")
                                curl_commands = None
                    else:
                        print("⚠️ No host dialogue found for HeyGen section")

                except Exception as heygen_error:
                    print(
                        f"⚠️ HeyGen section generation error: {heygen_error}")

            # NEW: Generate Demo Packages (skip unless checkbox checked)
            demo_packages = None
            if not quick_test and checkboxes.get("demo", False):
                try:
                    print("\n🎥 Generating Demo Packages...")
                    streamer.send_update("🎥 Creating demo packages...", 96)

                    from linedrive_azure.agents.openai_demo_agent_client import OpenAIDemoAgentClient

                    demo_client = OpenAIDemoAgentClient()
                    demo_result = demo_client.generate_demo_packages(
                        final_script_content,
                        max_tokens=12000,
                        audience=audience
                    )

                    if demo_result.get("success"):
                        demo_packages = demo_result["response"]
                        print(
                            f"✅ Demo packages created ({len(demo_packages)} characters)")
                        streamer.send_update("✅ Demo packages generated", 97)

                        # Extract HeyGen content from demos
                        from console_ui.text_processing import (
                            extract_demo_heygen_content,
                            format_demo_steps_plain
                        )
                        demo_heygen_content = extract_demo_heygen_content(
                            demo_packages)

                        if demo_heygen_content:
                            print(
                                f"✅ Demo HeyGen section prepared ({len(demo_heygen_content)} characters)")

                        # Format demo packages - remove numbers/bullets from steps
                        formatted_demo_packages = format_demo_steps_plain(
                            demo_packages)

                        # Append demo packages to final script
                        demo_section = f"\n\n{'=' * 80}\n"
                        demo_section += "# 🎥 DEMO PACKAGES\n"
                        demo_section += f"{'=' * 80}\n\n"
                        demo_section += formatted_demo_packages

                        final_script_with_tools += demo_section
                        print("✅ Demo packages appended to script")
                    else:
                        print(
                            f"⚠️ Demo generation failed: {demo_result.get('error', 'Unknown error')}")

                except Exception as demo_error:
                    print(f"⚠️ Demo generation error: {demo_error}")
            else:
                print("⚡ Quick test: Skipping demo package generation")

            # Store the original script before flow analysis
            original_script_before_flow = final_script_with_tools
            flow_analysis_report = None

            # NEW: Flow & Repetition Analysis (only if checkbox checked)
            # IMPORTANT: This runs AFTER all content is assembled
            if checkboxes.get("flow_analysis", False):
                try:
                    print("\n🔄 Analyzing script for repetition and flow...")
                    streamer.send_update(
                        "🔄 Analyzing script for repetition and flow...", 96)

                    from linedrive_azure.agents import ScriptRepeatAndFlowAgentClient

                    flow_agent = ScriptRepeatAndFlowAgentClient()
                    flow_result = flow_agent.analyze_and_improve_flow(
                        script_content=final_script_with_tools,
                        script_title=topic,
                        target_audience=audience,
                        timeout=300
                    )

                    logger.info(
                        f"📊 Flow agent result: success={flow_result.get('success')}, keys={list(flow_result.keys())}")

                    if flow_result.get("success", False):
                        improved_script = flow_result.get(
                            "improved_script", "")
                        repetition_analysis = flow_result.get(
                            "repetition_analysis", "")
                        flow_improvements = flow_result.get(
                            "flow_analysis", "")

                        logger.info(
                            f"📊 Improved script length: {len(improved_script)}, rep analysis: {len(repetition_analysis)}, flow: {len(flow_improvements)}")

                        if improved_script:
                            # Build analysis report
                            analysis_section = f"\n\n{'=' * 80}\n"
                            analysis_section += "# 🔄 FLOW & REPETITION ANALYSIS\n"
                            analysis_section += f"{'=' * 80}\n\n"

                            if repetition_analysis:
                                analysis_section += "## Repetition Analysis\n\n"
                                analysis_section += f"{repetition_analysis}\n\n"

                            if flow_improvements:
                                analysis_section += "## Flow Improvements\n\n"
                                analysis_section += f"{flow_improvements}\n\n"

                            # Store analysis for separate display
                            flow_analysis_report = analysis_section

                            # Update final_script_with_tools to improved version (WITHOUT analysis appended)
                            final_script_with_tools = improved_script

                            streamer.send_update(
                                f"✅ Flow analysis complete - script improved ({len(improved_script)} chars)",
                                97
                            )
                            print(
                                f"✅ Flow analysis complete - script improved ({len(improved_script)} chars)")
                        else:
                            logger.warning(
                                "⚠️ Flow agent returned empty improved script")
                            print("⚠️ Flow agent returned empty improved script")
                    else:
                        logger.warning(
                            f"⚠️ Flow analysis failed - result: {flow_result}")
                        print(f"⚠️ Flow analysis failed")

                except Exception as flow_error:
                    logger.error(f"❌ Flow analysis error: {flow_error}")
                    print(f"⚠️ Flow analysis error: {flow_error}")
            else:
                print("⚡ Skipping flow & repetition analysis (not selected)")

            # NEW: Generate YouTube Upload Details (AFTER flow analysis)
            print(
                f"🔍 DEBUG: Checking youtube_details - checkbox: {checkboxes.get('youtube_details', False)}, quick_test: {quick_test}")
            if not quick_test and checkboxes.get("youtube_details", False):
                print("✅ DEBUG: YouTube Details condition TRUE - calling agent")
                try:
                    print("\n📺 Generating YouTube Upload Details...")
                    streamer.send_update(
                        "📺 Generating YouTube upload metadata...", 97)

                    from linedrive_azure.agents import YouTubeUploadDetailsAgentClient

                    youtube_agent = YouTubeUploadDetailsAgentClient()
                    youtube_result = youtube_agent.generate_upload_details(
                        script_content=final_script_content,
                        script_title=topic,
                        target_audience=audience,
                        video_length=video_length,
                        primary_keywords=None,
                        channel_focus=None,
                        timeout=180
                    )

                    if youtube_result.get("success", False):
                        youtube_upload_details = youtube_result.get(
                            "upload_details", "")
                        print(
                            f"✅ YouTube upload details generated ({len(youtube_upload_details)} characters)")
                        streamer.send_update(
                            "✅ YouTube metadata generated", 98)

                        # Append YouTube details to script
                        youtube_section = f"\n\n{'=' * 80}\n"
                        youtube_section += "# 📺 YOUTUBE UPLOAD DETAILS\n"
                        youtube_section += f"{'=' * 80}\n\n"
                        youtube_section += youtube_upload_details

                        final_script_with_tools += youtube_section
                        print("✅ YouTube details appended to script")
                    else:
                        print(
                            f"⚠️ YouTube generation failed: {youtube_result.get('error')}")

                except Exception as youtube_error:
                    print(f"⚠️ YouTube upload details error: {youtube_error}")
            else:
                print("⚡ Quick test: Skipping YouTube details generation")

            # NEW: Generate Emotional Thumbnails (only if checkbox checked)
            thumbnail_results = None
            if checkboxes.get("thumbnails", False):
                try:
                    thumbnail_hook_text_options = _extract_thumbnail_hook_text_options(
                        hook_result=result,
                        script_text=final_script_with_tools
                    )
                    thumbnail_hook_text = thumbnail_hook_text_options[
                        0] if thumbnail_hook_text_options else ""
                    print("\n" + "="*70)
                    print("🖼️  THUMBNAIL GENERATION STARTING")
                    print("="*70)
                    print(f"📝 Topic: {resolved_script_title}")
                    print(
                        f"📄 Script length: {len(final_script_content)} chars")
                    print(
                        f"🎥 YouTube details: {'Available' if youtube_upload_details else 'None'}")
                    if thumbnail_hook_text_options:
                        print(
                            f"🏷️ Thumbnail hook text options: {thumbnail_hook_text_options}")
                        streamer.send_update(
                            f"🏷️ Thumbnail hook options: {', '.join(thumbnail_hook_text_options)}", 98
                        )
                    else:
                        print(
                            "⚠️ Thumbnail hook text not found; using script title fallback")
                        streamer.send_update(
                            "⚠️ Thumbnail hook text not found; using script title", 98
                        )

                    expected_thumbnails = 6 * \
                        len(thumbnail_hook_text_options) if thumbnail_hook_text_options else 6
                    streamer.send_update(
                        f"🖼️ Generating {expected_thumbnails} thumbnail variations...", 98)

                    from tools.media.emotional_thumbnail_generator import (
                        EmotionalThumbnailGenerator,
                    )

                    print(f"\n🔧 Initializing EmotionalThumbnailGenerator...")
                    # Let generator get API key from environment (same as test page)
                    api_key = os.getenv("GOOGLE_API_KEY", "")
                    print(
                        f"   Environment API key: {api_key[:20]}... (length: {len(api_key)})")

                    # Route thumbnails into configured output_dir/thumbnails when available.
                    _thumb_out_dir = run_output_dir / "thumbnails"
                    # No api_key param - use environment
                    thumbnail_gen = EmotionalThumbnailGenerator(
                        output_dir=_thumb_out_dir)
                    print(f"✅ Generator initialized successfully")
                    print(f"   Template: {thumbnail_gen.template_path}")
                    print(f"   Output: {thumbnail_gen.output_dir}")
                    print(
                        f"   API Key set: {'Set' if thumbnail_gen.api_key else 'MISSING'}")
                    print(
                        f"   API Key matches: {thumbnail_gen.api_key == api_key}")

                    print(f"\n🎬 Calling generate_all_thumbnails()...")
                    # always fresh — never reuse a stale cancelled event
                    thumb_cancel_evt = _threading.Event()
                    thumbnail_cancel_events[session_id] = thumb_cancel_evt
                    thumbnail_results = thumbnail_gen.generate_all_thumbnails(
                        script_title=resolved_script_title,
                        script_content=final_script_content,
                        youtube_upload_details=youtube_upload_details,
                        headline_text=thumbnail_hook_text or None,
                        headline_options=thumbnail_hook_text_options or None,
                        progress_callback=lambda msg: streamer.send_update(
                            msg, 98),
                        cancel_check=thumb_cancel_evt.is_set,
                    )

                    print(f"\n🔍 THUMBNAIL GENERATION RESULTS:")
                    print(f"   Type: {type(thumbnail_results)}")
                    print(f"   Is None: {thumbnail_results is None}")

                    if thumbnail_results:
                        print(f"   Keys: {list(thumbnail_results.keys())}")
                        output_dir = thumbnail_results.get("output_dir")
                        variations = thumbnail_results.get("variations") or []
                        if variations:
                            if output_dir:
                                print(
                                    f"   📁 Thumbnails saved to: {output_dir}")
                                streamer.send_update(
                                    f"📁 Thumbnails saved to: {output_dir}", 99
                                )
                            print(f"   Variations count: {len(variations)}")
                            print(
                                f"\n✅ Generated {len(variations)} thumbnail variations")

                            # Show details of each variation
                            for i, var in enumerate(variations, 1):
                                print(
                                    f"      #{i}: {var.get('emotion')} - {var.get('filename')}")

                            streamer.send_update(
                                f"✅ Generated {len(variations)} thumbnails", 99
                            )
                            if thumbnail_results.get("cancelled"):
                                streamer.send_update(
                                    f"🛑 Thumbnail generation stopped — keeping {len(variations)} thumbnails so far",
                                    99,
                                )
                        else:
                            print(
                                "   ⚠️ Thumbnail directory created, but no thumbnails were generated")
                            if output_dir:
                                print(f"   📁 Directory: {output_dir}")
                                streamer.send_update(
                                    f"⚠️ No thumbnails generated (directory only): {output_dir}", 99
                                )
                            attempted = thumbnail_results.get(
                                "total_attempted", 0)
                            streamer.send_update(
                                f"❌ Thumbnail generation failed (0/{attempted})", 99
                            )
                            if thumbnail_results.get("error"):
                                streamer.send_update(
                                    f"❌ Thumbnail API error: {thumbnail_results.get('error')}", 99
                                )
                            print(f"   Full results: {thumbnail_results}")
                    else:
                        print("   ⚠️ thumbnail_results is None")

                    print("="*70 + "\n")

                except Exception as thumb_error:
                    import traceback
                    err_msg = f"{type(thumb_error).__name__}: {thumb_error}"
                    print("\n" + "="*70)
                    print("❌ THUMBNAIL GENERATION ERROR")
                    print("="*70)
                    print(f"Error type: {type(thumb_error).__name__}")
                    print(f"Error message: {thumb_error}")
                    print("\n📋 Full traceback:")
                    traceback.print_exc()
                    print("="*70 + "\n")
                    streamer.send_update(
                        f"❌ Thumbnail generation error: {err_msg}", 99)
                finally:
                    thumbnail_cancel_events.pop(session_id, None)

            saved_markdown_path, saved_docx_path = await _save_script_md_and_docx(
                resolved_script_title,
                final_script_with_tools,
            )

            streamer.result = {
                "success": True,
                "enhanced_script": final_script_with_tools,
                "script_title": resolved_script_title,
                "demo_packages": demo_packages,
                "youtube_details": youtube_upload_details,
                "thumbnail_results": thumbnail_results,
                "broll_images": broll_images,
                "grok_videos": grok_videos,
                "broll_table": broll_table,
                "broll_rows": broll_rows,
                "word_count": len(final_script_with_tools.split()),
                "reading_time": f"{len(final_script_with_tools.split()) // 150} min",
                "audience": audience,
                "production_type": production_type,
                "comparison_file": result.get("comparison_file") if isinstance(result, dict) else None,
                "chapter_comparisons": result.get("chapter_comparisons") if isinstance(result, dict) else None,
                "flow_original_script": original_script_before_flow if checkboxes.get("flow_analysis") else None,
                "flow_improved_script": final_script_with_tools if checkboxes.get("flow_analysis") and flow_analysis_report else None,
                "flow_analysis_report": flow_analysis_report,
                "edl_content": edl_content,
                "edl_filename": edl_filename,
                "curl_commands": curl_commands,
                "markdown_path": saved_markdown_path,
                "docx_path": saved_docx_path,
            }

            # Debug: Log what's in the result
            print(
                f"🔍 DEBUG: Result object curl_commands = {curl_commands is not None} ({len(curl_commands) if curl_commands else 0} chars)")
            print(f"🔍 DEBUG: Result keys: {list(streamer.result.keys())}")

            # Send final completion message (SIMPLE)
            print("🎯 Sending completion message...")
            print(f"DEBUG: About to send completion for session {session_id}")
            print(
                f"DEBUG: Streamer done status before completion: {streamer.done}")

            # Send the completion message
            streamer.send_update("✅ Script creation complete!", 100)

            print("✅ Completion message sent successfully")
            print(
                f"DEBUG: Streamer done status after completion: {streamer.done}")
            print("🎉 Script creation completed successfully!")

        else:
            error_msg = result.get("error", "Unknown workflow error")
            streamer.send_update(f"❌ Script creation failed: {error_msg}", -1)
            streamer.result = {"success": False, "error": error_msg}
            print(f"❌ Script creation failed: {error_msg}")
            logger.error(
                f"❌ Script creation failed for session {session_id}: {error_msg}")

    except Exception as e:
        error_msg = f"Error: {str(e)}"
        logger.error(f"💥 Exception during script creation: {error_msg}")
        if session_id in progress_streams:
            streamer = progress_streams[session_id]
            streamer.send_update(f"❌ {error_msg}", -1)
            streamer.result = {"success": False, "error": error_msg}


async def process_existing_script(
    session_id, script_content, audience, tone,
    video_length, checkboxes, heygen_template_id="",
    heygen_api_key="", heygen_voice_id="", grok_api_key=""
):
    """Process existing script with progress updates"""
    import re

    logger.info(f"📝 SCRIPT PROCESSING STARTED: session={session_id}")

    if session_id not in progress_streams:
        logger.error(f"❌ Session {session_id} not found!")
        return

    streamer = progress_streams[session_id]

    try:
        # Initialize EDL variables
        edl_content = None
        edl_filename = None

        # Calculate total steps based on checkboxes
        total_steps = sum([
            checkboxes.get("hook_summary", False),
            checkboxes.get("youtube_details", False),
            checkboxes.get("broll", False),
            checkboxes.get("broll_images", False),
            checkboxes.get("heygen", False),
            checkboxes.get("curl", False),
            checkboxes.get("demo", False),
            checkboxes.get("thumbnails", False),
            checkboxes.get("flow_analysis", False),
            checkboxes.get("grok_videos", False)
        ])

        if total_steps == 0:
            streamer.send_update("❌ No processing options selected", -1)
            streamer.result = {"success": False,
                               "error": "No options selected"}
            return

        completed_steps = 0
        progress_per_step = 90 / total_steps  # Reserve 10% for completion

        streamer.send_update("🚀 Initializing script processing...", 5)

        # Extract title from script
        script_title = "Untitled Script"
        logger.info("📰 Extracting title from script content...")

        # 1) Prefer an explicit "Heading:" line (first one wins), but skip
        #    chapter-style headings like "Heading: Chapter 1 - ..." so the
        #    real script title (top-of-doc) can be used instead.
        heading_match = None
        for _hm in re.finditer(r'^\s*Heading:\s*(.+)$', script_content, re.MULTILINE | re.IGNORECASE):
            _val = _hm.group(1).strip().strip('*').strip()
            if re.match(r'^chapter\s+\d+\b', _val, re.IGNORECASE):
                continue
            heading_match = _hm
            break
        if heading_match:
            script_title = heading_match.group(1).strip().strip('*').strip()
            logger.info(f"📰 ✅ Found 'Heading:' title: '{script_title}'")
        else:
            # 2) Fall back to first markdown H1, but skip known analysis/report headings
            _SKIP_TITLE_PATTERNS = (
                'flow analysis', 'flow analysis report', 'analysis report',
                'hook summary', 'b-roll', 'broll', 'thumbnail', 'demo package',
                'youtube details', 'heygen', 'curl commands',
                'opening hook', 'hook options',
            )
            for m in re.finditer(r'^#+\s+(.+)$', script_content, re.MULTILINE):
                candidate = m.group(1).strip()
                normalized = re.sub(r'[^a-z0-9 ]+', ' ',
                                    candidate.lower()).strip()
                if any(p in normalized for p in _SKIP_TITLE_PATTERNS):
                    continue
                # Skip chapter headings like 'Chapter 1 - ...' or 'Chapter 2:'
                if re.match(r'^chapter\s+\d+\b', normalized):
                    continue
                script_title = candidate
                logger.info(f"📰 ✅ Found H1 title: '{script_title}'")
                break
            else:
                lines = script_content.split('\n')
                for raw_line in lines[:15]:
                    line = raw_line.strip()
                    # Strip surrounding markdown bold/italic wrappers so a line
                    # like '**Direct Video - Story Power with AI**' is usable.
                    line = re.sub(r'^[\*_]+\s*', '', line)
                    line = re.sub(r'\s*[\*_]+$', '', line)
                    line = line.strip()
                    if not line or line.startswith(('#', '-', '[')):
                        continue
                    if all(c in '_=-~' for c in line):
                        continue
                    if len(line) > 150:
                        continue
                    # Skip chapter headings
                    if re.match(r'^chapter\s+\d+\b', line, re.IGNORECASE):
                        continue
                    script_title = line
                    logger.info(f"📰 ✅ Using line as title: '{script_title}'")
                    break

                script_title = re.sub(
                    r'^\s*#\s*', '', (script_title or '').strip())
                script_title = re.sub(
                    r'^\s*Direct\s+Video\s*-\s*', '', script_title, flags=re.IGNORECASE)
                script_title = script_title or "Untitled Script"
                logger.info(f"📰 ✅ Using line as title: '{script_title}'")

        # 0) FINAL OVERRIDE: an explicit "Title: ..." line in the script always
        # wins over Heading: / H1 / first-line fallbacks. This keeps the
        # DaVinci project name and run output folder consistent with the
        # human-authored Title line (e.g. "Title: Homeschool Heroes with AI").
        _explicit_title_match = re.search(
            r'^[ \t]*\**[ \t]*Title[ \t]*\**[ \t]*:[ \t]*\**[ \t]*(.+?)[ \t]*\**[ \t]*$',
            script_content,
            re.MULTILINE | re.IGNORECASE,
        )
        if _explicit_title_match:
            _explicit_title = _explicit_title_match.group(
                1).strip().strip('*').strip()
            if _explicit_title:
                script_title = _explicit_title
                logger.info(
                    f"📰 ⭐ Overriding with explicit 'Title:' line: '{script_title}'")

        # Ensure run_output_dir is always set (regardless of which title branch was taken)
        run_output_dir = _get_run_output_dir(script_title, create=True)
        logger.info(f"📁 Using run folder: {run_output_dir}")

        # Clean the script
        streamer.send_update("🧹 Cleaning script formatting...", 10)
        cleaned_script = script_content

        # Remove separator lines
        cleaned_script = re.sub(r'^[_\-=~]+\s*$', '', cleaned_script,
                                flags=re.MULTILINE)
        cleaned_script = re.sub(r'\n{3,}', '\n\n', cleaned_script)

        # Remove emojis
        cleaned_script = re.sub(r'🎬|📺|📊|🎥|🎯|💡|✨|🚀', '', cleaned_script)

        # Remove markdown formatting
        cleaned_script = re.sub(r'\*\*([^*]+)\*\*', r'\1', cleaned_script)
        cleaned_script = re.sub(r'\*([^*]+)\*', r'\1', cleaned_script)
        cleaned_script = re.sub(r'__([^_]+)__', r'\1', cleaned_script)
        cleaned_script = re.sub(r'_([^_]+)_', r'\1', cleaned_script)

        # Apply clean formatting
        cleaned_script = re.sub(r'^Host:\s*$', r'**Host:**\n\n',
                                cleaned_script, flags=re.MULTILINE)
        cleaned_script = re.sub(r'^VISUAL CUE:\s*(.+)$', r'\n**VISUAL CUE:** \1',
                                cleaned_script, flags=re.MULTILINE)

        # Remove the title from the beginning of cleaned_script to avoid duplication
        # since we'll prepend it separately with proper formatting
        if script_title and script_title in cleaned_script:
            # Remove markdown title format (# Title)
            cleaned_script = re.sub(
                r'^#\s+' + re.escape(script_title) + r'\s*\n*',
                '',
                cleaned_script,
                count=1,
                flags=re.MULTILINE
            )
            # Also remove plain title at the start
            if cleaned_script.strip().startswith(script_title):
                cleaned_script = cleaned_script.strip()[
                    len(script_title):].lstrip()

        # Initialize output variables (will build final_output after all sections are generated)
        final_output = ""
        hook_text = ""
        summary_text = ""
        flow_analysis = ""
        heygen_script = None
        curl_commands = None
        hook_result = {}  # Initialize to empty dict to prevent undefined variable errors
        thumbnail_hook_text = ""
        thumbnail_hook_text_options: list[str] = []

        # Generate Hook & Summary if requested
        if checkboxes.get("hook_summary", False):
            streamer.send_update("🎯 Generating Hook & Summary...", 15)
            try:
                from linedrive_azure.agents import HookAndSummaryAgentClient

                hook_agent = HookAndSummaryAgentClient()
                hook_result = hook_agent.generate_hook_and_summary(
                    script_content=cleaned_script,
                    script_title=script_title,
                    target_audience=audience,
                    tone=tone,
                    video_length=video_length,
                    timeout=180
                )

                if hook_result.get("success", False):
                    hook_text = hook_result.get("hook", "")
                    summary_text = hook_result.get("summary", "")
                    flow_analysis = hook_result.get("flow_analysis", "")
                    thumbnail_hook_text_options = _extract_thumbnail_hook_text_options(
                        hook_result=hook_result
                    )
                    thumbnail_hook_text = thumbnail_hook_text_options[
                        0] if thumbnail_hook_text_options else ""
                    logger.info(
                        f"🔍 DEBUG: process_existing_script hook_result keys: {list(hook_result.keys())}")
                    if thumbnail_hook_text_options:
                        logger.info(
                            f"🏷️ Thumbnail hook text options detected: {thumbnail_hook_text_options}")
                        print(
                            f"🏷️ Thumbnail hook text options detected: {thumbnail_hook_text_options}")
                        streamer.send_update(
                            f"🏷️ Thumbnail hook options: {', '.join(thumbnail_hook_text_options)}", 15
                        )
                    else:
                        logger.warning(
                            "⚠️ Thumbnail hook text missing in Hook-and-Summary response")
                        raw_response = hook_result.get(
                            "full_response", "") or hook_result.get("raw_response", "")
                        print(
                            "⚠️ Thumbnail hook text missing in Hook-and-Summary response")
                        print(
                            f"🔍 DEBUG: THUMBNAIL_HOOK_TEXT marker present: {'THUMBNAIL_HOOK_TEXT' in raw_response}")
                        print(
                            f"🔍 DEBUG: THUMBNAIL HOOK section present: {'THUMBNAIL HOOK' in raw_response.upper()}")
                        streamer.send_update(
                            "⚠️ Thumbnail hook text missing in Hook-and-Summary response",
                            15,
                        )
                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        f"✅ Hook & Summary generated "
                        f"({len(hook_text)} + {len(summary_text)} chars)",
                        int(progress)
                    )
                else:
                    logger.warning(f"⚠️ Hook & summary failed")
            except Exception as e:
                logger.error(f"❌ Hook & summary error: {e}")

        # Build script with hooks and summary
        hook1_text = ""
        hook2_text = ""
        hook3_text = ""
        opening_statement = ""

        if hook_result.get("success", False):
            # Extract all three hooks from the result
            # Use hook1 or fallback to hook (for backward compatibility)
            hook1_text = hook_result.get("hook1", hook_result.get("hook", ""))
            hook2_text = hook_result.get("hook2", "")
            hook3_text = hook_result.get("hook3", "")
            opening_statement = hook_result.get("opening_statement", "")

            # DEBUG: Log what we actually got
            logger.info(
                f"🔍 DEBUG: hook_result keys: {list(hook_result.keys())}")
            logger.info(
                f"🔍 DEBUG: hook1_text length: {len(hook1_text) if hook1_text else 0}")
            logger.info(
                f"🔍 DEBUG: hook2_text length: {len(hook2_text) if hook2_text else 0}")
            logger.info(
                f"🔍 DEBUG: hook3_text length: {len(hook3_text) if hook3_text else 0}")

            logger.info(
                f"📌 Adding {3 if hook3_text else (2 if hook2_text else 1)} hook option(s) to output")

            final_output += "**🎬 OPENING HOOK OPTIONS**\n\n"
            final_output += "*Choose one of these three hooks for your video:*\n\n"

            # Add Hook Option 1
            final_output += "**OPTION 1:**\n\n"
            final_output += f"**Host:**\n\n{hook1_text}\n\n"

            # Add Hook Option 2 if available
            if hook2_text:
                final_output += "---\n\n**OPTION 2:**\n\n"
                final_output += f"**Host:**\n\n{hook2_text}\n\n"

            # Add Hook Option 3 if available
            if hook3_text:
                final_output += "---\n\n**OPTION 3:**\n\n"
                final_output += f"**Host:**\n\n{hook3_text}\n\n"

            final_output += "---\n\n"

            # Add opening statement after hooks
            if opening_statement:
                final_output += "**📺 OPENING STATEMENT**\n\n"
                final_output += f"**Host:**\n\n{opening_statement}\n\n"
                final_output += "---\n\n"

        # Add script title with line break and bold formatting
        final_output += f"\n**{script_title}**\n\n"

        final_output += cleaned_script
        logger.info(
            f"📌 After adding cleaned_script, final_output length: {len(final_output)} chars")

        if summary_text:
            logger.info(
                f"📌 Adding summary to output ({len(summary_text)} chars)")
            final_output += "\n\n---\n\n**📝 CONCLUSION/SUMMARY**\n\n"
            final_output += f"**Host:**\n\n{summary_text}\n\n"

        if flow_analysis:
            logger.info(
                f"📌 Adding flow analysis to output ({len(flow_analysis)} chars)")
            final_output += f"\n\n{'=' * 80}\n# 📊 FLOW ANALYSIS\n"
            final_output += f"{'=' * 80}\n\n{flow_analysis}"

        # Generate YouTube Upload Details if requested
        youtube_details = None
        if checkboxes.get("youtube_details", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("📺 Generating YouTube Upload Details...",
                                 int(current_progress))
            try:
                from linedrive_azure.agents import YouTubeUploadDetailsAgentClient

                youtube_agent = YouTubeUploadDetailsAgentClient()
                youtube_result = youtube_agent.generate_upload_details(
                    script_content=cleaned_script,
                    script_title=script_title,
                    target_audience=audience,
                    video_length=video_length,
                    timeout=180
                )

                if youtube_result.get("success", False):
                    youtube_details = youtube_result.get("upload_details", "")
                    youtube_section = f"\n\n{'=' * 80}\n"
                    youtube_section += "# 📺 YOUTUBE UPLOAD DETAILS\n"
                    youtube_section += f"{'=' * 80}\n\n{youtube_details}"
                    final_output += youtube_section
                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        f"✅ YouTube details generated ({len(youtube_details)} chars)",
                        int(progress)
                    )
            except Exception as e:
                logger.error(f"❌ YouTube details error: {e}")

        # Generate B-roll Table if requested
        broll_table = None
        broll_rows = []
        if checkboxes.get("broll", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("📊 Generating B-roll Search Terms Table...",
                                 int(current_progress))
            try:
                from linedrive_azure.agents import ScriptBRollAgentClient

                broll_agent = ScriptBRollAgentClient()
                broll_result = broll_agent.generate_broll_table_with_timecodes(
                    script_content=cleaned_script,
                    script_title=script_title,
                    words_per_minute=150,
                    # Increased to 5 minutes for larger tables (40-60+ rows)
                    timeout=300
                )

                if broll_result.get("success", False):
                    broll_table = broll_result.get("table", "")
                    parsed_data = broll_result.get("parsed_data", [])
                    broll_rows = parsed_data

                    broll_section = f"\n\n{'=' * 80}\n"
                    broll_section += "# 📊 B-ROLL SEARCH TERMS TABLE\n"
                    broll_section += f"{'=' * 80}\n\n{broll_table}"
                    final_output += broll_section

                    # Generate EDL file
                    if parsed_data:
                        try:
                            import time
                            timestamp = time.strftime("%Y%m%d_%H%M%S")
                            edl_filename = f"broll_markers_{timestamp}.edl"
                            edl_dir = run_output_dir / "MDL"
                            edl_dir.mkdir(parents=True, exist_ok=True)
                            edl_path = edl_dir / edl_filename

                            edl_result = broll_agent.create_edl_markers(
                                broll_data=parsed_data,
                                output_file=str(edl_path),
                                frame_rate='24'
                            )

                            if edl_result.get("success"):
                                # Read EDL content for frontend display
                                try:
                                    with open(edl_path, 'r', encoding='utf-8') as f:
                                        edl_content = f.read()
                                    logger.info(
                                        f"✅ EDL content read ({len(edl_content)} bytes)")
                                except Exception as read_error:
                                    logger.error(
                                        f"⚠️ Could not read EDL file: {read_error}")

                                edl_info = f"\n\n**EDL File:** `{edl_filename}`"
                                edl_info += f" ({edl_result.get('marker_count')}"
                                edl_info += " markers)\n"
                                final_output += edl_info
                        except Exception as edl_error:
                            logger.error(f"❌ EDL error: {edl_error}")

                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        f"✅ B-roll table generated ({len(broll_table)} chars)",
                        int(progress)
                    )
                else:
                    logger.warning(f"⚠️ B-roll agent returned success=False")
                    streamer.send_update(
                        "⚠️ B-roll table generation failed", int(current_progress))
            except Exception as e:
                logger.error(f"❌ B-roll table error: {e}")
                import traceback
                traceback.print_exc()
                streamer.send_update(
                    f"❌ B-roll table error: {str(e)}", int(current_progress))

        # Generate B-roll Images using Gemini (ONLY if broll_images checkbox checked AND broll_table was successfully generated)
        broll_images = None
        # If the user only checked "B-roll Images" (without "B-roll Table"), try to reuse a B-roll
        # table that's already embedded in the loaded script so they don't have to re-run the agent.
        if checkboxes.get("broll_images", False) and not broll_table:
            try:
                extracted_table = _extract_broll_table_from_script(
                    cleaned_script)
                if extracted_table:
                    broll_table = extracted_table
                    current_progress = 15 + \
                        (completed_steps * progress_per_step)
                    streamer.send_update(
                        "📋 Found existing B-roll table in script — using it for image generation",
                        int(current_progress)
                    )
                    logger.info(
                        f"📋 Reusing existing B-roll table from script ({len(broll_table)} chars)"
                    )
                else:
                    current_progress = 15 + \
                        (completed_steps * progress_per_step)
                    streamer.send_update(
                        "⚠️ B-roll Images selected but no B-roll table found in script — also enable 'B-roll Table' to generate one.",
                        int(current_progress)
                    )
            except Exception as ex_err:
                logger.warning(
                    f"⚠️ Could not extract B-roll table from script: {ex_err}")

        if checkboxes.get("broll_images", False) and broll_table:
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("🎨 Generating B-roll images with AI...",
                                 int(current_progress))
            try:
                from tools.media.broll_image_generator import BRollImageGenerator

                broll_gen = BRollImageGenerator()

                # Per-session cancel hook so the user can stop mid-generation and keep partial results.
                cancel_evt = broll_image_cancel_events.setdefault(
                    session_id, _threading.Event())

                # Notify the UI of every variation in real time so the user sees progress per image.
                step_progress_base = current_progress
                step_progress_span = max(1.0, progress_per_step * 0.95)

                def _broll_progress(message: str, current: int, total: int, image_info=None):
                    try:
                        ratio = (current / total) if total else 0
                        pct = int(
                            step_progress_base + (step_progress_span * min(1.0, max(0.0, ratio))))
                        streamer.send_update(message, pct)
                    except Exception as cb_err:  # never let UI plumbing break image gen
                        logger.warning(
                            f"⚠️ B-roll progress callback error: {cb_err}")
                    # Live-copy each saved image into the user's configured output dir.
                    try:
                        if image_info and image_info.get("success"):
                            src = image_info.get(
                                "filename") or image_info.get("filepath")
                            dst = _copy_to_output_subfolder(
                                src, "images", script_title=script_title)
                            if dst:
                                streamer.send_update(f"💾 Saved → {dst}", pct)
                    except Exception as copy_err:
                        logger.warning(
                            f"⚠️ live copy of B-roll image failed: {copy_err}")

                # Generate images from the B-roll table
                # max_images now refers to max ENTRIES (each gets 3 variations)
                # Set to None to generate ALL entries, or set a limit if needed
                _broll_out_dir = run_output_dir / "images"
                image_results = broll_gen.generate_all_broll_images(
                    broll_table=broll_table,
                    script_title=script_title,
                    max_images=None,  # Generate all entries with 3 variations each
                    progress_callback=_broll_progress,
                    cancel_check=cancel_evt.is_set,
                    output_dir=_broll_out_dir,
                )

                if image_results.get("success", False) or image_results.get("cancelled"):
                    broll_images = image_results.get("images", [])
                    entries_count = image_results.get("total_entries", 0)
                    variations = image_results.get("variations_per_entry", 3)
                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    if image_results.get("cancelled"):
                        logger.info(
                            f"🛑 B-roll image generation cancelled by user — keeping {len(broll_images)} images")
                        streamer.send_update(
                            f"🛑 B-roll image generation stopped — keeping {len(broll_images)} images so far",
                            int(progress)
                        )
                    else:
                        logger.info(
                            f"✅ Generated {len(broll_images)} B-roll images ({entries_count} entries × {variations} variations)")
                        streamer.send_update(
                            f"✅ Generated {len(broll_images)} B-roll images ({entries_count} entries × {variations} variations)",
                            int(progress)
                        )
                else:
                    logger.warning(
                        f"⚠️ B-roll images generation failed: {image_results.get('error')}")
            except Exception as e:
                logger.error(f"❌ B-roll images error: {e}")
                import traceback
                traceback.print_exc()
            finally:
                # Always discard the cancel event so a later run starts fresh.
                broll_image_cancel_events.pop(session_id, None)

        # Defer Grok AI generation until user selects rows from B-roll table
        # Auto-generate broll_table if needed and not already done
        grok_videos = []
        if checkboxes.get("grok_videos", False):
            if not broll_table:
                # Prefer reusing a B-roll table already embedded in the script before regenerating one.
                try:
                    extracted_table = _extract_broll_table_from_script(
                        cleaned_script)
                except Exception:
                    extracted_table = ""
                if extracted_table:
                    broll_table = extracted_table
                    logger.info(
                        f"📋 Grok Videos: reusing existing B-roll table from script ({len(broll_table)} chars)"
                    )
                    streamer.send_update(
                        "📋 Grok Videos: using existing B-roll table from script",
                        int(15 + (completed_steps * progress_per_step))
                    )
            if not broll_table:
                logger.info(
                    "🤖 Grok Videos: broll_table not available, auto-generating...")
                streamer.send_update("🤖 Grok Videos: generating B-roll table first...",
                                     int(15 + (completed_steps * progress_per_step)))
                try:
                    from linedrive_azure.agents import ScriptBRollAgentClient
                    broll_agent = ScriptBRollAgentClient()
                    broll_result = broll_agent.generate_broll_table_with_timecodes(
                        script_content=cleaned_script,
                        script_title=script_title,
                        words_per_minute=150,
                        timeout=300
                    )
                    if broll_result.get("success", False):
                        broll_table = broll_result.get("table", "")
                        broll_rows = broll_result.get("parsed_data", [])
                        logger.info(
                            f"✅ Auto-generated broll_table ({len(broll_table)} chars) for Grok Videos")
                    else:
                        logger.warning(
                            "⚠️ Auto broll_table generation failed for Grok Videos")
                except Exception as auto_broll_err:
                    logger.error(f"❌ Auto broll_table error: {auto_broll_err}")
        if checkboxes.get("grok_videos", False) and broll_table:
            total_rows = len(broll_rows) if broll_rows else 0
            streamer.send_update(
                f"⏸️ Select B-roll rows for Grok generation ({total_rows} available)",
                int(15 + (completed_steps * progress_per_step))
            )

        # Generate HeyGen Ready Script if requested
        if checkboxes.get("heygen", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("🎬 Generating HeyGen Ready section...",
                                 int(current_progress))
            try:
                from console_ui.text_processing import extract_heygen_host_script

                # Extract from the cleaned source script ONLY (not final_output) so the
                # HeyGen text never includes the prepended OPENING HOOK OPTIONS block,
                # the OPTION 1/2/3 "Host:" hooks, the opening statement, or any
                # generated headers added below.
                heygen_script = extract_heygen_host_script(cleaned_script)
                # Fallback: if no Host: markers found, use cleaned script directly
                if not heygen_script:
                    logger.info(
                        "⚠️ No Host: markers found - using full script for HeyGen section")
                    heygen_script = cleaned_script
                if heygen_script:
                    heygen_section = f"\n\n{'=' * 80}\n"
                    heygen_section += "# 🎬 HEYGEN READY SCRIPT\n"
                    heygen_section += f"{'=' * 80}\n\n{heygen_script}"
                    final_output += heygen_section
                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        f"✅ HeyGen script extracted ({len(heygen_script)} chars)",
                        int(progress)
                    )
            except Exception as e:
                logger.error(f"❌ HeyGen section error: {e}")
                heygen_script = None

        # Generate curl commands if requested
        if checkboxes.get("curl", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("🚀 Generating HeyGen curl commands...",
                                 int(current_progress))
            try:
                from console_ui.text_processing import (
                    extract_heygen_host_script,
                    generate_heygen_curl_commands
                )

                if not heygen_script:
                    heygen_script = extract_heygen_host_script(cleaned_script)

                # Fallback: if no Host: markers found, use the cleaned script directly
                if not heygen_script:
                    logger.info(
                        "⚠️ No Host: markers found - using full script for curl generation")
                    heygen_script = cleaned_script

                if heygen_script:
                    heygen_with_header = (
                        f"# 🎬 HEYGEN READY SCRIPT\n{'=' * 80}\n\n{heygen_script}"
                    )
                    curl_commands = generate_heygen_curl_commands(
                        heygen_with_header, script_title,
                        heygen_api_key, heygen_template_id,
                        heygen_voice_id
                    )
                    if curl_commands:
                        # Store curl commands separately (don't append to script)
                        num_commands = curl_commands.count(
                            'curl --request POST')
                        completed_steps += 1
                        progress = 15 + (completed_steps * progress_per_step)
                        streamer.send_update(
                            f"✅ Generated {num_commands} curl commands",
                            int(progress)
                        )
                        _live_curl_dst = _save_curl_commands_live(
                            curl_commands, script_title)
                        if _live_curl_dst:
                            streamer.send_update(
                                f"💾 Saved curl commands → {_live_curl_dst}",
                                int(progress)
                            )
                    else:
                        curl_commands = None
            except Exception as e:
                logger.error(f"❌ Curl generation error: {e}")
                curl_commands = None

        # Generate Demo Package if requested
        if checkboxes.get("demo", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("🔧 Generating demo package...",
                                 int(current_progress))
            try:
                from linedrive_azure.agents.openai_demo_agent_client import (
                    OpenAIDemoAgentClient
                )

                demo_agent = OpenAIDemoAgentClient()
                demo_result = demo_agent.generate_demo_package(
                    script_content=script_content,
                    script_title=script_title,
                    timeout=180
                )

                if demo_result.get("success", False):
                    demo_packages = demo_result.get("demo_packages", "")
                    demo_section = f"\n\n{'=' * 80}\n# 🔧 DEMO PACKAGE\n"
                    demo_section += f"{'=' * 80}\n\n{demo_packages}"
                    final_output += demo_section
                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        f"✅ Demo package generated ({len(demo_packages)} chars)",
                        int(progress)
                    )
            except Exception as e:
                logger.error(f"❌ Demo package error: {e}")

        # Generate Thumbnails if requested
        thumbnail_results = None
        if checkboxes.get("thumbnails", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            expected_thumbnails = 6 * \
                len(thumbnail_hook_text_options) if thumbnail_hook_text_options else 6
            streamer.send_update(
                f"🖼️ Generating {expected_thumbnails} thumbnail variations...",
                int(current_progress),
            )
            try:
                from tools.media.emotional_thumbnail_generator import (
                    EmotionalThumbnailGenerator,
                )

                _thumb_out_dir = run_output_dir / "thumbnails"
                thumbnail_gen = EmotionalThumbnailGenerator(
                    output_dir=_thumb_out_dir)
                if thumbnail_hook_text_options:
                    logger.info(
                        f"🏷️ Script processing will use thumbnail hook text options instead of script title: {thumbnail_hook_text_options}")
                else:
                    logger.warning(
                        "⚠️ Script processing did not find thumbnail hook text; using script title fallback")
                # always fresh — never reuse a stale cancelled event
                thumb_cancel_evt = _threading.Event()
                thumbnail_cancel_events[session_id] = thumb_cancel_evt
                thumbnail_results = thumbnail_gen.generate_all_thumbnails(
                    script_title=script_title,
                    script_content=cleaned_script,
                    headline_text=thumbnail_hook_text or None,
                    headline_options=thumbnail_hook_text_options or None,
                    progress_callback=lambda msg: streamer.send_update(
                        msg, int(current_progress)
                    ),
                    cancel_check=thumb_cancel_evt.is_set,
                )
                logger.info(
                    f"🔍 DEBUG: script processing headline options passed to thumbnail generator = {repr(thumbnail_hook_text_options or None)}")

                variations = (thumbnail_results or {}).get("variations") or []
                output_dir = (thumbnail_results or {}).get("output_dir")

                if variations:
                    if output_dir:
                        logger.info(
                            f"📁 Thumbnails saved to: {output_dir}")
                        streamer.send_update(
                            f"📁 Thumbnails saved to: {output_dir}",
                            int(current_progress)
                        )
                    thumb_section = f"\n\n{'=' * 80}\n"
                    thumb_section += "# 🖼️ EMOTIONAL THUMBNAIL VARIATIONS\n"
                    thumb_section += f"{'=' * 80}\n\n"
                    for i, var in enumerate(variations, 1):
                        thumb_section += f"## Variation #{i}: {var.get('emotion')}\n"
                        thumb_section += f"- **Text:** {var.get('text')}\n"
                        thumb_section += f"- **Expression:** {var.get('expression')}\n"
                        thumb_section += f"- **File:** {var.get('filename')}\n\n"
                    final_output += thumb_section
                    completed_steps += 1
                    progress = 15 + (completed_steps * progress_per_step)
                    streamer.send_update(
                        f"✅ Generated {len(variations)} thumbnails",
                        int(progress)
                    )
                    if (thumbnail_results or {}).get("cancelled"):
                        streamer.send_update(
                            f"🛑 Thumbnail generation stopped — keeping {len(variations)} thumbnails so far",
                            int(progress)
                        )
                    api_error = (thumbnail_results or {}).get("error")
                    if api_error:
                        logger.warning(
                            f"⚠️ Thumbnail API returned errors during generation: {api_error}")
                        streamer.send_update(
                            f"⚠️ Thumbnail API error during generation: {api_error}",
                            int(current_progress)
                        )
                elif output_dir:
                    logger.warning(
                        f"⚠️ Thumbnail directory created but no files generated: {output_dir}")
                    streamer.send_update(
                        f"⚠️ No thumbnails generated (directory only): {output_dir}",
                        int(current_progress)
                    )
                    attempted = (thumbnail_results or {}).get(
                        "total_attempted", 0)
                    streamer.send_update(
                        f"❌ Thumbnail generation failed (0/{attempted})",
                        int(current_progress)
                    )
                    if (thumbnail_results or {}).get("error"):
                        streamer.send_update(
                            f"❌ Thumbnail API error: {(thumbnail_results or {}).get('error')}",
                            int(current_progress)
                        )
            except Exception as e:
                logger.error(f"❌ Thumbnail generation error: {e}")
                streamer.send_update(
                    f"❌ Thumbnail generation error: {e}", int(current_progress))
            finally:
                thumbnail_cancel_events.pop(session_id, None)

        # Generate Flow & Repetition Analysis if requested
        # IMPORTANT: This should run AFTER all other content is assembled
        if checkboxes.get("flow_analysis", False):
            current_progress = 15 + (completed_steps * progress_per_step)
            streamer.send_update("🔄 Analyzing script for repetition and flow...",
                                 int(current_progress))
            try:
                from linedrive_azure.agents import ScriptRepeatAndFlowAgentClient

                flow_agent = ScriptRepeatAndFlowAgentClient()
                flow_result = flow_agent.analyze_and_improve_flow(
                    script_content=final_output,
                    script_title=script_title,
                    target_audience=audience,
                    timeout=300
                )

                logger.info(
                    f"📊 Flow agent result: success={flow_result.get('success')}, keys={list(flow_result.keys())}")

                if flow_result.get("success", False):
                    # Replace final_output with the improved version
                    improved_script = flow_result.get("improved_script", "")
                    repetition_analysis = flow_result.get(
                        "repetition_analysis", "")
                    flow_improvements = flow_result.get("flow_analysis", "")

                    logger.info(
                        f"📊 Improved script length: {len(improved_script)}, rep analysis: {len(repetition_analysis)}, flow: {len(flow_improvements)}")

                    if improved_script:
                        # Add analysis section first
                        analysis_section = f"\n\n{'=' * 80}\n"
                        analysis_section += "# 🔄 FLOW & REPETITION ANALYSIS\n"
                        analysis_section += f"{'=' * 80}\n\n"

                        if repetition_analysis:
                            analysis_section += "## Repetition Analysis\n\n"
                            analysis_section += f"{repetition_analysis}\n\n"

                        if flow_improvements:
                            analysis_section += "## Flow Improvements\n\n"
                            analysis_section += f"{flow_improvements}\n\n"

                        # Replace the script with improved version and add analysis at end
                        final_output = improved_script + analysis_section

                        completed_steps += 1
                        progress = 15 + (completed_steps * progress_per_step)
                        streamer.send_update(
                            f"✅ Flow analysis complete - script improved ({len(improved_script)} chars)",
                            int(progress)
                        )
                    else:
                        logger.warning(
                            "⚠️ Flow agent returned empty improved script")
                else:
                    logger.warning(
                        f"⚠️ Flow analysis failed - result: {flow_result}")
            except Exception as e:
                logger.error(f"❌ Flow analysis error: {e}")

        # Complete — save files and set result BEFORE sending done=True so the
        # SSE handler always finds a valid streamer.result when it reads on progress==100.
        saved_markdown_path, saved_docx_path = await _save_script_md_and_docx(
            script_title,
            final_output,
        )

        streamer.result = {
            "success": True,
            "enhanced_script": final_output,  # SSE handler expects "enhanced_script" key
            "script": final_output,  # Also include "script" for compatibility
            "script_title": script_title,
            "thumbnail_results": thumbnail_results,
            "thumbnail_hook_text": thumbnail_hook_text,
            "thumbnail_hook_text_options": thumbnail_hook_text_options,
            "edl_content": edl_content,
            "edl_filename": edl_filename,
            "curl_commands": curl_commands,
            "broll_images": broll_images,
            "grok_videos": grok_videos,
            "broll_table": broll_table,
            "broll_rows": broll_rows,
            "youtube_details": youtube_details,
            "markdown_path": saved_markdown_path,
            "docx_path": saved_docx_path,
        }
        logger.info("✅ Script processing completed successfully")
        # Send done=True AFTER result is set so the SSE handler always finds a valid result.
        streamer.send_update("✅ Script processing complete!", 100)

    except Exception as e:
        error_msg = f"Error: {str(e)}"
        logger.error(f"💥 Exception during script processing: {error_msg}")
        streamer.send_update(f"❌ {error_msg}", -1)
        streamer.result = {"success": False, "error": error_msg}


async def test_mode_simulation(session_id):
    """Quick test mode - simulates script creation in 10 seconds"""
    import time
    import asyncio

    logger.info(f"⚡ Starting test mode simulation for {session_id}")

    if session_id not in progress_streams:
        logger.error(f"❌ No progress stream for session {session_id}")
        return

    streamer = progress_streams[session_id]

    try:
        # Simulate the 7-chapter workflow with fast updates
        test_chapters = [
            "Introduction: Setting the Stage",
            "Chapter 1: The Problem",
            "Chapter 2: Understanding the Context",
            "Chapter 3: Exploring Solutions",
            "Chapter 4: Implementation Strategy",
            "Chapter 5: Real-World Applications",
            "Conclusion: Looking Forward"
        ]

        # Send initial message
        streamer.send_update("🚀 Starting test script creation...", 5)
        await asyncio.sleep(1)

        # Simulate writing each chapter (1 second per chapter)
        for i, chapter in enumerate(test_chapters):
            progress = 10 + (i * 12)  # Distribute progress from 10-94%
            streamer.send_update(f"📝 Writing {chapter}...", progress)
            await asyncio.sleep(1)

            # Show completion of chapter
            char_count = 1000 + (i * 200)  # Fake character counts
            streamer.send_update(f"✅ {chapter} completed ({char_count} chars)",
                                 progress + 2)
            await asyncio.sleep(0.5)

        # Final processing steps
        streamer.send_update("✨ Enhancing script formatting...", 96)
        await asyncio.sleep(1)

        streamer.send_update("🔍 Adding tool recommendations...", 98)
        await asyncio.sleep(1)

        # Create a simple test result
        test_script = """Welcome to AI with Roz, I'm Roz's AI Clone. 

This is a test script created in fast mode for debugging purposes.

Chapter 1: The Problem
This chapter discusses the main challenges...

Chapter 2: Understanding the Context  
Here we explore the broader implications...

[Continue with remaining chapters...]

===============================
TOOLS
===============================

• OpenAI GPT-4 - AI writing assistant
• Canva - Design and presentation tools
• YouTube Studio - Video optimization"""

        # Store test result
        streamer.result = {
            "success": True,
            "enhanced_script": test_script,
            "word_count": len(test_script.split()),
            "reading_time": f"{len(test_script.split()) // 150} min",
            "audience": "test audience",
            "production_type": "test",
        }

        # Send completion message
        streamer.send_update("✅ Test script creation complete!", 100)
        logger.info(f"✅ Test mode completed for session {session_id}")

    except Exception as e:
        error_msg = f"Test mode error: {str(e)}"
        logger.error(f"❌ Test mode error for {session_id}: {error_msg}")
        streamer.send_update(f"❌ {error_msg}", -1)
        streamer.result = {"success": False, "error": error_msg}


@app.route("/test-sse")
def test_sse():
    """Test SSE completion behavior without running full script creation"""
    return render_template("test_sse.html")


@app.route("/test-progress/<session_id>")
def test_progress_stream(session_id):
    """Test SSE endpoint that simulates script creation completion"""
    import time

    def generate():
        try:
            # Simulate progress messages
            messages = [
                ("🚀 Starting test...", 10),
                ("📝 Processing step 1...", 30),
                ("📝 Processing step 2...", 60),
                ("✅ Almost done...", 90),
                ("✅ Test complete!", 100)
            ]

            for message, progress in messages:
                update = {
                    "message": message,
                    "progress": progress,
                    "timestamp": time.time(),
                    "done": progress == 100
                }

                logger.info(f"🔄 TEST SSE sending: {json.dumps(update)}")
                yield f"data: {json.dumps(update)}\n\n"

                if progress == 100:
                    logger.info(f"🎯 TEST SSE completion message sent")
                    # Don't send closing message - client will close connection
                    break

                time.sleep(1)  # 1 second between messages

        except Exception as e:
            logger.error(f"TEST SSE error: {e}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
        finally:
            logger.info(f"🔚 TEST SSE stream ended")

    response = Response(generate(), mimetype="text/event-stream")
    response.headers['Cache-Control'] = 'no-cache'
    response.headers['Connection'] = 'keep-alive'
    response.headers['X-Accel-Buffering'] = 'no'
    return response


@app.route("/")
def index():
    """Serve the main ScriptCraft web interface"""
    # Resolve API keys with precedence: saved settings file > environment variable > empty.
    # Environment variables can be set in ~/.zshrc (or any shell rc) for one-time setup:
    #   export XAI_API_KEY="xai-..."
    #   export HEYGEN_API_KEY="sk_V2_..."
    #   export HEYGEN_VOICE_ID="..."
    #   export GOOGLE_API_KEY="..."
    settings = {}
    try:
        settings = _load_scriptcraft_settings()
    except Exception as e:
        logger.warning(f"⚠️ Could not preload settings for template: {e}")

    def _resolve(setting_key: str, env_var: str) -> str:
        return (
            (settings.get(setting_key) or "").strip()
            or os.getenv(env_var, "").strip()
        )

    return render_template(
        "index.html",
        version=VERSION,
        agent_mode=(os.environ.get("FOUNDRY_API_MODE") or "v2").lower(),
        saved_grok_api_key=_resolve("grok_api_key", "XAI_API_KEY"),
        saved_heygen_api_key=_resolve("heygen_api_key", "HEYGEN_API_KEY"),
        saved_heygen_voice_id=_resolve("heygen_voice_id", "HEYGEN_VOICE_ID"),
        saved_google_api_key=_resolve("google_api_key", "GOOGLE_API_KEY"),
    )


@app.route("/api/agent-mode", methods=["GET", "POST"])
def agent_mode_api():
    """Get or set the active Foundry agent API mode (v1=classic Assistants, v2=new Foundry)."""
    if request.method == "POST":
        data = request.get_json(silent=True) or {}
        requested = (data.get("mode") or "").strip().lower()
        if requested not in ("v1", "v2"):
            return jsonify({"success": False, "error": "mode must be 'v1' or 'v2'"}), 400
        os.environ["FOUNDRY_API_MODE"] = requested
        logger.info(f"\U0001F500 Agent API mode switched to: {requested}")
    current = (os.environ.get("FOUNDRY_API_MODE") or "v2").lower()
    return jsonify({"success": True, "mode": current})


@app.route("/test")
def test_page():
    """Serve the test page for debugging SSE streaming"""
    return render_template("test.html")


@app.route("/test-capture")
def test_capture_page():
    """Test page to debug console capture and threading issues"""
    return """
    <!DOCTYPE html>
    <html>
    <head><title>Console Capture Test</title></head>
    <body>
        <h1>Console Capture Test Harness</h1>
        <button onclick="runTest()">Run Console Capture Test</button>
        <div id="progress"></div>
        <div id="result"></div>
        <script>
        function runTest() {
            fetch('/test-capture-process', {method: 'POST'})
            .then(r => r.json())
            .then(data => {
                document.getElementById('progress').innerHTML = 'Test started...';
                watchProgress(data.session_id);
            });
                resp = make_response(render_template(
                    "index.html",
                    version=VERSION,
                    agent_mode=(os.environ.get("FOUNDRY_API_MODE") or "v2").lower(),
                    saved_grok_api_key=_resolve("grok_api_key", "XAI_API_KEY"),
                    saved_heygen_api_key=_resolve("heygen_api_key", "HEYGEN_API_KEY"),
                    saved_heygen_voice_id=_resolve("heygen_voice_id", "HEYGEN_VOICE_ID"),
                    saved_google_api_key=_resolve("google_api_key", "GOOGLE_API_KEY"),
                ))
                resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
                resp.headers['Pragma'] = 'no-cache'
                resp.headers['Expires'] = '0'
                return resp
        }
                    document.getElementById('result').innerHTML = 'Test completed!';
                }
            };
        }
        </script>
    </body>
    </html>
    """


@app.route('/test-capture-process', methods=['POST'])
def test_capture_process():
    """Test the console capture mechanism without full workflow"""
    session_id = str(uuid.uuid4())
    progress_streams[session_id] = ProgressStreamer(session_id)

    # Start the test in background
    thread = threading.Thread(
        target=run_test_capture,
        args=(session_id,),
        daemon=True
    )
    thread.start()
    running_tasks[session_id] = thread

    return jsonify({'session_id': session_id, 'status': 'started'})


def run_test_capture(session_id):
    """Test console capture and the exact point where it hangs"""
    import time
    import sys

    logger.info(f"🧪 TEST CAPTURE STARTED: session={session_id}")

    try:
        if session_id not in progress_streams:
            logger.error(f"❌ Session {session_id} not found!")
            return

        streamer = progress_streams[session_id]

        # Set up console capture exactly like the real workflow
        console_capture = ConsoleCapture(streamer)
        original_stdout = sys.stdout

        try:
            # Redirect stdout to capture console output
            sys.stdout = console_capture

            # Test 1: Basic console output
            print("🧪 TEST 1: Basic console output")
            streamer.send_update("🧪 Basic console output test", 10)
            time.sleep(1)

            # Test 2: Import text processing (where it hangs)
            print("🧪 TEST 2: Importing text processing modules")
            streamer.send_update("🧪 Testing text processing import", 20)

            try:
                from console_ui.text_processing import (
                    enhance_script_with_bold_tools,
                    extract_tool_links_and_info,
                )
                print("✅ Text processing modules imported successfully")
                streamer.send_update("✅ Text processing import success", 30)
            except Exception as e:
                print(f"❌ Text processing import failed: {e}")
                streamer.send_update(
                    f"❌ Text processing import failed: {e}", 30)
                raise

            # Test 3: Call enhance_script_with_bold_tools (where it might hang)
            print("🧪 TEST 3: Testing enhance_script_with_bold_tools")
            streamer.send_update("🧪 Testing script enhancement", 40)

            test_script = "This is a test script with **bold** and [ChatGPT](https://chatgpt.com)."

            try:
                enhanced_script = enhance_script_with_bold_tools(test_script)
                print(f"✅ Script enhanced: {len(enhanced_script)} chars")
                streamer.send_update("✅ Script enhancement success", 50)
            except Exception as e:
                print(f"❌ Script enhancement failed: {e}")
                streamer.send_update(f"❌ Script enhancement failed: {e}", 50)
                raise

            # Test 4: Call extract_tool_links_and_info
            print("🧪 TEST 4: Testing extract_tool_links_and_info")
            streamer.send_update("🧪 Testing tool link extraction", 60)

            try:
                tool_links = extract_tool_links_and_info(enhanced_script)
                print(f"✅ Tool links extracted: {len(tool_links)} chars")
                streamer.send_update("✅ Tool link extraction success", 70)
            except Exception as e:
                print(f"❌ Tool link extraction failed: {e}")
                streamer.send_update(f"❌ Tool link extraction failed: {e}", 70)
                raise

            # Test 5: Minimal rapid output test to isolate deadlock
            print("🧪 TEST 5: Minimal rapid output test")
            streamer.send_update("🧪 Starting minimal rapid test", 80)

            # Test just 3 rapid outputs with debugging
            for i in range(3):
                print(f"   📝 Before output {i+1}")
                streamer.send_update(f"🧪 Processing output {i+1}/3", 85 + i*2)
                print(f"   📝 After output {i+1}")
                time.sleep(0.5)  # Longer sleep to prevent issues

            print("✅ Minimal rapid test completed")
            logger.info(
                "🔍 DEBUG: About to send minimal rapid test success update")
            streamer.send_update("✅ Minimal rapid test success", 90)
            logger.info("🔍 DEBUG: Minimal rapid test success update sent")

            # Final test
            print("🧪 ALL TESTS COMPLETED SUCCESSFULLY!")
            logger.info("🔍 DEBUG: About to send all tests completed update")
            streamer.send_update("🧪 All tests completed!", 100)
            logger.info("🔍 DEBUG: All tests completed update sent")

            # Store success result
            logger.info("🔍 DEBUG: About to set streamer.result")
            streamer.result = {
                "success": True,
                "message": "All console capture tests passed",
                "enhanced_script": enhanced_script,
                "tool_links": tool_links
            }
            logger.info("🔍 DEBUG: streamer.result set successfully")

        finally:
            # Always restore stdout
            logger.info("🔍 DEBUG: Entering finally block")
            sys.stdout = original_stdout
            logger.info("🔍 DEBUG: stdout restored, function should exit now")

    except Exception as e:
        error_msg = f"Test failed: {str(e)}"
        logger.error(f"💥 Test exception: {error_msg}")
        if session_id in progress_streams:
            streamer = progress_streams[session_id]
            streamer.send_update(f"❌ {error_msg}", -1)
            streamer.result = {"success": False, "error": error_msg}


@app.route("/test_progress/<session_id>/<int:progress>/<message>")
def test_progress(session_id, progress, message):
    """Test endpoint to manually send progress updates"""
    if session_id in progress_streams:
        streamer = progress_streams[session_id]
        streamer.send_update(message, progress)
        logger.info(f"🧪 Test progress sent: {message} -> {progress}%")
        return jsonify({
            "success": True,
            "message": f"Sent: {message} -> {progress}%"
        })
    return jsonify({
        "success": False,
        "error": "Session not found"
    }), 404


@app.route("/test_create")
def test_create():
    """Create a test session with hardcoded progress messages"""
    import uuid
    import time
    import threading

    # Generate session ID
    session_id = str(uuid.uuid4())

    # Create progress streamer
    progress_streams[session_id] = ProgressStreamer(session_id)

    logger.info(f"🧪 Starting test session: {session_id}")

    def send_test_messages():
        streamer = progress_streams[session_id]
        time.sleep(1)

        # Send a series of test messages
        test_messages = [
            (10, "🚀 Starting test..."),
            (25, "📝 Writing Chapter 1/7: Test Chapter"),
            (40, "✅ Chapter 1 completed (1234 chars)"),
            (55, "📝 Writing Chapter 2/7: Another Test Chapter"),
            (70, "✅ Chapter 2 completed (2345 chars)"),
            (85, "🔍 Reviewing content..."),
            (100, "✅ Test completed successfully!")
        ]

        for progress, message in test_messages:
            streamer.send_update(message, progress)
            logger.info(f"🧪 Sent: {message} -> {progress}%")
            time.sleep(2)  # 2 second delay between messages

    # Start test in background
    thread = threading.Thread(target=send_test_messages, daemon=True)
    running_tasks[session_id] = thread
    thread.start()

    return jsonify({
        "success": True,
        "session_id": session_id,
        "message": "Test session started",
        "test_url": f"/progress/{session_id}"
    })


@app.route("/debug/stdout")
def debug_stdout():
    """Debug endpoint to check stdout status"""
    return jsonify({
        "stdout_type": str(type(sys.stdout)),
        "is_console_capture": isinstance(sys.stdout, ConsoleCapture),
        "active_sessions": len(progress_streams),
        "running_tasks": len(running_tasks)
    })


@app.route("/create", methods=["POST"])
def create():
    """Create a new script with progress streaming"""
    # Get form data
    data = request.get_json() or {}
    topic = data.get("topic", "AI in daily life")
    audience = data.get("audience", "general")
    tone = data.get("tone", "professional")
    video_length = data.get("video_length", "medium")
    production_type = data.get("production_type", "standard")
    goals = data.get("goals", "educational")
    test_mode = data.get("test_mode", False)  # Quick test mode flag
    quick_test = data.get("quick_test", False)  # 1-chapter quick test

    # NEW: Get checkbox selections
    checkboxes = data.get("checkboxes", {})

    # Get HeyGen parameters
    heygen_template_id = data.get("heygen_template_id", "")
    heygen_api_key = data.get("heygen_api_key", "")
    heygen_voice_id = data.get("heygen_voice_id", "")
    grok_api_key = data.get("grok_api_key", "")

    # If "All" is checked, enable everything
    if checkboxes.get("all", False):
        checkboxes = {
            "script": True,
            "hook_summary": True,
            "youtube_details": True,
            "heygen": True,
            "curl": True,
            "broll": True,
            "broll_images": True,
            "thumbnails": True,
            "demo": True,
            "flow_analysis": True,
            "all": True
        }
    # If no checkboxes provided, default to script only
    elif not checkboxes:
        checkboxes = {"script": True}

    # Generate session ID
    session_id = str(uuid.uuid4())

    # Create progress streamer
    progress_streams[session_id] = ProgressStreamer(session_id)

    logger.info(f"🎬 Starting script creation: {session_id}")
    logger.info(f"📝 Topic: {topic}, Audience: {audience}, Tone: {tone}")
    logger.info(f"📋 Checkboxes: {checkboxes}")
    if quick_test:
        logger.info(f"⚡ QUICK TEST MODE: 1 chapter only")

    # Start script creation in background thread
    def run_script_creation():
        try:
            if test_mode:
                # Quick test mode - simulate progress and complete in 10 seconds
                logger.info(f"⚡ Running in test mode for session {session_id}")
                asyncio.run(test_mode_simulation(session_id))
            else:
                # Override video_length for quick test
                actual_length = "1 minute (150-200 words)" if quick_test else video_length
                asyncio.run(process_script_creation(
                    session_id, topic, audience, tone, actual_length,
                    production_type, goals, quick_test, checkboxes,
                    heygen_template_id, heygen_api_key,
                    heygen_voice_id, grok_api_key
                ))
        except Exception as e:
            logger.error(f"❌ Script creation error: {e}")
            if session_id in progress_streams:
                progress_streams[session_id].send_update(
                    f"❌ Error: {str(e)}", 100
                )

    thread = threading.Thread(target=run_script_creation, daemon=True)
    running_tasks[session_id] = thread
    thread.start()

    return jsonify({
        "success": True,
        "session_id": session_id,
        "message": "Script creation started"
    })


@app.route("/progress/<session_id>")
def progress_stream(session_id):
    """Server-Sent Events endpoint for progress updates"""
    logger.info(f"🌊 SSE connection established for session {session_id}")

    if session_id not in progress_streams:
        return "Session not found", 404

    streamer = progress_streams[session_id]

    def generate():
        try:
            while True:  # Changed from while not streamer.done
                try:
                    # Get update with timeout
                    update = streamer.queue.get(timeout=30)
                    logger.info(f"🔄 SSE sending: {update}")

                    # Parse the update to check if it's the completion message
                    update_data = json.loads(update)

                    # If this is the completion message (progress=100), add the script content
                    if update_data.get("done", False) or update_data.get("progress") == 100:
                        logger.info(
                            f"🎯 Processing completion message for session {session_id}")

                        # Create a completion message with status and script
                        if streamer.result and streamer.result.get("success"):
                            # Prepare thumbnail data for frontend
                            thumbnail_results = streamer.result.get(
                                "thumbnail_results", {})

                            print("\n" + "="*70)
                            print("📦 PREPARING COMPLETION MESSAGE")
                            print("="*70)
                            print(f"Session: {session_id}")
                            print(
                                f"Thumbnail results type: {type(thumbnail_results)}")
                            print(
                                f"Thumbnail results is None: {thumbnail_results is None}")

                            thumbnails = _collect_all_thumbnail_entries(
                                streamer.result.get(
                                    "script_title", "Untitled Script"),
                                thumbnail_results,
                            )

                            if thumbnail_results and thumbnail_results.get("variations"):
                                print(f"✅ Found variations in thumbnail_results")
                                print(
                                    f"   Generated this run: {len(thumbnail_results.get('variations') or [])}")
                            elif thumbnail_results:
                                print(
                                    f"⚠️ No variations found in thumbnail_results")
                                print(
                                    f"   Available keys: {list(thumbnail_results.keys())}")

                            print(
                                f"\n✅ Prepared {len(thumbnails)} thumbnail objects for frontend (directory-wide)")

                            print("="*70 + "\n")

                            # DEBUG: Check chapter_comparisons data
                            chapter_comps = streamer.result.get(
                                "chapter_comparisons")
                            print(
                                f"🔍 DEBUG: chapter_comparisons type: {type(chapter_comps)}")
                            if chapter_comps:
                                print(
                                    f"🔍 DEBUG: chapter_comparisons length: {len(chapter_comps)}")
                                print(
                                    f"🔍 DEBUG: First comparison keys: {list(chapter_comps[0].keys()) if len(chapter_comps) > 0 else 'N/A'}")
                            else:
                                print(
                                    f"🔍 DEBUG: chapter_comparisons is None or empty")

                            completion_msg = {
                                "status": "complete",
                                "message": update_data.get("message", "✅ Script creation complete!"),
                                "script": streamer.result.get("enhanced_script", ""),
                                "script_title": streamer.result.get("script_title", "Untitled Script"),
                                "demo_packages": streamer.result.get("demo_packages"),
                                "word_count": streamer.result.get("word_count", 0),
                                "reading_time": streamer.result.get("reading_time", "N/A"),
                                "thumbnails": thumbnails,
                                "thumbnail_results": thumbnail_results,
                                "comparison_file": streamer.result.get("comparison_file"),
                                "chapter_comparisons": streamer.result.get("chapter_comparisons"),
                                "flow_original_script": streamer.result.get("flow_original_script"),
                                "flow_improved_script": streamer.result.get("flow_improved_script"),
                                "flow_analysis_report": streamer.result.get("flow_analysis_report"),
                                "edl_content": streamer.result.get("edl_content"),
                                "edl_filename": streamer.result.get("edl_filename"),
                                "curl_commands": streamer.result.get("curl_commands"),
                                "markdown_path": streamer.result.get("markdown_path"),
                                "docx_path": streamer.result.get("docx_path"),
                                "broll_images": streamer.result.get("broll_images"),
                                "grok_videos": streamer.result.get("grok_videos"),
                                "broll_table": streamer.result.get("broll_table"),
                                "broll_rows": streamer.result.get("broll_rows"),
                                "youtube_details": streamer.result.get("youtube_details"),
                            }

                            print("\n" + "="*70)
                            print("📤 SENDING COMPLETION MESSAGE TO FRONTEND")
                            print("="*70)
                            print(f"Status: {completion_msg['status']}")
                            print(
                                f"Script length: {len(completion_msg['script'])} chars")
                            print(
                                f"Thumbnails array length: {len(completion_msg['thumbnails'])}")
                            if completion_msg['thumbnails']:
                                print("\nThumbnails being sent:")
                                for i, t in enumerate(completion_msg['thumbnails'], 1):
                                    print(
                                        f"   #{i}: {t['emotion']} - {t['filename']}")
                            else:
                                print("⚠️ No thumbnails in completion message!")

                            # Debug B-roll images
                            broll_imgs = completion_msg.get('broll_images')
                            if broll_imgs:
                                print(
                                    f"\n🎨 B-roll images being sent: {len(broll_imgs)} images")
                                for i, img in enumerate(broll_imgs, 1):
                                    print(
                                        f"   #{i}: {img.get('search_term')} - {img.get('filename')}")
                            else:
                                print(
                                    "\n⚠️ No B-roll images in completion message!")

                            print("="*70 + "\n")

                            logger.info(
                                f"✅ Sending completion with script ({len(completion_msg.get('script', ''))} chars) and {len(thumbnails)} thumbnails")
                        else:
                            # Error case
                            error_msg = streamer.result.get(
                                "error", "Unknown error") if streamer.result else "No result available"
                            completion_msg = {
                                "status": "error",
                                "error": error_msg
                            }
                            logger.error(
                                f"❌ Sending error completion: {error_msg}")

                        yield f"data: {json.dumps(completion_msg)}\n\n"
                        logger.info(
                            f"🎯 SSE completion message sent for session {session_id}")
                        break
                    else:
                        # Regular progress update
                        progress_msg = {
                            "status": "progress",
                            "message": update_data.get("message", "Processing..."),
                            "progress": update_data.get("progress", 0)
                        }
                        yield f"data: {json.dumps(progress_msg)}\n\n"

                except queue.Empty:
                    # Only check done status if queue is empty
                    if streamer.done:
                        logger.info(
                            f"🏁 SSE ending - streamer marked done for session {session_id}")
                        break
                    # Send keepalive
                    logger.debug(f"💓 SSE keepalive for session {session_id}")
                    yield f"data: {json.dumps({'status': 'keepalive'})}\n\n"
                    continue

        except Exception as e:
            # Suppress "Broken pipe" errors - these happen when client disconnects
            import errno
            if isinstance(e, (BrokenPipeError, IOError)) and getattr(e, 'errno', None) == errno.EPIPE:
                logger.debug(
                    f"🔌 Client disconnected (broken pipe) for session {session_id}")
            else:
                logger.error(f"SSE error for session {session_id}: {e}")
                yield f"data: {json.dumps({'status': 'error', 'error': str(e)})}\n\n"
        finally:
            logger.info(f"🔚 SSE stream ended for session {session_id}")

    response = Response(generate(), mimetype="text/event-stream")
    response.headers['Cache-Control'] = 'no-cache'
    response.headers['Connection'] = 'keep-alive'
    response.headers['X-Accel-Buffering'] = 'no'  # Disable nginx buffering
    return response


@app.route("/result/<session_id>")
def get_result(session_id):
    """Get final result"""
    logger.info(f"🔍 Result requested for session {session_id}")

    if session_id in progress_streams:
        streamer = progress_streams[session_id]
        logger.info(
            f"📊 Session {session_id} found. Result available: {bool(streamer.result)}")
        logger.info(f"📊 Session {session_id} done status: {streamer.done}")

        if streamer.result:
            logger.info(f"✅ Result found for session {session_id}")
            result_size = len(str(streamer.result)) if streamer.result else 0
            logger.info(f"📏 Result size: {result_size} chars")
            return jsonify(streamer.result)
        else:
            logger.warning(
                f"⚠️ No result available yet for session {session_id}")
            return jsonify({"error": "Result not ready yet, please try again"}), 202
    else:
        logger.error(f"❌ Session {session_id} not found in progress_streams")
        logger.info(f"📋 Available sessions: {list(progress_streams.keys())}")

    return jsonify({"error": "Result not available"}), 404


@app.route("/process-script", methods=["POST"])
def process_script():
    """Process existing script to generate YouTube details, B-roll, HeyGen, etc."""
    logger.info("=== PROCESS SCRIPT REQUEST RECEIVED ===")

    try:
        data = request.json
        script_content = data.get("script", "")
        # Note: topic field is only used for script CREATION, not processing
        audience = data.get("audience", "general audience")
        video_length = data.get("videoLength", "5-8 minutes")
        tone = data.get("tone", "informative")
        checkboxes = data.get("checkboxes", {})
        heygen_template_id = data.get("heygen_template_id", "")
        heygen_api_key = data.get("heygen_api_key", "")
        heygen_voice_id = data.get("heygen_voice_id", "")
        grok_api_key = data.get("grok_api_key", "")

        if not script_content.strip():
            logger.warning("❌ Empty script received")
            return jsonify({"success": False, "error": "Please provide a script to process"})

        logger.info(f"📋 Checkboxes: {checkboxes}")
        if heygen_template_id:
            logger.info(f"🎬 HeyGen Template: {heygen_template_id}")

        # Generate session ID
        session_id = str(uuid.uuid4())

        # Create progress streamer
        progress_streams[session_id] = ProgressStreamer(session_id)

        logger.info(f"🎬 Starting script processing: {session_id}")

        # Start script processing in background thread
        def run_script_processing():
            try:
                asyncio.run(process_existing_script(
                    session_id, script_content, audience, tone,
                    video_length, checkboxes, heygen_template_id,
                    heygen_api_key, heygen_voice_id, grok_api_key
                ))
            except Exception as e:
                logger.error(f"❌ Script processing error: {e}")
                if session_id in progress_streams:
                    progress_streams[session_id].send_update(
                        f"❌ Error: {str(e)}", -1
                    )

        thread = threading.Thread(target=run_script_processing, daemon=True)
        running_tasks[session_id] = thread
        thread.start()

        return jsonify({
            "success": True,
            "session_id": session_id,
            "message": "Script processing started"
        })

    except Exception as e:
        error_msg = f"Script processing request failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({"success": False, "error": error_msg})


@app.route("/export_markdown", methods=["POST"])
def export_markdown():
    """Export script content as Markdown file"""
    logger.info("=== MARKDOWN EXPORT REQUEST RECEIVED ===")

    try:
        data = request.json
        script_content = data.get("script_content", "")
        title = data.get("title", "AI_Script")

        if not script_content:
            logger.warning("❌ No script content provided for Markdown export")
            return jsonify(
                {"success": False, "error": "No script content to export"}
            )

        logger.info(f"📝 Exporting script to Markdown: {title}")

        # Clean up title for filename
        safe_title = "".join(
            c for c in title if c.isalnum() or c in (' ', '-', '_')
        ).strip()
        safe_title = safe_title.replace(' ', '_')

        # Create Markdown content with proper formatting
        markdown_content = f"# {title}\n\n"
        markdown_content += f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n"
        markdown_content += "---\n\n"
        markdown_content += script_content

        # Create in-memory file
        markdown_bytes = markdown_content.encode('utf-8')
        markdown_io = BytesIO(markdown_bytes)

        logger.info(
            f"✅ Markdown file created successfully ({len(markdown_bytes)} bytes)"
        )

        # Send file
        response = send_file(
            markdown_io,
            mimetype='text/markdown',
            as_attachment=True,
            download_name=f'{safe_title}.md'
        )

        return response

    except Exception as e:
        error_msg = f"Markdown export failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({"success": False, "error": error_msg})


@app.route("/download_comparison", methods=["POST"])
def download_comparison():
    """Download the chapter comparison report file"""
    logger.info("=== COMPARISON FILE DOWNLOAD REQUEST RECEIVED ===")

    try:
        data = request.json
        file_path = data.get("file_path", "")

        if not file_path:
            logger.warning("❌ No file path provided for comparison download")
            return jsonify(
                {"success": False, "error": "No file path provided"}
            ), 400

        # Security: Ensure the file is in the current directory and is a comparison file
        import os
        file_name = os.path.basename(file_path)
        if not file_name.startswith("chapter_comparison_") or not file_name.endswith(".md"):
            logger.warning(f"❌ Invalid file name: {file_name}")
            return jsonify(
                {"success": False, "error": "Invalid comparison file"}
            ), 400

        # Check if file exists
        if not os.path.exists(file_path):
            logger.warning(f"❌ File not found: {file_path}")
            return jsonify(
                {"success": False, "error": "Comparison file not found"}
            ), 404

        logger.info(f"📊 Sending comparison file: {file_name}")

        # Send the file
        return send_file(
            file_path,
            mimetype='text/markdown',
            as_attachment=True,
            download_name=file_name
        )

    except Exception as e:
        error_msg = f"Comparison file download failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({"success": False, "error": error_msg}), 500


@app.route("/export_heygen_curl", methods=["POST"])
def export_heygen_curl():
    """Extract and export HeyGen curl commands from script"""
    logger.info("=== HEYGEN CURL EXPORT REQUEST RECEIVED ===")

    try:
        data = request.json
        script_content = data.get("script_content", "")

        if not script_content:
            logger.warning(
                "❌ No script content provided for HeyGen curl export")
            return jsonify({"error": "No script content to export"}), 400

        # Extract HeyGen section
        heygen_section_start = script_content.find("# 🎬 HEYGEN READY SCRIPT")

        if heygen_section_start == -1:
            logger.warning("❌ No HeyGen section found in script")
            return jsonify({"error": "No HeyGen section found in this script. Make sure the script has been generated with HeyGen content."}), 404

        # Find curl commands section (try multiple markers)
        curl_markers = [
            "# 🚀 HEYGEN API CURL COMMANDS",
            "CURL COMMANDS:",
            "# HEYGEN API CURL COMMANDS"
        ]

        curl_start = -1
        marker_found = None
        for marker in curl_markers:
            pos = script_content.find(marker, heygen_section_start)
            if pos != -1:
                curl_start = pos
                marker_found = marker
                break

        if curl_start == -1:
            logger.warning("❌ No curl commands found in HeyGen section")
            msg = "No curl commands found in the HeyGen section."
            return jsonify({"error": msg}), 404

        # Extract curl commands (skip the marker line, get content after)
        curl_commands = script_content[curl_start:]

        # Curl commands now run to the end of the script (after section order fix)
        # No need to look for Demo Packages section as it comes BEFORE HeyGen now
        curl_commands = curl_commands.strip()

        if not curl_commands:
            logger.warning("❌ Curl commands section is empty")
            return jsonify({"error": "HeyGen curl commands section is empty."}), 404

        logger.info(
            f"✅ Extracted {len(curl_commands)} characters of curl commands")

        # Create in-memory text file
        curl_bytes = curl_commands.encode('utf-8')
        curl_io = BytesIO(curl_bytes)

        # Send file
        response = send_file(
            curl_io,
            mimetype='text/plain',
            as_attachment=True,
            download_name='heygen_curl_commands.txt'
        )

        logger.info("✅ HeyGen curl commands exported successfully")
        return response

    except Exception as e:
        error_msg = f"HeyGen curl export failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({"error": error_msg}), 500


@app.route("/api/heygen/templates", methods=["POST"])
def heygen_list_templates():
    """Proxy endpoint to fetch HeyGen templates (avoids CORS)"""
    try:
        data = request.get_json()
        api_key = data.get("api_key", "") if data else ""
        if not api_key:
            return jsonify({"error": "No API key provided"}), 400

        resp = requests.get(
            "https://api.heygen.com/v2/templates",
            headers={"accept": "application/json", "x-api-key": api_key},
            timeout=15,
        )
        resp.raise_for_status()
        return jsonify(resp.json())

    except requests.exceptions.HTTPError as e:
        status = e.response.status_code if e.response is not None else 500
        logger.error(f"HeyGen API error: {status} - {e}")
        return jsonify({"error": f"HeyGen API returned {status}"}), status
    except Exception as e:
        logger.error(f"HeyGen template fetch failed: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/heygen/template/<template_id>", methods=["POST"])
def heygen_template_details(template_id):
    """Proxy endpoint to fetch details for a single HeyGen template.
    Merges data from the list endpoint (name, thumbnail) and the
    detail endpoint (variables)."""
    try:
        data = request.get_json()
        api_key = data.get("api_key", "") if data else ""
        if not api_key:
            return jsonify({"error": "No API key provided"}), 400

        headers = {"accept": "application/json", "x-api-key": api_key}

        # Detail endpoint gives variables
        detail_resp = requests.get(
            f"https://api.heygen.com/v2/template/{template_id}",
            headers=headers, timeout=15,
        )
        detail_resp.raise_for_status()
        detail_data = detail_resp.json().get("data", {})

        # List endpoint gives name and thumbnail
        list_resp = requests.get(
            "https://api.heygen.com/v2/templates",
            headers=headers, timeout=15,
        )
        list_resp.raise_for_status()
        templates = list_resp.json().get("data", {}).get("templates", [])
        summary = next((t for t in templates if t.get(
            "template_id") == template_id), {})

        merged = {**summary, **detail_data, "template_id": template_id}
        return jsonify({"error": None, "data": merged})

    except requests.exceptions.HTTPError as e:
        status = e.response.status_code if e.response is not None else 500
        logger.error(f"HeyGen template detail error: {status} - {e}")
        return jsonify({"error": f"HeyGen API returned {status}"}), status
    except Exception as e:
        logger.error(f"HeyGen template detail fetch failed: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/heygen/voices", methods=["POST"])
def heygen_list_voices():
    """Proxy endpoint to fetch HeyGen voices (avoids CORS)"""
    try:
        data = request.get_json()
        api_key = data.get("api_key", "") if data else ""
        if not api_key:
            return jsonify({"error": "No API key provided"}), 400

        resp = requests.get(
            "https://api.heygen.com/v2/voices",
            headers={"accept": "application/json", "x-api-key": api_key},
            timeout=15,
        )
        resp.raise_for_status()
        return jsonify(resp.json())

    except requests.exceptions.HTTPError as e:
        status = e.response.status_code if e.response is not None else 500
        logger.error(f"HeyGen voices API error: {status} - {e}")
        return jsonify({"error": f"HeyGen API returned {status}"}), status
    except Exception as e:
        logger.error(f"HeyGen voices fetch failed: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/heygen/test-generate", methods=["POST"])
def heygen_test_generate():
    """Proxy endpoint to test HeyGen video generation from template."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400

        api_key = data.get("api_key", "")
        template_id = data.get("template_id", "")
        request_body = data.get("request_body", {})

        if not api_key:
            return jsonify({"error": "No API key provided"}), 400
        if not template_id:
            return jsonify({"error": "No template ID provided"}), 400

        url = f"https://api.heygen.com/v2/template/{template_id}/generate"
        headers = {
            "accept": "application/json",
            "content-type": "application/json",
            "x-api-key": api_key,
        }

        logger.info(f"🧪 Test generate: POST {url}")
        logger.info(f"🧪 Request body: {json.dumps(request_body)[:500]}")

        resp = requests.post(url, headers=headers,
                             json=request_body, timeout=30)

        logger.info(f"🧪 Response status: {resp.status_code}")
        logger.info(f"🧪 Response body: {resp.text[:500]}")

        # Return the full response with status so frontend can show it
        try:
            resp_json = resp.json()
        except Exception:
            resp_json = {"raw_response": resp.text}

        return jsonify(resp_json), resp.status_code

    except requests.exceptions.Timeout:
        logger.error("HeyGen test-generate timed out")
        return jsonify({"error": "Request timed out after 30 seconds"}), 504
    except Exception as e:
        logger.error(f"HeyGen test-generate failed: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/grok/test-video", methods=["POST"])
def grok_test_video_generate():
    """Generate a single Grok test video from a prompt and return a local preview path."""
    try:
        data = request.get_json() or {}
        prompt = (data.get("prompt") or "").strip()
        duration = int(data.get("duration") or 6)
        aspect_ratio = data.get("aspect_ratio") or "16:9"
        resolution = data.get("resolution") or "480p"

        if not prompt:
            return jsonify({"success": False, "error": "Prompt is required"}), 400

        duration = max(1, min(15, duration))

        xai_api_key = (data.get("api_key") or "").strip(
        ) or os.getenv("XAI_API_KEY", "").strip()
        if not xai_api_key or xai_api_key == "your-xai-api-key-here":
            return jsonify({
                "success": False,
                "error": "Grok API key is missing. Enter it in the Web GUI Grok Key field (or set XAI_API_KEY)."
            }), 400

        import xai_sdk
        import certifi
        import datetime as dt

        logger.info("🤖 Grok test video request started")
        logger.info(f"🤖 Prompt: {prompt[:180]}")
        logger.info(
            f"🤖 Settings: duration={duration}, aspect_ratio={aspect_ratio}, resolution={resolution}")

        client = xai_sdk.Client(api_key=xai_api_key)
        response = client.video.generate(
            prompt=prompt,
            model="grok-imagine-video",
            duration=duration,
            aspect_ratio=aspect_ratio,
            resolution=resolution,
        )

        broll_videos_dir = Path.home() / "Dev" / "brollvideos"
        broll_videos_dir.mkdir(parents=True, exist_ok=True)

        timestamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_prompt = "".join(c if c.isalnum() else "_" for c in prompt)[
            :40].strip("_")
        if not safe_prompt:
            safe_prompt = "test_video"
        filename = f"grok_test_{timestamp}_{safe_prompt}.mp4"
        local_path = broll_videos_dir / filename

        with requests.get(response.url, stream=True, timeout=180, verify=certifi.where()) as dl_resp:
            dl_resp.raise_for_status()
            with open(local_path, "wb") as out_file:
                for chunk in dl_resp.iter_content(chunk_size=8192):
                    if chunk:
                        out_file.write(chunk)

        logger.info(f"✅ Grok test video saved: {local_path}")

        return jsonify({
            "success": True,
            "filename": filename,
            "video_path": f"/broll-videos/{filename}",
            "source_url": response.url,
            "duration": getattr(response, "duration", duration),
            "model": getattr(response, "model", "grok-imagine-video"),
            "respect_moderation": getattr(response, "respect_moderation", None),
        })

    except Exception as e:
        logger.error(f"❌ Grok test video generation failed: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/cancel-broll-images/<session_id>", methods=["POST"])
def api_cancel_broll_images(session_id):
    """Signal the in-flight B-roll image generation loop to stop ASAP and return what was produced so far."""
    try:
        evt = broll_image_cancel_events.get(session_id)
        if evt is None:
            # Pre-create the event in case the request races image generation startup.
            evt = _threading.Event()
            broll_image_cancel_events[session_id] = evt
        evt.set()
        logger.info(
            f"🛑 B-roll image generation cancel requested for session {session_id}")
        return jsonify({"success": True, "session_id": session_id})
    except Exception as e:
        logger.error(f"❌ Failed to cancel B-roll image generation: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/cancel-thumbnails/<session_id>", methods=["POST"])
def api_cancel_thumbnails(session_id):
    """Signal in-flight thumbnail generation to stop and keep partial results."""
    try:
        evt = thumbnail_cancel_events.get(session_id)
        if evt is None:
            # Pre-create in case cancel arrives before generation fully starts.
            evt = _threading.Event()
            thumbnail_cancel_events[session_id] = evt
        evt.set()
        logger.info(
            f"🛑 Thumbnail generation cancel requested for session {session_id}")
        return jsonify({"success": True, "session_id": session_id})
    except Exception as e:
        logger.error(f"❌ Failed to cancel thumbnail generation: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/extract-docx", methods=["POST"])
def api_extract_docx():
    """Extract plain text from an uploaded .docx file so the front-end can load Word scripts."""
    try:
        upload = request.files.get("file")
        if upload is None or not upload.filename:
            return jsonify({"success": False, "error": "No file uploaded"}), 400

        name_lower = upload.filename.lower()
        if not (name_lower.endswith(".docx") or name_lower.endswith(".doc")):
            return jsonify({"success": False, "error": "Only .docx files are supported"}), 400
        if name_lower.endswith(".doc") and not name_lower.endswith(".docx"):
            return jsonify({"success": False, "error": "Legacy .doc not supported. Please save as .docx first."}), 400

        try:
            from docx import Document  # python-docx
        except ImportError as e:
            return jsonify({"success": False, "error": f"python-docx not installed: {e}"}), 500

        import io
        data = upload.read()
        doc = Document(io.BytesIO(data))

        def _para_text(para):
            """Extract paragraph text, preserving soft line-breaks (<w:br>) as \\n."""
            parts = []
            for node in para._p.iter():
                local = node.tag.split(
                    "}")[-1] if "}" in node.tag else node.tag
                if local == "br":
                    br_type = node.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "")
                    if br_type not in ("page", "column"):
                        parts.append("\n")
                elif local == "t":
                    parts.append(node.text or "")
            return "".join(parts).rstrip()

        lines = []
        for para in doc.paragraphs:
            raw = _para_text(para)
            style = (para.style.name or "") if para.style else ""
            # A paragraph may itself contain soft-return-separated sub-lines
            sub_lines = raw.split("\n")
            first = True
            for sub in sub_lines:
                sub = sub.rstrip()
                if first and sub and style.startswith("Heading"):
                    try:
                        level = int(style.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    level = max(1, min(level, 6))
                    lines.append(("#" * level) + " " + sub)
                else:
                    lines.append(sub)
                first = False

        # Append simple table extraction (tab-separated rows).
        for table in doc.tables:
            lines.append("")
            for row in table.rows:
                cells = [(c.text or "").strip().replace("\n", " ")
                         for c in row.cells]
                lines.append("\t".join(cells))

        text = "\n".join(lines).strip() + "\n"
        return jsonify({"success": True, "text": text, "filename": upload.filename, "length": len(text)})
    except Exception as e:
        logger.error(f"❌ /api/extract-docx failed: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/resolve/inspect", methods=["GET"])
def api_resolve_inspect():
    """Inspect the currently-loaded DaVinci Resolve timeline.

    Returns every clip on every video & audio track including B-roll,
    adjustment clips, fusion comp names, and the property bag for each
    timeline item (transform, composite mode, retime, etc.).
    """
    try:
        from resolve_inspector import inspect_current_timeline
        result = inspect_current_timeline()
        status = 200 if result.get("success") else 400
        return jsonify(result), status
    except Exception as e:
        logger.error(f"❌ /api/resolve/inspect failed: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/resolve/test-swipe", methods=["POST"])
def api_resolve_test_swipe():
    """Fast feedback loop for the V3 broll swipe transition.

    Body JSON:
        {
            "video_path": "/abs/path/to/sample.mp4"   (required),
            "track_index": 3                          (optional, default 3),
            "slide_seconds": 0.4                      (optional, default 0.4),
            "timeline_name": "swipe_test"             (optional)
        }

    Behavior:
        1. Connect to running Resolve, get/create a project.
        2. Import the video into a `swipe_test` bin.
        3. Create a fresh empty timeline (name suffixed with timestamp).
        4. Append the clip to V1 *and* to the configured track (default V3)
           so the swipe (which runs on V3 by default) has something to bite.
        5. Call `_apply_wipe_to_v3_clips(...)` from davinci_resolve_api.

    Returns the wipe-result dict so the UI can show diagnostics.
    """
    payload = request.get_json(silent=True) or {}
    video_path = (payload.get("video_path") or "").strip()
    track_index = int(payload.get("track_index") or 3)
    slide_seconds = float(payload.get("slide_seconds") or 0.4)
    base_name = (payload.get("timeline_name")
                 or "swipe_test").strip() or "swipe_test"

    if not video_path:
        return jsonify({"success": False, "error": "video_path is required"}), 400
    video_path = os.path.expanduser(video_path)
    if not os.path.isfile(video_path):
        return jsonify({"success": False, "error": f"video_path not found: {video_path}"}), 400

    try:
        # Connect to Resolve.
        resolve_script_api = (
            "/Library/Application Support/Blackmagic Design/"
            "DaVinci Resolve/Developer/Scripting"
        )
        resolve_script_lib = (
            "/Applications/DaVinci Resolve/DaVinci Resolve.app/"
            "Contents/Libraries/Fusion/fusionscript.so"
        )
        os.environ["RESOLVE_SCRIPT_API"] = resolve_script_api
        os.environ["RESOLVE_SCRIPT_LIB"] = resolve_script_lib
        modules_path = os.path.join(resolve_script_api, "Modules")
        if modules_path not in sys.path:
            sys.path.append(modules_path)

        import DaVinciResolveScript as dvr_script  # type: ignore
        resolve = dvr_script.scriptapp("Resolve")
        if resolve is None:
            return jsonify({"success": False,
                            "error": "Could not connect to DaVinci Resolve. "
                                     "Is it running?"}), 500

        pm = resolve.GetProjectManager()
        project = pm.GetCurrentProject()
        if project is None:
            # Try to create / open a scratch project.
            project = pm.CreateProject("ScriptCraft_SwipeTest") or \
                pm.LoadProject("ScriptCraft_SwipeTest")
            if project is None:
                return jsonify({"success": False,
                                "error": "No current project and could not "
                                         "create 'ScriptCraft_SwipeTest'"}), 500

        media_pool = project.GetMediaPool()
        root_folder = media_pool.GetRootFolder()

        # Make / find a swipe_test bin.
        bin_name = "swipe_test"
        target_bin = None
        try:
            for sub in (root_folder.GetSubFolderList() or []):
                if sub.GetName() == bin_name:
                    target_bin = sub
                    break
        except Exception:
            pass
        if target_bin is None:
            try:
                media_pool.SetCurrentFolder(root_folder)
                target_bin = media_pool.AddSubFolder(root_folder, bin_name)
            except Exception:
                target_bin = root_folder
        try:
            media_pool.SetCurrentFolder(target_bin)
        except Exception:
            pass

        # Import the video.
        imported = media_pool.ImportMedia([video_path]) or []
        if not imported:
            return jsonify({"success": False,
                            "error": f"ImportMedia returned nothing for {video_path}"}), 500
        clip = imported[0]

        # Create a fresh timeline.
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        timeline_name = f"{base_name}_{ts}"
        timeline = media_pool.CreateEmptyTimeline(timeline_name)
        if timeline is None:
            return jsonify({"success": False,
                            "error": f"CreateEmptyTimeline('{timeline_name}') returned None"}), 500
        project.SetCurrentTimeline(timeline)

        # Make sure the target track exists.
        try:
            current_video_tracks = timeline.GetTrackCount("video") or 1
            while current_video_tracks < track_index:
                timeline.AddTrack("video")
                current_video_tracks = timeline.GetTrackCount(
                    "video") or current_video_tracks + 1
        except Exception as _te:
            logger.warning(f"⚠️ Could not ensure V{track_index} exists: {_te}")

        # Place the clip exactly once on the target track.
        appended_target = media_pool.AppendToTimeline([{
            "mediaPoolItem": clip,
            "startFrame": 0,
            "endFrame": int(clip.GetClipProperty("Frames") or 240) - 1,
            "trackIndex": track_index,
        }]) or []
        logger.info(
            f"🧪 swipe-test: appended V{track_index}={len(appended_target)}")

        # Apply the swipe.
        from davinci_resolve_api import _apply_wipe_to_v3_clips
        wipe_result = _apply_wipe_to_v3_clips(
            timeline,
            track_index=track_index,
            slide_seconds=slide_seconds,
        )

        return jsonify({
            "success": bool(wipe_result.get("success")),
            "timeline": timeline_name,
            "video_path": video_path,
            "track_index": track_index,
            "slide_seconds": slide_seconds,
            "appended_target": len(appended_target),
            "wipe": wipe_result,
        })
    except Exception as e:
        logger.error(f"❌ /api/resolve/test-swipe failed: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/resolve/probe-fusion", methods=["GET"])
def api_resolve_probe_fusion():
    """For a given timeline item unique_id, dump every Fusion comp's tool list.

    Use this to see whether a paid template (e.g. mTuber 4 Tiles Swap) exposes
    its inputs to the scripting API or is encrypted.
    """
    try:
        unique_id = request.args.get("unique_id", "").strip()
        if not unique_id:
            return jsonify({"success": False, "error": "unique_id query param is required"}), 400
        from resolve_inspector import probe_fusion_comp
        result = probe_fusion_comp(unique_id)
        status = 200 if result.get("success") else 400
        return jsonify(result), status
    except Exception as e:
        logger.error(f"❌ /api/resolve/probe-fusion failed: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/resolve/save-comp-preset", methods=["POST"])
def api_resolve_save_comp_preset():
    """Save the Fusion composition from one timeline clip as a reusable preset.

    Body JSON: { "unique_id": "...", "preset_name": "default_broll" (optional) }

    The preset file (default: scriptcraft-app-v2/fusion_presets/default_broll.setting)
    is what Create DaVinci Project will auto-apply to every V3 broll clip.
    """
    try:
        from resolve_inspector import save_comp_preset
        payload = request.get_json(silent=True) or {}
        unique_id = (payload.get("unique_id") or "").strip()
        if not unique_id:
            return jsonify({"success": False, "error": "unique_id is required"}), 400
        preset_name = (payload.get("preset_name") or "").strip() or None
        result = save_comp_preset(unique_id, preset_name=preset_name)
        return jsonify(result), (200 if result.get("success") else 400)
    except Exception as e:
        logger.error(f"❌ /api/resolve/save-comp-preset failed: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/keys", methods=["GET"])
def api_keys_get():
    """Return all saved API keys (Google, HeyGen, Grok). Falls back to env vars."""
    try:
        settings = _load_scriptcraft_settings()

        def _v(key, env):
            return (settings.get(key) or "").strip() or os.getenv(env, "").strip()
        return jsonify({
            "success": True,
            "keys": {
                "google_api_key": _v("google_api_key", "GOOGLE_API_KEY"),
                "heygen_api_key": _v("heygen_api_key", "HEYGEN_API_KEY"),
                "grok_api_key": _v("grok_api_key", "XAI_API_KEY"),
            },
        })
    except Exception as e:
        logger.error(f"❌ Failed to load API keys: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/keys", methods=["POST"])
def api_keys_save():
    """Persist Google / HeyGen / Grok keys and update process env so backend modules see them."""
    try:
        data = request.get_json(silent=True) or {}
        settings = _load_scriptcraft_settings()

        mapping = {
            "google_api_key": "GOOGLE_API_KEY",
            "heygen_api_key": "HEYGEN_API_KEY",
            "grok_api_key": "XAI_API_KEY",
        }
        updated = []
        for setting_key, env_var in mapping.items():
            if setting_key in data:
                val = (data.get(setting_key) or "").strip()
                if val:
                    settings[setting_key] = val
                    os.environ[env_var] = val
                else:
                    settings.pop(setting_key, None)
                    os.environ.pop(env_var, None)
                updated.append(setting_key)

        if not updated:
            return jsonify({"success": False, "error": "No keys provided"}), 400

        if not _save_scriptcraft_settings(settings):
            return jsonify({"success": False, "error": "Could not save settings"}), 500

        logger.info(f"🔑 API keys updated: {updated}")
        return jsonify({"success": True, "updated": updated})
    except Exception as e:
        logger.error(f"❌ Failed to save API keys: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/output-dir", methods=["GET"])
def api_output_dir_get():
    """Return the currently configured base output directory (or empty string)."""
    try:
        settings = _load_scriptcraft_settings()
        return jsonify({
            "success": True,
            "output_dir": (settings.get("output_dir") or "").strip(),
        })
    except Exception as e:
        logger.error(f"❌ Failed to load output_dir: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/output-dir", methods=["POST"])
def api_output_dir_save():
    """Persist the base output directory path without creating it immediately."""
    try:
        data = request.get_json(silent=True) or {}
        raw = (data.get("output_dir") or "").strip()
        settings = _load_scriptcraft_settings()

        if not raw:
            settings.pop("output_dir", None)
            _save_scriptcraft_settings(settings)
            return jsonify({"success": True, "output_dir": "", "cleared": True})

        expanded = Path(raw).expanduser().resolve()
        if expanded.exists() and not expanded.is_dir():
            return jsonify({
                "success": False,
                "error": "Path exists but is not a directory",
            }), 400

        settings["output_dir"] = str(expanded)
        if not _save_scriptcraft_settings(settings):
            return jsonify({"success": False, "error": "Could not save settings"}), 500

        logger.info(f"📁 Output directory set: {expanded}")
        return jsonify({"success": True, "output_dir": str(expanded)})
    except Exception as e:
        logger.error(f"❌ Failed to save output_dir: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/browse-dirs", methods=["GET"])
def api_browse_dirs():
    """List immediate subdirectories of a path for the in-app folder browser.

    Query params:
      path: starting absolute or ~-prefixed path. Empty/missing -> $HOME.
      show_hidden: '1' to include dotfile directories.
    Returns:
      {success, path, parent, entries: [{name, path}], home, can_create}
    """
    try:
        raw = (request.args.get("path") or "").strip()
        show_hidden = request.args.get("show_hidden") == "1"
        home = str(Path.home())

        if not raw:
            target = Path.home()
        else:
            target = Path(raw).expanduser()

        try:
            target = target.resolve()
        except Exception:
            return jsonify({"success": False, "error": f"Invalid path: {raw}"}), 400

        if not target.exists():
            return jsonify({"success": False, "error": f"Path does not exist: {target}"}), 404
        if not target.is_dir():
            return jsonify({"success": False, "error": f"Not a directory: {target}"}), 400

        entries = []
        try:
            for child in sorted(target.iterdir(), key=lambda p: p.name.lower()):
                try:
                    if not child.is_dir():
                        continue
                    if not show_hidden and child.name.startswith("."):
                        continue
                    entries.append({"name": child.name, "path": str(child)})
                except (PermissionError, OSError):
                    continue
        except PermissionError:
            return jsonify({"success": False, "error": f"Permission denied: {target}"}), 403

        parent = str(target.parent) if target.parent != target else None
        return jsonify({
            "success": True,
            "path": str(target),
            "parent": parent,
            "home": home,
            "entries": entries,
        })
    except Exception as e:
        logger.error(f"❌ /api/browse-dirs failed: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/create-dir", methods=["POST"])
def api_create_dir():
    """Create a new subdirectory inside the given parent path. Used by the folder browser."""
    try:
        data = request.get_json(silent=True) or {}
        parent_raw = (data.get("parent") or "").strip()
        name_raw = (data.get("name") or "").strip()
        if not parent_raw or not name_raw:
            return jsonify({"success": False, "error": "parent and name are required"}), 400
        # Reject path separators in folder name
        if "/" in name_raw or "\\" in name_raw or name_raw in (".", ".."):
            return jsonify({"success": False, "error": "Invalid folder name"}), 400

        parent = Path(parent_raw).expanduser().resolve()
        if not parent.is_dir():
            return jsonify({"success": False, "error": f"Parent is not a directory: {parent}"}), 400

        target = parent / name_raw
        target.mkdir(parents=False, exist_ok=True)
        return jsonify({"success": True, "path": str(target)})
    except Exception as e:
        logger.error(f"❌ /api/create-dir failed: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/save-outputs", methods=["POST"])
def api_save_outputs():
    """Copy the current run's artifacts into {output_parent}/{safe_script_title}.

    Layout:
        {run_dir}/script/<title>.md
        {run_dir}/images/<broll image files>
        {run_dir}/broll/<grok video files>
        {run_dir}/MDL/<edl file>
        {run_dir}/thumbnails/<thumbnail files>
        {run_dir}/curls/<title>_curls.json
    """
    import shutil
    import re as _re

    try:
        data = request.get_json(silent=True) or {}
        title_raw = (data.get("title") or "script").strip() or "script"
        safe_title = _re.sub(r'[^A-Za-z0-9._-]+', '_',
                             title_raw).strip('_') or "script"
        base_path = _get_run_output_dir(title_raw, create=True)

        script_text = data.get("script") or ""
        edl_filename = (data.get("edl_filename") or "").strip()
        edl_content = data.get("edl_content") or ""
        curl_commands = data.get("curl_commands")
        broll_images = data.get("broll_images") or []
        grok_videos = data.get("grok_videos") or []
        thumbnails = data.get("thumbnails") or []

        report = {
            "script": None,
            "edl": None,
            "curls": None,
            "images": [], "images_skipped": [],
            "videos": [], "videos_skipped": [],
            "thumbnails": [], "thumbnails_skipped": [],
        }

        # --- script ---
        if script_text:
            script_dir = base_path / "script"
            script_dir.mkdir(parents=True, exist_ok=True)
            script_path = script_dir / f"{safe_title}.md"
            script_path.write_text(script_text, encoding="utf-8")
            report["script"] = str(script_path)

        # --- EDL (user calls it MDL) ---
        if edl_content or edl_filename:
            mdl_dir = base_path / "MDL"
            mdl_dir.mkdir(parents=True, exist_ok=True)
            target_name = Path(
                edl_filename).name if edl_filename else f"{safe_title}.edl"
            edl_path = mdl_dir / target_name
            if edl_content:
                edl_path.write_text(edl_content, encoding="utf-8")
            else:
                src = _resolve_media_file(target_name, 'edl')
                if src:
                    shutil.copy2(src, edl_path)
                else:
                    edl_path = None
            if edl_path:
                report["edl"] = str(edl_path)

        # --- curl commands ---
        if curl_commands:
            curl_dir = base_path / "curls"
            curl_dir.mkdir(parents=True, exist_ok=True)
            if isinstance(curl_commands, str):
                curl_path = curl_dir / f"{safe_title}_curls.sh"
                curl_path.write_text(curl_commands, encoding="utf-8")
            else:
                curl_path = curl_dir / f"{safe_title}_curls.json"
                curl_path.write_text(json.dumps(
                    curl_commands, indent=2), encoding="utf-8")
            report["curls"] = str(curl_path)

        # --- images / videos / thumbnails: pull names out of structured items ---
        def _names_of(items):
            out = []
            for it in (items or []):
                if isinstance(it, str):
                    out.append(it)
                elif isinstance(it, dict):
                    fn = it.get("filename") or it.get(
                        "filepath") or it.get("path")
                    if fn:
                        out.append(Path(fn).name)
            return out

        image_names = _names_of(broll_images)
        if image_names:
            img_dir = base_path / "images"
            img_dir.mkdir(parents=True, exist_ok=True)
            for fn in image_names:
                src = _resolve_media_file(fn, 'image')
                if src:
                    dst = img_dir / src.name
                    shutil.copy2(src, dst)
                    report["images"].append(str(dst))
                else:
                    report["images_skipped"].append(fn)

        video_names = _names_of(grok_videos)
        if video_names:
            vid_dir = base_path / "broll"
            vid_dir.mkdir(parents=True, exist_ok=True)
            for fn in video_names:
                src = _resolve_media_file(fn, 'video')
                if src:
                    dst = vid_dir / src.name
                    shutil.copy2(src, dst)
                    report["videos"].append(str(dst))
                else:
                    report["videos_skipped"].append(fn)

        thumb_names = _names_of(thumbnails)
        if thumb_names:
            thumb_dir = base_path / "thumbnails"
            thumb_dir.mkdir(parents=True, exist_ok=True)
            for fn in thumb_names:
                src = _resolve_media_file(fn, 'thumbnail')
                if src:
                    dst = thumb_dir / src.name
                    shutil.copy2(src, dst)
                    report["thumbnails"].append(str(dst))
                else:
                    report["thumbnails_skipped"].append(fn)

        total_saved = (
            (1 if report["script"] else 0)
            + (1 if report["edl"] else 0)
            + (1 if report["curls"] else 0)
            + len(report["images"]) + len(report["videos"]) +
            len(report["thumbnails"])
        )
        total_skipped = (
            len(report["images_skipped"])
            + len(report["videos_skipped"])
            + len(report["thumbnails_skipped"])
        )
        logger.info(
            f"💾 save-outputs → run_dir={base_path} | saved={total_saved} skipped={total_skipped}"
        )
        return jsonify({
            "success": True,
            "base_dir": str(base_path),
            "report": report,
            "total_saved": total_saved,
            "total_skipped": total_skipped,
        })
    except Exception as e:
        logger.error(f"❌ /api/save-outputs failed: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/grok/key", methods=["GET"])
def grok_get_saved_key():
    """Return saved Grok API key from server settings (if any)."""
    try:
        settings = _load_scriptcraft_settings()
        key = (settings.get("grok_api_key") or "").strip()
        return jsonify({"success": True, "api_key": key})
    except Exception as e:
        logger.error(f"❌ Failed to load saved Grok key: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/grok/key", methods=["POST"])
def grok_save_key():
    """Persist Grok API key in server settings so it survives browser/session changes."""
    try:
        data = request.get_json() or {}
        api_key = (data.get("api_key") or "").strip()
        if not api_key:
            return jsonify({"success": False, "error": "api_key is required"}), 400

        settings = _load_scriptcraft_settings()
        settings["grok_api_key"] = api_key
        if not _save_scriptcraft_settings(settings):
            return jsonify({"success": False, "error": "Could not save settings"}), 500

        return jsonify({"success": True})
    except Exception as e:
        logger.error(f"❌ Failed to save Grok key: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


def _grok_get_output_dir(script_title: Optional[str] = None) -> Path:
    """Resolve directory for Grok video output under the active run folder."""
    if script_title:
        base = _get_run_output_dir(script_title, create=True)
    else:
        base = _get_output_parent_dir()
    target = base / "broll"
    target.mkdir(parents=True, exist_ok=True)
    return target


def _grok_emit(session_id: str, payload: dict) -> None:
    q = grok_video_streams.get(session_id)
    if q is not None:
        try:
            q.put_nowait(payload)
        except Exception:
            pass


def _grok_generate_worker(session_id: str, selected_rows: list, api_key: str, script_title: str = "") -> None:
    """Background worker: generate Grok videos one-by-one, streaming progress and honoring cancel."""
    cancel_evt = grok_video_cancel_events.setdefault(
        session_id, _threading.Event())
    out_dir = _grok_get_output_dir(script_title or None)
    videos: list = []
    failures: list = []
    total = len(selected_rows)
    cancelled = False

    try:
        import xai_sdk  # type: ignore
        import certifi  # type: ignore
        client = xai_sdk.Client(api_key=api_key)
    except Exception as e:
        _grok_emit(session_id, {
            "type": "error",
            "message": f"❌ Failed to init Grok client: {e}",
        })
        _grok_emit(session_id, {
            "type": "done", "cancelled": False, "videos": [], "failures": [],
            "generated_count": 0, "failed_count": 0, "total_requested": total,
            "output_dir": str(out_dir),
        })
        grok_video_results[session_id] = {
            "videos": [], "failures": [], "cancelled": False}
        return

    _grok_emit(session_id, {
        "type": "start",
        "message": f"🎬 Starting Grok video generation: {total} clip(s) → {out_dir}",
        "total": total,
        "output_dir": str(out_dir),
    })

    for i, row in enumerate(selected_rows):
        if cancel_evt.is_set():
            cancelled = True
            break

        prompt = (row.get("description") or row.get(
            "search_term") or "").strip()
        term = row.get("search_term", f"vid{i}") or f"vid{i}"
        timecode = row.get("timecode", "")

        if not prompt:
            failures.append(
                {"index": i, "error": "Missing prompt", "row": row})
            _grok_emit(session_id, {
                "type": "video_failed",
                "index": i, "current": i + 1, "total": total,
                "search_term": term, "timecode": timecode,
                "error": "Missing prompt",
                "message": f"⚠️ [{i+1}/{total}] Skipped '{term}' — missing prompt",
            })
            continue

        _grok_emit(session_id, {
            "type": "video_start",
            "index": i, "current": i + 1, "total": total,
            "search_term": term, "timecode": timecode, "description": prompt,
            "message": f"🎨 [{i+1}/{total}] Generating '{term}'…",
        })

        try:
            response = client.video.generate(
                prompt=prompt,
                model="grok-imagine-video",
                duration=6,
                aspect_ratio="16:9",
                resolution="480p",
            )

            if cancel_evt.is_set():
                cancelled = True
                break

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_term = "".join(c if c.isalnum() else "_" for c in term)[:30]
            filename = f"grok_{timestamp}_{i}_{safe_term}.mp4"
            local_path = out_dir / filename

            with requests.get(response.url, stream=True, timeout=180, verify=certifi.where()) as dl_resp:
                dl_resp.raise_for_status()
                with open(local_path, "wb") as out_file:
                    for chunk in dl_resp.iter_content(chunk_size=8192):
                        if cancel_evt.is_set():
                            break
                        if chunk:
                            out_file.write(chunk)

            if cancel_evt.is_set():
                # Partial download — drop the incomplete file.
                try:
                    if local_path.exists():
                        local_path.unlink()
                except Exception:
                    pass
                cancelled = True
                break

            video_entry = {
                "timecode": timecode,
                "search_term": term,
                "description": prompt,
                "url": response.url,
                "filename": str(local_path),
            }
            videos.append(video_entry)

            _grok_emit(session_id, {
                "type": "video_ready",
                "index": i, "current": i + 1, "total": total,
                "video": video_entry,
                "message": f"✅ [{i+1}/{total}] Saved '{term}' → {local_path}",
            })
        except Exception as row_err:
            failures.append({
                "index": i,
                "timecode": timecode,
                "search_term": term,
                "error": str(row_err),
            })
            _grok_emit(session_id, {
                "type": "video_failed",
                "index": i, "current": i + 1, "total": total,
                "search_term": term, "timecode": timecode,
                "error": str(row_err),
                "message": f"❌ [{i+1}/{total}] Failed '{term}': {row_err}",
            })

    grok_video_results[session_id] = {
        "videos": videos,
        "failures": failures,
        "cancelled": cancelled,
        "output_dir": str(out_dir),
    }

    final_msg = (
        f"🛑 Cancelled — kept {len(videos)} of {total} videos"
        if cancelled
        else f"✅ Grok videos complete: {len(videos)}/{total} (failed: {len(failures)})"
    )
    _grok_emit(session_id, {
        "type": "done",
        "cancelled": cancelled,
        "videos": videos,
        "failures": failures,
        "generated_count": len(videos),
        "failed_count": len(failures),
        "total_requested": total,
        "output_dir": str(out_dir),
        "message": final_msg,
    })


@app.route("/api/grok/generate-selected", methods=["POST"])
def grok_generate_selected_videos():
    """Start background Grok video generation for selected B-roll rows. Streams via SSE."""
    try:
        data = request.get_json() or {}
        selected_rows = data.get("selected_rows") or []
        api_key = (data.get("api_key") or "").strip(
        ) or os.getenv("XAI_API_KEY", "").strip()
        script_title = (data.get("script_title") or "").strip()
        session_id = (data.get("session_id")
                      or "").strip() or str(uuid.uuid4())

        if not api_key:
            return jsonify({"success": False, "error": "Grok API key is required"}), 400
        if not selected_rows:
            return jsonify({"success": False, "error": "No rows selected"}), 400

        # Reset per-session state
        grok_video_cancel_events[session_id] = _threading.Event()
        grok_video_streams[session_id] = _queue.Queue()
        grok_video_results.pop(session_id, None)

        thread = _threading.Thread(
            target=_grok_generate_worker,
            args=(session_id, selected_rows, api_key, script_title),
            daemon=True,
        )
        thread.start()

        return jsonify({
            "success": True,
            "session_id": session_id,
            "total": len(selected_rows),
            "stream_url": f"/api/grok/progress/{session_id}",
            "cancel_url": f"/api/grok/cancel/{session_id}",
        })
    except Exception as e:
        logger.error(f"❌ Grok selected generation failed to start: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/grok/progress/<session_id>")
def grok_progress_stream(session_id):
    """SSE stream of per-video progress events for a Grok generation session."""
    q = grok_video_streams.get(session_id)
    if q is None:
        return "Session not found", 404

    def generate():
        try:
            while True:
                try:
                    payload = q.get(timeout=300)
                except _queue.Empty:
                    yield ": keepalive\n\n"
                    continue
                yield f"data: {json.dumps(payload)}\n\n"
                if payload.get("type") == "done":
                    break
        finally:
            # Leave results in grok_video_results for inspection; clean stream + cancel event.
            grok_video_streams.pop(session_id, None)
            grok_video_cancel_events.pop(session_id, None)

    return Response(generate(), mimetype="text/event-stream", headers={
        "Cache-Control": "no-cache",
        "X-Accel-Buffering": "no",
        "Connection": "keep-alive",
    })


@app.route("/api/grok/cancel/<session_id>", methods=["POST"])
def grok_cancel(session_id):
    """Signal an in-flight Grok generation to stop after the current step. Already-saved videos are kept."""
    try:
        evt = grok_video_cancel_events.get(session_id)
        if evt is None:
            evt = _threading.Event()
            grok_video_cancel_events[session_id] = evt
        evt.set()
        logger.info(
            f"🛑 Grok video generation cancel requested for session {session_id}")
        return jsonify({"success": True, "session_id": session_id})
    except Exception as e:
        logger.error(f"❌ Failed to cancel Grok generation: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/test_word")
def test_word_page():
    """Test page for Word export functionality"""
    return render_template("test_word_export.html")


@app.route("/export_word", methods=["POST"])
def export_word():
    """Export script content to Word document"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No data provided"})

        script_content = data.get("script_content", "")
        title = data.get("title", "AI Script")

        if not script_content:
            return jsonify({"success": False, "error": "No script content to export"})

        logger.info(
            f"📄 Exporting script to Word: {len(script_content)} characters")

        # Import Word processing libraries directly
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            logger.info("✅ Word processing libraries imported successfully")
        except ImportError as e:
            logger.error(f"❌ Failed to import docx libraries: {e}")
            return jsonify({"success": False, "error": "python-docx not available"})

        # Create temporary file path
        import tempfile
        temp_dir = Path(tempfile.gettempdir())
        temp_file = temp_dir / f"script_{uuid.uuid4().hex[:8]}.docx"

        # Convert script to Word document directly
        try:
            logger.info("📝 Creating Word document...")

            # Create a new Word document
            doc = Document()

            # Process the script content line by line
            lines = script_content.split('\n')
            logger.info(f"🔍 Processing {len(lines)} lines")

            i = 0
            while i < len(lines):
                line = lines[i].rstrip()

                # Skip empty lines
                if not line:
                    i += 1
                    continue

                # Handle markdown tables - detect header row with pipes
                if '|' in line:
                    # Look ahead to check if next line is a separator
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].rstrip()
                        # Check if next line is a table separator (pipes, dashes, colons, spaces)
                        if '|' in next_line and all(c in '|-: \t' for c in next_line.replace('|', '')):
                            # This is a markdown table!
                            logger.info(f"📊 Detected table at line {i}")

                            # Collect header row
                            header_line = line
                            i += 1  # Skip to separator line
                            i += 1  # Skip separator, move to first data row

                            # Collect all data rows (lines with pipes)
                            data_lines = []
                            while i < len(lines):
                                data_line = lines[i].rstrip()
                                if '|' in data_line:
                                    data_lines.append(data_line)
                                    i += 1
                                else:
                                    break

                            # Parse the table
                            # Split on pipes and clean up
                            header_cells = [cell.strip()
                                            for cell in header_line.split('|')]
                            # Remove empty cells from start/end if table has leading/trailing pipes
                            if header_cells and not header_cells[0]:
                                header_cells = header_cells[1:]
                            if header_cells and not header_cells[-1]:
                                header_cells = header_cells[:-1]

                            # Parse data rows
                            data_rows = []
                            for data_line in data_lines:
                                cells = [cell.strip()
                                         for cell in data_line.split('|')]
                                # Remove empty cells from start/end
                                if cells and not cells[0]:
                                    cells = cells[1:]
                                if cells and not cells[-1]:
                                    cells = cells[:-1]
                                if cells:  # Only add non-empty rows
                                    data_rows.append(cells)

                            # Create Word table
                            if header_cells and data_rows:
                                logger.info(
                                    f"📊 Creating Word table: {len(header_cells)} columns, {len(data_rows)} rows")

                                table = doc.add_table(
                                    rows=1 + len(data_rows), cols=len(header_cells))
                                table.style = 'Light Grid Accent 1'

                                # Add headers
                                for col_idx, header in enumerate(header_cells):
                                    if col_idx < len(table.rows[0].cells):
                                        cell = table.rows[0].cells[col_idx]
                                        cell.text = header
                                        # Make header bold
                                        for paragraph in cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.bold = True

                                # Add data rows
                                for row_idx, row_data in enumerate(data_rows, start=1):
                                    for col_idx, cell_text in enumerate(row_data):
                                        if col_idx < len(header_cells) and row_idx < len(table.rows):
                                            table.rows[row_idx].cells[col_idx].text = cell_text

                                # Add spacing after table
                                doc.add_paragraph()
                                logger.info(f"✅ Table created successfully")
                            else:
                                logger.warning(
                                    f"⚠️ Table parsing failed: headers={len(header_cells)}, rows={len(data_rows)}")

                            continue

                # Handle horizontal rules (---)
                elif line.strip() in ['---', '___', '***']:
                    doc.add_paragraph('_' * 50)
                    i += 1
                    continue

                # Handle headers (lines starting with #)
                elif line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('#').strip()

                    if level == 1:  # Main title
                        paragraph = doc.add_heading(text, level=1)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif level == 2:  # Chapter/Section
                        paragraph = doc.add_heading(text, level=2)
                    else:  # Sub-sections
                        paragraph = doc.add_heading(text, level=3)
                    i += 1
                    continue

                # Handle bold text (**text**)
                elif '**' in line:
                    paragraph = doc.add_paragraph()
                    parts = line.split('**')
                    for idx, part in enumerate(parts):
                        if idx % 2 == 0:  # Normal text
                            if part:
                                paragraph.add_run(part)
                        else:  # Bold text
                            if part:
                                run = paragraph.add_run(part)
                                run.bold = True
                    i += 1
                    continue

                # Handle visual cues (lines with 🎬 emoji or [Visual:])
                elif ('🎬' in line or '[Visual:' in line or
                      (line.startswith('[') and ']' in line)):
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.italic = True
                    paragraph.space_before = Pt(6)
                    paragraph.space_after = Pt(6)
                    i += 1
                    continue

                # Regular paragraph
                else:
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(line)
                    i += 1

            # Save document to temporary file
            doc.save(str(temp_file))
            logger.info(f"💾 Word document saved to: {temp_file}")

            if not temp_file.exists():
                return jsonify({"success": False, "error": "Failed to create Word document"})

            # Read the file and return as binary response
            with open(temp_file, 'rb') as f:
                file_content = f.read()

            # Clean up temp file
            temp_file.unlink()

            logger.info(f"✅ Word export successful: {len(file_content)} bytes")

            # Return file as binary response
            response = make_response(file_content)
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename="{title}.docx"'
            return response

        except Exception as e:
            logger.error(f"❌ Word conversion failed: {e}")
            return jsonify({"success": False, "error": f"Word conversion failed: {str(e)}"})

    except Exception as e:
        error_msg = f"Export to Word failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({"success": False, "error": error_msg})


@app.route("/test-thumbnails")
def test_thumbnails():
    """Test page to view thumbnail gallery with existing thumbnails"""
    try:
        thumbnail_dir = Path.home() / "Dev" / "Thumbnails"

        # Get all PNG files
        if thumbnail_dir.exists():
            png_files = list(thumbnail_dir.glob("*.png"))
            # Take the most recent 6 files
            png_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
            png_files = png_files[:6]

            # Create thumbnail data
            thumbnails = []
            emotions = ["ANGRY", "SHOCKED", "SCARED",
                        "EXCITED", "SKEPTICAL", "DETERMINED"]
            for i, file_path in enumerate(png_files):
                thumbnails.append({
                    "emotion": emotions[i] if i < len(emotions) else f"VARIANT_{i+1}",
                    "text": f"Test thumbnail {i+1}",
                    "filename": file_path.name
                })

            return render_template("test_thumbnails.html", thumbnails=thumbnails)
        else:
            return "Thumbnail directory not found", 404
    except Exception as e:
        return f"Error: {str(e)}", 500


@app.route("/thumbnail-test")
def thumbnail_test():
    """Interactive thumbnail generation test page"""
    return render_template("thumbnail_test.html")


def _thumbnail_test_worker(session_id, topic, headlines, variations_per_option):
    """Background worker: generates test thumbnails and streams progress events."""
    q = thumbnail_test_streams.get(session_id)

    def _emit(event_type, **kwargs):
        if q is None:
            return
        payload = {"type": event_type, **kwargs, "ts": time.time()}
        try:
            q.put(payload)
        except Exception:
            pass

    def _log(msg):
        print(msg)
        _emit("log", message=str(msg))

    try:
        _log(f"🎬 Generating thumbnails for: {topic}")
        if headlines:
            _log(f"🏷️ Hooks ({len(headlines)}): {headlines}")
        if variations_per_option:
            _log(f"🔢 Variations per hook: {variations_per_option}")

        from tools.media.emotional_thumbnail_generator import (
            EmotionalThumbnailGenerator,
        )

        thumbnail_gen = EmotionalThumbnailGenerator()
        _log(f"✅ Generator initialized")
        _log(f"   Template: {thumbnail_gen.template_path}")
        _log(f"   Output:   {thumbnail_gen.output_dir}")

        thumb_cancel_evt = _threading.Event()
        thumbnail_cancel_events[session_id] = thumb_cancel_evt

        thumbnail_results = thumbnail_gen.generate_all_thumbnails(
            script_title=topic,
            script_content="Test script content",
            youtube_upload_details=None,
            headline_options=headlines or None,
            variations_per_option=variations_per_option,
            progress_callback=_log,
            cancel_check=thumb_cancel_evt.is_set,
        )

        if thumbnail_results and thumbnail_results.get("variations"):
            variations = thumbnail_results["variations"]
            output_dir = thumbnail_results.get("output_dir")
            _log(f"✅ Generated {len(variations)} thumbnails")
            if output_dir:
                _log(f"📁 Saved to: {output_dir}")

            thumbnail_test_results[session_id] = {
                "success": True,
                "thumbnails": variations,
                "output_dir": output_dir,
                "count": len(variations),
                "topic": topic,
            }
            _emit("done", success=True, thumbnails=variations,
                  output_dir=output_dir, count=len(variations),
                  download_url=f"/api/thumbnails/download/{session_id}")
        else:
            err = (thumbnail_results or {}).get(
                "error") or "No thumbnails generated"
            output_dir = (thumbnail_results or {}).get("output_dir")
            _log(f"❌ {err}")
            thumbnail_test_results[session_id] = {
                "success": False,
                "error": err,
                "output_dir": output_dir,
                "thumbnails": [],
            }
            _emit("done", success=False, error=err,
                  output_dir=output_dir, thumbnails=[])

    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        _log(f"❌ Thumbnail worker error: {e}")
        _log(tb)
        thumbnail_test_results[session_id] = {
            "success": False,
            "error": str(e),
            "thumbnails": [],
        }
        _emit("done", success=False, error=str(e), thumbnails=[])
    finally:
        thumbnail_cancel_events.pop(session_id, None)


@app.route("/generate-test-thumbnails", methods=["POST"])
def generate_test_thumbnails():
    """Start a thumbnail-generation session; progress streams via /api/thumbnails/progress/<id>."""
    try:
        data = request.json or {}
        topic = data.get("topic", "AI Testing")
        raw_headlines = data.get("headlines") or []
        headlines = [str(h).strip()
                     for h in raw_headlines if str(h or "").strip()][:3]
        variations_per_option = int(
            data.get("variations_per_option") or 0) or None

        session_id = str(uuid.uuid4())
        thumbnail_test_streams[session_id] = _queue.Queue()
        thumbnail_test_results.pop(session_id, None)

        print(f"\n{'='*70}")
        print(f"🎬 GENERATING TEST THUMBNAILS (session={session_id})")
        print(f"📝 Topic: {topic}")
        if headlines:
            print(f"🏷️ Headlines: {headlines}")
        if variations_per_option:
            print(f"🔢 Variations per headline: {variations_per_option}")
        print(f"{'='*70}\n")

        thread = _threading.Thread(
            target=_thumbnail_test_worker,
            args=(session_id, topic, headlines, variations_per_option),
            daemon=True,
        )
        thread.start()

        return jsonify({
            "success": True,
            "session_id": session_id,
            "stream_url": f"/api/thumbnails/progress/{session_id}",
            "download_url": f"/api/thumbnails/download/{session_id}",
        })
    except Exception as e:
        print(f"\n❌ Error starting test thumbnails: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route("/api/thumbnails/progress/<session_id>")
def thumbnails_progress_stream(session_id):
    """SSE stream of log/done events for a thumbnail generation session."""
    q = thumbnail_test_streams.get(session_id)
    if q is None:
        return "Session not found", 404

    def generate():
        try:
            while True:
                try:
                    payload = q.get(timeout=300)
                except _queue.Empty:
                    yield ": keepalive\n\n"
                    continue
                yield f"data: {json.dumps(payload)}\n\n"
                if payload.get("type") == "done":
                    break
        finally:
            thumbnail_test_streams.pop(session_id, None)

    return Response(generate(), mimetype="text/event-stream", headers={
        "Cache-Control": "no-cache",
        "X-Accel-Buffering": "no",
        "Connection": "keep-alive",
    })


@app.route("/api/thumbnails/download/<session_id>")
def thumbnails_download_zip(session_id):
    """Download all thumbnails generated for a session as a single ZIP."""
    info = thumbnail_test_results.get(session_id)
    if not info or not info.get("thumbnails"):
        return jsonify({"error": "Session not found or no thumbnails"}), 404

    import io
    import zipfile

    buf = io.BytesIO()
    added = 0
    with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for var in info["thumbnails"]:
            fp = var.get("filepath")
            fn = var.get("filename") or (Path(fp).name if fp else None)
            if not fp or not fn:
                continue
            try:
                p = Path(fp)
                if p.exists():
                    zf.write(p, arcname=fn)
                    added += 1
            except Exception as e:
                logger.warning(f"Skipping thumbnail {fp}: {e}")

    if added == 0:
        return jsonify({"error": "No thumbnail files found on disk"}), 404

    buf.seek(0)
    safe_topic = re.sub(r'[^\w\-]+', '_', (info.get("topic")
                        or "thumbnails")).strip('_') or "thumbnails"
    zip_name = f"{safe_topic}_thumbnails_{session_id[:8]}.zip"
    return send_file(
        buf,
        mimetype="application/zip",
        as_attachment=True,
        download_name=zip_name,
    )


@app.route("/videos")
def public_videos():
    """Public read-only Video Gallery (finished videos). No toolbar/forms."""
    html = """<!doctype html>
<html lang="en"><head>
<meta charset="utf-8" />
<meta name="viewport" content="width=device-width,initial-scale=1" />
<meta name="robots" content="noindex,nofollow" />
<title>Video Gallery</title>
<style>
  :root { color-scheme: light dark; }
  body { margin:0; font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,sans-serif; background:#0b1220; color:#e5e7eb; }
  header { padding:14px 22px; border-bottom:1px solid #1f2937; display:flex; align-items:center; gap:12px; flex-wrap:wrap; position:sticky; top:0; background:#0b1220; z-index:10; }
  header h1 { font-size:18px; margin:0; font-weight:600; }
  header .count { color:#9ca3af; font-size:13px; }
  main { padding:20px; }
  .grid { display:grid; grid-template-columns:repeat(auto-fill,minmax(320px,1fr)); gap:18px; }
  .card { background:#111827; border:1px solid #1f2937; border-radius:10px; overflow:hidden; display:flex; flex-direction:column; }
  .card video { width:100%; height:auto; display:block; aspect-ratio:16/9; object-fit:cover; background:#000; }
  .card .body { padding:10px 12px; font-size:13px; }
  .card .title { font-weight:600; color:#f3f4f6; margin-bottom:6px; line-height:1.3; word-break:break-word; }
  .card .meta { color:#9ca3af; font-size:12px; }
  .empty, .err { color:#9ca3af; padding:24px; }
  .err { color:#fca5a5; }
  footer { padding:14px 22px; color:#6b7280; font-size:12px; border-top:1px solid #1f2937; }
</style></head>
<body>
  <header>
    <h1>🎥 Video Gallery</h1>
    <span id="count" class="count">Loading…</span>
  </header>
  <main>
    <div id="status" class="empty">Loading videos…</div>
    <div id="grid" class="grid" style="display:none"></div>
  </main>
  <footer>Read-only shared gallery.</footer>
<script>
(async function(){
  const grid = document.getElementById('grid');
  const status = document.getElementById('status');
  const countEl = document.getElementById('count');
  function fmtSize(n){ if(!n) return ''; const u=['B','KB','MB','GB']; let i=0; let v=n; while(v>=1024&&i<u.length-1){v/=1024;i++;} return v.toFixed(v>=100?0:1)+' '+u[i]; }
  function fmtDate(t){ if(!t) return ''; try { return new Date(t*1000).toLocaleDateString(); } catch(e){ return ''; } }
  try {
    const r = await fetch('/api/finished-videos');
    const data = await r.json();
    if (!data.success) {
      status.className = 'err';
      status.textContent = '⚠️ ' + (data.error || 'Failed to load videos');
      countEl.textContent = '';
      return;
    }
    const videos = data.videos || [];
    countEl.textContent = videos.length + ' video' + (videos.length===1?'':'s');
    if (videos.length === 0) {
      status.textContent = 'No videos available.';
      return;
    }
    status.style.display = 'none';
    grid.style.display = '';
    for (const v of videos) {
      const card = document.createElement('div');
      card.className = 'card';
      const title = v.title || v.filename || 'Untitled';
      const meta = [fmtSize(v.size), fmtDate(v.mtime), (v.ext||'').toUpperCase()].filter(Boolean).join(' · ');
      card.innerHTML = `
        <video controls preload="metadata" playsinline src="${v.stream_url}"></video>
        <div class="body">
          <div class="title">${title}</div>
          <div class="meta">${meta}</div>
        </div>`;
      grid.appendChild(card);
    }
  } catch (e) {
    status.className = 'err';
    status.textContent = '⚠️ ' + e.message;
  }
})();
</script>
</body></html>"""
    return Response(html, mimetype="text/html")


@app.route("/thumbnails/<filename>")
def serve_thumbnail(filename):
    """Serve generated thumbnail images"""
    try:
        search_roots = [
            _get_output_parent_dir(),
            Path.home() / "Dev" / "Thumbnails",
        ]

        file_path = None
        for root in search_roots:
            if not root.exists():
                continue
            matches = list(root.rglob(filename))
            if matches:
                file_path = matches[0]
                break

        if file_path and file_path.exists() and file_path.suffix.lower() in ['.png', '.jpg', '.jpeg']:
            mimetype = 'image/png'
            if file_path.suffix.lower() in ['.jpg', '.jpeg']:
                mimetype = 'image/jpeg'
            return send_file(file_path, mimetype=mimetype)
        return jsonify({"error": "Thumbnail not found"}), 404
    except Exception as e:
        logger.error(f"Error serving thumbnail: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/broll-images/<filename>")
def serve_broll_image(filename):
    """Serve generated B-roll images"""
    try:
        search_roots = [
            _get_output_parent_dir(),
            Path.home() / "Dev" / "brollimages",
        ]

        file_path = None
        for root in search_roots:
            if not root.exists():
                continue
            matches = list(root.rglob(filename))
            if matches:
                file_path = matches[0]
                break

        if file_path and file_path.exists() and file_path.suffix.lower() in ['.png', '.jpg', '.jpeg', '.webp']:
            # Determine mimetype based on extension
            mimetype = 'image/png'
            if file_path.suffix.lower() in ['.jpg', '.jpeg']:
                mimetype = 'image/jpeg'
            elif file_path.suffix.lower() == '.webp':
                mimetype = 'image/webp'

            return send_file(file_path, mimetype=mimetype)
        return jsonify({"error": "B-roll image not found"}), 404
    except Exception as e:
        logger.error(f"Error serving B-roll image: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/broll-videos/<filename>")
def serve_broll_video(filename):
    """Serve generated Grok B-roll videos from {output_dir}/broll or legacy ~/Dev/brollvideos."""
    try:
        safe_name = Path(filename).name
        candidates = []
        base = _get_output_base_dir()
        if base is not None:
            candidates.append(base / "broll" / safe_name)
            for run_broll in base.glob("*/broll"):
                candidates.append(run_broll / safe_name)
        candidates.append(Path.home() / "Dev" / "brollvideos" / safe_name)
        for file_path in candidates:
            if file_path.exists() and file_path.suffix.lower() == '.mp4':
                return send_file(file_path, mimetype='video/mp4')
        return jsonify({"error": "B-roll video not found"}), 404
    except Exception as e:
        logger.error(f"Error serving B-roll video: {e}")
        return jsonify({"error": str(e)}), 500


def _resolve_media_file(filename: str, kind: str):
    """Resolve a media filename to an absolute path. kind: 'image' | 'video' | 'thumbnail' | 'edl'."""
    # Strip any path components for safety
    safe_name = Path(filename).name
    if not safe_name or safe_name in ('.', '..'):
        return None

    if kind == 'video':
        candidates = []
        base = _get_output_base_dir()
        if base is not None:
            candidates.append(base / "broll" / safe_name)
            for run_broll in base.glob("*/broll"):
                candidates.append(run_broll / safe_name)
        candidates.append(Path.home() / "Dev" / "brollvideos" / safe_name)
        for candidate in candidates:
            if candidate.exists() and candidate.suffix.lower() == '.mp4':
                return candidate
        return None

    if kind == 'thumbnail':
        search_roots = [
            _get_output_parent_dir(),
            Path.home() / "Dev" / "Thumbnails",
        ]
        allowed_ext = {'.png', '.jpg', '.jpeg'}
        for root in search_roots:
            if not root.exists():
                continue
            matches = list(root.rglob(safe_name))
            if matches and matches[0].suffix.lower() in allowed_ext:
                return matches[0]
        return None

    if kind == 'edl':
        base = _get_output_base_dir()
        if base is not None:
            direct = base / "MDL" / safe_name
            if direct.exists() and direct.suffix.lower() == '.edl':
                return direct
            for run_mdl in base.glob("*/MDL"):
                candidate = run_mdl / safe_name
                if candidate.exists() and candidate.suffix.lower() == '.edl':
                    return candidate

        # Legacy fallback: previous flow wrote EDLs in cwd.
        candidate = Path.cwd() / safe_name
        if candidate.exists() and candidate.suffix.lower() == '.edl':
            return candidate
        # Also search the repo root just in case.
        repo_root = Path(__file__).resolve().parent.parent
        for root in (repo_root, repo_root.parent):
            c = root / safe_name
            if c.exists() and c.suffix.lower() == '.edl':
                return c
        return None

    # images: search the same roots used by serve_broll_image
    search_roots = [
        _get_output_parent_dir(),
        Path.home() / "Dev" / "brollimages",
    ]
    allowed_ext = {'.png', '.jpg', '.jpeg', '.webp'}
    for root in search_roots:
        if not root.exists():
            continue
        matches = list(root.rglob(safe_name))
        if matches:
            match = matches[0]
            if match.suffix.lower() in allowed_ext:
                return match
    return None


@app.route("/api/download-media-zip", methods=["POST"])
def download_media_zip():
    """Bundle selected B-roll images and/or Grok videos into a ZIP for download.

    Request JSON:
      {
        "title": "Optional script title (used in zip filename)",
        "images": ["filename1.png", ...],   # optional
        "videos": ["filename1.mp4", ...]    # optional
      }
    """
    import zipfile
    import re as _re

    try:
        data = request.get_json(silent=True) or {}
        title = (data.get("title") or "media").strip() or "media"
        images = data.get("images") or []
        videos = data.get("videos") or []

        if not images and not videos:
            return jsonify({"success": False, "error": "No files requested"}), 400

        # Sanitize title for filename
        safe_title = _re.sub(r'[^A-Za-z0-9._-]+', '_',
                             title).strip('_') or "media"
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        zip_name = f"{safe_title}_media_{ts}.zip"

        buf = BytesIO()
        added = 0
        skipped = []
        with zipfile.ZipFile(buf, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
            for fn in images:
                path = _resolve_media_file(fn, 'image')
                if path is None:
                    skipped.append(fn)
                    continue
                zf.write(path, arcname=f"broll-images/{path.name}")
                added += 1
            for fn in videos:
                path = _resolve_media_file(fn, 'video')
                if path is None:
                    skipped.append(fn)
                    continue
                zf.write(path, arcname=f"grok-videos/{path.name}")
                added += 1

        if added == 0:
            return jsonify({
                "success": False,
                "error": "None of the requested files could be found on disk",
                "skipped": skipped,
            }), 404

        if skipped:
            logger.warning(
                f"⚠️ Media zip skipped {len(skipped)} missing files: {skipped[:5]}{'…' if len(skipped) > 5 else ''}")

        buf.seek(0)
        logger.info(
            f"📦 Built media zip '{zip_name}' with {added} files ({len(buf.getvalue())} bytes)")
        response = make_response(buf.getvalue())
        response.headers['Content-Type'] = 'application/zip'
        response.headers['Content-Disposition'] = f'attachment; filename="{zip_name}"'
        return response
    except Exception as e:
        logger.error(f"❌ download_media_zip failed: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/create_resolve_project", methods=["POST"])
def create_resolve_project():
    """Create a new DaVinci Resolve project with automated setup"""
    logger.info("=== DAVINCI RESOLVE PROJECT CREATION REQUEST ===")

    try:
        data = request.json
        script_title = data.get("script_title", "AI_Video_Project")
        edl_content = data.get("edl_content")
        edl_filename = data.get("edl_filename")
        generate_subtitles = bool(data.get("generate_subtitles", True))

        logger.info(f"📊 Script Title: {script_title}")
        logger.info(f"📄 EDL Filename: {edl_filename}")
        logger.info(
            f"📄 EDL Content Length: {len(edl_content) if edl_content else 0}")

        # Ensure generated Grok videos/images are in expected DaVinci project folders.
        try:
            sync_info = _sync_generated_media_to_project(
                _get_run_output_dir(script_title, create=True)
            )
            logger.info(
                f"🎞️ Pre-Resolve media sync: {sync_info['videos_copied']} video(s), "
                f"{sync_info['images_copied']} image(s) copied"
            )
        except Exception as sync_err:
            logger.warning(
                f"⚠️ Media sync before create_resolve_project failed: {sync_err}")

        # Save EDL content to temp file if provided
        edl_file_path = None
        if edl_content:
            import tempfile
            # Create temp file with .edl extension
            temp_fd, edl_file_path = tempfile.mkstemp(suffix='.edl', text=True)
            try:
                with os.fdopen(temp_fd, 'w') as f:
                    f.write(edl_content)
                logger.info(f"💾 Saved EDL to temp file: {edl_file_path}")
            except Exception as write_error:
                logger.error(f"❌ Failed to write EDL temp file: {write_error}")
                edl_file_path = None

        project_root = _get_run_output_dir(script_title, create=True)
        broll_media_path = str(project_root / "bRoll")
        images_media_path = str(project_root / "images")

        # Import the DaVinci Resolve API module
        from davinci_resolve_api import create_resolve_project as create_project

        # Create the project with EDL file path
        result = create_project(
            script_title,
            edl_file_path,
            broll_folder=broll_media_path,
            images_folder=images_media_path,
        )

        # Clean up temp file
        if edl_file_path and os.path.exists(edl_file_path):
            try:
                os.unlink(edl_file_path)
                logger.info(f"🗑️ Deleted temp EDL file: {edl_file_path}")
            except Exception as cleanup_error:
                logger.warning(
                    f"⚠️ Failed to delete temp EDL file: {cleanup_error}")

        if result.get("success"):
            logger.info(f"✅ Project created: {result.get('project_name')}")
            return jsonify(result)
        else:
            logger.error(f"❌ Project creation failed: {result.get('error')}")
            return jsonify(result), 500

    except ImportError as import_error:
        error_msg = f"Failed to import DaVinci Resolve API: {str(import_error)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500
    except Exception as e:
        error_msg = f"Unexpected error: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/check_aroll_videos", methods=["POST"])
def check_aroll_videos():
    """Check if aRoll folder has any videos"""
    logger.info("=== CHECKING AROLL VIDEOS ===")

    def _is_video_file(filename: str) -> bool:
        return filename.lower().endswith((".mp4", ".mov", ".m4v"))

    def _get_default_aroll_source():
        """Return (folder_path, sorted_video_files) for default A-roll test clips.

        Search order:
        1. SCRIPTCRAFT_DEFAULT_AROLL_DIR env var (if set)
        2. ~/Dev/Davinci/Template/aRoll
        3. ~/Dev/Davinci/Template/Raw
        """
        candidate_dirs = []

        env_dir = (os.environ.get(
            "SCRIPTCRAFT_DEFAULT_AROLL_DIR") or "").strip()
        if env_dir:
            candidate_dirs.append(os.path.expanduser(env_dir))

        candidate_dirs.append(os.path.expanduser(
            "~/Dev/Davinci/Template/aRoll"))
        candidate_dirs.append(os.path.expanduser("~/Dev/Davinci/Template/Raw"))

        for folder in candidate_dirs:
            if not os.path.isdir(folder):
                continue
            files = sorted(
                [f for f in os.listdir(folder) if _is_video_file(f)])
            if files:
                return folder, files

        return None, []

    try:
        data = request.json
        script_title = data.get("script_title", "")

        if not script_title:
            return jsonify({
                "success": False,
                "error": "No script title provided"
            }), 400

        project_path = _get_run_output_dir(script_title, create=False)
        aroll_path = str(project_path / "aRoll")

        logger.info(f"📁 Script title: '{script_title}'")
        logger.info(f"📁 Run folder: '{project_path}'")
        logger.info(f"📁 Checking aroll path: {aroll_path}")
        logger.info(f"📁 Path exists: {os.path.exists(aroll_path)}")

        if not os.path.exists(aroll_path):
            # List what's actually in the output parent path to help debug
            base_path = _get_output_parent_dir()
            if os.path.exists(base_path):
                logger.info(f"📁 Base path exists, folders found:")
                for folder in os.listdir(base_path):
                    logger.info(f"   - {folder}")
            else:
                logger.info(f"📁 Base path does not exist: {base_path}")

            logger.info(f"📁 aRoll folder does not exist yet")

            default_source, default_videos = _get_default_aroll_source()
            return jsonify({
                "success": True,
                "video_count": 0,
                "videos": [],
                "default_available": len(default_videos) > 0,
                "default_video_count": len(default_videos),
                "default_source": default_source,
                "default_videos": default_videos,
            })

        # Count local project video files in aRoll folder
        video_files = [f for f in os.listdir(aroll_path) if _is_video_file(f)]

        logger.info(f"✅ Found {len(video_files)} video(s) in aRoll folder:")
        for vf in video_files:
            logger.info(f"   - {vf}")

        default_source, default_videos = _get_default_aroll_source()

        return jsonify({
            "success": True,
            "video_count": len(video_files),
            "videos": video_files,
            "default_available": len(default_videos) > 0,
            "default_video_count": len(default_videos),
            "default_source": default_source,
            "default_videos": default_videos,
        })

    except Exception as e:
        error_msg = f"Error checking aRoll videos: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/create_resolve_with_videos", methods=["POST"])
def create_resolve_with_videos():
    """Create a DaVinci Resolve project with aRoll videos added to timeline"""
    logger.info("=== DAVINCI RESOLVE PROJECT CREATION WITH VIDEOS ===")

    try:
        data = request.json
        script_title = data.get("script_title", "AI_Video_Project")
        edl_content = data.get("edl_content")
        edl_filename = data.get("edl_filename")
        generate_subtitles = bool(data.get("generate_subtitles", True))

        logger.info(f"📊 Script Title: {script_title}")

        # Ensure generated Grok videos/images are in expected DaVinci project folders.
        try:
            sync_info = _sync_generated_media_to_project(
                _get_run_output_dir(script_title, create=True)
            )
            logger.info(
                f"🎞️ Pre-Resolve media sync: {sync_info['videos_copied']} video(s), "
                f"{sync_info['images_copied']} image(s) copied"
            )
        except Exception as sync_err:
            logger.warning(
                f"⚠️ Media sync before create_resolve_with_videos failed: {sync_err}")

        def _is_video_file(filename: str) -> bool:
            return filename.lower().endswith((".mp4", ".mov", ".m4v"))

        def _get_default_aroll_source():
            candidate_dirs = []

            env_dir = (os.environ.get(
                "SCRIPTCRAFT_DEFAULT_AROLL_DIR") or "").strip()
            if env_dir:
                candidate_dirs.append(os.path.expanduser(env_dir))

            candidate_dirs.append(os.path.expanduser(
                "~/Dev/Davinci/Template/aRoll"))
            candidate_dirs.append(os.path.expanduser(
                "~/Dev/Davinci/Template/Raw"))

            for folder in candidate_dirs:
                if not os.path.isdir(folder):
                    continue
                files = sorted(
                    [f for f in os.listdir(folder) if _is_video_file(f)])
                if files:
                    return folder, files

            return None, []

        use_default_aroll = bool(data.get("use_default_aroll", False))

        project_path = _get_run_output_dir(script_title, create=False)
        aroll_path = str(project_path / "aRoll")

        # Get list of video files
        video_files = []
        if os.path.exists(aroll_path):
            video_files = [f for f in os.listdir(
                aroll_path) if _is_video_file(f)]
            logger.info(f"📹 Found {len(video_files)} video(s) in aroll folder")

        # Optional fallback for testing DaVinci flow without downloaded HeyGen aRoll.
        used_default_aroll = False
        default_source = None
        if use_default_aroll and len(video_files) == 0:
            default_source, default_videos = _get_default_aroll_source()
            if default_source and default_videos:
                logger.info(
                    f"🧪 Using default A-roll test clips from: {default_source} "
                    f"({len(default_videos)} videos)"
                )
                aroll_path = default_source
                video_files = default_videos
                used_default_aroll = True
            else:
                logger.warning(
                    "⚠️ use_default_aroll requested, but no default clips were found"
                )

        # Sort videos by chapter/part order
        def parse_chapter_info(filename):
            """Extract chapter and part numbers from filename.

            Matches any filename containing Ch{N} and p{N}, e.g.:
              Ch1p1, Ch1p2, Ch6p1         (no separator)
              Ch1-Pt1, Ch1-pt1, Ch2-p2   (dash + optional 't')
              Ch1_Pt1                     (underscore)
              Title-Ch3p2.mp4            (title prefix, any separator)
              heygen_...-Ch1p1_id.mp4    (heygen with ID suffix)
              Ch1p1b                     (b-duplicate sorts after Ch1p1)

            'AI with Roz' intro files always sort first.
            Unknown files sort to the end.
            """
            import re
            name = filename

            # Exit clip ("AI with Roz Exit Clip.mov") always goes LAST,
            # after every chapter clip and any unknown clips.
            if re.search(r'AI\s+with\s+Roz.*Exit', name, re.IGNORECASE):
                return (99999, 0)

            # Other "AI with Roz" intro clips always go first
            if re.search(r'AI\s+with\s+Roz', name, re.IGNORECASE):
                return (0, 0)

            # Universal pattern: Ch{X} then optional separator then p/P then optional t/T then {Y} then optional b
            # Covers: Ch1p1, Ch1-p1, Ch1-Pt1, Ch1-pt1, Ch1_Pt1, Ch6p1, Ch2-p2, etc.
            match = re.search(r'Ch(\d+)[-_]?[Pp][Tt]?(\d+)(b?)', name)
            if match:
                chapter_num = int(match.group(1))
                part_num = int(match.group(2))
                # 'b' duplicate sorts just after the main part
                part_frac = 0.5 if match.group(3).lower() == 'b' else 0
                return (chapter_num, part_num + part_frac)

            # Default to end of list
            return (999, 999)

        # Sort videos by chapter order
        sorted_videos = sorted(video_files, key=parse_chapter_info)
        logger.info(f"📋 Sorted video order: {sorted_videos}")

        # Save EDL content to temp file if provided
        edl_file_path = None
        if edl_content:
            import tempfile
            temp_fd, edl_file_path = tempfile.mkstemp(suffix='.edl', text=True)
            try:
                with os.fdopen(temp_fd, 'w') as f:
                    f.write(edl_content)
                logger.info(f"💾 Saved EDL to temp file: {edl_file_path}")
            except Exception as write_error:
                logger.error(f"❌ Failed to write EDL temp file: {write_error}")
                edl_file_path = None

        # Import the DaVinci Resolve API module (force reload to get latest changes)
        import importlib
        import sys
        if 'davinci_resolve_api' in sys.modules:
            logger.info("🔄 Reloading davinci_resolve_api module...")
            import davinci_resolve_api
            importlib.reload(davinci_resolve_api)
            from davinci_resolve_api import create_resolve_project_with_videos
            logger.info("✅ Module reloaded successfully")
        else:
            logger.info(
                "📦 Loading davinci_resolve_api module for first time...")
            from davinci_resolve_api import create_resolve_project_with_videos

        logger.info(f"🎬 Calling create_resolve_project_with_videos...")
        logger.info(f"   Script title: {script_title}")
        logger.info(f"   EDL file: {edl_file_path}")
        logger.info(f"   aRoll path: {aroll_path}")
        logger.info(f"   Videos to add: {len(sorted_videos)}")

        broll_media_path = str(project_path / "bRoll")
        images_media_path = str(project_path / "images")

        # Create the project with videos
        result = create_resolve_project_with_videos(
            script_title,
            edl_file_path,
            aroll_path,
            sorted_videos,
            broll_media_path,
            images_media_path,
            generate_subtitles,
        )

        if isinstance(result, dict):
            result["used_default_aroll"] = used_default_aroll
            result["aroll_source"] = aroll_path
            if used_default_aroll:
                result["default_aroll_source"] = default_source

        # Clean up temp file
        if edl_file_path and os.path.exists(edl_file_path):
            try:
                os.unlink(edl_file_path)
                logger.info(f"🗑️ Deleted temp EDL file: {edl_file_path}")
            except Exception as cleanup_error:
                logger.warning(
                    f"⚠️ Failed to delete temp EDL file: {cleanup_error}")

        if result.get("success"):
            logger.info(
                f"✅ Project created with {len(sorted_videos)} videos: {result.get('project_name')}")
            return jsonify(result)
        else:
            logger.error(f"❌ Project creation failed: {result.get('error')}")
            return jsonify(result), 500

    except ImportError as import_error:
        error_msg = f"Failed to import DaVinci Resolve API: {str(import_error)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500
    except Exception as e:
        error_msg = f"Unexpected error: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/execute_curl", methods=["POST"])
def execute_curl():
    """Execute a curl command and return the result"""
    logger.info("=== CURL COMMAND EXECUTION REQUEST ===")

    try:
        data = request.json
        curl_command = data.get("curl_command", "")

        if not curl_command:
            return jsonify({
                "success": False,
                "error": "No curl command provided"
            }), 400

        logger.info(
            f"📤 Executing curl command (first 100 chars): {curl_command[:100]}...")

        # Execute the curl command using subprocess
        import subprocess
        import shlex

        # Parse the curl command
        # Remove leading/trailing whitespace and newlines
        curl_command = curl_command.strip()

        # Execute the command
        try:
            result = subprocess.run(
                curl_command,
                shell=True,
                capture_output=True,
                text=True,
                timeout=30
            )

            # Parse the response
            if result.returncode == 0:
                # Try to parse JSON response
                try:
                    import json
                    response_data = json.loads(result.stdout)

                    # Extract job ID if available (HeyGen specific)
                    job_id = None
                    if isinstance(response_data, dict):
                        data_obj = response_data.get('data') or {}
                        job_id = (data_obj.get('video_id') if isinstance(data_obj, dict) else None) or \
                            response_data.get('video_id') or \
                            response_data.get('job_id') or \
                            response_data.get('id')

                    logger.info(f"✅ Curl command executed successfully")
                    if job_id:
                        logger.info(f"📋 Job ID: {job_id}")

                    return jsonify({
                        "success": True,
                        "response": response_data,
                        "job_id": job_id,
                        "raw_output": result.stdout
                    })
                except json.JSONDecodeError:
                    # Not JSON, return raw output
                    logger.info(f"✅ Curl command executed (non-JSON response)")
                    return jsonify({
                        "success": True,
                        "raw_output": result.stdout
                    })
            else:
                error_msg = result.stderr or "Command failed with no error message"
                logger.error(f"❌ Curl command failed: {error_msg}")
                return jsonify({
                    "success": False,
                    "error": error_msg,
                    "return_code": result.returncode
                }), 500

        except subprocess.TimeoutExpired:
            logger.error("❌ Curl command timed out after 30 seconds")
            return jsonify({
                "success": False,
                "error": "Command timed out after 30 seconds"
            }), 500

    except Exception as e:
        error_msg = f"Error executing curl command: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/setup_project", methods=["POST"])
def setup_project():
    """Set up project folder structure and copy template files"""
    logger.info("=== PROJECT SETUP REQUEST ===")

    try:
        data = request.json
        script_title = data.get("script_title", "")
        script_content = data.get("script_content", "")

        if not script_title:
            return jsonify({
                "success": False,
                "error": "No script title provided"
            }), 400

        import os
        import shutil
        from pathlib import Path

        # Define paths
        safe_title = _safe_title_for_paths(script_title)
        project_path = _get_run_output_dir(script_title, create=True)
        template_path = Path.home() / "Dev" / "Davinci" / "Template"

        logger.info(f"📁 Creating project folder: {project_path}")

        # Create project directory if it doesn't exist
        # IMPORTANT: Do NOT delete existing project - preserve downloaded videos!
        if project_path.exists():
            logger.info(f"✅ Project folder already exists: {project_path}")
            logger.info(
                "   Preserving existing aRoll videos and other content")
        else:
            project_path.mkdir(parents=True, exist_ok=True)
            logger.info(f"✅ Created new project folder")

        # Ensure subfolders exist (without deleting existing content)
        aroll_path = project_path / "aRoll"
        aroll_path.mkdir(exist_ok=True)
        broll_path = project_path / "bRoll"
        broll_path.mkdir(exist_ok=True)
        images_path = project_path / "images"
        images_path.mkdir(exist_ok=True)
        script_path = project_path / "script"
        script_path.mkdir(exist_ok=True)

        # Copy template files ONLY if project is new (no aroll videos exist)
        aroll_videos = list(aroll_path.glob("*.mp4"))
        if template_path.exists() and len(aroll_videos) == 0:
            logger.info(f"📋 Copying template from: {template_path}")
            file_count = sum(
                1 for _ in template_path.rglob('*') if _.is_file())
            logger.info(f"📦 Found {file_count} files to copy...")

            copied_count = 0
            for item in template_path.iterdir():
                # Never copy the heygen folder into the project.
                if item.name.lower() == "heygen":
                    logger.info(
                        f"⏭️ Skipping heygen folder (excluded from project copy)")
                    continue
                dest = project_path / item.name
                if item.is_dir():
                    # Count files in directory
                    dir_files = sum(1 for _ in item.rglob('*') if _.is_file())
                    logger.info(
                        f"📁 Copying directory: {item.name} ({dir_files} files)..."
                    )
                    shutil.copytree(item, dest, dirs_exist_ok=True)
                    copied_count += dir_files
                else:
                    shutil.copy2(item, dest)
                    copied_count += 1

                # Log progress every 100 files
                if copied_count % 100 == 0:
                    logger.info(
                        f"📊 Progress: {copied_count}/{file_count} files copied")

            logger.info(f"✅ Template files copied ({copied_count} total)")
        elif len(aroll_videos) > 0:
            logger.info(
                f"ℹ️ Skipping template copy - {len(aroll_videos)} "
                f"aRoll videos already exist"
            )
        else:
            logger.warning(
                f"⚠️ Template path not found: {template_path}"
            )

        # Create script subdirectory
        script_dir = project_path / "script"
        script_dir.mkdir(exist_ok=True)

        # Sync generated media into DaVinci project structure.
        # - Grok videos: broll -> bRoll
        # - Generated images already produced in project/images are kept there
        # - Legacy fallback copies from ~/Dev/brollimages when needed
        media_sync = _sync_generated_media_to_project(project_path)
        logger.info(
            f"🎞️ Media sync complete: {media_sync['videos_copied']} video(s) -> bRoll, "
            f"{media_sync['images_copied']} image(s) copied"
        )

        # Save script as Word document if content provided
        if script_content:
            import sys
            import asyncio

            # Import word processing from console_ui
            sys.path.insert(
                0, str(Path(__file__).parent.parent / "console_ui")
            )
            from word_processing import convert_markdown_to_word

            script_file = script_dir / f"{safe_title}.docx"
            logger.info(f"📝 Creating script document: {script_file}")

            # Convert and save (handle async function)
            try:
                asyncio.run(convert_markdown_to_word(
                    markdown_content=script_content,
                    output_file_path=str(script_file),
                    template_path=None,
                    title=script_title
                ))
                logger.info(f"✅ Script saved: {script_file}")
            except Exception as word_error:
                logger.error(f"❌ Word document error: {word_error}")
                # Continue anyway - project still set up

        return jsonify({
            "success": True,
            "project_path": str(project_path),
            "script_dir": str(script_dir),
            "media_sync": media_sync,
        })

    except Exception as e:
        error_msg = f"Error setting up project: {str(e)}"
        logger.error(f"❌ {error_msg}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/check_video_status", methods=["POST"])
def check_video_status():
    """Check HeyGen video generation status and download when complete"""
    logger.info("=== VIDEO STATUS CHECK REQUEST ===")

    try:
        data = request.json
        video_id = data.get("video_id", "")
        api_key = data.get("api_key", "")

        if not video_id:
            return jsonify({
                "success": False,
                "error": "No video ID provided"
            }), 400

        if not api_key:
            return jsonify({
                "success": False,
                "error": "No API key provided"
            }), 400

        logger.info(f"📹 Checking status for video ID: {video_id}")

        import requests

        # Check video status
        base_url = "https://api.heygen.com/v1/video_status.get"
        video_status_url = f"{base_url}?video_id={video_id}"
        headers = {
            "X-Api-Key": api_key
        }

        try:
            response = requests.get(
                video_status_url, headers=headers, timeout=10
            )
            response_data = response.json()

            if response.status_code != 200:
                error_msg = response_data.get("message", "Unknown error")
                logger.error(f"❌ HeyGen API error: {error_msg}")
                return jsonify({
                    "success": False,
                    "error": f"HeyGen API error: {error_msg}"
                }), 500

            status = response_data.get("data", {}).get("status", "unknown")
            logger.info(f"📊 Video status: {status}")

            result = {
                "success": True,
                "status": status,
                "video_id": video_id
            }

            if status == "completed":
                video_url = response_data.get("data", {}).get("video_url")
                thumb_url = response_data.get("data", {}).get(
                    "thumbnail_url"
                )

                if video_url:
                    result["video_url"] = video_url
                    result["thumbnail_url"] = thumb_url
                    logger.info(f"✅ Video completed! URL: {video_url}")
                else:
                    msg = "⚠️ Video marked as completed but no URL"
                    logger.warning(msg)

            elif status == "failed":
                error = response_data.get("data", {}).get(
                    "error", "Unknown error"
                )
                result["error"] = error
                logger.error(f"❌ Video generation failed: {error}")

            return jsonify(result)

        except requests.RequestException as req_error:
            error_msg = f"Network error: {str(req_error)}"
            logger.error(f"❌ {error_msg}")
            return jsonify({
                "success": False,
                "error": error_msg
            }), 500

    except Exception as e:
        error_msg = f"Error checking video status: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/download_heygen_video", methods=["POST"])
def download_heygen_video():
    """Download completed HeyGen video to project folder"""
    logger.info("=== VIDEO DOWNLOAD REQUEST ===")

    try:
        data = request.json
        video_url = data.get("video_url", "")
        video_id = data.get("video_id", "")
        chapter_name = data.get("chapter_name", "video")
        script_title = data.get("script_title", "")

        if not video_url:
            return jsonify({
                "success": False,
                "error": "No video URL provided"
            }), 400

        logger.info(f"⬇️ Downloading video: {video_id}")

        import requests
        from pathlib import Path

        try:
            # Download video content
            video_response = requests.get(video_url, timeout=60)

            if video_response.status_code != 200:
                status_code = video_response.status_code
                logger.error(f"❌ Failed to download video: {status_code}")
                return jsonify({
                    "success": False,
                    "error": f"Failed to download video: {status_code}"
                }), 500

            # Create safe names
            safe_chars = (' ', '-', '_')
            safe_chapter_name = "".join(
                c for c in chapter_name
                if c.isalnum() or c in safe_chars
            ).rstrip()

            safe_title = "".join(
                c for c in script_title
                if c.isalnum() or c in safe_chars
            ).rstrip().replace(' ', '_') if script_title else "default"

            filename = f"{safe_chapter_name}_{video_id[:8]}.mp4"

            # Save to project aroll folder if script title provided
            if script_title:
                project_path = _get_run_output_dir(script_title, create=True)
                aroll_path = project_path / "aRoll"

                # Create project folders if they don't exist
                aroll_path.mkdir(parents=True, exist_ok=True)
                broll_path = project_path / "bRoll"
                broll_path.mkdir(parents=True, exist_ok=True)
                images_path = project_path / "images"
                images_path.mkdir(parents=True, exist_ok=True)

                # Save video file
                video_file = aroll_path / filename
                with open(video_file, 'wb') as f:
                    f.write(video_response.content)

                logger.info(
                    f"✅ Video saved to project: {video_file}"
                )

                return jsonify({
                    "success": True,
                    "filename": filename,
                    "path": str(video_file)
                })
            else:
                # No script title - return as download
                from io import BytesIO
                video_buffer = BytesIO(video_response.content)
                video_buffer.seek(0)

                logger.info(
                    f"✅ Video downloaded successfully: {filename}"
                )

                return send_file(
                    video_buffer,
                    mimetype='video/mp4',
                    as_attachment=True,
                    download_name=filename
                )

        except requests.RequestException as req_error:
            error_msg = f"Network error: {str(req_error)}"
            logger.error(f"❌ {error_msg}")
            return jsonify({
                "success": False,
                "error": error_msg
            }), 500

    except Exception as e:
        error_msg = f"Error downloading video: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({
            "success": False,
            "error": error_msg
        }), 500


@app.route("/test-comparisons")
def test_comparisons():
    """Test endpoint to preview the chapter comparisons tab without running full workflow"""

    # Mock chapter comparison data
    mock_comparisons = [
        {
            'chapter_num': 1,
            'original': '''Welcome to our channel! Today we're diving into an amazing topic that will change how you think about productivity.

Let's start with the basics and understand why this matters.''',
            'revised': '''Welcome to our channel! Today we're exploring a game-changing approach to productivity that actually works.

Let's dive into the fundamentals and discover why this is so important for your success.''',
            'feedback': '''Reviewer Feedback:
- Opening hook needs more impact
- Changed "amazing topic" to "game-changing approach" for stronger language
- Adjusted second paragraph for better flow'''
        },
        {
            'chapter_num': 2,
            'original': '''The first key principle is understanding your goals. You need to know what you want to achieve.

Without clear goals, you'll struggle to make progress.''',
            'revised': '''The first crucial principle is crystal-clear goal setting. You must define exactly what success looks like.

Without specific, measurable goals, you're just wandering without direction.''',
            'feedback': '''Reviewer Feedback:
- Made language more specific and actionable
- Changed "understanding" to "crystal-clear goal setting"
- Strengthened the consequence statement'''
        },
        {
            'chapter_num': 3,
            'original': '''Now let's talk about time management. It's important to schedule your day properly.

Make sure you allocate time for important tasks.''',
            'revised': '''Here's where time management becomes your superpower. Block out your day strategically, not randomly.

Protect prime hours for your most critical, high-impact work.''',
            'feedback': '''Reviewer Feedback:
- Added energy to the opening with "superpower"
- Changed passive "schedule" to active "block out"
- Made the advice more specific and actionable'''
        }
    ]

    # Mock thumbnails for testing
    mock_thumbnails = [
        {
            'id': 'test-thumb-1',
            'url': 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMTI4MCIgaGVpZ2h0PSI3MjAiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+PHJlY3Qgd2lkdGg9IjEyODAiIGhlaWdodD0iNzIwIiBmaWxsPSIjMWU0MGFmIi8+PHRleHQgeD0iNTAlIiB5PSI1MCUiIGZvbnQtZmFtaWx5PSJBcmlhbCIgZm9udC1zaXplPSI4MCIgZmlsbD0id2hpdGUiIHRleHQtYW5jaG9yPSJtaWRkbGUiIGR5PSIuM2VtIj5UZXN0IFRodW1ibmFpbCAxPC90ZXh0Pjwvc3ZnPg==',
            'style': 'bold'
        }
    ]

    # Render the main template
    return render_template('index.html',
                           test_mode=True,
                           test_comparisons=mock_comparisons,
                           test_thumbnails=mock_thumbnails)


# ── DaVinci Resolve Audio Render ──────────────────────────────────────────────

def _resolve_env_setup():
    """Set RESOLVE_SCRIPT_API/LIB env vars and ensure Modules path is in sys.path."""
    resolve_script_api = (
        "/Library/Application Support/Blackmagic Design/"
        "DaVinci Resolve/Developer/Scripting"
    )
    os.environ["RESOLVE_SCRIPT_API"] = resolve_script_api
    os.environ["RESOLVE_SCRIPT_LIB"] = (
        "/Applications/DaVinci Resolve/DaVinci Resolve.app/"
        "Contents/Libraries/Fusion/fusionscript.so"
    )
    modules_path = os.path.join(resolve_script_api, "Modules")
    if modules_path not in sys.path:
        sys.path.append(modules_path)


@app.route("/api/resolve/list-timelines", methods=["GET"])
def api_resolve_list_timelines():
    """Return the current DaVinci Resolve project and its timelines."""
    try:
        _resolve_env_setup()
        import DaVinciResolveScript as dvr_script  # type: ignore
        resolve = dvr_script.scriptapp("Resolve")
        if resolve is None:
            return jsonify({"success": False, "error": "DaVinci Resolve is not running or not reachable."}), 500

        pm = resolve.GetProjectManager()
        project = pm.GetCurrentProject()
        if project is None:
            return jsonify({"success": False, "error": "No project is open in DaVinci Resolve."}), 400

        project_name = project.GetName()
        timelines = []
        for i in range(1, project.GetTimelineCount() + 1):
            tl = project.GetTimelineByIndex(i)
            if tl:
                timelines.append(tl.GetName())

        current_tl = project.GetCurrentTimeline()
        current_tl_name = current_tl.GetName() if current_tl else None

        return jsonify({
            "success": True,
            "project": project_name,
            "timelines": timelines,
            "current_timeline": current_tl_name,
        })
    except Exception as e:
        logger.error(
            f"❌ /api/resolve/list-timelines failed: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


def _audio_render_worker(session_id: str, timeline_name: str, output_path: str):
    """Background worker that runs the Resolve audio export and posts progress to the session queue."""
    q = audio_render_streams.get(session_id)
    if q is None:
        return

    def push(msg, progress=None, status="progress", **extra):
        payload = {"type": status, "message": msg}
        if progress is not None:
            payload["progress"] = progress
        if extra:
            payload.update(extra)
        q.put(payload)

    try:
        push(f"🔌 Connecting to DaVinci Resolve…", 5)
        _resolve_env_setup()
        import DaVinciResolveScript as dvr_script  # type: ignore
        resolve = dvr_script.scriptapp("Resolve")
        if resolve is None:
            push("❌ DaVinci Resolve is not running or not reachable.", status="error")
            audio_render_results[session_id] = {
                "success": False, "error": "Resolve not running"}
            return

        pm = resolve.GetProjectManager()
        project = pm.GetCurrentProject()
        if project is None:
            push("❌ No project is open in DaVinci Resolve.", status="error")
            audio_render_results[session_id] = {
                "success": False, "error": "No project open"}
            return

        push(f"🎬 Switching to timeline '{timeline_name}'…", 10)

        # Find and set the timeline
        tl = None
        for i in range(1, project.GetTimelineCount() + 1):
            t = project.GetTimelineByIndex(i)
            if t and t.GetName() == timeline_name:
                tl = t
                break
        if tl is None:
            push(f"❌ Timeline '{timeline_name}' not found.", status="error")
            audio_render_results[session_id] = {
                "success": False, "error": f"Timeline '{timeline_name}' not found"}
            return
        project.SetCurrentTimeline(tl)

        push(f"⚙️ Configuring render settings (WAV / 48 kHz / 24-bit)…", 20)
        output_dir = str(Path(output_path).parent)
        # Strip extension from CustomName — Resolve appends it based on format
        output_filename = Path(output_path).stem
        Path(output_dir).mkdir(parents=True, exist_ok=True)

        # Try the "Audio Only" preset first (most reliable across Resolve versions)
        try:
            presets = project.GetRenderPresetList() or []
        except Exception:
            presets = []
        audio_preset = next(
            (p for p in presets if isinstance(p, str)
             and p.lower() in ("audio only", "audio_only")),
            None,
        )
        if audio_preset:
            try:
                project.LoadRenderPreset(audio_preset)
                push(f"   • Loaded render preset: {audio_preset}", 22)
            except Exception:
                pass

        # Set format+codec BEFORE SetRenderSettings (codec keys depend on format)
        try:
            project.SetCurrentRenderFormatAndCodec("wav", "LinearPCM")
        except Exception as _e:
            push(f"   • Note: SetCurrentRenderFormatAndCodec raised: {_e}", 22)

        # Minimal settings dict — Resolve rejects the entire dict if any key is
        # unknown for the current format. Audio-specific keys are already set
        # by the format/codec selection above.
        settings = {
            "SelectAllFrames": True,
            "TargetDir": output_dir,
            "CustomName": output_filename,
            "ExportVideo": False,
            "ExportAudio": True,
        }
        if not project.SetRenderSettings(settings):
            # Fallback: try without ExportVideo/ExportAudio (some Resolve versions reject these)
            push(
                "   • Initial SetRenderSettings rejected — retrying with minimal keys…", 23)
            minimal = {
                "SelectAllFrames": True,
                "TargetDir": output_dir,
                "CustomName": output_filename,
            }
            if not project.SetRenderSettings(minimal):
                push(
                    "❌ SetRenderSettings failed. Open Resolve → Deliver page, choose 'Audio Only' "
                    "preset manually once, then retry. Also verify the output directory exists and "
                    "is writable.",
                    status="error",
                )
                audio_render_results[session_id] = {
                    "success": False, "error": "SetRenderSettings failed"}
                return

        push(f"📋 Queueing render job…", 30)
        job_id = project.AddRenderJob()
        if not job_id:
            push(
                "❌ AddRenderJob failed. Check the Deliver page in Resolve.", status="error")
            audio_render_results[session_id] = {
                "success": False, "error": "AddRenderJob failed"}
            return

        if not project.StartRendering(job_id):
            project.DeleteRenderJobList()
            push("❌ StartRendering failed.", status="error")
            audio_render_results[session_id] = {
                "success": False, "error": "StartRendering failed"}
            return

        push(f"🎵 Render started — polling for completion…", 35)
        import time as _time
        deadline = _time.time() + 1800  # 30 min max
        job_status = ""
        while _time.time() < deadline:
            _time.sleep(3)
            status_info = project.GetRenderJobStatus(job_id) or {}
            job_status = status_info.get("JobStatus", "")
            completion_pct = status_info.get("CompletionPercentage", 0)
            sse_progress = 35 + int(completion_pct *
                                    0.6)  # maps 0-100% → 35-95%
            push(
                f"⏳ Rendering… {completion_pct:.0f}% ({job_status})", sse_progress)
            if job_status in ("Complete", "Cancelled", "Failed"):
                break

        if job_status != "Complete":
            push(f"❌ Render ended with status '{job_status}'.", status="error")
            audio_render_results[session_id] = {
                "success": False, "error": f"Render status: {job_status}"}
            return

        # Locate output file (Resolve appends extension; output_filename is stem only)
        candidates = sorted(Path(output_dir).glob(f"{output_filename}*.wav")) \
            or sorted(Path(output_dir).glob(f"{output_filename}*"))
        actual_file = str(
            candidates[0]) if candidates else f"{Path(output_dir) / output_filename}.wav"
        push(f"✅ Audio exported to: {actual_file}",
             100, status="done", output_file=actual_file)
        audio_render_results[session_id] = {
            "success": True, "output_file": actual_file}

    except Exception as e:
        logger.error(f"❌ _audio_render_worker failed: {e}", exc_info=True)
        push(f"❌ Unexpected error: {e}", status="error")
        audio_render_results[session_id] = {"success": False, "error": str(e)}


@app.route("/api/resolve/render-audio", methods=["POST"])
def api_resolve_render_audio_start():
    """Start an audio-only render job in DaVinci Resolve.

    Body JSON:
        {
            "timeline_name": "My Timeline",   (required)
            "script_title":  "My Script",     (used for output filename)
        }

    Returns:
        {"success": true, "session_id": "...", "stream_url": "..."}
    """
    try:
        data = request.get_json(silent=True) or {}
        timeline_name = (data.get("timeline_name") or "").strip()
        script_title = (data.get("script_title")
                        or "audio_export").strip() or "audio_export"

        if not timeline_name:
            return jsonify({"success": False, "error": "timeline_name is required"}), 400

        # Build output path: <output_dir>/<safe_title> (no extension; Resolve adds it)
        safe_title = re.sub(r'[^A-Za-z0-9._-]+', '_',
                            script_title).strip('_') or "audio"
        output_parent = _get_output_parent_dir()
        output_path = str(output_parent / safe_title)

        session_id = str(uuid.uuid4())
        audio_render_streams[session_id] = _queue.Queue()
        audio_render_results.pop(session_id, None)

        thread = _threading.Thread(
            target=_audio_render_worker,
            args=(session_id, timeline_name, output_path),
            daemon=True,
        )
        thread.start()

        logger.info(
            f"🎵 Audio render started: session={session_id} timeline='{timeline_name}' output='{output_path}'")
        return jsonify({
            "success": True,
            "session_id": session_id,
            "output_path": output_path,
            "stream_url": f"/api/resolve/render-audio/progress/{session_id}",
        })
    except Exception as e:
        logger.error(f"❌ /api/resolve/render-audio failed to start: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/resolve/render-audio/progress/<session_id>")
def api_resolve_render_audio_progress(session_id):
    """SSE stream of progress events for an in-flight audio render session."""
    q = audio_render_streams.get(session_id)
    if q is None:
        return "Session not found", 404

    def generate():
        try:
            while True:
                try:
                    payload = q.get(timeout=60)
                except _queue.Empty:
                    yield ": keepalive\n\n"
                    continue
                yield f"data: {json.dumps(payload)}\n\n"
                if payload.get("type") in ("done", "error"):
                    break
        finally:
            audio_render_streams.pop(session_id, None)

    return Response(generate(), mimetype="text/event-stream", headers={
        "Cache-Control": "no-cache",
        "X-Accel-Buffering": "no",
        "Connection": "keep-alive",
    })


# ---------------------------------------------------------------------------
# Whisper transcription
# ---------------------------------------------------------------------------

# Local openai-whisper CLI (preferred). Falls back to OpenAI Whisper API if
# the CLI is not installed.
_WHISPER_CLI_CANDIDATES = [
    "/Library/Frameworks/Python.framework/Versions/3.14/bin/whisper",
    "/opt/homebrew/bin/whisper",
    "whisper",
]


def _find_whisper_cli():
    import shutil
    for c in _WHISPER_CLI_CANDIDATES:
        path = shutil.which(c) or (c if Path(c).is_file() else None)
        if path:
            return path
    return None


@app.route("/api/transcribe-audio/upload", methods=["POST"])
def api_transcribe_audio_upload():
    """Accept a browser-uploaded audio/video file and start transcription.

    multipart/form-data:
        file:     the audio/video file (required)
        model:    whisper model name (optional, default 'medium')
        language: ISO language code (optional)

    Returns: { success, session_id, stream_url }
    Stream events identical to /api/transcribe-audio/progress/<session_id>.
    """
    import tempfile
    import uuid
    import shutil as _shutil
    try:
        if "file" not in request.files:
            return jsonify({"success": False, "error": "No file uploaded (expected form field 'file')"}), 400
        f = request.files["file"]
        if not f or not f.filename:
            return jsonify({"success": False, "error": "Empty file upload"}), 400

        model = (request.form.get("model") or "medium").strip() or "medium"
        language = (request.form.get("language") or "").strip()

        whisper_bin = _find_whisper_cli()
        if not whisper_bin:
            return jsonify({
                "success": False,
                "error": "whisper CLI not found. Install with: pip install openai-whisper",
            }), 500

        # Save upload to a per-session temp dir; cleaned up by the worker’s rmtree of out_dir
        # is separate — we keep the input file in its own temp dir and let the OS clean tmp.
        upload_dir = Path(tempfile.mkdtemp(prefix="whisper_in_"))
        # Preserve a safe filename + original extension
        safe_name = Path(f.filename).name.replace("/", "_").replace("\\", "_")
        if not safe_name:
            safe_name = "upload.bin"
        src = upload_dir / safe_name
        f.save(str(src))

        if src.stat().st_size == 0:
            _shutil.rmtree(upload_dir, ignore_errors=True)
            return jsonify({"success": False, "error": "Uploaded file is empty"}), 400

        session_id = str(uuid.uuid4())
        transcribe_streams[session_id] = _queue.Queue()
        transcribe_results.pop(session_id, None)

        def _worker_then_cleanup():
            try:
                _transcribe_worker(session_id, whisper_bin,
                                   str(src), model, language)
            finally:
                _shutil.rmtree(upload_dir, ignore_errors=True)

        thread = _threading.Thread(target=_worker_then_cleanup, daemon=True)
        thread.start()

        logger.info(
            f"🎙️ Transcribe-upload started: session={session_id} file='{src.name}' model={model} size={src.stat().st_size}")
        return jsonify({
            "success": True,
            "session_id": session_id,
            "stream_url": f"/api/transcribe-audio/progress/{session_id}",
            "filename": src.name,
        })
    except Exception as e:
        logger.error(
            f"❌ /api/transcribe-audio/upload failed: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


# ---------------------------------------------------------------------------
# Finished Videos Gallery
# ---------------------------------------------------------------------------
# When FINISHED_VIDEOS_BLOB_CONTAINER is set, list/stream from Azure Blob
# Storage instead of the local iCloud folder. Auth uses DefaultAzureCredential
# (system-assigned managed identity in Container Apps; az login locally) and
# returns short-lived user-delegation SAS URLs so the browser streams the
# video directly from Azure with native HTTP Range support.
FINISHED_VIDEOS_ROOT = Path(
    os.path.expanduser(
        "~/Library/Mobile Documents/com~apple~CloudDocs/Desktop/Podcast/Videos/Final"
    )
).resolve()
FINISHED_VIDEOS_BLOB_ACCOUNT = os.environ.get(
    "FINISHED_VIDEOS_BLOB_ACCOUNT", "linedrivestorage"
)
FINISHED_VIDEOS_BLOB_CONTAINER = os.environ.get(
    "FINISHED_VIDEOS_BLOB_CONTAINER", ""
).strip()
FINISHED_VIDEOS_SAS_MINUTES = int(
    os.environ.get("FINISHED_VIDEOS_SAS_MINUTES", "120")
)
_VIDEO_EXTS = {".mp4", ".mov", ".m4v", ".webm", ".mkv"}
_MIME_MAP = {
    ".mp4": "video/mp4",
    ".m4v": "video/mp4",
    ".mov": "video/quicktime",
    ".webm": "video/webm",
    ".mkv": "video/x-matroska",
}
# Cached user-delegation key (refreshed ~hourly)
_udk_cache: "dict[str, object]" = {"key": None, "expires": 0.0}


def _finished_videos_blob_service():
    from azure.identity import DefaultAzureCredential
    from azure.storage.blob import BlobServiceClient
    account_url = f"https://{FINISHED_VIDEOS_BLOB_ACCOUNT}.blob.core.windows.net"
    return BlobServiceClient(account_url=account_url, credential=DefaultAzureCredential())


def _get_user_delegation_key():
    import time as _t
    now = _t.time()
    if _udk_cache["key"] and now < float(_udk_cache["expires"]) - 300:
        return _udk_cache["key"]
    from datetime import datetime, timedelta, timezone
    svc = _finished_videos_blob_service()
    start = datetime.now(timezone.utc) - timedelta(minutes=5)
    expiry = datetime.now(timezone.utc) + timedelta(hours=2)
    key = svc.get_user_delegation_key(
        key_start_time=start, key_expiry_time=expiry)
    _udk_cache["key"] = key
    _udk_cache["expires"] = expiry.timestamp()
    return key


def _blob_sas_url(blob_name: str) -> str:
    from datetime import datetime, timedelta, timezone
    from urllib.parse import quote as _q
    from azure.storage.blob import generate_blob_sas, BlobSasPermissions
    udk = _get_user_delegation_key()
    expiry = datetime.now(timezone.utc) + \
        timedelta(minutes=FINISHED_VIDEOS_SAS_MINUTES)
    sas = generate_blob_sas(
        account_name=FINISHED_VIDEOS_BLOB_ACCOUNT,
        container_name=FINISHED_VIDEOS_BLOB_CONTAINER,
        blob_name=blob_name,
        user_delegation_key=udk,
        permission=BlobSasPermissions(read=True),
        expiry=expiry,
    )
    encoded = "/".join(_q(seg) for seg in blob_name.split("/"))
    return (
        f"https://{FINISHED_VIDEOS_BLOB_ACCOUNT}.blob.core.windows.net/"
        f"{FINISHED_VIDEOS_BLOB_CONTAINER}/{encoded}?{sas}"
    )


def _list_finished_videos_from_blob():
    """Returns (items, source_label) by listing blobs in the configured container.

    Groups by the first path segment ('folder'); picks the first video blob in
    each group. Loose root-level video blobs become their own cards.
    """
    svc = _finished_videos_blob_service()
    container = svc.get_container_client(FINISHED_VIDEOS_BLOB_CONTAINER)
    folders: "dict[str, list]" = {}
    loose: list = []
    poster_stems: set = set()
    for b in container.list_blobs():
        name = b.name
        if not name or name.endswith("/"):
            continue
        ext = ("." + name.rsplit(".", 1)[-1].lower()) if "." in name else ""
        if ext == ".jpg":
            poster_stems.add(name[:-4])
            continue
        if ext not in _VIDEO_EXTS:
            continue
        if "/" in name:
            folder = name.split("/", 1)[0]
            folders.setdefault(folder, []).append(b)
        else:
            loose.append(b)

    def _poster_for(blob_name: str):
        stem = blob_name.rsplit(".", 1)[0]
        if stem in poster_stems:
            return _blob_sas_url(stem + ".jpg")
        return None

    items = []
    for folder in sorted(folders.keys(), key=str.lower):
        blobs = sorted(folders[folder], key=lambda x: x.name.lower())
        b = blobs[0]
        size = getattr(b, "size", 0) or 0
        mtime = 0.0
        try:
            if getattr(b, "last_modified", None):
                mtime = b.last_modified.timestamp()
        except Exception:  # noqa: BLE001
            mtime = 0.0
        ext = b.name.rsplit(".", 1)[-1].lower()
        items.append({
            "title": folder,
            "filename": b.name.rsplit("/", 1)[-1],
            "rel_path": b.name,
            "ext": ext,
            "size": size,
            "mtime": mtime,
            "stream_url": _blob_sas_url(b.name),
            "poster_url": _poster_for(b.name),
        })
    for b in sorted(loose, key=lambda x: x.name.lower()):
        size = getattr(b, "size", 0) or 0
        mtime = 0.0
        try:
            if getattr(b, "last_modified", None):
                mtime = b.last_modified.timestamp()
        except Exception:  # noqa: BLE001
            mtime = 0.0
        ext = b.name.rsplit(".", 1)[-1].lower()
        stem = b.name.rsplit(".", 1)[0]
        items.append({
            "title": stem,
            "filename": b.name,
            "rel_path": b.name,
            "ext": ext,
            "size": size,
            "mtime": mtime,
            "stream_url": _blob_sas_url(b.name),
            "poster_url": _poster_for(b.name),
        })
    label = (
        f"azure://{FINISHED_VIDEOS_BLOB_ACCOUNT}/"
        f"{FINISHED_VIDEOS_BLOB_CONTAINER}"
    )
    return items, label


@app.route("/api/finished-videos", methods=["GET"])
def api_finished_videos():
    """List finished videos (Azure Blob if configured, else local iCloud folder)."""
    from urllib.parse import quote
    try:
        if FINISHED_VIDEOS_BLOB_CONTAINER:
            try:
                items, label = _list_finished_videos_from_blob()
                return jsonify({"success": True, "root": label,
                                "source": "blob", "videos": items})
            except Exception as e:  # noqa: BLE001
                logger.error(f"❌ blob listing failed: {e}", exc_info=True)
                return jsonify({
                    "success": False,
                    "error": f"Azure blob listing failed: {e}",
                    "root": (f"azure://{FINISHED_VIDEOS_BLOB_ACCOUNT}/"
                             f"{FINISHED_VIDEOS_BLOB_CONTAINER}"),
                    "source": "blob",
                    "videos": [],
                }), 500

        root = FINISHED_VIDEOS_ROOT
        if not root.exists() or not root.is_dir():
            return jsonify({
                "success": False,
                "error": f"Finished videos folder not found: {root}",
                "root": str(root),
                "source": "local",
                "videos": [],
            }), 404

        items = []
        for entry in sorted(root.iterdir(), key=lambda p: p.name.lower()):
            if entry.name.startswith("."):
                continue
            if entry.is_dir():
                vids = sorted(
                    [p for p in entry.iterdir()
                     if p.is_file() and p.suffix.lower() in _VIDEO_EXTS
                     and not p.name.startswith(".")],
                    key=lambda p: p.name.lower(),
                )
                if not vids:
                    continue
                vid = vids[0]
                rel = vid.relative_to(root).as_posix()
                try:
                    size = vid.stat().st_size
                    mtime = vid.stat().st_mtime
                except OSError:
                    size, mtime = 0, 0
                items.append({
                    "title": entry.name,
                    "filename": vid.name,
                    "rel_path": rel,
                    "ext": vid.suffix.lower().lstrip("."),
                    "size": size,
                    "mtime": mtime,
                    "stream_url": f"/api/finished-videos/file?path={quote(rel)}",
                })
            elif entry.is_file() and entry.suffix.lower() in _VIDEO_EXTS:
                rel = entry.name
                try:
                    size = entry.stat().st_size
                    mtime = entry.stat().st_mtime
                except OSError:
                    size, mtime = 0, 0
                items.append({
                    "title": entry.stem,
                    "filename": entry.name,
                    "rel_path": rel,
                    "ext": entry.suffix.lower().lstrip("."),
                    "size": size,
                    "mtime": mtime,
                    "stream_url": f"/api/finished-videos/file?path={quote(rel)}",
                })

        return jsonify({"success": True, "root": str(root),
                        "source": "local", "videos": items})
    except Exception as e:
        logger.error(f"❌ /api/finished-videos failed: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/finished-videos/file", methods=["GET"])
def api_finished_videos_file():
    """Stream a local video file (used in local mode only)."""
    try:
        rel = request.args.get("path", "").strip()
        if not rel:
            return jsonify({"success": False, "error": "Missing 'path' query param"}), 400
        root = FINISHED_VIDEOS_ROOT
        candidate = (root / rel).resolve()
        try:
            candidate.relative_to(root)
        except ValueError:
            return jsonify({"success": False, "error": "Path escapes finished-videos root"}), 400
        if not candidate.is_file():
            return jsonify({"success": False, "error": "File not found"}), 404
        if candidate.suffix.lower() not in _VIDEO_EXTS:
            return jsonify({"success": False, "error": "Not a supported video type"}), 400
        mimetype = _MIME_MAP.get(
            candidate.suffix.lower(), "application/octet-stream")
        return send_file(str(candidate), mimetype=mimetype, conditional=True)
    except Exception as e:
        logger.error(f"❌ /api/finished-videos/file failed: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/finished-videos/upload", methods=["POST"])
def api_finished_videos_upload():
    """Start a background transcode + upload to the gallery.

    multipart/form-data fields:
      - file:   the video file (any common video container)
      - folder: optional gallery card name (defaults to file stem)

    Returns immediately with {success, job_id}. Poll GET
    /api/finished-videos/upload/status/<job_id> for progress + final result.
    """
    import subprocess
    import tempfile
    import shutil
    import uuid
    import re as _re
    import threading
    import time
    import json as _json
    from werkzeug.utils import secure_filename
    from azure.storage.blob import ContentSettings

    if not FINISHED_VIDEOS_BLOB_CONTAINER:
        return jsonify({
            "success": False,
            "error": "Uploads disabled: FINISHED_VIDEOS_BLOB_CONTAINER not set",
        }), 400

    if shutil.which("ffmpeg") is None:
        return jsonify({"success": False, "error": "ffmpeg not found on PATH"}), 500

    f = request.files.get("file")
    if not f or not f.filename:
        return jsonify({"success": False, "error": "No file uploaded"}), 400

    safe_name = secure_filename(f.filename) or f"upload-{uuid.uuid4().hex}.mp4"
    stem = Path(safe_name).stem or f"upload-{uuid.uuid4().hex}"

    folder_raw = (request.form.get("folder") or "").strip()
    if folder_raw:
        folder = _re.sub(r"[^A-Za-z0-9 ._-]+", "_",
                         folder_raw).strip("/. ") or stem
    else:
        folder = stem
    blob_name = f"{folder}/{stem}.mp4"

    # Persist the upload to a temp dir BEFORE returning (FileStorage dies with the request).
    tmpdir = Path(tempfile.mkdtemp(prefix="scriptcraft-upload-"))
    src_path = tmpdir / safe_name
    f.save(str(src_path))

    job_id = uuid.uuid4().hex
    with video_upload_jobs_lock:
        video_upload_jobs[job_id] = {
            "state": "queued",
            "percent": 0,
            "duration": None,
            "out_time": 0.0,
            "error": None,
            "video": None,
            "filename": safe_name,
            "blob_name": blob_name,
            "started": time.time(),
            "finished": None,
        }

    def _set(**kw):
        with video_upload_jobs_lock:
            video_upload_jobs[job_id].update(kw)

    def _worker():
        out_path = tmpdir / f"{stem}.mp4"
        ffmpeg_log = tmpdir / "ffmpeg.log"
        try:
            # 1) Probe duration so we can compute %.
            duration = None
            try:
                pr = subprocess.run(
                    ["ffprobe", "-v", "error", "-show_entries", "format=duration",
                     "-of", "default=noprint_wrappers=1:nokey=1", str(src_path)],
                    capture_output=True, text=True, timeout=60,
                )
                if pr.returncode == 0 and pr.stdout.strip():
                    duration = float(pr.stdout.strip())
                    _set(duration=duration)
            except Exception as pe:
                logger.warning(f"ffprobe failed (continuing without %): {pe}")

            # 2) Transcode. Progress key=value lines on stdout, errors to file.
            ffmpeg_cmd = [
                "ffmpeg", "-hide_banner", "-nostdin", "-loglevel", "error", "-y",
                "-i", str(src_path),
                "-vf", "scale='min(1920,iw)':'-2'",
                "-c:v", "libx264", "-preset", "veryfast", "-crf", "22",
                "-pix_fmt", "yuv420p",
                "-c:a", "aac", "-b:a", "192k",
                "-movflags", "+faststart",
                "-progress", "pipe:1", "-nostats",
                str(out_path),
            ]
            logger.info(
                f"📼 [{job_id[:8]}] transcoding {src_path.name} -> {blob_name} (dur={duration})")
            _set(state="transcoding")
            with ffmpeg_log.open("wb") as logf:
                proc = subprocess.Popen(
                    ffmpeg_cmd,
                    stdin=subprocess.DEVNULL,
                    stdout=subprocess.PIPE,
                    stderr=logf,
                    bufsize=1, text=True,
                )
                try:
                    assert proc.stdout is not None
                    for line in proc.stdout:
                        line = line.strip()
                        if not line or "=" not in line:
                            continue
                        k, _, v = line.partition("=")
                        if k == "out_time_ms":
                            try:
                                secs = int(v) / 1_000_000.0
                                pct = int((secs / duration) *
                                          100) if duration else 0
                                pct = max(0, min(99, pct))
                                _set(out_time=secs, percent=pct)
                            except Exception:
                                pass
                        elif k == "out_time_us":
                            try:
                                secs = int(v) / 1_000_000.0
                                pct = int((secs / duration) *
                                          100) if duration else 0
                                pct = max(0, min(99, pct))
                                _set(out_time=secs, percent=pct)
                            except Exception:
                                pass
                        elif k == "progress" and v == "end":
                            break
                finally:
                    rc = proc.wait()

            if rc != 0 or not out_path.exists() or out_path.stat().st_size == 0:
                try:
                    tail = ffmpeg_log.read_bytes(
                    )[-2000:].decode("utf-8", errors="replace").strip()
                except Exception:
                    tail = ""
                err = tail or f"ffmpeg exited {rc}"
                logger.error(f"❌ [{job_id[:8]}] transcode failed: {err}")
                _set(state="error", error=f"Transcode failed: {err[:500]}",
                     finished=time.time())
                return

            # 3) Generate poster (thumbnail) — single frame at ~3s (or 10% in for short clips).
            poster_path = tmpdir / f"{stem}.jpg"
            poster_blob = f"{folder}/{stem}.jpg"
            try:
                seek = 3.0
                if duration and duration > 0:
                    seek = min(3.0, max(0.5, duration * 0.1))
                pcmd = [
                    "ffmpeg", "-hide_banner", "-nostdin", "-loglevel", "error", "-y",
                    "-ss", f"{seek:.2f}", "-i", str(out_path),
                    "-vframes", "1",
                    "-vf", "scale='min(1280,iw)':'-2'",
                    "-q:v", "3",
                    str(poster_path),
                ]
                pres = subprocess.run(pcmd, stdin=subprocess.DEVNULL,
                                      capture_output=True, timeout=120)
                if pres.returncode != 0 or not poster_path.exists() or poster_path.stat().st_size == 0:
                    logger.warning(
                        f"⚠️ [{job_id[:8]}] poster generation failed (rc={pres.returncode}); continuing without thumbnail")
                    poster_path = None
            except Exception as pe:
                logger.warning(
                    f"⚠️ [{job_id[:8]}] poster generation error: {pe}")
                poster_path = None

            # 4) Upload to blob.
            _set(state="uploading", percent=99)
            svc = _finished_videos_blob_service()
            container = svc.get_container_client(
                FINISHED_VIDEOS_BLOB_CONTAINER)
            with out_path.open("rb") as fh:
                container.upload_blob(
                    name=blob_name,
                    data=fh,
                    overwrite=True,
                    content_settings=ContentSettings(content_type="video/mp4"),
                )
            if poster_path is not None:
                try:
                    with poster_path.open("rb") as ph:
                        container.upload_blob(
                            name=poster_blob,
                            data=ph,
                            overwrite=True,
                            content_settings=ContentSettings(
                                content_type="image/jpeg"),
                        )
                    logger.info(
                        f"🖼️ [{job_id[:8]}] uploaded poster {poster_blob}")
                except Exception as ue:
                    logger.warning(
                        f"⚠️ [{job_id[:8]}] poster upload failed: {ue}")
            size = out_path.stat().st_size
            logger.info(
                f"✅ [{job_id[:8]}] uploaded {blob_name} ({size} bytes)")
            _set(
                state="done",
                percent=100,
                finished=time.time(),
                video={
                    "title": folder,
                    "filename": f"{stem}.mp4",
                    "rel_path": blob_name,
                    "ext": "mp4",
                    "size": size,
                    "stream_url": _blob_sas_url(blob_name),
                    "poster_url": _blob_sas_url(poster_blob) if poster_path is not None else None,
                },
            )
        except Exception as e:
            logger.error(
                f"❌ [{job_id[:8]}] upload worker failed: {e}", exc_info=True)
            _set(state="error", error=str(e), finished=time.time())
        finally:
            try:
                shutil.rmtree(tmpdir, ignore_errors=True)
            except Exception:
                pass

    threading.Thread(
        target=_worker, name=f"video-upload-{job_id[:8]}", daemon=True).start()
    return jsonify({"success": True, "job_id": job_id, "filename": safe_name})


@app.route("/api/finished-videos/upload/status/<job_id>", methods=["GET"])
def api_finished_videos_upload_status(job_id: str):
    with video_upload_jobs_lock:
        job = video_upload_jobs.get(job_id)
        if not job:
            return jsonify({"success": False, "error": "Unknown job_id"}), 404
        # Return a shallow copy so the caller doesn't see live mutations.
        snap = dict(job)
    return jsonify({"success": True, "job": snap})


@app.route("/api/finished-videos/regenerate-posters", methods=["POST"])
def api_finished_videos_regenerate_posters():
    """Backfill poster JPGs for any video blobs that don't already have one.

    Streams the source video into ffmpeg via a temp file, extracts a single
    frame at ~3s, and uploads as `<stem>.jpg`. Skips blobs whose poster
    already exists. Returns a summary list.
    """
    import subprocess
    import tempfile
    import shutil as _shutil
    from azure.storage.blob import ContentSettings

    if not FINISHED_VIDEOS_BLOB_CONTAINER:
        return jsonify({"success": False, "error": "Blob container not configured"}), 400
    if _shutil.which("ffmpeg") is None:
        return jsonify({"success": False, "error": "ffmpeg not found on PATH"}), 500

    svc = _finished_videos_blob_service()
    container = svc.get_container_client(FINISHED_VIDEOS_BLOB_CONTAINER)
    existing_posters: set = set()
    videos: list = []
    for b in container.list_blobs():
        n = b.name or ""
        if n.endswith(".jpg"):
            existing_posters.add(n[:-4])
        else:
            ext = n.rsplit(".", 1)[-1].lower() if "." in n else ""
            if ("." + ext) in _VIDEO_EXTS:
                videos.append(n)

    created, skipped, failed = [], [], []
    tmp_root = Path(tempfile.mkdtemp(prefix="poster-backfill-"))
    try:
        for vname in videos:
            stem = vname.rsplit(".", 1)[0]
            if stem in existing_posters:
                skipped.append(vname)
                continue
            local_vid = tmp_root / Path(vname).name
            try:
                with local_vid.open("wb") as fh:
                    container.download_blob(vname).readinto(fh)
                poster_local = tmp_root / (Path(vname).stem + ".jpg")
                pcmd = [
                    "ffmpeg", "-hide_banner", "-nostdin", "-loglevel", "error", "-y",
                    "-ss", "3", "-i", str(local_vid),
                    "-vframes", "1",
                    "-vf", "scale='min(1280,iw)':'-2'",
                    "-q:v", "3",
                    str(poster_local),
                ]
                pres = subprocess.run(pcmd, stdin=subprocess.DEVNULL,
                                      capture_output=True, timeout=180)
                if pres.returncode != 0 or not poster_local.exists() or poster_local.stat().st_size == 0:
                    failed.append({"blob": vname, "error": "ffmpeg failed"})
                    continue
                with poster_local.open("rb") as ph:
                    container.upload_blob(
                        name=stem + ".jpg", data=ph, overwrite=True,
                        content_settings=ContentSettings(
                            content_type="image/jpeg"),
                    )
                created.append(stem + ".jpg")
                logger.info(f"🖼️ backfilled poster for {vname}")
            except Exception as e:
                logger.error(
                    f"❌ poster backfill failed for {vname}: {e}", exc_info=True)
                failed.append({"blob": vname, "error": str(e)})
            finally:
                try:
                    if local_vid.exists():
                        local_vid.unlink()
                except Exception:
                    pass
    finally:
        _shutil.rmtree(tmp_root, ignore_errors=True)

    return jsonify({
        "success": True,
        "created": created,
        "skipped": len(skipped),
        "failed": failed,
        "total_videos": len(videos),
    })


@app.route("/api/transcribe-audio", methods=["POST"])
def api_transcribe_audio():
    """Start a local-whisper transcription in the background.

    Body JSON:
        { "file_path": "/abs/path/to/audio.wav",
          "model":     "medium" (optional),
          "language":  "en"     (optional) }

    Returns: { success, session_id, stream_url }
    Stream events on /api/transcribe-audio/progress/<session_id>:
      {type:"progress", message:"...", line:"..."}
      {type:"segment",  start:"00:00.000", end:"00:05.000", text:"..."}
      {type:"done",     text:"<full transcript>", model:"medium"}
      {type:"error",    error:"..."}
    """
    try:
        data = request.get_json(silent=True) or {}
        file_path = (data.get("file_path") or "").strip()
        model = (data.get("model") or "medium").strip() or "medium"
        language = (data.get("language") or "").strip()

        if not file_path:
            return jsonify({"success": False, "error": "file_path is required"}), 400

        src = Path(file_path).expanduser()
        if not src.exists() or not src.is_file():
            return jsonify({"success": False, "error": f"File not found: {src}"}), 404

        whisper_bin = _find_whisper_cli()
        if not whisper_bin:
            return jsonify({
                "success": False,
                "error": "whisper CLI not found. Install with: pip install openai-whisper",
            }), 500

        session_id = str(uuid.uuid4())
        transcribe_streams[session_id] = _queue.Queue()
        transcribe_results.pop(session_id, None)

        thread = _threading.Thread(
            target=_transcribe_worker,
            args=(session_id, whisper_bin, str(src), model, language),
            daemon=True,
        )
        thread.start()

        logger.info(
            f"🎙️ Transcribe started: session={session_id} file='{src.name}' model={model}")
        return jsonify({
            "success": True,
            "session_id": session_id,
            "stream_url": f"/api/transcribe-audio/progress/{session_id}",
        })
    except Exception as e:
        logger.error(
            f"❌ /api/transcribe-audio failed to start: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


def _transcribe_worker(session_id: str, whisper_bin: str, file_path: str,
                       model: str, language: str):
    """Run whisper CLI as a subprocess and stream stdout lines as SSE events."""
    import subprocess
    import tempfile
    import shutil as _shutil
    import re as _re
    import time as _time

    q = transcribe_streams.get(session_id)

    def push(payload):
        if q is not None:
            try:
                q.put(payload)
            except Exception:
                pass

    out_dir = Path(tempfile.mkdtemp(prefix="whisper_out_"))
    cmd = [
        whisper_bin,
        file_path,
        "--model", model,
        "--output_format", "txt",
        "--output_dir", str(out_dir),
        # --verbose True (default) prints "[mm:ss.xxx --> mm:ss.xxx]  text" per segment
        "--verbose", "True",
    ]
    if language:
        cmd += ["--language", language]

    push({"type": "progress",
         "message": f"🎙️ Loading whisper model '{model}' (first run downloads weights)…"})
    logger.info(f"🎙️ whisper cmd: {' '.join(cmd)}")

    seg_re = _re.compile(r"^\[(\d+:\d+\.\d+)\s*-->\s*(\d+:\d+\.\d+)\]\s*(.*)$")

    def _to_hhmmss(ts: str) -> str:
        # ts is "MM:SS.mmm" or "HH:MM:SS.mmm" (whisper uses minutes:seconds when <1h)
        try:
            base = ts.split(".")[0]
            parts = base.split(":")
            if len(parts) == 2:
                h, m, s = 0, int(parts[0]), int(parts[1])
            elif len(parts) == 3:
                h, m, s = int(parts[0]), int(parts[1]), int(parts[2])
            else:
                return ts
            # normalize overflow
            m += s // 60
            s = s % 60
            h += m // 60
            m = m % 60
            return f"{h:02d}:{m:02d}:{s:02d}"
        except Exception:
            return ts

    collected_segments: list = []  # list of (start_hhmmss, text)

    env = dict(os.environ)
    env["PYTHONUNBUFFERED"] = "1"
    env["PYTHONIOENCODING"] = "utf-8"
    # Force UTF-8 locale so the child's stdio encoding (and our pipe decode)
    # don't fall back to ASCII and crash on tqdm's non-ASCII progress chars.
    env.setdefault("LC_ALL", "en_US.UTF-8")
    env.setdefault("LANG", "en_US.UTF-8")

    try:
        proc = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            encoding="utf-8",
            errors="replace",
            bufsize=0,  # unbuffered — we'll segment manually on \n and \r
            env=env,
        )
    except Exception as e:
        push({"type": "error", "error": f"failed to launch whisper: {e}"})
        _shutil.rmtree(out_dir, ignore_errors=True)
        return

    # Heartbeat thread — keeps the UI alive while whisper loads the model silently
    stop_heartbeat = _threading.Event()

    def _heartbeat():
        secs = 0
        while not stop_heartbeat.wait(5):
            secs += 5
            push({"type": "progress", "message": f"…still working ({secs}s elapsed)"})

    hb_thread = _threading.Thread(target=_heartbeat, daemon=True)
    hb_thread.start()

    try:
        assert proc.stdout is not None
        # Read char-by-char so we capture both \n (segment lines) and \r (tqdm updates)
        buf = []
        first_output_seen = False
        while True:
            ch = proc.stdout.read(1)
            if ch == "" and proc.poll() is not None:
                break
            if ch in ("\n", "\r"):
                line = "".join(buf).strip()
                buf = []
                if not line:
                    continue
                if not first_output_seen:
                    first_output_seen = True
                    # First real output — model is loaded
                    stop_heartbeat.set()
                m = seg_re.match(line)
                if m:
                    start_hms = _to_hhmmss(m.group(1))
                    end_hms = _to_hhmmss(m.group(2))
                    seg_text = m.group(3).strip()
                    collected_segments.append((start_hms, seg_text))
                    push({
                        "type": "segment",
                        "start": start_hms,
                        "end": end_hms,
                        "text": seg_text,
                        "line": line,
                    })
                else:
                    push({"type": "progress", "message": line, "line": line})
            else:
                buf.append(ch)

        # Drain any trailing buffered text
        if buf:
            tail = "".join(buf).strip()
            if tail:
                push({"type": "progress", "message": tail, "line": tail})

        stop_heartbeat.set()
        rc = proc.wait()
        if rc != 0:
            push({"type": "error", "error": f"whisper exited with code {rc}"})
            _shutil.rmtree(out_dir, ignore_errors=True)
            return

        txt_files = sorted(out_dir.glob("*.txt"))
        if not txt_files:
            push({"type": "error", "error": "whisper produced no .txt output"})
            _shutil.rmtree(out_dir, ignore_errors=True)
            return

        plain_text = txt_files[0].read_text(
            encoding="utf-8", errors="replace").strip()
        # Build a timestamped transcript in the format the YouTube agent expects:
        #   [HH:MM:SS] spoken text
        if collected_segments:
            text = "\n".join(f"[{ts}] {seg}" for ts,
                             seg in collected_segments if seg).strip()
        else:
            text = plain_text
        transcribe_results[session_id] = {
            "success": True, "text": text, "model": model}
        push({"type": "done", "text": text, "model": model,
             "segments": len(collected_segments)})
    except Exception as e:
        logger.error(f"❌ transcribe worker crash: {e}", exc_info=True)
        push({"type": "error", "error": str(e)})
    finally:
        stop_heartbeat.set()
        _shutil.rmtree(out_dir, ignore_errors=True)


@app.route("/api/transcribe-audio/progress/<session_id>")
def api_transcribe_audio_progress(session_id):
    """SSE stream of transcription events."""
    q = transcribe_streams.get(session_id)
    if q is None:
        return "Session not found", 404

    def generate():
        try:
            while True:
                try:
                    payload = q.get(timeout=60)
                except _queue.Empty:
                    yield ": keepalive\n\n"
                    continue
                yield f"data: {json.dumps(payload)}\n\n"
                if payload.get("type") in ("done", "error"):
                    break
        finally:
            transcribe_streams.pop(session_id, None)

    return Response(generate(), mimetype="text/event-stream", headers={
        "Cache-Control": "no-cache",
        "X-Accel-Buffering": "no",
        "Connection": "keep-alive",
    })


if __name__ == "__main__":
    # Use environment variable for port, default to 8080 for local, 5007 for container
    port = int(os.environ.get("PORT", "8080"))

    logger.info("🎬 Starting ScriptCraft Web GUI (Universal Version)...")
    logger.info("📦 Version: %s", VERSION)
    logger.info("🚀 Server starting on http://0.0.0.0:%d", port)

    if port == 5007:
        logger.info("🐳 Running in Azure Container mode")
    else:
        logger.info("💻 Running in Local mode")

    app.run(host="0.0.0.0", port=port, debug=False)
