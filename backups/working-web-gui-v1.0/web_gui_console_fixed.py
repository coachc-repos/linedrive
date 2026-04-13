#!/usr/bin/env python3
"""
ScriptCraft Web GUI - Clean Console Capture Version
Fixed implementation that displays real-time EnhancedAutoGenSystem messages in browser
"""

from flask import Flask, render_template, request, jsonify, Response, make_response
from pathlib import Path
import logging
import threading
import asyncio
import uuid
import queue
import time
import json
import sys
import re


VERSION = "15.18-console-capture-fix"


# Set up logging
logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Add paths for imports
current_dir = Path(__file__).parent
linedrive_dir = current_dir.parent
sys.path.insert(0, str(linedrive_dir))

# Import text processing functions

app = Flask(__name__)
progress_streams = {}
results = {}
running_tasks = {}


class ConsoleCapture:
    """Captures console output and redirects to progress streamer"""

    def __init__(self, streamer):
        self.streamer = streamer
        self.original_stdout = sys.stdout
        self.buffer = []
        self.long_operation_start = None
        self.last_chapter_revision = None

    def write(self, message):
        """Write to both original stdout and the progress streamer"""
        if not message.strip():
            return

        # Debug: Add marker to identify when write method is called
        if "TEST 5" in message or "Minimal rapid" in message:
            self.original_stdout.write(f"🔍 WRITE DEBUG: {message.strip()}\n")
            self.original_stdout.flush()

        # Write to original stdout first 
        self.original_stdout.write(message)
        self.original_stdout.flush()

        # Filter Azure SDK noise that can cause blocking
        if self._should_filter_azure_sdk_logging(message):
            return

        # Extract progress and send to web interface 
        progress = self._extract_progress(message)
        if progress is not None:
            if self.streamer:
                try:
                    self.streamer.send_update(message.strip(), progress)
                except Exception as e:
                    self.original_stdout.write(f"⚠️ Stream update failed: {e}\n")
                    self.original_stdout.flush()
        else:
            # Send regular message without progress
            if self.streamer:
                try:
                    self.streamer.send_update(message.strip())
                except Exception as e:
                    self.original_stdout.write(f"⚠️ Stream message failed: {e}\n")
                    self.original_stdout.flush()

    def _track_long_operations(self, message):
        """Track when long operations start for timing purposes"""
        import time
        
        # Track when chapter revision starts
        if "Reviewing Chapter" in message and "revision failed" not in message:
            chapter_match = re.search(r'Chapter (\d+)', message)
            if chapter_match:
                chapter_num = int(chapter_match.group(1))
                self.last_chapter_revision = {
                    'chapter': chapter_num,
                    'start_time': time.time()
                }
                
        # Track other long operations
        elif any(phrase in message for phrase in [
            "Script Reviewer", "Script review", "Reviewing", "Revising"
        ]) and "failed" not in message and "completed" not in message:
            if self.long_operation_start is None:
                self.long_operation_start = time.time()

    def _enhance_message_with_timing(self, primary_text, fallback_text):
        """Add timing information to progress messages"""
        import time
        
        msg = primary_text if len(primary_text) > 10 else fallback_text
        
        # Add elapsed time for chapter revisions
        if self.last_chapter_revision and "revised" in msg:
            elapsed = time.time() - self.last_chapter_revision['start_time']
            if elapsed > 30:  # Only show timing after 30 seconds
                chapter_num = self.last_chapter_revision['chapter']
                mins, secs = divmod(int(elapsed), 60)
                if mins > 0:
                    msg = f"{msg} (Chapter {chapter_num} took {mins}m {secs}s)"
                else:
                    msg = f"{msg} (Chapter {chapter_num} took {secs}s)"
                    
        # Add elapsed time for other long operations
        elif self.long_operation_start and any(phrase in msg for phrase in [
            "reviewing", "revising", "processing"
        ]):
            elapsed = time.time() - self.long_operation_start
            if elapsed > 60:  # Show timing after 1 minute
                mins, secs = divmod(int(elapsed), 60)
                msg = f"{msg} (running {mins}m {secs}s)"
        
        return msg

    def _should_filter_initialization_message(self, message):
        """Filter out initialization messages not relevant to script work"""
        initialization_patterns = [
            "TournamentAgentClient initialized",
            "TournamentAgent model client created",
            "AITipsAgentClient initialized",
            "AITipsAgent model client created",
            "📡 AUTOGEN SYSTEM TRACE: Initializing agents",
            "🤖 AUTOGEN SYSTEM TRACE: Creating model clients",
            "🎯 AUTOGEN SYSTEM TRACE: System fully initialized"
        ]

        return any(pattern in message for pattern in initialization_patterns)

    def _should_filter_azure_sdk_logging(self, message):
        """Filter out Azure AI SDK internal logging that can block operations"""
        # Only filter very specific Azure SDK internal messages that cause blocking
        blocking_patterns = [
            "INFO - ManagedIdentityCredential will use IMDS",
            "INFO - Request URL:",
            "Request method:",
            "Request headers:",
            "No body was attached to the request",
            "A body is sent with the request", 
            "INFO - Response status:",
            "Response headers:",
            "INFO - DefaultAzureCredential acquired a token",
            "User-Agent': 'azsdk-python-identity",
            "User-Agent': 'AIProjectClient azsdk-python-ai-agents",
            "Authorization': 'REDACTED'",
            "x-ms-client-request-id':",
            "Content-Type': 'application/json'",
            "openai-processing-ms':",
            "Date': 'Fri, 26 Sep 2025",
            "No environment configuration found."
        ]

        # Keep important workflow messages even if they contain Azure terms
        important_workflow_patterns = [
            "STEP 1:", "STEP 2:", "STEP 3:", "STEP 4:",
            "Topic enhanced", "Chapter", "Script", "Writing", "Reviewing", "Revising",
            "completed", "revised", "Assembling", "WORKFLOW", "DEBUG:"
        ]
        
        # Don't filter if it's an important workflow message
        if any(pattern in message for pattern in important_workflow_patterns):
            return False
            
        return any(pattern in message for pattern in blocking_patterns)

    def _extract_progress(self, message):
        """Extract progress percentage based on console message patterns"""
        import re

        # Initialization messages (5-15%)
        if "Creating EnhancedAutoGenSystem" in message:
            return 10
        elif "All 5 agents initialized" in message or "initialized successfully" in message:
            return 15

        # Topic planning phase (20-35%)
        if "STEP 1: Topic Enhancement" in message or "Topic Assistant" in message:
            return 25
        elif "Topic enhanced" in message and "chars" in message:
            return 35

        # Chapter writing phase (35-85%)
        elif "STEP 2: Chapter-by-Chapter Script Writing" in message:
            return 40
        elif "Writing Chapter" in message:
            # Extract chapter numbers like "Writing Chapter 5/7"
            match = re.search(r'Chapter (\d+)/(\d+)', message)
            if match:
                current, total = int(match.group(1)), int(match.group(2))
                progress = 45 + int(current * 40 / total)
                return min(progress, 85)
            return 50
        elif "completed" in message and "chars" in message:
            # "✅ Chapter 4 completed (3773 chars)"
            match = re.search(r'Chapter (\d+)', message)
            if match:
                chapter_num = int(match.group(1))
                progress = 45 + int(chapter_num * 40 / 7)  # Assume 7 chapters
                return min(progress, 85)
            return 70

        # Review phase (85-98%)
        elif "STEP 3: Script Review" in message or "Reviewing" in message:
            return 90
        elif "revised" in message and "chars" in message:
            # "✅ Chapter 5 revised (3444 chars)" - but don't max out at 94% for final chapter
            match = re.search(r'Chapter (\d+)', message)
            if match:
                chapter_num = int(match.group(1))
                if chapter_num >= 7:  # Final chapter
                    return 96  # Leave room for final steps
                else:
                    return 93
            return 93
        elif "Assembling" in message and "chapters" in message:
            # "🔗 Assembling 7 revised chapters..."
            return 97

        # Post-processing steps (97-99%)
        elif "Script enhanced with bold tool formatting" in message:
            return 98
        elif "Extracted" in message and "tool references" in message:
            return 99
        elif "Sending completion message" in message:
            return 99

        # Completion messages
        elif "SEQUENTIAL WORKFLOW COMPLETED" in message:
            return 99
        elif "Script creation complete" in message:
            return 100

        return None

    def flush(self):
        self.original_stdout.flush()

    def fileno(self):
        """Return file descriptor for compatibility"""
        return self.original_stdout.fileno()

    def isatty(self):
        """Return whether this is a terminal"""
        return self.original_stdout.isatty()

    def readable(self):
        """Return whether file is readable"""
        return False

    def writable(self):
        """Return whether file is writable"""
        return True

    def seekable(self):
        """Return whether file is seekable"""
        return False

    def __getattr__(self, name):
        """Delegate unknown attributes to original stdout"""
        return getattr(self.original_stdout, name)


class ProgressStreamer:
    """Handles Server-Sent Events progress streaming"""

    def __init__(self, session_id: str):
        self.session_id = session_id
        self.queue = queue.Queue()
        self.result = None
        self.done = False

    def send_update(self, message: str, progress: int = None):
        """Send a progress update"""
        if progress is None:
            progress = 50

        update = {
            "message": message,
            "progress": progress,
            "timestamp": time.time(),
            "done": progress == 100
        }

        try:
            logger.info(f"🔍 send_update: putting message in queue for session {self.session_id}")
            self.queue.put(json.dumps(update))
            logger.info(f"🔍 send_update: message queued successfully")
            if progress == 100:
                logger.info(f"🔍 send_update: marking done=True for session {self.session_id}")
                self.done = True
            logger.info(f"🔍 send_update: method completed successfully")
        except Exception as e:
            logger.error(f"Failed to send update: {e}")


async def process_script_creation(session_id, topic, audience,
                                  video_length, production_type, goals):
    """Clean script creation with only console capture"""
    logger.info(f"🎬 SCRIPT CREATION STARTED: session={session_id}")

    try:
        if session_id not in progress_streams:
            logger.error(f"❌ Session {session_id} not found!")
            return

        streamer = progress_streams[session_id]

        # Initial progress
        streamer.send_update("🚀 Initializing script creation system...", 5)

        # Import the system
        from linedrive_azure.agents.enhanced_autogen_system import (
            EnhancedAutoGenSystem
        )

        # Set up console capture
        console_capture = ConsoleCapture(streamer)
        original_stdout = sys.stdout

        # Start a background thread to provide periodic progress updates during long waits
        def progress_updater():
            import time
            last_update = time.time()
            while not streamer.done:
                time.sleep(30)  # Check every 30 seconds
                current_time = time.time()
                
                # If no progress update in the last 2 minutes, send a "still working" message
                if current_time - last_update > 120 and not streamer.done:
                    elapsed_mins = int((current_time - last_update) / 60)
                    streamer.send_update(
                        f"⏳ Still processing... ({elapsed_mins}+ minutes elapsed)", 
                        None  # Don't change progress percentage
                    )
                    last_update = current_time
                    
        progress_thread = threading.Thread(target=progress_updater, daemon=True)
        progress_thread.start()

        try:
            # Redirect stdout to capture console output
            sys.stdout = console_capture

            # Send a test message to verify streaming works
            streamer.send_update(
                "🧪 Console capture test - starting script creation", 5)

            # Initialize system (this will be captured)
            print("🔧 Creating EnhancedAutoGenSystem...")
            system = EnhancedAutoGenSystem(verbose=True)
            print("🔧 All 5 agents initialized successfully")

            # Run the workflow with timeout to prevent hanging
            print("🚀 Starting script workflow with 15-minute timeout...")
            
            # DEBUG: Add detailed logging for hang investigation
            print("🔍 DEBUG: About to call async workflow method")
            print(f"🔍 DEBUG: Current thread: {threading.current_thread().name}")
            print(f"🔍 DEBUG: ConsoleCapture active: {isinstance(sys.stdout, ConsoleCapture)}")
            
            try:
                print("🔍 DEBUG: Calling asyncio.wait_for with async method")
                result = await asyncio.wait_for(
                    system.run_complete_script_workflow_sequential(
                        script_topic=topic,
                        topic_description="",
                        audience=audience,
                        tone="conversational and educational",
                        script_length=video_length,
                    ),
                    timeout=900  # 15 minutes timeout
                )
                print("🔍 DEBUG: Async workflow completed successfully")
                    
            except asyncio.TimeoutError:
                print("🔍 DEBUG: AsyncIO timeout occurred")
                print("⚠️ Workflow timeout - using fallback approach")
                streamer.send_update(
                    "⚠️ Script creation taking longer than expected, "
                    "trying alternative approach...", 85
                )
                
                # Fallback: Create a simpler script without revision
                result = {
                    "success": False,
                    "error": "Workflow timeout - chapter revision taking too long",
                    "timeout": True
                }
            except asyncio.TimeoutError:
                print("⚠️ Workflow timeout - using fallback approach")
                streamer.send_update(
                    "⚠️ Script creation taking longer than expected, "
                    "trying alternative approach...", 85
                )
                
                # Fallback: Create a simpler script without revision
                result = {
                    "success": False,
                    "error": "Workflow timeout - chapter revision taking too long",
                    "timeout": True
                }

        finally:
            # Always restore stdout
            sys.stdout = original_stdout

        # Debug: Check what we actually got
        print(f"🔍 DEBUG: Result type: {type(result)}")
        if isinstance(result, dict):
            print(f"🔍 DEBUG: Result keys: {list(result.keys())}")
        else:
            print(f"🔍 DEBUG: Result value: {result}")
            # Convert string result to error format
            result = {"success": False, "error": str(result)}

        if result.get("success"):
            # EXACT COPY OF WORKING CONSOLE UI LOGIC
            print(f"\n🎉 COMPLETE 4-AGENT WORKFLOW FINISHED!")
            print("=" * 60)
            print(f"📋 Topic: {topic}")
            print(f"👥 Audience: {audience}")
            print(f"💬 Tone: conversational and educational")
            print(f"⏱️ Length: {video_length}")
            print("-" * 60)
            print(f"✅ Topic Enhanced by Topic Assistant")
            print(f"✅ Script Created by Script Writer")
            print(f"✅ Script Reviewed by Script Reviewer")
            print(f"✅ Sequential Workflow Completed Successfully")

            # Add introduction to script content (EXACT COPY)
            introduction = """Welcome to AI with Roz, I'm Roz's AI Clone. She is a high-tech sales leader, bonafide tech nerd and busy Mom of two really incredible kids. She has spent her career in high-tech working to this moment and beyond and we are here to guide you through it.

---

"""

            # Combine introduction with script content (EXACT COPY)
            raw_script_content = introduction + result["script_content"]

            # Enhanced script with bold tool formatting (EXACT COPY - NO TIMEOUTS!)
            from console_ui.text_processing import (
                enhance_script_with_bold_tools,
                extract_tool_links_and_info,
            )

            enhanced_script_content = enhance_script_with_bold_tools(
                raw_script_content)
            print("✅ Script enhanced with bold tool formatting")
            # Force progress update for enhancement step
            streamer.send_update("✅ Script enhanced with formatting", 98)

            # Use enhanced script as final content (EXACT COPY)
            final_script_content = enhanced_script_content

            # Extract tool links for YouTube description (EXACT COPY - NO TIMEOUTS!)
            tool_links = extract_tool_links_and_info(final_script_content)
            tool_count = len(tool_links.splitlines())
            print(f"📊 Extracted {tool_count} tool references")
            # Force progress update for tool extraction
            streamer.send_update(
                f"📊 Extracted {tool_count} tool references", 99)

            # Add tool links to the script content for YouTube description (EXACT COPY)
            final_script_with_tools = final_script_content + "\n\n" + "=" * 60 + "\n"
            final_script_with_tools += "� YOUTUBE VIDEO DESCRIPTION\n"
            final_script_with_tools += "=" * 60 + "\n"
            final_script_with_tools += (
                "Copy the section below for your YouTube video description:\n\n"
            )
            final_script_with_tools += f"🎥 {topic}\n\n"
            final_script_with_tools += tool_links
            final_script_with_tools += "\n\n🔔 Don't forget to SUBSCRIBE for more AI tools and productivity tips!"
            final_script_with_tools += "\n� What tools would you like to see featured next? Drop a comment below!"

            # Store result (SIMPLIFIED)
            streamer.result = {
                "success": True,
                "enhanced_script": final_script_with_tools,
                "word_count": len(final_script_with_tools.split()),
                "reading_time": f"{len(final_script_with_tools.split()) // 150} min",
                "audience": audience,
                "production_type": production_type,
            }

            # Send final completion message (SIMPLE)
            print("🎯 Sending completion message...")
            print(f"DEBUG: About to send completion for session {session_id}")
            print(
                f"DEBUG: Streamer done status before completion: {streamer.done}")

            # Send the completion message
            streamer.send_update("✅ Script creation complete!", 100)

            print("✅ Completion message sent successfully")
            print(
                f"DEBUG: Streamer done status after completion: {streamer.done}")
            print("🎉 Script creation completed successfully!")

        else:
            error_msg = result.get("error", "Unknown workflow error")
            streamer.send_update(f"❌ Script creation failed: {error_msg}", -1)
            streamer.result = {"success": False, "error": error_msg}
            print(f"❌ Script creation failed: {error_msg}")
            logger.error(
                f"❌ Script creation failed for session {session_id}: {error_msg}")

    except Exception as e:
        error_msg = f"Error: {str(e)}"
        logger.error(f"💥 Exception during script creation: {error_msg}")
        if session_id in progress_streams:
            streamer = progress_streams[session_id]
            streamer.send_update(f"❌ {error_msg}", -1)
            streamer.result = {"success": False, "error": error_msg}


async def test_mode_simulation(session_id):
    """Quick test mode - simulates script creation in 10 seconds"""
    import time
    import asyncio

    logger.info(f"⚡ Starting test mode simulation for {session_id}")

    if session_id not in progress_streams:
        logger.error(f"❌ No progress stream for session {session_id}")
        return

    streamer = progress_streams[session_id]

    try:
        # Simulate the 7-chapter workflow with fast updates
        test_chapters = [
            "Introduction: Setting the Stage",
            "Chapter 1: The Problem",
            "Chapter 2: Understanding the Context",
            "Chapter 3: Exploring Solutions",
            "Chapter 4: Implementation Strategy",
            "Chapter 5: Real-World Applications",
            "Conclusion: Looking Forward"
        ]

        # Send initial message
        streamer.send_update("🚀 Starting test script creation...", 5)
        await asyncio.sleep(1)

        # Simulate writing each chapter (1 second per chapter)
        for i, chapter in enumerate(test_chapters):
            progress = 10 + (i * 12)  # Distribute progress from 10-94%
            streamer.send_update(f"📝 Writing {chapter}...", progress)
            await asyncio.sleep(1)

            # Show completion of chapter
            char_count = 1000 + (i * 200)  # Fake character counts
            streamer.send_update(f"✅ {chapter} completed ({char_count} chars)",
                                 progress + 2)
            await asyncio.sleep(0.5)

        # Final processing steps
        streamer.send_update("✨ Enhancing script formatting...", 96)
        await asyncio.sleep(1)

        streamer.send_update("🔍 Adding tool recommendations...", 98)
        await asyncio.sleep(1)

        # Create a simple test result
        test_script = """Welcome to AI with Roz, I'm Roz's AI Clone. 

This is a test script created in fast mode for debugging purposes.

Chapter 1: The Problem
This chapter discusses the main challenges...

Chapter 2: Understanding the Context  
Here we explore the broader implications...

[Continue with remaining chapters...]

===============================
TOOLS
===============================

• OpenAI GPT-4 - AI writing assistant
• Canva - Design and presentation tools
• YouTube Studio - Video optimization"""

        # Store test result
        streamer.result = {
            "success": True,
            "enhanced_script": test_script,
            "word_count": len(test_script.split()),
            "reading_time": f"{len(test_script.split()) // 150} min",
            "audience": "test audience",
            "production_type": "test",
        }

        # Send completion message
        streamer.send_update("✅ Test script creation complete!", 100)
        logger.info(f"✅ Test mode completed for session {session_id}")

    except Exception as e:
        error_msg = f"Test mode error: {str(e)}"
        logger.error(f"❌ Test mode error for {session_id}: {error_msg}")
        streamer.send_update(f"❌ {error_msg}", -1)
        streamer.result = {"success": False, "error": error_msg}


@app.route("/test-sse")
def test_sse():
    """Test SSE completion behavior without running full script creation"""
    return render_template("test_sse.html")


@app.route("/test-progress/<session_id>")
def test_progress_stream(session_id):
    """Test SSE endpoint that simulates script creation completion"""
    import time

    def generate():
        try:
            # Simulate progress messages
            messages = [
                ("🚀 Starting test...", 10),
                ("📝 Processing step 1...", 30),
                ("📝 Processing step 2...", 60),
                ("✅ Almost done...", 90),
                ("✅ Test complete!", 100)
            ]

            for message, progress in messages:
                update = {
                    "message": message,
                    "progress": progress,
                    "timestamp": time.time(),
                    "done": progress == 100
                }

                logger.info(f"🔄 TEST SSE sending: {json.dumps(update)}")
                yield f"data: {json.dumps(update)}\n\n"

                if progress == 100:
                    logger.info(f"🎯 TEST SSE completion message sent")
                    # Don't send closing message - client will close connection
                    break

                time.sleep(1)  # 1 second between messages

        except Exception as e:
            logger.error(f"TEST SSE error: {e}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
        finally:
            logger.info(f"🔚 TEST SSE stream ended")

    response = Response(generate(), mimetype="text/event-stream")
    response.headers['Cache-Control'] = 'no-cache'
    response.headers['Connection'] = 'keep-alive'
    response.headers['X-Accel-Buffering'] = 'no'
    return response


@app.route("/")
def index():
    """Serve the main ScriptCraft web interface"""
    return render_template("index.html")


@app.route("/test")
def test_page():
    """Serve the test page for debugging SSE streaming"""
    return render_template("test.html")


@app.route("/test-capture")
def test_capture_page():
    """Test page to debug console capture and threading issues"""
    return """
    <!DOCTYPE html>
    <html>
    <head><title>Console Capture Test</title></head>
    <body>
        <h1>Console Capture Test Harness</h1>
        <button onclick="runTest()">Run Console Capture Test</button>
        <div id="progress"></div>
        <div id="result"></div>
        <script>
        function runTest() {
            fetch('/test-capture-process', {method: 'POST'})
            .then(r => r.json())
            .then(data => {
                document.getElementById('progress').innerHTML = 'Test started...';
                watchProgress(data.session_id);
            });
        }
        
        function watchProgress(sessionId) {
            const eventSource = new EventSource('/progress/' + sessionId);
            eventSource.onmessage = function(event) {
                const data = JSON.parse(event.data);
                document.getElementById('progress').innerHTML = 
                    `Progress: ${data.progress}% - ${data.message}`;
                if (data.done) {
                    eventSource.close();
                    document.getElementById('result').innerHTML = 'Test completed!';
                }
            };
        }
        </script>
    </body>
    </html>
    """


@app.route('/test-capture-process', methods=['POST'])
def test_capture_process():
    """Test the console capture mechanism without full workflow"""
    session_id = str(uuid.uuid4())
    progress_streams[session_id] = ProgressStreamer(session_id)
    
    # Start the test in background
    thread = threading.Thread(
        target=run_test_capture,
        args=(session_id,),
        daemon=True
    )
    thread.start()
    running_tasks[session_id] = thread
    
    return jsonify({'session_id': session_id, 'status': 'started'})


def run_test_capture(session_id):
    """Test console capture and the exact point where it hangs"""
    import time
    import sys
    
    logger.info(f"🧪 TEST CAPTURE STARTED: session={session_id}")
    
    try:
        if session_id not in progress_streams:
            logger.error(f"❌ Session {session_id} not found!")
            return

        streamer = progress_streams[session_id]
        
        # Set up console capture exactly like the real workflow
        console_capture = ConsoleCapture(streamer)
        original_stdout = sys.stdout

        try:
            # Redirect stdout to capture console output
            sys.stdout = console_capture
            
            # Test 1: Basic console output
            print("🧪 TEST 1: Basic console output")
            streamer.send_update("🧪 Basic console output test", 10)
            time.sleep(1)
            
            # Test 2: Import text processing (where it hangs)
            print("🧪 TEST 2: Importing text processing modules")
            streamer.send_update("🧪 Testing text processing import", 20)
            
            try:
                from console_ui.text_processing import (
                    enhance_script_with_bold_tools,
                    extract_tool_links_and_info,
                )
                print("✅ Text processing modules imported successfully")
                streamer.send_update("✅ Text processing import success", 30)
            except Exception as e:
                print(f"❌ Text processing import failed: {e}")
                streamer.send_update(f"❌ Text processing import failed: {e}", 30)
                raise
            
            # Test 3: Call enhance_script_with_bold_tools (where it might hang)
            print("🧪 TEST 3: Testing enhance_script_with_bold_tools")
            streamer.send_update("🧪 Testing script enhancement", 40)
            
            test_script = "This is a test script with **bold** and [ChatGPT](https://chatgpt.com)."
            
            try:
                enhanced_script = enhance_script_with_bold_tools(test_script)
                print(f"✅ Script enhanced: {len(enhanced_script)} chars")
                streamer.send_update("✅ Script enhancement success", 50)
            except Exception as e:
                print(f"❌ Script enhancement failed: {e}")
                streamer.send_update(f"❌ Script enhancement failed: {e}", 50)
                raise
            
            # Test 4: Call extract_tool_links_and_info
            print("🧪 TEST 4: Testing extract_tool_links_and_info")
            streamer.send_update("🧪 Testing tool link extraction", 60)
            
            try:
                tool_links = extract_tool_links_and_info(enhanced_script)
                print(f"✅ Tool links extracted: {len(tool_links)} chars")
                streamer.send_update("✅ Tool link extraction success", 70)
            except Exception as e:
                print(f"❌ Tool link extraction failed: {e}")
                streamer.send_update(f"❌ Tool link extraction failed: {e}", 70)
                raise
            
            # Test 5: Minimal rapid output test to isolate deadlock
            print("🧪 TEST 5: Minimal rapid output test")
            streamer.send_update("🧪 Starting minimal rapid test", 80)
            
            # Test just 3 rapid outputs with debugging
            for i in range(3):
                print(f"   📝 Before output {i+1}")
                streamer.send_update(f"🧪 Processing output {i+1}/3", 85 + i*2)
                print(f"   📝 After output {i+1}")
                time.sleep(0.5)  # Longer sleep to prevent issues
            
            print("✅ Minimal rapid test completed")
            logger.info("🔍 DEBUG: About to send minimal rapid test success update")
            streamer.send_update("✅ Minimal rapid test success", 90)
            logger.info("🔍 DEBUG: Minimal rapid test success update sent")
            
            # Final test
            print("🧪 ALL TESTS COMPLETED SUCCESSFULLY!")
            logger.info("🔍 DEBUG: About to send all tests completed update")
            streamer.send_update("🧪 All tests completed!", 100)
            logger.info("🔍 DEBUG: All tests completed update sent")
            
            # Store success result
            logger.info("🔍 DEBUG: About to set streamer.result")
            streamer.result = {
                "success": True,
                "message": "All console capture tests passed",
                "enhanced_script": enhanced_script,
                "tool_links": tool_links
            }
            logger.info("🔍 DEBUG: streamer.result set successfully")
            
        finally:
            # Always restore stdout
            logger.info("🔍 DEBUG: Entering finally block")
            sys.stdout = original_stdout
            logger.info("🔍 DEBUG: stdout restored, function should exit now")
            
    except Exception as e:
        error_msg = f"Test failed: {str(e)}"
        logger.error(f"💥 Test exception: {error_msg}")
        if session_id in progress_streams:
            streamer = progress_streams[session_id]
            streamer.send_update(f"❌ {error_msg}", -1)
            streamer.result = {"success": False, "error": error_msg}


@app.route("/test_progress/<session_id>/<int:progress>/<message>")
def test_progress(session_id, progress, message):
    """Test endpoint to manually send progress updates"""
    if session_id in progress_streams:
        streamer = progress_streams[session_id]
        streamer.send_update(message, progress)
        logger.info(f"🧪 Test progress sent: {message} -> {progress}%")
        return jsonify({
            "success": True,
            "message": f"Sent: {message} -> {progress}%"
        })
    return jsonify({
        "success": False,
        "error": "Session not found"
    }), 404


@app.route("/test_create")
def test_create():
    """Create a test session with hardcoded progress messages"""
    import uuid
    import time
    import threading

    # Generate session ID
    session_id = str(uuid.uuid4())

    # Create progress streamer
    progress_streams[session_id] = ProgressStreamer(session_id)

    logger.info(f"🧪 Starting test session: {session_id}")

    def send_test_messages():
        streamer = progress_streams[session_id]
        time.sleep(1)

        # Send a series of test messages
        test_messages = [
            (10, "🚀 Starting test..."),
            (25, "📝 Writing Chapter 1/7: Test Chapter"),
            (40, "✅ Chapter 1 completed (1234 chars)"),
            (55, "📝 Writing Chapter 2/7: Another Test Chapter"),
            (70, "✅ Chapter 2 completed (2345 chars)"),
            (85, "🔍 Reviewing content..."),
            (100, "✅ Test completed successfully!")
        ]

        for progress, message in test_messages:
            streamer.send_update(message, progress)
            logger.info(f"🧪 Sent: {message} -> {progress}%")
            time.sleep(2)  # 2 second delay between messages

    # Start test in background
    thread = threading.Thread(target=send_test_messages, daemon=True)
    running_tasks[session_id] = thread
    thread.start()

    return jsonify({
        "success": True,
        "session_id": session_id,
        "message": "Test session started",
        "test_url": f"/progress/{session_id}"
    })


@app.route("/debug/stdout")
def debug_stdout():
    """Debug endpoint to check stdout status"""
    return jsonify({
        "stdout_type": str(type(sys.stdout)),
        "is_console_capture": isinstance(sys.stdout, ConsoleCapture),
        "active_sessions": len(progress_streams),
        "running_tasks": len(running_tasks)
    })


@app.route("/create", methods=["POST"])
def create():
    """Create a new script with progress streaming"""
    # Get form data
    data = request.get_json() or {}
    topic = data.get("topic", "AI in daily life")
    audience = data.get("audience", "general")
    video_length = data.get("video_length", "medium")
    production_type = data.get("production_type", "standard")
    goals = data.get("goals", "educational")
    test_mode = data.get("test_mode", False)  # Quick test mode flag

    # Generate session ID
    session_id = str(uuid.uuid4())

    # Create progress streamer
    progress_streams[session_id] = ProgressStreamer(session_id)

    logger.info(f"🎬 Starting script creation: {session_id}")
    logger.info(f"📝 Topic: {topic}, Audience: {audience}")

    # Start script creation in background thread
    def run_script_creation():
        try:
            if test_mode:
                # Quick test mode - simulate progress and complete in 10 seconds
                logger.info(f"⚡ Running in test mode for session {session_id}")
                asyncio.run(test_mode_simulation(session_id))
            else:
                asyncio.run(process_script_creation(
                    session_id, topic, audience, video_length,
                    production_type, goals
                ))
        except Exception as e:
            logger.error(f"❌ Script creation error: {e}")
            if session_id in progress_streams:
                progress_streams[session_id].send_update(
                    f"❌ Error: {str(e)}", 100
                )

    thread = threading.Thread(target=run_script_creation, daemon=True)
    running_tasks[session_id] = thread
    thread.start()

    return jsonify({
        "success": True,
        "session_id": session_id,
        "message": "Script creation started"
    })


@app.route("/progress/<session_id>")
def progress_stream(session_id):
    """Server-Sent Events endpoint for progress updates"""
    logger.info(f"🌊 SSE connection established for session {session_id}")

    if session_id not in progress_streams:
        return "Session not found", 404

    streamer = progress_streams[session_id]

    def generate():
        try:
            while True:  # Changed from while not streamer.done
                try:
                    # Get update with timeout
                    update = streamer.queue.get(timeout=30)
                    logger.info(f"🔄 SSE sending: {update}")
                    yield f"data: {update}\n\n"

                    # Check if done AFTER sending the message
                    update_data = json.loads(update)
                    if update_data.get("done", False):
                        logger.info(
                            f"🎯 SSE completion message sent for session {session_id}")
                        # Break AFTER sending the final message
                        break

                except queue.Empty:
                    # Only check done status if queue is empty
                    if streamer.done:
                        logger.info(f"🏁 SSE ending - streamer marked done for session {session_id}")
                        break
                    # Send keepalive
                    logger.debug(f"💓 SSE keepalive for session {session_id}")
                    yield f"data: {json.dumps({'keepalive': True})}\n\n"
                    continue

        except Exception as e:
            logger.error(f"SSE error for session {session_id}: {e}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
        finally:
            logger.info(f"🔚 SSE stream ended for session {session_id}")

    response = Response(generate(), mimetype="text/event-stream")
    response.headers['Cache-Control'] = 'no-cache'
    response.headers['Connection'] = 'keep-alive'
    response.headers['X-Accel-Buffering'] = 'no'  # Disable nginx buffering
    return response


@app.route("/result/<session_id>")
def get_result(session_id):
    """Get final result"""
    logger.info(f"🔍 Result requested for session {session_id}")

    if session_id in progress_streams:
        streamer = progress_streams[session_id]
        logger.info(
            f"📊 Session {session_id} found. Result available: {bool(streamer.result)}")
        logger.info(f"📊 Session {session_id} done status: {streamer.done}")

        if streamer.result:
            logger.info(f"✅ Result found for session {session_id}")
            result_size = len(str(streamer.result)) if streamer.result else 0
            logger.info(f"📏 Result size: {result_size} chars")
            return jsonify(streamer.result)
        else:
            logger.warning(
                f"⚠️ No result available yet for session {session_id}")
            return jsonify({"error": "Result not ready yet, please try again"}), 202
    else:
        logger.error(f"❌ Session {session_id} not found in progress_streams")
        logger.info(f"📋 Available sessions: {list(progress_streams.keys())}")

    return jsonify({"error": "Result not available"}), 404


@app.route("/polish", methods=["POST"])
def polish():
    """Polish existing script content using AI - EXACT COPY OF CONSOLE UI LOGIC"""
    logger.info("=== SCRIPT POLISH REQUEST RECEIVED ===")

    try:
        data = request.json
        script_content = data.get("script", "")
        audience = data.get("audience", "general")
        production_type = data.get("production_type", "video")

        if not script_content.strip():
            logger.warning("❌ Empty script received for polishing")
            return jsonify(
                {"success": False, "error": "Please enter a script to polish"}
            )

        logger.info(
            f"📝 Polishing script for {audience} audience, {production_type} type")

        # EXACT COPY OF CONSOLE UI LOGIC - Step 1: Polish the script
        print("\n🪄 STEP 1: Script Polishing")
        print("-" * 40)
        print("🚀 Initializing script polisher agent...")

        try:
            from linedrive_azure.agents.script_polisher_agent_client import (
                ScriptPolisherAgentClient,
            )

            polisher_agent = ScriptPolisherAgentClient()
            print("✅ Script polisher agent initialized")

            print(
                "🔄 Polishing script with chapters, visual cues, and tool integration...")
            polish_result = polisher_agent.polish_script(
                raw_script=script_content,
                script_title="AI Script (Web GUI)",
                target_audience=audience,
                production_type=production_type,
                timeout=300,
            )

            if polish_result.get("success", False):
                polished_script = polish_result.get("response", "")
                print(
                    f"✅ Script polished successfully ({len(polished_script)} characters)")

                # EXACT COPY - Enhanced script with bold tool formatting
                from console_ui.text_processing import (
                    enhance_script_with_bold_tools,
                    extract_tool_links_and_info,
                )

                enhanced_script_content = enhance_script_with_bold_tools(
                    polished_script)
                print("✅ Script enhanced with bold tool formatting")

                # EXACT COPY - Extract tool links for YouTube description
                tool_links = extract_tool_links_and_info(
                    enhanced_script_content)
                print(
                    f"📊 Extracted {len(tool_links.splitlines())} tool references")

                # EXACT COPY - Add tool links to the script content for YouTube description
                final_script_with_tools = enhanced_script_content + "\n\n" + "=" * 60 + "\n"
                final_script_with_tools += "📺 YOUTUBE VIDEO DESCRIPTION\n"
                final_script_with_tools += "=" * 60 + "\n"
                final_script_with_tools += (
                    "Copy the section below for your YouTube video description:\n\n"
                )
                final_script_with_tools += tool_links
                final_script_with_tools += "\n\n🔔 Don't forget to SUBSCRIBE for more AI tools and productivity tips!"
                final_script_with_tools += "\n💬 What tools would you like to see featured next? Drop a comment below!"

                print("🎉 Script polishing completed successfully!")

                return jsonify({
                    "success": True,
                    "enhanced_script": final_script_with_tools,
                    "word_count": len(final_script_with_tools.split()),
                    "reading_time": f"{len(final_script_with_tools.split()) // 150} min",
                    "audience": audience,
                    "production_type": production_type,
                })

            else:
                error_msg = f"Script polishing failed: {polish_result.get('error', 'Unknown error')}"
                print(f"❌ {error_msg}")
                return jsonify({"success": False, "error": error_msg})

        except Exception as e:
            error_msg = f"Error initializing script polisher: {e}"
            print(f"❌ {error_msg}")
            return jsonify({"success": False, "error": error_msg})

    except Exception as e:
        error_msg = f"Script polish request failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({"success": False, "error": error_msg})


@app.route("/export_word", methods=["POST"])
def export_word():
    """Export script content to Word document"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No data provided"})

        script_content = data.get("script_content", "")
        title = data.get("title", "AI Script")

        if not script_content:
            return jsonify({"success": False, "error": "No script content to export"})

        logger.info(
            f"📄 Exporting script to Word: {len(script_content)} characters")

        # Import Word processing libraries directly
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            logger.info("✅ Word processing libraries imported successfully")
        except ImportError as e:
            logger.error(f"❌ Failed to import docx libraries: {e}")
            return jsonify({"success": False, "error": "python-docx not available"})

        # Create temporary file path
        import tempfile
        temp_dir = Path(tempfile.gettempdir())
        temp_file = temp_dir / f"script_{uuid.uuid4().hex[:8]}.docx"

        # Convert script to Word document directly
        try:
            logger.info("📝 Creating Word document...")

            # Create a new Word document
            doc = Document()

            # Process the script content line by line
            lines = script_content.split('\n')
            logger.info(f"🔍 Processing {len(lines)} lines")

            for line in lines:
                line = line.rstrip()

                # Skip empty lines
                if not line:
                    continue

                # Handle headers (lines starting with #)
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('#').strip()

                    if level == 1:  # Main title
                        paragraph = doc.add_heading(text, level=1)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif level == 2:  # Chapter/Section
                        paragraph = doc.add_heading(text, level=2)
                    else:  # Sub-sections
                        paragraph = doc.add_heading(text, level=3)
                    continue

                # Handle bold text (**text**)
                elif '**' in line:
                    paragraph = doc.add_paragraph()
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:  # Normal text
                            if part:
                                paragraph.add_run(part)
                        else:  # Bold text
                            if part:
                                run = paragraph.add_run(part)
                                run.bold = True
                    continue

                # Handle visual cues (lines with 🎬 emoji or [Visual:])
                elif ('🎬' in line or '[Visual:' in line or
                      (line.startswith('[') and ']' in line)):
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.italic = True
                    paragraph.space_before = Pt(6)
                    paragraph.space_after = Pt(6)
                    continue

                # Regular paragraph
                else:
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(line)

            # Save document to temporary file
            doc.save(str(temp_file))
            logger.info(f"💾 Word document saved to: {temp_file}")

            if not temp_file.exists():
                return jsonify({"success": False, "error": "Failed to create Word document"})

            # Read the file and return as binary response
            with open(temp_file, 'rb') as f:
                file_content = f.read()

            # Clean up temp file
            temp_file.unlink()

            logger.info(f"✅ Word export successful: {len(file_content)} bytes")

            # Return file as binary response
            response = make_response(file_content)
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename="{title}.docx"'
            return response

        except Exception as e:
            logger.error(f"❌ Word conversion failed: {e}")
            return jsonify({"success": False, "error": f"Word conversion failed: {str(e)}"})

    except Exception as e:
        error_msg = f"Export to Word failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return jsonify({"success": False, "error": error_msg})


if __name__ == "__main__":
    logger.info("🎬 Starting ScriptCraft Web GUI Clean Version...")
    logger.info("📦 Version: %s", VERSION)
    logger.info("🚀 Server starting on http://localhost:8080")

    app.run(host="0.0.0.0", port=8080, debug=False)
